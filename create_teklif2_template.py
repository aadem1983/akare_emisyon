from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def create_teklif2_template():
    """TEKLİF - 2 FİYAT.docx şablonunu oluşturur"""
    
    doc = Document()
    
    # Sayfa ayarları
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4 height
    section.page_width = Inches(8.27)    # A4 width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Başlık
    heading = doc.add_heading('ÖLÇÜM METOTLARI VE ÜCRETLENDİRME', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.runs[0]
    heading_run.font.size = Pt(14)
    heading_run.font.bold = True
    
    doc.add_paragraph()  # Boşluk
    
    # Tablo oluştur - 5 sütun
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Sütun genişliklerini ayarla (ADET sütunu daha dar)
    # Toplam genişlik: ~6.5 inch (A4 sayfa genişliği - kenar boşlukları)
    widths = [Inches(2.0), Inches(2.0), Inches(0.5), Inches(1.0), Inches(1.0)]  # ADET: 0.5 inch
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    
    # Başlık satırı
    header_cells = table.rows[0].cells
    headers = ['ÖLÇÜLECEK PARAMETRE', 'ÖLÇÜM METODU', 'ADET', 'BİRİM FİYATI', 'TOP. FİYAT']
    
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        
        # Başlık hücresi formatı
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        
        # Başlık arka plan rengi (açık gri)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9D9D9')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Örnek parametre satırları (3 adet - kod bunları dolduracak)
    for _ in range(3):
        row = table.add_row()
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]  # Sütun genişliğini uygula
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run('')
            run.font.size = Pt(10)
    
    # Özet satırları ekle
    # TOPLAM satırı
    row_toplam = table.add_row()
    row_toplam.cells[0].text = 'TOPLAM:'
    row_toplam.cells[0].merge(row_toplam.cells[3])  # İlk 4 hücreyi birleştir
    for cell in row_toplam.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run('')
        run.font.bold = True
        run.font.size = Pt(10)
    
    # İSKONTO satırı
    row_iskonto = table.add_row()
    row_iskonto.cells[0].text = 'İSKONTO (TL):'
    row_iskonto.cells[0].merge(row_iskonto.cells[3])
    for cell in row_iskonto.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run('')
        run.font.bold = True
        run.font.size = Pt(10)
    
    # TOPLAM TUTAR satırı
    row_net = table.add_row()
    row_net.cells[0].text = 'TOPLAM TUTAR (TL):'
    row_net.cells[0].merge(row_net.cells[3])
    for cell in row_net.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run('')
        run.font.bold = True
        run.font.size = Pt(10)
        
        # Arka plan rengi (açık yeşil)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D4EDDA')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Footer ekle
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # Footer paragrafını temizle
        for p in footer.paragraphs:
            p.clear()
        
        # Footer tablosu ekle (1 satır, 3 sütun)
        footer_table = footer.add_table(1, 3)
        footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Sütun genişlikleri
        widths = [Inches(2.0), Inches(3.5), Inches(2.0)]
        for idx, cell in enumerate(footer_table.rows[0].cells):
            cell.width = widths[idx]
        
        # Sol hücre: Sayı placeholder
        cell_left = footer_table.rows[0].cells[0]
        p_left = cell_left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_left = p_left.add_run("Sayı:{{TEKLIF_NO}}")
        run_left.font.bold = True
        run_left.font.color.rgb = RGBColor(0, 0, 128)
        run_left.font.size = Pt(8)
        
        # Orta hücre: Firma bilgileri
        cell_center = footer_table.rows[0].cells[1]
        p_center = cell_center.paragraphs[0]
        p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run1 = p_center.add_run("AKARE ÇEVRE LABORATUVAR VE DAN. HİZM. TİC.LTD.ŞTİ Kirazlıyalı Mah. Süleyman Demirel Cad. No:28/A")
        run1.font.bold = True
        run1.font.color.rgb = RGBColor(0, 0, 128)
        run1.font.size = Pt(8)
        
        p_center.add_run("\n")
        
        run2 = p_center.add_run("Körfez V.D 013 065 1290 Körfez-KOCAELİ")
        run2.font.bold = True
        run2.font.color.rgb = RGBColor(0, 0, 128)
        run2.font.size = Pt(8)
        
        p_center.add_run("\n")
        
        run3 = p_center.add_run("info@akarecevre.com  www.akarecevre.com")
        run3.font.bold = True
        run3.font.color.rgb = RGBColor(0, 0, 128)
        run3.font.size = Pt(8)
        
        # Sağ hücre: Form kodu ve sayfa numarası
        cell_right = footer_table.rows[0].cells[2]
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run_form = p_right.add_run("AÇ.F.102/Rev04/14.08.2025  ")
        run_form.font.bold = True
        run_form.font.color.rgb = RGBColor(0, 0, 128)
        run_form.font.size = Pt(8)
        
        # Sayfa numarası field
        run_page = p_right.add_run()
        run_page.font.bold = True
        run_page.font.color.rgb = RGBColor(0, 0, 128)
        run_page.font.size = Pt(8)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_page._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        run_page._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_page._r.append(fldChar2)
        
        run_page.add_text('/')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        run_page._r.append(fldChar3)
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = 'NUMPAGES'
        run_page._r.append(instrText2)
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run_page._r.append(fldChar4)
        
        # Tablo kenarlıklarını kaldır
        for cell in footer_table.rows[0].cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)
    
    # Dosyayı kaydet
    output_path = os.path.join('static', 'images', 'TEKLİF - 2 FİYAT.docx')
    doc.save(output_path)
    print(f"✅ Şablon oluşturuldu (footer ile): {output_path}")
    return output_path

if __name__ == '__main__':
    create_teklif2_template()
