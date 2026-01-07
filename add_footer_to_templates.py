"""Word şablonlarına footer ekler (placeholder sistemi ile)"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def add_footer_to_template(template_path):
    """Word şablonuna footer ekler"""
    try:
        doc = Document(template_path)
        
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            
            # Mevcut içeriği temizle
            for p in list(footer.paragraphs):
                p._element.getparent().remove(p._element)
            for t in list(footer.tables):
                t._element.getparent().remove(t._element)
            
            # 1 satır, 3 sütun tablo
            table = footer.add_table(1, 3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Sütun genişlikleri
            widths = [Inches(2.0), Inches(3.5), Inches(2.0)]
            for idx, cell in enumerate(table.rows[0].cells):
                cell.width = widths[idx]
            
            # Sol hücre: Sayı placeholder
            cell_left = table.rows[0].cells[0]
            p_left = cell_left.paragraphs[0]
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_left = p_left.add_run("Sayı:{{TEKLIF_NO}}")
            run_left.font.bold = True
            run_left.font.color.rgb = RGBColor(0, 0, 128)
            run_left.font.size = Pt(9)
            
            # Orta hücre: Firma bilgileri
            cell_center = table.rows[0].cells[1]
            p_center = cell_center.paragraphs[0]
            p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run1 = p_center.add_run("AKARE ÇEVRE LABORATUVAR VE DAN. HİZM. TİC.LTD.ŞTİ Kirazlıyalı Mah. Süleyman Demirel Cad. No:28/A")
            run1.font.bold = True
            run1.font.color.rgb = RGBColor(0, 0, 128)
            run1.font.size = Pt(9)
            
            p_center.add_run("\n")
            
            run2 = p_center.add_run("Körfez V.D 013 065 1290 Körfez-KOCAELİ")
            run2.font.bold = True
            run2.font.color.rgb = RGBColor(0, 0, 128)
            run2.font.size = Pt(9)
            
            p_center.add_run("\n")
            
            run3 = p_center.add_run("info@akarecevre.com  www.akarecevre.com")
            run3.font.bold = True
            run3.font.color.rgb = RGBColor(0, 0, 128)
            run3.font.size = Pt(9)
            
            # Sağ hücre: Form kodu ve sayfa placeholder
            cell_right = table.rows[0].cells[2]
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            run_form = p_right.add_run("AÇ.F.102/Rev04/14.08.2025  ")
            run_form.font.bold = True
            run_form.font.color.rgb = RGBColor(0, 0, 128)
            run_form.font.size = Pt(9)
            
            # Sayfa numarası field
            run_page = p_right.add_run()
            run_page.font.bold = True
            run_page.font.color.rgb = RGBColor(0, 0, 128)
            run_page.font.size = Pt(9)
            
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run_page._r.append(fldChar1)
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            run_page._r.append(instrText)
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run_page._r.append(fldChar2)
            
            run_page.add_text('/')
            
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'begin')
            run_page._r.append(fldChar3)
            instrText2 = OxmlElement('w:instrText')
            instrText2.text = 'NUMPAGES'
            run_page._r.append(instrText2)
            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')
            run_page._r.append(fldChar4)
            
            # Tablo kenarlıklarını kaldır
            for cell in table.rows[0].cells:
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
        
        # Kaydet
        doc.save(template_path)
        print(f"✅ Footer eklendi: {template_path}")
        return True
        
    except Exception as e:
        print(f"❌ Hata ({template_path}): {e}")
        return False

if __name__ == '__main__':
    templates_dir = os.path.join('static', 'images')
    
    templates = [
        'TEKLİF-1 GİRİŞ.docx',
        'TEKLİF - 2 FİYAT.docx',
        'TEKLİF - 3 GENEL HUKUM.docx'
    ]
    
    for template_name in templates:
        template_path = os.path.join(templates_dir, template_name)
        if os.path.exists(template_path):
            add_footer_to_template(template_path)
        else:
            print(f"⚠️ Şablon bulunamadı: {template_path}")
    
    print("\n✅ Tüm şablonlar güncellendi!")
