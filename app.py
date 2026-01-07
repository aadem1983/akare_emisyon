import os
import json
import gc
from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, send_file, make_response
from werkzeug.utils import secure_filename
import requests
import uuid
from uuid import uuid4
import tempfile
from datetime import datetime
from io import BytesIO
from copy import deepcopy

# Lazy loading - sadece gerektiğinde yükle
pandas_loaded = False
matplotlib_loaded = False
plt = None
pd = None
np = None
mdates = None
Rectangle = None

def load_pandas():
    global pd, pandas_loaded
    if not pandas_loaded:
        import pandas as pd
        pandas_loaded = True
    return pd

def load_matplotlib():
    global plt, np, mdates, Rectangle, matplotlib_loaded
    if not matplotlib_loaded:
        import matplotlib
        matplotlib.use('Agg')  # GUI olmadan çalışması için
        import matplotlib.pyplot as plt
        import matplotlib.dates as mdates
        import numpy as np
        from matplotlib.patches import Rectangle
        matplotlib_loaded = True
    return plt, np, mdates, Rectangle
# Conditional imports for optional dependencies - Lazy loading
DOCX_AVAILABLE = False
WEASYPRINT_AVAILABLE = False

def load_docx():
    global DOCX_AVAILABLE
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        DOCX_AVAILABLE = True
        return Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    except ImportError as e:
        DOCX_AVAILABLE = False
        print(f"HATA: python-docx kütüphanesi yüklü değil: {e}")
        return None, None, None, None, None, None, None

app = Flask(__name__, template_folder='templates', static_folder='static')
app.secret_key = os.environ.get('SECRET_KEY', 'your_secret_key_development')  # Production'da ortam değişkeninden al

# Production-safe Flask ayarları
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0
is_dev = os.environ.get('FLASK_ENV') == 'development'
app.config['TEMPLATES_AUTO_RELOAD'] = True if is_dev else False
app.config['DEBUG'] = True if is_dev else (os.environ.get('FLASK_DEBUG', 'false').lower() == 'true')
app.jinja_env.auto_reload = True if is_dev else False

def allowed_file(filename, allowed_extensions):
    """Dosya uzantısının izin verilen uzantılar arasında olup olmadığını kontrol eder."""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in allowed_extensions

DATA_DIR = os.environ.get('DATA_DIR')
if DATA_DIR:
    try:
        os.makedirs(DATA_DIR, exist_ok=True)
    except Exception as _e:
        print(f"DATA_DIR oluşturulamadı: {DATA_DIR}, hata: {_e}")

def data_path(filename: str) -> str:
    # If DATA_DIR is not set, use a stable absolute path under the project root.
    # Otherwise, relative paths depend on current working directory and data can appear to "reset" after restart.
    return os.path.join(DATA_DIR, filename) if DATA_DIR else os.path.join(app.root_path, filename)

USERS_FILE = data_path('users.json')
EMISSIONS_FILE = data_path('emissions.json')
PARAMETERS_FILE = data_path('parameters.json')
MEASUREMENTS_FILE = data_path('measurements.json')
FIRMA_OLCUM_FILE = data_path('firma_olcum.json')
FIRMA_KAYIT_FILE = data_path('firma_kayit.json')
SAHA_OLC_FILE = data_path('saha_olc.json')
BACA_BILGILERI_FILE = data_path('baca_bilgileri.json')
PARAMETRE_OLCUM_FILE = data_path('parametre_olcum.json')
TEKLIF_FILE = data_path('teklif.json')
USED_TEKLIF_NUMBERS_FILE = data_path('used_teklif_numbers.json')
BACA_PARALAR_FILE = data_path('baca_paralar.json')
PARAMETRE_SAHABIL_FILE = data_path('parametre_sahabil.json')
PARAMETRE_FIELDS_FILE = data_path('parametre_fields.json')
ASGARI_FIYATLAR_FILE = data_path('asgari_fiyatlar.json')
ASGARI_FIYAT_UI_STATE_FILE = data_path('asgari_fiyat_ui_state.json')
TEKLIF_PARAMETRE_SECIM_UI_STATE_FILE = data_path('teklif_parametre_secim_ui_state.json')
PAR_SAHA_HEADERS_FILE = data_path('par_saha_header_groups.json')
FORMS_FILE = data_path('forms.json')
IL_ILCE_FILE = data_path('il-ilce.json')

def ensure_data_files():
    # İlk çalıştırmada mevcut repo kökünden DATA_DIR'e veri kopyala
    if not DATA_DIR:
        return
    try:
        # İç içe import ile global import tekrarlarını önle
        import shutil
        base_files = [
            'users.json', 'emissions.json', 'parameters.json', 'measurements.json',
            'firma_olcum.json', 'firma_kayit.json', 'saha_olc.json', 'baca_bilgileri.json',
            'parametre_olcum.json', 'teklif.json', 'used_teklif_numbers.json',
            'parametre_sahabil.json', 'forms.json', 'parametre_fields.json', 'baca_paralar.json',
            'par_saha_header_groups.json'
        ]
        for fname in base_files:
            src = os.path.join(app.root_path, fname)  # proje kökü
            dst = data_path(fname)
            # DATA_DIR'de yoksa ve repo kökünde varsa kopyala
            if not os.path.exists(dst) and os.path.exists(src):
                try:
                    shutil.copy2(src, dst)
                    print(f"DATA_DIR bootstrap: {src} -> {dst}")
                except Exception as e:
                    print(f"DATA_DIR kopyalama hatası {src} -> {dst}: {e}")
    except Exception as e:
        print(f"DATA_DIR başlangıç kopyalama hatası: {e}")

# Başlangıçta tek seferlik kontrol
ensure_data_files()

def load_users():
    """Kullanıcıları JSON dosyasından yükler."""
    if not os.path.exists(USERS_FILE):
        # Dosya yoksa, varsayılan admin kullanıcısı ile oluştur
        default_users = {'admin': {'password': '1111', 'role': 'admin'}}
        save_users(default_users)
        return default_users
    with open(USERS_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)

def save_users(users_data):
    """Kullanıcıları JSON dosyasına kaydeder."""
    with open(USERS_FILE, 'w', encoding='utf-8') as f:
        json.dump(users_data, f, indent=4, ensure_ascii=False)

def load_emissions():
    """Emisyon verilerini JSON dosyasından yükler."""
    if not os.path.exists(EMISSIONS_FILE):
        return []
    with open(EMISSIONS_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)

def save_emissions(emissions_data):
    """Emisyon verilerini JSON dosyasına kaydeder."""
    with open(EMISSIONS_FILE, 'w', encoding='utf-8') as f:
        json.dump(emissions_data, f, indent=4, ensure_ascii=False)



def load_parameters():
    """Parametreleri JSON dosyasından yükler."""
    try:
        if not os.path.exists(PARAMETERS_FILE):
            print(f"load_parameters: file not found: {PARAMETERS_FILE}")
            return []
        with open(PARAMETERS_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
        try:
            print(f"load_parameters: loaded {len(data) if isinstance(data, list) else 'n/a'} items from {PARAMETERS_FILE}")
        except Exception:
            pass
        return data
    except Exception as e:
        print(f"load_parameters error: {e} file={PARAMETERS_FILE}")
        return []

def _atomic_write_json(file_path: str, data_obj, indent: int = 4, ensure_ascii: bool = False) -> bool:
    try:
        dir_name = os.path.dirname(file_path)
        if dir_name:
            os.makedirs(dir_name, exist_ok=True)
        import tempfile
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(mode='w', delete=False, encoding='utf-8', dir=(dir_name or None), suffix='.tmp') as tf:
                tmp_path = tf.name
                json.dump(data_obj, tf, indent=indent, ensure_ascii=ensure_ascii)
                tf.flush()
                try:
                    os.fsync(tf.fileno())
                except Exception:
                    pass
            os.replace(tmp_path, file_path)
            return True
        finally:
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass
    except Exception as e:
        print(f"atomic_write_json error: {e} file={file_path}")
        return False

def save_parameters(parameters_data):
    """Parametreleri JSON dosyasına kaydeder."""
    ok = _atomic_write_json(PARAMETERS_FILE, parameters_data, indent=4, ensure_ascii=False)
    try:
        if ok:
            print(f"save_parameters: saved {len(parameters_data) if isinstance(parameters_data, list) else 'n/a'} items to {PARAMETERS_FILE}")
    except Exception:
        pass
    return bool(ok)


@app.route('/api/debug/storage')
def api_debug_storage():
    """Debug endpoint: shows resolved storage paths and file stats."""
    try:
        def _stat(p):
            try:
                if not p:
                    return None
                if not os.path.exists(p):
                    return {'exists': False, 'path': p}
                st = os.stat(p)
                return {
                    'exists': True,
                    'path': p,
                    'size': st.st_size,
                    'mtime': datetime.fromtimestamp(st.st_mtime).isoformat()
                }
            except Exception as _e:
                return {'exists': False, 'path': p, 'error': str(_e)}

        return jsonify({
            'success': True,
            'DATA_DIR': DATA_DIR,
            'app_root_path': app.root_path,
            'files': {
                'PARAMETERS_FILE': _stat(PARAMETERS_FILE),
                'FIRMA_KAYIT_FILE': _stat(FIRMA_KAYIT_FILE),
                'TEKLIF_FILE': _stat(TEKLIF_FILE),
                'ASGARI_FIYATLAR_FILE': _stat(ASGARI_FIYATLAR_FILE),
                'PARAMETRE_FIELDS_FILE': _stat(PARAMETRE_FIELDS_FILE)
            }
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/debug/parameters')
def api_debug_parameters():
    """Debug endpoint: shows parameters.json path, file stats and loaded count."""
    try:
        def _stat(p):
            try:
                if not p:
                    return None
                if not os.path.exists(p):
                    return {'exists': False, 'path': p}
                st = os.stat(p)
                return {
                    'exists': True,
                    'path': p,
                    'size': st.st_size,
                    'mtime': datetime.fromtimestamp(st.st_mtime).isoformat()
                }
            except Exception as _e:
                return {'exists': False, 'path': p, 'error': str(_e)}

        params = load_parameters()
        sample_ids = []
        try:
            for p in (params or [])[:20]:
                if isinstance(p, dict) and 'id' in p:
                    sample_ids.append(p.get('id'))
        except Exception:
            pass

        return jsonify({
            'success': True,
            'DATA_DIR': DATA_DIR,
            'app_root_path': app.root_path,
            'PARAMETERS_FILE': _stat(PARAMETERS_FILE),
            'loaded_count': len(params) if isinstance(params, list) else None,
            'sample_ids': sample_ids
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

def load_measurements():
    """Ölçüm verilerini JSON dosyasından yükler."""
    if not os.path.exists(MEASUREMENTS_FILE):
        return []
    with open(MEASUREMENTS_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)

def save_measurements(measurements_data):
    """Ölçüm verilerini JSON dosyasına kaydeder."""
    with open(MEASUREMENTS_FILE, 'w', encoding='utf-8') as f:
        json.dump(measurements_data, f, indent=4, ensure_ascii=False)

def load_firma_olcum():
    """Firma ölçüm verilerini JSON dosyasından yükler."""
    try:
        if not os.path.exists(FIRMA_OLCUM_FILE):
            return []
        with open(FIRMA_OLCUM_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if not isinstance(data, list):
                return []
            return data
    except Exception as e:
        print(f"Firma ölçüm verileri yüklenirken hata: {e}")
        return []

def save_firma_olcum(firma_olcum_data):
    """Firma ölçüm verilerini JSON dosyasına kaydeder."""
    try:
        if not isinstance(firma_olcum_data, list):
            print("Hata: firma_olcum_data liste olmalıdır")
            return False
        with open(FIRMA_OLCUM_FILE, 'w', encoding='utf-8') as f:
            json.dump(firma_olcum_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Firma ölçüm verileri kaydedilirken hata: {e}")
        return False

def load_firma_kayit():
    """Firma kayıt verilerini JSON dosyasından yükler."""
    try:
        if not os.path.exists(FIRMA_KAYIT_FILE):
            return []
        with open(FIRMA_KAYIT_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if not isinstance(data, list):
                return []
            return data
    except Exception as e:
        print(f"Firma kayıt verileri yüklenirken hata: {e}")
        return []

def save_firma_kayit(firma_kayit_data):
    """Firma kayıt verilerini JSON dosyasına kaydeder."""
    try:
        if not isinstance(firma_kayit_data, list):
            print("Hata: firma_kayit_data liste olmalıdır")
            return False
        with open(FIRMA_KAYIT_FILE, 'w', encoding='utf-8') as f:
            json.dump(firma_kayit_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Firma kayıt verileri kaydedilirken hata: {e}")
        return False

def load_saha_olc():
    """Saha ölçüm verilerini JSON dosyasından yükler."""
    try:
        if not os.path.exists(SAHA_OLC_FILE):
            return []
        with open(SAHA_OLC_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if not isinstance(data, list):
                return []
            return data
    except Exception as e:
        print(f"Saha ölçüm verileri yüklenirken hata: {e}")
        return []

def save_saha_olc(saha_olc_data):
    """Saha ölçüm verilerini JSON dosyasına kaydeder."""
    try:
        if not isinstance(saha_olc_data, list):
            print("Hata: saha_olc_data liste olmalıdır")
            return False
        with open(SAHA_OLC_FILE, 'w', encoding='utf-8') as f:
            json.dump(saha_olc_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Saha ölçüm verileri kaydedilirken hata: {e}")
        return False

def load_baca_bilgileri():
    """Baca bilgilerini JSON dosyasından yükler."""
    try:
        if not os.path.exists(BACA_BILGILERI_FILE):
            return []
        with open(BACA_BILGILERI_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if not isinstance(data, list):
                return []
            return data
    except Exception as e:
        print(f"Baca bilgileri yüklenirken hata: {e}")
        return []

def save_baca_bilgileri(baca_bilgileri_data):
    """Baca bilgilerini JSON dosyasına kaydeder."""
    try:
        if not isinstance(baca_bilgileri_data, list):
            print("Hata: baca_bilgileri_data liste olmalıdır")
            return False
        with open(BACA_BILGILERI_FILE, 'w', encoding='utf-8') as f:
            json.dump(baca_bilgileri_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Baca bilgileri kaydedilirken hata: {e}")
        return False

def load_parametre_olcum():
    """Parametre ölçüm verilerini JSON dosyasından yükler."""
    try:
        if not os.path.exists(PARAMETRE_OLCUM_FILE):
            return []
        with open(PARAMETRE_OLCUM_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if not isinstance(data, list):
                return []
            return data
    except Exception as e:
        print(f"Parametre ölçüm verileri yüklenirken hata: {e}")
        return []

def save_parametre_olcum(parametre_olcum_data):
    """Parametre ölçüm verilerini JSON dosyasına kaydeder."""
    try:
        if not isinstance(parametre_olcum_data, list):
            print("Hata: parametre_olcum_data liste olmalıdır")
            return False
        with open(PARAMETRE_OLCUM_FILE, 'w', encoding='utf-8') as f:
            json.dump(parametre_olcum_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Parametre ölçüm verileri kaydedilirken hata: {e}")
        return False

def load_teklif():
    """Teklif verilerini JSON dosyasından yükler."""
    try:
        if not os.path.exists(TEKLIF_FILE):
            return []
        with open(TEKLIF_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if not isinstance(data, list):
                return []
            return data
    except Exception as e:
        print(f"Teklif verileri yüklenirken hata: {e}")
        return []

def save_teklif(teklif_data):
    """Teklif verilerini JSON dosyasına kaydeder."""
    try:
        if not isinstance(teklif_data, list):
            print("Hata: teklif_data liste olmalıdır")
            return False
        with open(TEKLIF_FILE, 'w', encoding='utf-8') as f:
            json.dump(teklif_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Teklif verileri kaydedilirken hata: {e}")
        return False

def load_used_teklif_numbers():
    """Kullanılmış teklif numaralarını yükler (silinen teklifler dahil)"""
    try:
        used_numbers_file = USED_TEKLIF_NUMBERS_FILE
        if not os.path.exists(used_numbers_file):
            return set()
        with open(used_numbers_file, 'r', encoding='utf-8') as f:
            return set(json.load(f))
    except Exception as e:
        print(f"Kullanılmış teklif numaraları yüklenirken hata: {e}")
        return set()

def save_used_teklif_numbers(used_numbers):
    """Kullanılmış teklif numaralarını kaydeder"""
    try:
        used_numbers_file = USED_TEKLIF_NUMBERS_FILE
        with open(used_numbers_file, 'w', encoding='utf-8') as f:
            json.dump(list(used_numbers), f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"Kullanılmış teklif numaraları kaydedilirken hata: {e}")
        return False

def generate_teklif_no():
    """Yeni teklif numarası oluşturur (YYYY/TE-XXX formatında) - BENZERSİZ GARANTİLİ"""
    try:
        current_year = datetime.now().year
        teklif_prefix = f"{current_year}/TE-"
        
        # Mevcut teklifleri yükle
        teklifler = load_teklif()
        
        # Mevcut tekliflerdeki numaraları topla (basit sistem)
        used_numbers = set()
        for teklif in teklifler:
            teklif_no = teklif.get('teklif_no', '')
            if teklif_no:
                used_numbers.add(teklif_no)

        # Sadece bu yıla ait teklif numaralarını dikkate al
        used_numbers = {n for n in used_numbers if isinstance(n, str) and n.startswith(teklif_prefix)}
        
        # Bu yıl için kullanılmamış en küçük numarayı bul (3 haneli format)
        number = 1
        while True:
            new_teklif_no = f"{teklif_prefix}{number:03d}"
            if new_teklif_no not in used_numbers:
                return new_teklif_no
            number += 1
            
            # Güvenlik için maksimum 999'a kadar dene
            if number > 999:
                raise Exception("Teklif numarası limiti aşıldı!")
        
    except Exception as e:
        print(f"Teklif numarası oluşturulurken hata: {e}")
    # Hata durumunda varsayılan format (3 haneli)
    current_year = datetime.now().year
    return f"{current_year}/TE-001"

def reserve_teklif_no():
    """Teklif numarasını rezerve eder (henüz kullanılmamış, sadece rezerve)"""
    try:
        current_year = datetime.now().year
        teklif_prefix = f"{current_year}/TE-"
        
        # Mevcut teklifleri yükle
        teklifler = load_teklif()
        
        # Kullanılmış tüm numaraları yükle
        all_used_numbers = load_used_teklif_numbers()
        
        # Mevcut tekliflerdeki numaraları da ekle
        for teklif in teklifler:
            teklif_no = teklif.get('teklif_no', '')
            if teklif_no:
                all_used_numbers.add(teklif_no)

        # Sadece bu yıla ait teklif numaralarını dikkate al
        used_numbers_year = {n for n in all_used_numbers if isinstance(n, str) and n.startswith(teklif_prefix)}
        
        # Bu yıl için kullanılmamış en küçük numarayı bul ve rezerve et (3 haneli format)
        number = 1
        while True:
            new_teklif_no = f"{teklif_prefix}{number:03d}"
            if new_teklif_no not in used_numbers_year:
                # Bu numarayı rezerve et (kullanılmış olarak işaretle)
                all_used_numbers.add(new_teklif_no)
                save_used_teklif_numbers(all_used_numbers)
                return new_teklif_no
            number += 1
            
            # Güvenlik için maksimum 999'a kadar dene
            if number > 999:
                raise Exception("Teklif numarası limiti aşıldı!")
                
    except Exception as e:
        print(f"Teklif numarası rezerve edilirken hata: {e}")
        # Hata durumunda varsayılan numara (3 haneli)
        current_year = datetime.now().year
        return f"{current_year}/TE-001"

def release_teklif_no(teklif_no):
    """Rezerve edilmiş teklif numarasını serbest bırakır (vazgeçme durumunda)"""
    try:
        used_numbers = load_used_teklif_numbers()
        
        # Eğer bu numara kullanılmış numaralar listesindeyse, çıkar
        if teklif_no in used_numbers:
            used_numbers.remove(teklif_no)
            save_used_teklif_numbers(used_numbers)
            return True
        else:
            return False
            
    except Exception as e:
        print(f"Teklif numarası serbest bırakılırken hata: {e}")
        return False

def migrate_existing_teklif_numbers():
    """Mevcut teklif numaralarını kullanılmış numaralar dosyasına ekler (tek seferlik)"""
    try:
        teklifler = load_teklif()
        used_numbers = load_used_teklif_numbers()
        
        # Mevcut teklif numaralarını ekle
        for teklif in teklifler:
            teklif_no = teklif.get('teklif_no', '')
            if teklif_no:
                used_numbers.add(teklif_no)
        
        # Kaydet
        save_used_teklif_numbers(used_numbers)
        
    except Exception as e:
        return

def convert_teklif_numbers_to_new_format():
    """Eski teklif numaralarını (TE26-001) yeni formata (2026/TE-001) dönüştürür"""
    try:
        teklifler = load_teklif()
        updated_count = 0
        
        for teklif in teklifler:
            teklif_no = teklif.get('teklif_no', '')
            if teklif_no and teklif_no.startswith('TE') and '-' in teklif_no:
                # TE26-001 -> 2026/TE-001 formatına dönüştür
                parts = teklif_no.split('-')
                if len(parts) == 2:
                    prefix = parts[0]  # TE26
                    number = parts[1]  # 001
                    
                    # Yıl suffix'ini çıkar (TE26 -> 26)
                    if prefix.startswith('TE') and len(prefix) == 4:
                        year_suffix = prefix[2:]  # 26
                        
                        # 2 haneli yılı 4 haneli yıla çevir
                        year_int = int(year_suffix)
                        if year_int >= 0 and year_int <= 99:
                            # 00-99 arası: 2000-2099 olarak kabul et
                            full_year = 2000 + year_int
                            
                            # Yeni format: YYYY/TE-XXX
                            new_number = f"{full_year}/TE-{number.zfill(3)}"
                            teklif['teklif_no'] = new_number
                            updated_count += 1
        
        if updated_count > 0:
            # Güncellenmiş teklifleri kaydet
            save_teklif(teklifler)
            
            # Kullanılmış numaralar listesini de güncelle
            used_numbers = load_used_teklif_numbers()
            new_used_numbers = set()
            
            for num in used_numbers:
                if num.startswith('TE') and '-' in num:
                    parts = num.split('-')
                    if len(parts) == 2:
                        prefix = parts[0]
                        number = parts[1]
                        
                        if prefix.startswith('TE') and len(prefix) == 4:
                            year_suffix = prefix[2:]
                            year_int = int(year_suffix)
                            if year_int >= 0 and year_int <= 99:
                                full_year = 2000 + year_int
                                new_num = f"{full_year}/TE-{number.zfill(3)}"
                                new_used_numbers.add(new_num)
                            else:
                                new_used_numbers.add(num)
                        else:
                            new_used_numbers.add(num)
                    else:
                        new_used_numbers.add(num)
                else:
                    new_used_numbers.add(num)
            
            save_used_teklif_numbers(new_used_numbers)
            
    except Exception as e:
        print(f"Teklif numarası dönüştürme hatası: {e}")

def resequence_teklif_numbers():
    """Mevcut teklifleri 001'den başlayarak yeniden sıralar"""
    try:
        teklifler = load_teklif()
        current_year = datetime.now().year
        year_suffix = str(current_year)[-2:]
        
        if not teklifler:
            print("Sıralanacak teklif bulunamadı")
            return
        
        # Teklifleri tarihe göre sırala (en eskiden en yeniye)
        teklifler.sort(key=lambda x: x.get('teklif_tarihi', ''))
        
        # Yeni numaraları ata
        new_used_numbers = set()
        for i, teklif in enumerate(teklifler, 1):
            old_number = teklif.get('teklif_no', '')
            new_number = f'TE{year_suffix}-{i:03d}'
            
            teklif['teklif_no'] = new_number
            new_used_numbers.add(new_number)
            
            pass
        
        # Güncellenmiş teklifleri kaydet
        save_teklif(teklifler)
        
        # Kullanılmış numaralar listesini güncelle
        save_used_teklif_numbers(new_used_numbers)
        
        pass
        
    except Exception as e:
        print(f"Teklif sıralama hatası: {e}")

def format_tarih_gg_aa_yyyy(tarih_str):
    """YYYY-MM-DD formatındaki tarihi GG.AA.YYYY formatına çevirir"""
    if not tarih_str:
        return ''
    try:
        # YYYY-MM-DD formatını parse et
        tarih_obj = datetime.strptime(tarih_str, '%Y-%m-%d')
        # GG.AA.YYYY formatına çevir
        return tarih_obj.strftime('%d.%m.%Y')
    except:
        return tarih_str

# Template'de kullanılacak yardımcı fonksiyonları Flask'a ekle
app.jinja_env.globals.update(format_tarih_gg_aa_yyyy=format_tarih_gg_aa_yyyy)

@app.route('/api/formlar/genel-hukumler', methods=['GET'])
def get_genel_hukumler():
    """Formlar sayfasındaki genel hükümler metnini döndürür"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        # Varsayılan genel hükümler metni
        default_text = """<strong>GENEL HÜKÜMLER</strong>
<ul>
    <li><strong>Fiyatlara KDV dahil değildir.</strong></li>
    <li><strong>Teklifimizin geçerlilik süresi 30 gündür.</strong></li>
    <li><strong>(*) işaretli parametrelerin ölçümü ve analizleri yetkili taşeron işbirliği laboratuvarı tarafından yapılacaktır.</strong></li>
    <li><strong>Fiyat tarifemiz; Çevre, Şehircilik ve İklim Değişikliği Bakanlığı'nın 27.12.2024 tarihinde yayınlanan duyurusu itibarıyla, 01.01.2025 tarihi itibariyle geçerli olacak şekil tarifi öz sınırları kaynaklamıştır.</strong></li>
    <li><strong>Yetkili Ölçüm ve Analiz Laboratuvarlarından 02.12.2024 tarihinden sonra Fiyat Tarifesi 13. Maddesi gereği, ödeme rapor tesliminden önce para yapılacaktır.</strong></li>
    <li><strong>Ölçüm esnasında firmanızdan kaynaklanan bir hangi bir sebepten dolayı ölçüm yapılamaması durumunda ücret iadeesi yapılmayacaktır.</strong></li>
    <li><strong>Tüm anlaşmazlıklarda İstanbul (Çağlayan) Adliyeleri yetkilidir.</strong></li>
    <li><strong>İş Güvenliği Firmanız sorumluluğundadır.</strong></li>
    <li><strong>Numuneler 30 gün saklanacaktır.</strong></li>
    <li><strong>Raporlara itiraz süresi teslim tarihinden itibaren 15 gündür.</strong></li>
    <li><strong>Raporlar 15 iş günü içinde teslim edilecektir.</strong></li>
    <li><strong>Laboratuvarımız TS EN ISO/IEC 17025 standardına uygun çalışmaktadır.</strong></li>
    <li><strong>Karar kuralları: Ölçüm sonuçları belirsizlik hesaplamaları ile birlikte değerlendirilir.</strong></li>
</ul>"""
        
        return jsonify({'success': True, 'content': default_text})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})


@app.route('/api/asgari_fiyatlar/save_table', methods=['POST'])
def api_save_asgari_fiyatlar_table():
    """Asgari fiyat tablosunu (çoklu yıl sütunları) toplu kaydeder."""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})

    try:
        data = request.get_json(silent=True) or {}
        years = data.get('years')
        rows = data.get('rows')

        if not isinstance(years, list) or not years:
            return jsonify({'success': False, 'message': 'years alanı liste olmalı'}), 400
        if not isinstance(rows, list):
            return jsonify({'success': False, 'message': 'rows alanı liste olmalı'}), 400

        cleaned = []

        def _to_float(v):
            if v is None:
                return None
            s = str(v).strip()
            if s == '':
                return None
            try:
                return float(s.replace(',', '.'))
            except Exception:
                return None

        for r in rows:
            if not isinstance(r, dict):
                continue
            kapsam = str(r.get('kapsam', '')).strip()
            parametre = str(r.get('parametre', '')).strip()
            metot = str(r.get('metot', '')).strip()
            yillar = r.get('yillar', [])
            if not parametre:
                continue
            if not isinstance(yillar, list):
                yillar = []

            if not kapsam:
                kapsam = 'EMİSYON'

            yillik = {}
            for i, y in enumerate(years):
                val = yillar[i] if i < len(yillar) else None
                fval = _to_float(val)
                if fval is not None:
                    yillik[str(y)] = fval

            cleaned.append({
                'kapsam': kapsam,
                'parametre': parametre,
                'metot': metot,
                'yillik': yillik
            })

        ok = save_asgari_fiyatlar(cleaned)
        return jsonify({'success': bool(ok), 'message': 'Asgari fiyat tablosu kaydedildi.' if ok else 'Kaydetme hatası'})

    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'}), 500


@app.route('/api/asgari_fiyatlar/full', methods=['GET'])
def api_get_asgari_fiyatlar_full():
    """Asgari fiyatların ham halini (yıllık alanlarıyla) döndürür."""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'}), 401
    try:
        raw = load_asgari_fiyatlar()
        return jsonify({'success': True, 'data': raw})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'}), 500


def load_asgari_fiyat_ui_state():
    try:
        if not os.path.exists(ASGARI_FIYAT_UI_STATE_FILE):
            return {}
        with open(ASGARI_FIYAT_UI_STATE_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except Exception as e:
        print(f"Asgari fiyat UI state yüklenirken hata: {e}")
        return {}


def save_asgari_fiyat_ui_state(data):
    try:
        if not isinstance(data, dict):
            data = {}
        return bool(_atomic_write_json(ASGARI_FIYAT_UI_STATE_FILE, data, indent=2, ensure_ascii=False))
    except Exception as e:
        print(f"Asgari fiyat UI state kaydedilirken hata: {e}")
        return False


def load_teklif_parametre_secim_ui_state():
    try:
        if not os.path.exists(TEKLIF_PARAMETRE_SECIM_UI_STATE_FILE):
            return {}
        with open(TEKLIF_PARAMETRE_SECIM_UI_STATE_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except Exception as e:
        print(f"Teklif parametre seçim UI state yüklenirken hata: {e}")
        return {}


def save_teklif_parametre_secim_ui_state(data):
    try:
        if not isinstance(data, dict):
            data = {}
        return bool(_atomic_write_json(TEKLIF_PARAMETRE_SECIM_UI_STATE_FILE, data, indent=2, ensure_ascii=False))
    except Exception as e:
        print(f"Teklif parametre seçim UI state kaydedilirken hata: {e}")
        return False


@app.route('/api/ui_state/asgari_fiyat', methods=['GET'])
def api_get_ui_state_asgari_fiyat():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'}), 401
    try:
        st = load_asgari_fiyat_ui_state()
        return jsonify({'success': True, 'data': st})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'}), 500


@app.route('/api/ui_state/asgari_fiyat', methods=['POST'])
def api_set_ui_state_asgari_fiyat():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'}), 401
    try:
        payload = request.get_json(silent=True) or {}
        if not isinstance(payload, dict):
            return jsonify({'success': False, 'message': 'Geçersiz veri'}), 400
        # Beklenen yapı: {col_widths: {"3": 200, ...}}
        col_widths = payload.get('col_widths')
        if col_widths is not None and not isinstance(col_widths, dict):
            return jsonify({'success': False, 'message': 'col_widths sözlük olmalı'}), 400
        current = load_asgari_fiyat_ui_state()
        current['col_widths'] = col_widths or {}
        current['updated_at'] = datetime.now().isoformat()
        ok = save_asgari_fiyat_ui_state(current)
        return jsonify({'success': bool(ok)})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'}), 500


@app.route('/api/ui_state/teklif_parametre_secim', methods=['GET'])
def api_get_ui_state_teklif_parametre_secim():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'}), 401
    try:
        st = load_teklif_parametre_secim_ui_state()
        return jsonify({'success': True, 'data': st})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'}), 500


@app.route('/api/ui_state/teklif_parametre_secim', methods=['POST'])
def api_set_ui_state_teklif_parametre_secim():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'}), 401
    try:
        payload = request.get_json(silent=True) or {}
        if not isinstance(payload, dict):
            return jsonify({'success': False, 'message': 'Geçersiz veri'}), 400
        col_widths = payload.get('col_widths')
        if col_widths is not None and not isinstance(col_widths, dict):
            return jsonify({'success': False, 'message': 'col_widths sözlük olmalı'}), 400
        current = load_teklif_parametre_secim_ui_state()
        current['col_widths'] = col_widths or {}
        current['updated_at'] = datetime.now().isoformat()
        ok = save_teklif_parametre_secim_ui_state(current)
        return jsonify({'success': bool(ok)})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'}), 500

def format_date_with_day(dt_str):
    try:
        if not dt_str or dt_str is None:
            return ''
        dt_str = str(dt_str).strip()
        if not dt_str:
            return ''
        # Eğer zaten GG.AA.YY (GUN) formatındaysa, olduğu gibi döndür
        if '(' in dt_str and ')' in dt_str:
            return dt_str
        # Eğer zaten GG.AA.YY formatındaysa, gün ekle
        if '.' in dt_str and len(dt_str) == 8:
            dt = datetime.strptime(dt_str, '%d.%m.%y')
        # Sadece tarih formatı (YYYY-MM-DD)
        elif len(dt_str) == 10 and dt_str.count('-') == 2:
            dt = datetime.strptime(dt_str, '%Y-%m-%d')
        # Tarih-saat formatı (YYYY-MM-DDTHH:MM veya YYYY-MM-DD HH:MM)
        elif 'T' in dt_str:
            dt = datetime.strptime(dt_str, '%Y-%m-%dT%H:%M')
        else:
            dt = datetime.strptime(dt_str, '%Y-%m-%d %H:%M')
        # Türkçe kısa günler
        gunler = ['PZT', 'SAL', 'ÇAR', 'PER', 'CUM', 'CTS', 'PAZ']
        gun_index = dt.weekday()  # Pazartesi=0, Pazar=6
        gun_kisa = gunler[gun_index]
        return dt.strftime('%d.%m.%y') + f' ({gun_kisa})'
    except Exception as e:
        print(f"Tarih formatı hatası: {dt_str}, Hata: {e}")
        return str(dt_str) if dt_str else ''

# Yetki kontrolü fonksiyonları
def can_read(role):
    """Rol 1, 2, 3 veya admin ise okuma yetkisi var"""
    return role in ['1', '2', '3', 'admin']

def can_write(role):
    """Rol 2, 3 veya admin ise yazma yetkisi var"""
    return role in ['2', '3', 'admin']

def can_edit(role):
    """Rol 3 veya admin ise düzenleme yetkisi var"""
    return role in ['3', 'admin']

def can_delete(role):
    """Rol 3 veya admin ise silme yetkisi var"""
    return role in ['3', 'admin']

# Uygulama başlangıcında verileri yükle
users = load_users()

# İl-ilçe verilerini yükle
CITIES_DATA = []
try:
    # Verileri yerel dosyadan okumayı dene
    with open('il-ilce.json', 'r', encoding='utf-8') as f:
        CITIES_DATA = json.load(f)
except FileNotFoundError:
    # Yerel dosya yoksa URL'den çek
    try:
        response = requests.get('https://raw.githubusercontent.com/furkan-dogu/Turkiye-Sehir-ve-Ilceleri/main/il-ilce.json')
        response.raise_for_status() # HTTP hatalarını kontrol et
        CITIES_DATA = response.json()
        # Gelen veriyi kalıcı veri klasörüne kaydet
        with open(IL_ILCE_FILE, 'w', encoding='utf-8') as f:
            json.dump(CITIES_DATA, f, ensure_ascii=False, indent=2)
    except requests.exceptions.RequestException as e:
        print(f"Uyarı: İl-ilçe verileri yüklenemedi. Hata: {e}")

@app.route('/login', methods=['GET', 'POST'])
def login():
    error = None
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        # Kullanıcı adını küçük harfe çevir ve büyük/küçük harf duyarlılığını kaldır
        username_lower = username.lower()
        
        # Kullanıcıları yükle ve kullanıcı adını küçük harfe çevirerek kontrol et
        current_users = load_users()
        user = None
        original_username = None
        
        # Kullanıcı adını büyük/küçük harf duyarlılığı olmadan ara
        for stored_username, user_data in current_users.items():
            if stored_username.lower() == username_lower:
                user = user_data
                original_username = stored_username
                break
        
        if user and user['password'] == password:
            session['logged_in'] = True
            session['username'] = original_username  # Orijinal kullanıcı adını kullan
            session['role'] = user['role']
            return redirect(url_for('index'))

        error = 'Geçersiz kullanıcı adı veya şifre'

    return render_template('login.html', error=error)

@app.route('/')
def index():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    emissions = load_emissions()
    
    # Tarihleri GG.AA.YY formatında formatla
    for emission in emissions:
        if 'tarih' in emission:
            emission['tarih'] = format_date_with_day(emission['tarih'])
    
    return render_template('index.html', username=session.get('username'), role=session.get('role'), emissions=emissions)

@app.route('/logout')
def logout():
    session.pop('logged_in', None)
    session.pop('username', None)
    session.pop('role', None)
    return redirect(url_for('login'))

@app.route('/pivot')
def pivot():
    # Geçici olarak giriş kontrolünü devre dışı bırak
    # if not session.get('logged_in'):
    #     return redirect(url_for('login'))
    # if session.get('role') != 'admin':
    #     return redirect(url_for('index'))
    return render_template('pivot.html', username='admin', role='admin')

@app.route('/api/pivot/summary')
def pivot_summary():
    if not session.get('logged_in'):
        return jsonify({'error': 'auth'}), 401
    if session.get('role') != 'admin':
        return jsonify({'error': 'forbidden'}), 403
    try:
        start = request.args.get('start')
        end = request.args.get('end')
        if not start or not end:
            return jsonify({'error': 'start/end required'}), 400
        from datetime import datetime
        def in_range(dt):
            try:
                d = datetime.fromisoformat(dt.replace('Z', '+00:00')).date()
            except Exception:
                try:
                    d = datetime.strptime(dt, '%Y-%m-%d').date()
                except Exception:
                    return False
            return datetime.strptime(start, '%Y-%m-%d').date() <= d <= datetime.strptime(end, '%Y-%m-%d').date()
        def _norm(sval: str) -> str:
            m = (sval or '').strip().lower()
            tr_map = str.maketrans({'ı':'i','İ':'i','ş':'s','Ş':'s','ç':'c','Ç':'c','ğ':'g','Ğ':'g','ü':'u','Ü':'u','ö':'o','Ö':'o'})
            return m.translate(tr_map)
        teklifler = load_teklif()
        filt_teklif = [t for t in teklifler if t.get('teklif_tarihi') and in_range(t.get('teklif_tarihi'))]
        toplam_teklif_adedi = len(filt_teklif)
        toplam_teklif_tutari = sum(float(t.get('netToplam', 0) or 0) for t in filt_teklif)
        accepted_list, rejected_list = [], []
        for t in filt_teklif:
            status = _norm(t.get('teklif_durumu',''))
            if any(k in status for k in ['kabul','onay']):
                accepted_list.append(t)
            elif any(k in status for k in ['red','ret','iptal','olumsuz','kabul edilmedi','kabul edilmez']):
                rejected_list.append(t)
        kabul_adet = len(accepted_list)
        red_adet = len(rejected_list)
        kapsam_ici = []
        kapsam_disi = []
        for t in accepted_list:
            tip = _norm(t.get('teklif_tipi',''))
            if 'kapsam' in tip and 'ici' in tip:
                kapsam_ici.append(t)
            elif 'kapsam' in tip and ('disi' in tip or 'dış' in tip):
                kapsam_disi.append(t)
        kapsam_ici_adet = len(kapsam_ici)
        kapsam_disi_adet = len(kapsam_disi)
        kapsam_ici_tutar = sum(float(t.get('netToplam', 0) or 0) for t in kapsam_ici)
        kapsam_disi_tutar = sum(float(t.get('netToplam', 0) or 0) for t in kapsam_disi)
        from collections import defaultdict
        parametre_ozet = defaultdict(lambda: {'adet': 0, 'toplam': 0.0})
        for t in filt_teklif:
            for pr in t.get('parametreler', []) or []:
                p_ad = pr.get('parametre') or 'Bilinmiyor'
                try:
                    p_adet = int(pr.get('adet', 0) or 0)
                except Exception:
                    p_adet = 0
                try:
                    p_top = float(pr.get('topFiyat', 0) or 0)
                except Exception:
                    p_top = 0.0
                parametre_ozet[p_ad]['adet'] += p_adet
                parametre_ozet[p_ad]['toplam'] += p_top
        parametre_list = [{'parametre': k, 'adet': v['adet'], 'toplam': round(v['toplam'], 2)} for k, v in sorted(parametre_ozet.items(), key=lambda x: x[0].lower())]
        try:
            baca_list = load_baca_bilgileri()
        except Exception:
            baca_list = []
        baca_index = {}
        for b in baca_list:
            if not (b.get('firma_adi') and b.get('olcum_kodu') and b.get('baca_adi')):
                continue
            key = (b.get('firma_adi'), b.get('olcum_kodu'), b.get('baca_adi'))
            baca_index[key] = (b.get('personel_adi') or '').strip() or 'Bilinmiyor'
        personel_set = set()
        matrix = defaultdict(lambda: defaultdict(int))
        for o in [o for o in load_parametre_olcum() if o.get('created_at') and in_range(o.get('created_at')[:10])]:
            param = (o.get('parametre_adi') or 'Bilinmiyor').strip() or 'Bilinmiyor'
            key = (o.get('firma_adi'), o.get('olcum_kodu'), o.get('baca_adi'))
            personel = (o.get('personel_adi') or '').strip()
            if not personel:
                personel = baca_index.get(key, 'Bilinmiyor')
            if not personel:
                personel = 'Bilinmiyor'
            personel_set.add(personel)
            matrix[param][personel] += 1
        personeller = sorted(list(personel_set), key=lambda x: x.lower())
        matrix_rows = []
        for param, counts in sorted(matrix.items(), key=lambda x: x[0].lower()):
            row = {'parametre': param, 'counts': {p: counts.get(p, 0) for p in personeller}, 'toplam': sum(counts.values())}
            matrix_rows.append(row)
        toplam_by_person = defaultdict(int)
        for counts in matrix.values():
            for p, v in counts.items():
                toplam_by_person[p] += v
        personel_list = [{'personel': p, 'adet': toplam_by_person[p]} for p in sorted(toplam_by_person.keys(), key=lambda x: (-toplam_by_person[x], x.lower()))]
        return jsonify({'summary': {'toplam_teklif_adedi': toplam_teklif_adedi, 'toplam_teklif_tutari': round(toplam_teklif_tutari, 2), 'kapsam_ici_adet': kapsam_ici_adet, 'kapsam_ici_tutar': round(kapsam_ici_tutar, 2), 'kapsam_disi_adet': kapsam_disi_adet, 'kapsam_disi_tutar': round(kapsam_disi_tutar, 2), 'kabul_adet': kabul_adet, 'red_adet': red_adet, 'toplam_olcum_adedi': len([o for o in load_parametre_olcum() if o.get('created_at') and in_range(o.get('created_at')[:10])])}, 'parametreler': parametre_list, 'personeller': personel_list, 'personel_parametre': {'personel_headers': personeller, 'rows': matrix_rows}})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/admin', methods=['GET', 'POST'])
def admin():
    # Protect this route
    if session.get('role') != 'admin':
        return redirect(url_for('index'))

    current_users = load_users()

    if request.method == 'POST':
        action = request.form.get('action')

        if action == 'add_user':
            username = request.form['username']
            password = request.form['password']
            surname = request.form.get('surname', '')
            gorev = request.form.get('gorev', '')
            role = request.form['role']
            
            # İmza dosyasını işle
            imza_filename = None
            if 'imza' in request.files and request.files['imza'].filename:
                imza_file = request.files['imza']
                if imza_file and allowed_file(imza_file.filename, {'png', 'jpg', 'jpeg', 'gif'}):
                    # Dosya adını güvenli hale getir
                    filename = secure_filename(imza_file.filename)
                    imza_filename = f"{username}_{filename}"
                    
                    # Static/images/signatures klasörünü oluştur
                    import os
                    signatures_dir = os.path.join('static', 'images', 'signatures')
                    os.makedirs(signatures_dir, exist_ok=True)
                    
                    # Dosyayı kaydet
                    imza_file.save(os.path.join(signatures_dir, imza_filename))
            
            if username and password and username not in current_users:
                current_users[username] = {
                    'password': password, 
                    'role': role,
                    'surname': surname,
                    'gorev': gorev,
                    'imza': imza_filename
                }
                save_users(current_users)
                flash(f'Kullanıcı "{username}" başarıyla eklendi!', 'success')
            else:
                flash('Kullanıcı adı zaten mevcut veya eksik bilgi!', 'danger')

        elif action == 'update_user':
            username = request.form.get('edit_username') or request.form['username']
            if username != 'admin' and username in current_users:
                # Şifre sadece girilmişse güncelle
                if request.form.get('password'):
                    current_users[username]['password'] = request.form['password']
                current_users[username]['role'] = request.form['role']
                current_users[username]['surname'] = request.form.get('surname', '')
                current_users[username]['gorev'] = request.form.get('gorev', '')
                
                # İmza dosyasını işle
                if 'imza' in request.files and request.files['imza'].filename:
                    imza_file = request.files['imza']
                    if imza_file and allowed_file(imza_file.filename, {'png', 'jpg', 'jpeg', 'gif'}):
                        # Dosya adını güvenli hale getir
                        filename = secure_filename(imza_file.filename)
                        imza_filename = f"{username}_{filename}"
                        
                        # Static/images/signatures klasörünü oluştur
                        import os
                        signatures_dir = os.path.join('static', 'images', 'signatures')
                        os.makedirs(signatures_dir, exist_ok=True)
                        
                        # Eski imzayı sil
                        if current_users[username].get('imza'):
                            old_imza_path = os.path.join(signatures_dir, current_users[username]['imza'])
                            if os.path.exists(old_imza_path):
                                os.remove(old_imza_path)
                        
                        # Yeni dosyayı kaydet
                        imza_file.save(os.path.join(signatures_dir, imza_filename))
                        current_users[username]['imza'] = imza_filename
                
                save_users(current_users)
                flash(f'Kullanıcı "{username}" başarıyla güncellendi!', 'success')
            else:
                flash('Kullanıcı güncellenemedi!', 'danger')

        elif action == 'delete_user':
            username = request.form['username']
            if username != 'admin' and username in current_users:
                del current_users[username]
                save_users(current_users)
                flash(f'Kullanıcı "{username}" başarıyla silindi!', 'success')
            else:
                flash('Kullanıcı silinemedi!', 'danger')
        
        return redirect(url_for('admin'))

    # For GET request, pass the list of users to the template
    return render_template('admin.html', username=session.get('username'), users=current_users)

@app.route('/add_emission', methods=['GET', 'POST'])
def add_emission():
    if not session.get('logged_in') or not can_write(session.get('role')):
        return redirect(url_for('index')) # Yetkisiz erişim

    if request.method == 'POST':
        emissions = load_emissions()
        new_emission = {
            'id': str(uuid4()),
            'tesis_adi': request.form['tesis_adi'],
            'tarih': request.form['tarih'],
            'parametre': request.form['parametre'],
            'sonuc': request.form['sonuc'],
            'birim': request.form['birim']
        }
        emissions.append(new_emission)
        save_emissions(emissions)
        return redirect(url_for('index'))

    return render_template('add_emission.html', username=session.get('username'))

@app.route('/edit_emission/<int:emission_id>', methods=['GET', 'POST'])
def edit_emission(emission_id):
    if not session.get('logged_in') or not can_edit(session.get('role')):
        return redirect(url_for('index'))

    emissions = load_emissions()
    emission_to_edit = next((e for e in emissions if e['id'] == emission_id), None)

    if not emission_to_edit:
        return redirect(url_for('index')) # Kayıt bulunamadı

    if request.method == 'POST':
        emission_to_edit['tesis_adi'] = request.form['tesis_adi']
        emission_to_edit['tarih'] = request.form['tarih']
        emission_to_edit['parametre'] = request.form['parametre']
        emission_to_edit['sonuc'] = request.form['sonuc']
        emission_to_edit['birim'] = request.form['birim']
        save_emissions(emissions)
        return redirect(url_for('index'))

    return render_template('edit_emission.html', username=session.get('username'), emission=emission_to_edit)

@app.route('/delete_emission/<int:emission_id>')
def delete_emission(emission_id):
    if not session.get('logged_in') or not can_delete(session.get('role')):
        return redirect(url_for('index'))

    emissions = load_emissions()
    emissions_to_keep = [e for e in emissions if e['id'] != emission_id]
    save_emissions(emissions_to_keep)
    return redirect(url_for('index'))



@app.route('/parametre')
def parametre():
    if not session.get('logged_in') or not can_read(session.get('role')):
        return redirect(url_for('login'))
    current_parameters = load_parameters()
    return render_template('parametre.html', parameters=current_parameters, username=session.get('username'), role=session.get('role'))

@app.route('/import_parameters', methods=['POST'])
def import_parameters():
    if 'username' not in session:
        return redirect(url_for('login'))

    if 'excel_file' not in request.files:
        flash('Dosya bölümü bulunamadı!', 'danger')
        return redirect(url_for('parametre'))

    file = request.files['excel_file']

    if not file or not file.filename or file.filename == '':
        flash('Dosya seçilmedi!', 'danger')
        return redirect(url_for('parametre'))

    if (file.filename and (file.filename.endswith('.xlsx') or file.filename.endswith('.xls'))):
        try:
            pd = load_pandas()
            df = pd.read_excel(file)
            
            # İstenen sütun sırası
            expected_columns = [
                'Sıra',
                'Parametre Adı',
                'Metot',
                'İzo Oran',
                'Nozzle',
                '1. İmp',
                '2. İmp',
                '3. İmp',
                '4. İmp',
                'L/DAK',
                'T.HAC',
                'LOQ',
                'KK',
                '-3S',
                '-2S',
                '+2S',
                '+3S'
            ]
            
            # Sadece zorunlu sütunları kontrol et
            required_columns = ['Parametre Adı', 'Metot']
            if not all(col in df.columns for col in required_columns):
                flash(f'Excel dosyasında gerekli sütunlar bulunamadı. En az "Parametre Adı" ve "Metot" sütunları olmalıdır.', 'danger')
                return redirect(url_for('parametre'))

            # NaN değerleri boş string ile değiştir
            df.fillna('', inplace=True)

            new_parameters = []
            for index, row in df.iterrows():
                param_dict = {}
                
                # Sadece beklenen sütunları al, diğerlerini atla
                for col in expected_columns:
                    if col in df.columns:
                        cell_value = row[col]
                        if pd.isna(cell_value) or cell_value == '' or cell_value is None:
                            param_dict[col] = ''
                        else:
                            param_dict[col] = str(cell_value)
                    else:
                        param_dict[col] = ''
                
                # Sıra sütununu parametre verilerinden çıkar
                if 'Sıra' in param_dict:
                    del param_dict['Sıra']
                
                param_dict['id'] = str(uuid4())
                new_parameters.append(param_dict)
            
            save_parameters(new_parameters)

            flash('Parametreler başarıyla Excel dosyasından yüklendi.', 'success')
        except Exception as e:
            flash(f'Dosya okunurken bir hata oluştu: {e}', 'danger')
        
        return redirect(url_for('parametre'))
    else:
        flash('Geçersiz dosya formatı. Lütfen .xlsx veya .xls uzantılı bir dosya yükleyin.', 'danger')
        return redirect(url_for('parametre'))


@app.route('/export_parameters')
def export_parameters():
    if 'username' not in session:
        return redirect(url_for('login'))

    parameters = load_parameters()
    if not parameters:
        flash('Dışarı aktarılacak veri bulunmuyor.', 'warning')
        return redirect(url_for('parametre'))

    pd = load_pandas()
    df = pd.DataFrame(parameters)
    # 'Lab LOQ' sütunu varsa kaldır
    if 'Lab LOQ' in df.columns:
        df = df.drop(columns=['Lab LOQ'])
    if 'id' in df.columns:
        df = df.drop(columns=['id'])

    # İstenen sıralamaya göre sütunları yeniden düzenle
    desired_order = [
        'Parametre Adı',
        'Metot', 
        'İzo Oran',
        'Nozzle',
        '1. İmp',
        '2. İmp',
        '3. İmp',
        '4. İmp',
        'L/DAK',
        'T.HAC',
        'LOQ',
        'KK',
        '-3S',
        '-2S',
        '+2S',
        '+3S'
    ]
    for col in desired_order:
        if col not in df.columns:
            df[col] = ''
    df = df[desired_order]

    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        df.to_excel(tmp.name, index=False, sheet_name='Parametreler', engine='openpyxl')
        tmp.flush()
        
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='parametreler.xlsx'
        )





@app.route('/parametre/add', methods=['GET', 'POST'])
def add_parameter():
    if not session.get('logged_in') or not can_write(session.get('role')):
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        new_param = {
            'id': str(uuid4()),
            'Parametre Adı': request.form.get('Parametre Adı'),
            'Metot': request.form.get('Metot'),
            'İzo Oran': request.form.get('İzo Oran'),
            'LOQ': request.form.get('LOQ'),
            'KK': request.form.get('KK'),
            '1. İmp': request.form.get('1. İmp'),
            '2. İmp': request.form.get('2. İmp'),
            '3. İmp': request.form.get('3. İmp'),
            '4. İmp': request.form.get('4. İmp'),
            'L/DAK': request.form.get('L/DAK'),
            'Nozzle': request.form.get('Nozzle'),
            'T.HAC': request.form.get('T.HAC'),
            '-3S': request.form.get('-3S'),
            '-2S': request.form.get('-2S'),
            '+2S': request.form.get('+2S'),
            '+3S': request.form.get('+3S')
        }
        
        parameters = load_parameters()
        parameters.append(new_param)
        save_parameters(parameters)
        flash('Yeni parametre başarıyla eklendi.', 'success')
        return redirect(url_for('parametre'))

    return render_template('add_edit_parameter.html', username=session.get('username'))

@app.route('/parametre/edit/<parameter_id>', methods=['GET', 'POST'])
def edit_parameter(parameter_id):
    if not session.get('logged_in') or not can_edit(session.get('role')):
        return redirect(url_for('login'))

    parameters = load_parameters()
    param_to_edit = next((p for p in parameters if p['id'] == parameter_id), None)

    if not param_to_edit:
        flash('Parametre bulunamadı.', 'danger')
        return redirect(url_for('parametre'))

    if request.method == 'POST':
        param_to_edit['Parametre Adı'] = request.form.get('Parametre Adı')
        param_to_edit['Metot'] = request.form.get('Metot')
        param_to_edit['İzo Oran'] = request.form.get('İzo Oran')
        param_to_edit['LOQ'] = request.form.get('LOQ')
        param_to_edit['KK'] = request.form.get('KK')
        param_to_edit['1. İmp'] = request.form.get('1. İmp')
        param_to_edit['2. İmp'] = request.form.get('2. İmp')
        param_to_edit['3. İmp'] = request.form.get('3. İmp')
        param_to_edit['4. İmp'] = request.form.get('4. İmp')
        param_to_edit['L/DAK'] = request.form.get('L/DAK')
        param_to_edit['Nozzle'] = request.form.get('Nozzle')
        param_to_edit['T.HAC'] = request.form.get('T.HAC')
        param_to_edit['-3S'] = request.form.get('-3S')
        param_to_edit['-2S'] = request.form.get('-2S')
        param_to_edit['+2S'] = request.form.get('+2S')
        param_to_edit['+3S'] = request.form.get('+3S')
        
        save_parameters(parameters)
        flash('Parametre başarıyla güncellendi.', 'success')
        return redirect(url_for('parametre'))

    return render_template('add_edit_parameter.html', parameter=param_to_edit, username=session.get('username'))

@app.route('/parametre/delete/<parameter_id>', methods=['POST'])
def delete_parameter(parameter_id):
    if not session.get('logged_in') or not can_delete(session.get('role')):
        return redirect(url_for('login'))

    parameters = load_parameters()
    parameters_to_keep = [p for p in parameters if p['id'] != parameter_id]

    if len(parameters) == len(parameters_to_keep):
        flash('Silinecek parametre bulunamadı.', 'danger')
    else:
        save_parameters(parameters_to_keep)
        flash('Parametre başarıyla silindi.', 'success')
    
    return redirect(url_for('parametre'))



@app.route('/api/ilceler/<il_adi>')
def api_ilceler(il_adi):
    # Gelen il adına göre ilçeleri bul ve sadece ilçe adlarını döndür
    for il in CITIES_DATA:
        if il['il_adi'].upper() == il_adi.upper():
            ilceler = [ilce['ilce_adi'] for ilce in il['ilceler']]
            return jsonify(sorted(ilceler))
    return jsonify([])



@app.route('/olcum_olustur')
def olcum_olustur():
    if not session.get('logged_in') or not can_read(session.get('role')):
        return redirect(url_for('login'))
    measurements = load_measurements()
    # Ölçüm başlangıç ve bitiş tarihlerini formatla
    for m in measurements:
        m['olcum_baslangic_fmt'] = format_date_with_day(m.get('olcum_baslangic', ''))
        m['olcum_bitis_fmt'] = format_date_with_day(m.get('olcum_bitis', ''))
    # Sadece görevi "Saha" olan kullanıcıları filtrele
    current_users = load_users()
    saha_users = []
    for username, user_data in current_users.items():
        if user_data.get('gorev') == 'Saha':
            saha_users.append(username)
    user_list = sorted(saha_users)
    # Parametreleri yükle ve isimleri benzersiz yap
    parameters = load_parameters()
    unique_param_names = []
    seen = set()
    for p in parameters:
        name = p.get('Parametre Adı') or p.get('parametre_adi') or p.get('ad') or p.get('isim')
        if name and name not in seen:
            unique_param_names.append(name)
            seen.add(name)
    return render_template('olcum_olustur_simple.html', username=session.get('username'), measurements=measurements, users=user_list, unique_param_names=unique_param_names, CITIES_DATA=CITIES_DATA)

@app.route('/add_measurement', methods=['POST'])
def add_measurement():
    if not session.get('logged_in') or not can_write(session.get('role')):
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        
        # Çoklu personel seçimini al ve virgülle birleştir
        selected_personnel = request.form.getlist('olcumPersoneli')
        personnel_str = ', '.join(selected_personnel) if selected_personnel else ''
        
        new_measurement = {
            'id': str(uuid4()),
            'firma_adi': request.form.get('firmaAdi'),
            'il': request.form.get('il'),
            'ilce': request.form.get('ilce'),
            'il_ilce': f"{request.form.get('il')} - {request.form.get('ilce')}" if request.form.get('il') and request.form.get('ilce') else '',
            'olcum_kodu': request.form.get('olcumKodu'),
            'baca_sayisi': request.form.get('bacaSayisi', ''),
            'olcum_baslangic': request.form.get('olcumBaslangic'),
            'olcum_bitis': request.form.get('olcumBitis'),
            'olcum_baslangic_fmt': format_date_with_day(request.form.get('olcumBaslangic')),
            'olcum_bitis_fmt': format_date_with_day(request.form.get('olcumBitis')),
            'yetkili': request.form.get('yetkili', ''),
            'telefon_no': request.form.get('telefonNo', ''),
            'olcumPersoneli': personnel_str,
            'parametre': request.form.get('parametre', ''),
            'tarih': request.form.get('olcumBaslangic', '')[:10] if request.form.get('olcumBaslangic') else '',  # Sadece tarih kısmı
            'durum': request.form.get('durum', 'Aktif')
        }
        
        measurements = load_measurements()
        measurements.append(new_measurement)
        save_measurements(measurements)
        
        # AJAX/fetch ile gelirse JSON döndür
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.is_json:
            return jsonify({'success': True, 'measurement_id': new_measurement['id']})
        
        flash('Yeni ölçüm başarıyla eklendi!', 'success')
        return redirect(url_for('olcum_olustur'))
    
    return redirect(url_for('olcum_olustur'))

@app.route('/save_baca_data', methods=['POST'])
def save_baca_data():
    if not session.get('logged_in') or not can_write(session.get('role')):
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        data = request.get_json()
        olcum_id = data.get('olcum_id')
        baca_data = data.get('baca_data', [])
        
        # Saha verilerini kaydet (basitleştirilmiş)
        return jsonify({'success': True, 'message': 'Baca bilgileri kaydedildi'})
        
        return jsonify({'success': True, 'message': 'Baca bilgileri kaydedildi'})
        
    except Exception as e:
        print(f"Baca veri kaydetme hatası: {e}")
        return jsonify({'error': f'Kaydetme hatası: {str(e)}'}), 500

@app.route('/edit_measurement/<measurement_id>', methods=['POST'])
def edit_measurement(measurement_id):
    if not session.get('logged_in') or not can_edit(session.get('role')):
        return redirect(url_for('login'))
    measurements = load_measurements()
    measurement = next((m for m in measurements if m['id'] == measurement_id), None)
    if not measurement:
        flash('Ölçüm kaydı bulunamadı.', 'danger')
        return redirect(url_for('olcum_olustur'))
    
    # Çoklu personel seçimini al ve virgülle birleştir
    selected_personnel = request.form.getlist('olcumPersoneli')
    personnel_str = ', '.join(selected_personnel) if selected_personnel else ''
    
    # Formdan gelen verilerle güncelle
    measurement['firma_adi'] = request.form.get('firmaAdi')
    measurement['olcum_kodu'] = request.form.get('olcumKodu')
    measurement['baca_sayisi'] = request.form.get('bacaSayisi')
    measurement['olcum_baslangic'] = request.form.get('olcumBas')
    measurement['olcum_bitis'] = request.form.get('olcumBit')
    measurement['olcumPersoneli'] = personnel_str
    measurement['tarih'] = request.form.get('olcumBas', '')[:10] if request.form.get('olcumBas') else ''
    save_measurements(measurements)
    flash('Ölçüm başarıyla güncellendi!', 'success')
    return redirect(url_for('olcum_olustur'))

@app.route('/delete_measurement/<measurement_id>', methods=['POST'])
def delete_measurement(measurement_id):
    if not session.get('logged_in') or not can_delete(session.get('role')):
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        measurements = load_measurements()
        new_measurements = [m for m in measurements if m['id'] != measurement_id]
        
        if len(measurements) == len(new_measurements):
            return jsonify({'error': 'Silinecek ölçüm kaydı bulunamadı'}), 404
        else:
            save_measurements(new_measurements)
            return jsonify({'success': True, 'message': 'Ölçüm başarıyla silindi'}), 200
            
    except Exception as e:
        return jsonify({'error': f'Silme işlemi sırasında hata: {str(e)}'}), 500

@app.route('/delete_selected_measurements', methods=['POST'])
def delete_selected_measurements():
    if not session.get('logged_in') or not can_delete(session.get('role')):
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        data = request.get_json()
        selected_ids = data.get('ids', [])
        
        if not selected_ids:
            return jsonify({'error': 'Seçilen ölçüm bulunamadı'}), 400
        
        measurements = load_measurements()
        original_count = len(measurements)
        
        # Seçilen ID'leri sil
        new_measurements = [m for m in measurements if m['id'] not in selected_ids]
        
        if len(new_measurements) < original_count:
            save_measurements(new_measurements)
            deleted_count = original_count - len(new_measurements)
            return jsonify({'success': True, 'message': f'{deleted_count} ölçüm başarıyla silindi'}), 200
        else:
            return jsonify({'error': 'Silinecek ölçüm bulunamadı'}), 404
            
    except Exception as e:
        return jsonify({'error': f'Silme işlemi sırasında hata: {str(e)}'}), 500

@app.route('/get_measurement_details/<measurement_id>')
def get_measurement_details(measurement_id):
    try:
        if not session.get('logged_in') or not can_read(session.get('role')):
            return jsonify({'error': 'Unauthorized'}), 401
        
        measurements = load_measurements()
        measurement = next((m for m in measurements if m['id'] == measurement_id), None)
        
        if not measurement:
            return jsonify({'error': 'Ölçüm bulunamadı'}), 404
        
        # Tarihleri formatla
        measurement['olcum_baslangic_fmt'] = format_date_with_day(measurement.get('olcum_baslangic', ''))
        measurement['olcum_bitis_fmt'] = format_date_with_day(measurement.get('olcum_bitis', ''))
        

        
        # Baca bilgilerini al (basitleştirilmiş)
        baca_bilgileri = []
        
        # HTML içeriği oluştur
        html_content = f"""
        <div class="container-fluid">
            <div class="row">
                <div class="col-md-6">
                    <div class="card mb-3">
                        <div class="card-header bg-primary text-white">
                            <h6 class="mb-0"><i class="fas fa-info-circle me-2"></i>Ölçüm Bilgileri</h6>
                        </div>
                        <div class="card-body">
                            <table class="table table-sm">
                                <tr>
                                    <td><strong>Firma Adı:</strong></td>
                                    <td>{measurement.get('firma_adi', 'Belirtilmemiş')}</td>
                                </tr>
                                <tr>
                                    <td><strong>Ölçüm Kodu:</strong></td>
                                    <td>{measurement.get('olcum_kodu', 'Belirtilmemiş')}</td>
                                </tr>
                                <tr>
                                    <td><strong>Baca Sayısı:</strong></td>
                                    <td>{measurement.get('baca_sayisi', 'Belirtilmemiş')}</td>
                                </tr>
                                <tr>
                                    <td><strong>Ölçüm Başlangıç:</strong></td>
                                    <td>{measurement.get('olcum_baslangic_fmt', 'Belirtilmemiş')}</td>
                                </tr>
                                <tr>
                                    <td><strong>Ölçüm Bitiş:</strong></td>
                                    <td>{measurement.get('olcum_bitis_fmt', 'Belirtilmemiş')}</td>
                                </tr>
                                <tr>
                                    <td><strong>Personel:</strong></td>
                                    <td>{measurement.get('olcumPersoneli', 'Belirtilmemiş')}</td>
                                </tr>
                                <tr>
                                    <td><strong>Durum:</strong></td>
                                    <td><span class="badge bg-success">{measurement.get('durum', 'Aktif')}</span></td>
                                </tr>
                            </table>
                        </div>
                    </div>
                </div>
                <div class="col-md-6">
                    <div class="card mb-3">
                        <div class="card-header bg-info text-white">
                            <h6 class="mb-0"><i class="fas fa-building me-2"></i>Firma Bilgileri</h6>
                        </div>
                        <div class="card-body">
        """
        html_content += """
                            <div class="text-center text-muted">
                                <i class="fas fa-info-circle fa-2x mb-2"></i>
                                <p>Firma bilgileri artık kullanılmıyor</p>
                            </div>
        """
        html_content += """
                        </div>
                    </div>
                </div>
            </div>
        """
        
        # Baca bilgileri varsa ekle
        if baca_bilgileri:
            html_content += """
            <div class="row">
                <div class="col-12">
                    <div class="card">
                        <div class="card-header bg-success text-white">
                            <h6 class="mb-0"><i class="fas fa-chimney me-2"></i>Baca ve Parametre Bilgileri</h6>
                        </div>
                        <div class="card-body">
                            <div class="table-responsive">
                                <table class="table table-sm table-bordered">
                                    <thead class="table-light">
                                        <tr>
                                            <th style="width: 50px;">Sıra</th>
                                            <th>Baca Adı</th>
                                            <th>Parametreler</th>
                                        </tr>
                                    </thead>
                                    <tbody>
            """
            
            for index, baca in enumerate(baca_bilgileri, 1):
                parametreler_html = ""
                if baca.get('parametreler'):
                    # Parametre sayılarını hesapla
                    param_counts = {}
                    for param in baca['parametreler']:
                        param_counts[param] = param_counts.get(param, 0) + 1
                    
                    # Her parametreyi sayısıyla birlikte göster
                    for param, count in param_counts.items():
                        parametreler_html += f'<span class="badge bg-primary me-1" style="position: relative;">{param}<span style="position: absolute; top: -8px; right: -8px; background: #dc3545; color: white; border-radius: 50%; width: 18px; height: 18px; font-size: 0.7rem; display: flex; align-items: center; justify-content: center; font-weight: bold;">{count}</span></span>'
                else:
                    parametreler_html = '<span class="text-muted">Parametre seçilmemiş</span>'
                
                html_content += f"""
                                        <tr>
                                            <td class="text-center">
                                                <span class="badge bg-secondary" style="min-width: 25px; font-size: 0.85rem; font-weight: bold; padding: 2px 6px;">{index}</span>
                                            </td>
                                            <td><strong>{baca.get('baca_adi', 'Belirtilmemiş')}</strong></td>
                                            <td>{parametreler_html}</td>
                                        </tr>
                """
            
            html_content += """
                                    </tbody>
                                </table>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
            """
        else:
            html_content += """
            <div class="row">
                <div class="col-12">
                    <div class="card">
                        <div class="card-header bg-warning text-dark">
                            <h6 class="mb-0"><i class="fas fa-exclamation-triangle me-2"></i>Baca Bilgileri</h6>
                        </div>
                        <div class="card-body">
                            <div class="text-center text-muted">
                                <i class="fas fa-info-circle fa-2x mb-2"></i>
                                <p>Bu ölçüm için baca ve parametre kaydı bulunamadı.</p>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
            """
        
        html_content += """
        </div>
        """
        
        return jsonify({
            'success': True,
            'html': html_content
        })
    except Exception as e:
        print(f"Hata: {e}")
        return jsonify({'error': f'Server hatası: {str(e)}'}), 500

@app.route('/export_measurement/<measurement_id>')
def export_measurement(measurement_id):
    if 'username' not in session:
        return redirect(url_for('login'))
    
    measurements = load_measurements()
    measurement = next((m for m in measurements if m['id'] == measurement_id), None)
    
    if not measurement:
        flash('Ölçüm bulunamadı.', 'danger')
        return redirect(url_for('olcum_olustur'))
    
    # Ölçüm verilerini DataFrame'e çevir
    data = {
        'Firma Adı': [measurement['firma_adi']],
        'Ölçüm Kodu': [measurement['olcum_kodu']],
        'Baca Sayısı': [measurement['baca_sayisi']],
        'Ölçüm Başlangıç': [measurement['olcum_baslangic']],
        'Ölçüm Bitiş': [measurement['olcum_bitis']],
        'Personel': [measurement['olcumPersoneli']],
        'Durum': [measurement['durum']],
        'Tarih': [measurement['tarih']]
    }
    
    df = pd.DataFrame(data)
    
    import tempfile
    import os
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        df.to_excel(tmp.name, index=False, sheet_name='Ölçüm Verileri', engine='openpyxl')
        tmp.flush()
        
        filename = f"olcum_{measurement['olcum_kodu']}_{measurement['firma_adi']}.xlsx"
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )

@app.route('/export_all_measurements')
def export_all_measurements():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    measurements = load_measurements()
    if not measurements:
        flash('Dışarı aktarılacak ölçüm bulunmuyor.', 'warning')
        return redirect(url_for('olcum_olustur'))
    
    # Ölçüm verilerini DataFrame'e çevir
    data_list = []
    for m in measurements:
        data_list.append({
            'Firma Adı': m['firma_adi'],
            'Ölçüm Kodu': m['olcum_kodu'],
            'Baca Sayısı': m['baca_sayisi'],
            'Ölçüm Başlangıç': m['olcum_baslangic'],
            'Ölçüm Bitiş': m['olcum_bitis'],
            'Personel': m['olcumPersoneli'],
            'Durum': m['durum'],
            'Tarih': m['tarih']
        })
    
    df = pd.DataFrame(data_list)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        df.to_excel(tmp.name, index=False, sheet_name='Tüm Ölçümler', engine='openpyxl')
        tmp.flush()
        
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='tum_olcumler.xlsx'
        )

@app.route('/import_measurement_data', methods=['POST'])
def import_measurement_data():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    if 'excel_file' not in request.files:
        flash('Dosya seçilmedi.', 'danger')
        return redirect(url_for('olcum_olustur'))
    
    file = request.files['excel_file']
    if file.filename == '':
        flash('Dosya seçilmedi.', 'danger')
        return redirect(url_for('olcum_olustur'))
    
    if file and file.filename and file.filename.endswith(('.xlsx', '.xls')):
        try:
            df = pd.read_excel(file)
            measurements = load_measurements()
            
            # Excel'den gelen verileri ölçümlere uygula
            for index, row in df.iterrows():
                # Ölçüm kodu ile eşleşen kaydı bul
                matching_measurement = next((m for m in measurements if m['olcum_kodu'] == row.get('Ölçüm Kodu', '')), None)
                
                if matching_measurement:
                    # Verileri güncelle
                    matching_measurement['firma_adi'] = row.get('Firma Adı', matching_measurement['firma_adi'])
                    matching_measurement['baca_sayisi'] = str(row.get('Baca Sayısı', matching_measurement['baca_sayisi']))
                    matching_measurement['olcum_baslangic'] = str(row.get('Ölçüm Başlangıç', matching_measurement['olcum_baslangic']))
                    matching_measurement['olcum_bitis'] = str(row.get('Ölçüm Bitiş', matching_measurement['olcum_bitis']))
                    matching_measurement['olcumPersoneli'] = str(row.get('Personel', matching_measurement['olcumPersoneli']))
                    matching_measurement['durum'] = row.get('Durum', matching_measurement['durum'])
                    matching_measurement['tarih'] = str(row.get('Tarih', matching_measurement['tarih']))
            
            save_measurements(measurements)
            flash('Ölçüm verileri başarıyla güncellendi.', 'success')
            
        except Exception as e:
            flash(f'Dosya okunurken hata oluştu: {e}', 'danger')
    else:
        flash('Geçersiz dosya formatı. Lütfen .xlsx veya .xls uzantılı dosya yükleyin.', 'danger')
    
    return redirect(url_for('olcum_olustur'))

@app.route('/export_selected_measurements')
def export_selected_measurements():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    ids_param = request.args.get('ids', '')
    if not ids_param:
        flash('Seçilen ölçüm bulunamadı.', 'warning')
        return redirect(url_for('olcum_olustur'))
    
    selected_ids = ids_param.split(',')
    measurements = load_measurements()
    selected_measurements = [m for m in measurements if m['id'] in selected_ids]
    
    if not selected_measurements:
        flash('Seçilen ölçüm bulunamadı.', 'warning')
        return redirect(url_for('olcum_olustur'))
    
    # Seçilen ölçümleri DataFrame'e çevir
    data_list = []
    for m in selected_measurements:
        data_list.append({
            'Firma Adı': m['firma_adi'],
            'Ölçüm Kodu': m['olcum_kodu'],
            'Baca Sayısı': m['baca_sayisi'],
            'Ölçüm Başlangıç': m['olcum_baslangic'],
            'Ölçüm Bitiş': m['olcum_bitis'],
            'Personel': m['olcumPersoneli'],
            'Durum': m['durum'],
            'Tarih': m['tarih']
        })
    
    df = pd.DataFrame(data_list)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        df.to_excel(tmp.name, index=False, sheet_name='Seçilen Ölçümler', engine='openpyxl')
        tmp.flush()
        
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'secilen_olcumler_{len(selected_measurements)}_kayit.xlsx'
        )



@app.route('/export_selected_parameters')
def export_selected_parameters():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    ids_param = request.args.get('ids', '')
    selected_ids = ids_param.split(',') if ids_param else []
    parameters = load_parameters()
    selected_parameters = [p for p in parameters if p['id'] in selected_ids]
    
    columns = [
        'Parametre Adı', 'Metot', 'İzo Oran', 'Nozzle', '1. İmp', '2. İmp', '3. İmp', '4. İmp', 'L/DAK', 'T.HAC', 'LOQ', 'KK', '-3S', '-2S', '+2S', '+3S'
    ]
    # Her parametrede başlıkları normalize et ve Lab LOQ'yu LOQ'ya kopyala
    for p in selected_parameters:
        for col in columns:
            if col not in p:
                p[col] = ''
        # Eski veri desteği: 'Lab LOQ' varsa 'LOQ'ya kopyala
        if not p['LOQ'] and ('Lab LOQ' in p and p['Lab LOQ']):
            p['LOQ'] = p['Lab LOQ']
        # Yanlış anahtarları düzelt (ör: nozzle, NOZZLE, Nozzle vs.)
        for key in list(p.keys()):
            if key.lower().replace(' ', '') == 'nozzle':
                p['Nozzle'] = p[key]
            if key.lower().replace(' ', '') == 't.hac':
                p['T.HAC'] = p[key]
    if not selected_parameters:
        df = pd.DataFrame()
    else:
        df = pd.DataFrame(selected_parameters)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        df.to_excel(tmp.name, index=False, sheet_name='Seçilen Parametreler', engine='openpyxl')
        tmp.flush()
        
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'secilen_parametreler_{len(selected_parameters)}_kayit.xlsx'
        )

@app.route('/delete_selected_users', methods=['POST'])
def delete_selected_users():
    if 'username' not in session:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401
    
    data = request.get_json()
    usernames = data.get('usernames', []) if data else []
    
    if not usernames:
        return jsonify({'success': False, 'error': 'No usernames provided'}), 400
    
    current_users = load_users()
    deleted_count = 0
    
    for username in usernames:
        if username != 'admin' and username in current_users:
            del current_users[username]
            deleted_count += 1
    
    save_users(current_users)
    
    if deleted_count > 0:
        flash(f'{deleted_count} kullanıcı başarıyla silindi!', 'success')
    
    return jsonify({'success': True, 'deleted_count': deleted_count})

@app.route('/export_selected_users')
def export_selected_users():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    usernames_param = request.args.get('usernames', '')
    if not usernames_param:
        flash('Seçilen kullanıcı bulunamadı.', 'warning')
        return redirect(url_for('admin'))
    
    selected_usernames = usernames_param.split(',')
    users_data = load_users()
    selected_users = []
    
    for username in selected_usernames:
        if username in users_data:
            selected_users.append({
                'Kullanıcı Adı': username,
                'Rol': users_data[username]['role']
            })
    
    if not selected_users:
        flash('Seçilen kullanıcı bulunamadı.', 'warning')
        return redirect(url_for('admin'))
    
    # Seçilen kullanıcıları DataFrame'e çevir
    df = pd.DataFrame(selected_users)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        df.to_excel(tmp.name, index=False, sheet_name='Seçilen Kullanıcılar', engine='openpyxl')
        tmp.flush()
        
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'secilen_kullanicilar_{len(selected_users)}_kayit.xlsx'
        )

@app.route('/export_selected_emissions')
def export_selected_emissions():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    ids_param = request.args.get('ids', '')
    if not ids_param:
        flash('Seçilen emisyon bulunamadı.', 'warning')
        return redirect(url_for('index'))
    
    selected_ids = ids_param.split(',')
    emissions = load_emissions()
    selected_emissions = [e for e in emissions if e['id'] in selected_ids]
    
    if not selected_emissions:
        flash('Seçilen emisyon bulunamadı.', 'warning')
        return redirect(url_for('index'))
    
    # Seçilen emisyonları DataFrame'e çevir
    data_list = []
    for e in selected_emissions:
        data_list.append({
            'Tesis Adı': e['tesis_adi'],
            'Tarih': e['tarih'],
            'Parametre': e['parametre'],
            'Sonuç': e['sonuc'],
            'Birim': e['birim']
        })
    
    df = pd.DataFrame(data_list)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        df.to_excel(tmp.name, index=False, sheet_name='Seçilen Emisyonlar', engine='openpyxl')
        tmp.flush()
        
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'secilen_emisyonlar_{len(selected_emissions)}_kayit.xlsx'
        )

@app.route('/delete_selected_parameters', methods=['POST'])
def delete_selected_parameters():
    if 'username' not in session:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401
    data = request.get_json()
    ids = data.get('ids', []) if data else []
    if not ids:
        return jsonify({'success': False, 'error': 'No IDs provided'}), 400
    parameters = load_parameters()
    print('Silinecek parametre idleri:', ids)
    new_parameters = [p for p in parameters if str(p.get('id')) not in ids and p.get('id') not in ids]
    save_parameters(new_parameters)
    return jsonify({'success': True})

@app.route('/add_selected_parameters', methods=['POST'])
def add_selected_parameters():
    if 'username' not in session:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401
    
    data = request.get_json()
    ids = data.get('ids', []) if data else []
    
    if not ids:
        return jsonify({'success': False, 'error': 'No IDs provided'}), 400
    
    try:
        parameters = load_parameters()
        selected_parameters = [p for p in parameters if str(p.get('id')) in ids or p.get('id') in ids]
        
        if not selected_parameters:
            return jsonify({'success': False, 'error': 'Seçilen parametreler bulunamadı'}), 404
        
        # Seçilen parametreleri kopyala ve yeni ID'ler ver
        new_parameters = []
        for param in selected_parameters:
            new_param = param.copy()
            new_param['id'] = str(uuid4())  # Yeni benzersiz ID
            new_param['created_at'] = datetime.now().isoformat()
            new_parameters.append(new_param)
        
        # Yeni parametreleri mevcut listeye ekle
        parameters.extend(new_parameters)
        save_parameters(parameters)
        
        print(f'{len(new_parameters)} adet parametre başarıyla eklendi')
        return jsonify({'success': True, 'message': f'{len(new_parameters)} adet parametre başarıyla eklendi'})
        
    except Exception as e:
        print(f'Parametre ekleme hatası: {e}')
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/update_parameter_field', methods=['POST'])
def update_parameter_field():
    if 'username' not in session:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401
    
    data = request.get_json()
    parameter_id = data.get('parameter_id')
    field = data.get('field')
    value = data.get('value')
    
    if not parameter_id or not field:
        return jsonify({'success': False, 'error': 'Missing required fields'}), 400
    
    parameters = load_parameters()
    parameter = next((p for p in parameters if str(p.get('id')) == str(parameter_id)), None)
    
    if not parameter:
        return jsonify({'success': False, 'error': 'Parameter not found'}), 404
    
    # Yeni sütunları güncelle
    parameter[field] = value
    save_parameters(parameters)
    
    return jsonify({'success': True})

@app.route('/firma_kayit')
def firma_kayit():
    if not session.get('logged_in') or not can_read(session.get('role')):
        return redirect(url_for('login'))
    
    # Firma kayıt verilerini yükle
    firma_kayitlar = load_firma_kayit()
    
    # İstediğiniz illeri en üste taşı
    oncelikli_iller = ['KOCAELİ', 'SAKARYA', 'DÜZCE', 'BOLU', 'İSTANBUL', 'BURSA', 'BİLECİK', 'KÜTAHYA']
    
    # İlleri öncelik sırasına göre düzenle
    sirali_cities = []
    
    # Önce öncelikli illeri ekle
    for il_adi in oncelikli_iller:
        for il in CITIES_DATA:
            if il['il_adi'] == il_adi:
                sirali_cities.append(il)
                break
    
    # Sonra diğer illeri ekle
    for il in CITIES_DATA:
        if il['il_adi'] not in oncelikli_iller:
            sirali_cities.append(il)
    
    return render_template('firma_kayit.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         cities_data=sirali_cities,
                         firma_kayitlar=firma_kayitlar)

@app.route('/teklif')
def teklif():
    if not session.get('logged_in') or not can_read(session.get('role')):
        return redirect(url_for('login'))
    
    # Teklif verilerini yükle
    teklifler = load_teklif()
    
    # Firma listesini yükle (firma seçimi için)
    firma_kayitlar = load_firma_kayit()
    
    return render_template('teklif.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         teklifler=teklifler,
                         firma_kayitlar=firma_kayitlar)

@app.route('/api/teklif/add', methods=['POST'])
def add_teklif():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        
        # Yeni teklif oluştur
        yeni_teklif = {
            'id': str(uuid.uuid4()),
            'teklif_tipi': data.get('teklif_tipi', ''),
            'firma_adi': data.get('firma_adi', ''),
            'teklif_tarihi': data.get('teklif_tarihi', ''),
            'teklif_no': data.get('teklif_no', ''),
            'indirim_orani': data.get('indirim_orani', ''),
            'indirim_tipi': data.get('indirim_tipi', ''),  # TL veya %
            'parametreler': data.get('parametreler', []),
            'toplam': data.get('toplam', 0),
            'indirim': data.get('indirim', 0),
            'netToplam': data.get('netToplam', 0),
            'teklif_giris_metni': data.get('teklif_giris_metni', ''),
            'genel_hukumler': data.get('genel_hukumler', ''),
            'teklif_durumu': data.get('teklif_durumu', 'BEKLEMEDE'),
            'created_at': datetime.now().isoformat(),
            'updated_at': datetime.now().isoformat()
        }
        
        # Mevcut verileri yükle
        teklifler = load_teklif()
        teklifler.append(yeni_teklif)
        
        # Verileri kaydet
        if save_teklif(teklifler):
            return jsonify({'success': True, 'message': 'Teklif başarıyla kaydedildi', 'teklif': yeni_teklif})
        else:
            return jsonify({'success': False, 'message': 'Teklif kaydedilirken hata oluştu'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/teklif/next_number')
def get_next_teklif_number():
    """Bir sonraki teklif numarasını döndürür"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        next_number = generate_teklif_no()
        return jsonify({'success': True, 'teklif_no': next_number})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/teklif/reserve_number')
def reserve_teklif_number():
    """Teklif numarasını rezerve eder (vazgeçme durumunda serbest bırakılabilir)"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        reserved_number = reserve_teklif_no()
        return jsonify({'success': True, 'reserved_number': reserved_number})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/teklif/release_number', methods=['POST'])
def release_teklif_number():
    """Rezerve edilmiş teklif numarasını serbest bırakır"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        teklif_no = data.get('teklif_no', '')
        
        if not teklif_no:
            return jsonify({'success': False, 'message': 'Teklif numarası gerekli'})
        
        success = release_teklif_no(teklif_no)
        if success:
            return jsonify({'success': True, 'message': f'Teklif numarası serbest bırakıldı: {teklif_no}'})
        else:
            return jsonify({'success': False, 'message': f'Teklif numarası zaten serbest: {teklif_no}'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/asgari_fiyatlar')
def get_asgari_fiyatlar():
    """Asgari fiyat verilerini döndürür"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        current_year = datetime.now().year
        raw = load_asgari_fiyatlar()
        # UI uyumluluğu: {parametre, metot, fiyat} listesine dönüştür
        asgari_fiyatlar = []
        for item in raw:
            yillik = item.get('yillik', {}) or {}
            fiyat = yillik.get(str(current_year))
            if fiyat is None:
                # en yeni yıla fallback
                if yillik:
                    latest_year = sorted(yillik.keys())[-1]
                    fiyat = yillik.get(latest_year)
            asgari_fiyatlar.append({
                'parametre': item.get('parametre', ''),
                'metot': item.get('metot', ''),
                'fiyat': float(fiyat or 0)
            })
        return jsonify({'success': True, 'asgari_fiyatlar': asgari_fiyatlar, 'current_year': current_year})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/asgari_fiyatlar/add', methods=['POST'])
def add_asgari_fiyat():
    """Yeni asgari fiyat parametresi ekler"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        
        # Yeni parametre
        yeni_parametre = {
            'parametre': data.get('parametre', ''),
            'metot': data.get('metot', ''),
            'fiyat': float(data.get('fiyat', 0))
        }
        
        # Validasyon
        if not yeni_parametre['parametre']:
            return jsonify({'success': False, 'message': 'Parametre adı gerekli'})
        
        # Dosyadan oku, ekle ve kaydet
        data_file = load_asgari_fiyatlar()
        # Aynı parametre varsa güncelle
        updated = False
        for item in data_file:
            if item.get('parametre', '').upper() == yeni_parametre['parametre'].upper():
                item['metot'] = yeni_parametre.get('metot', item.get('metot', ''))
                yillik = item.get('yillik', {}) or {}
                yillik[str(datetime.now().year)] = yeni_parametre['fiyat']
                item['yillik'] = yillik
                updated = True
                break
        if not updated:
            data_file.append({
                'parametre': yeni_parametre['parametre'],
                'metot': yeni_parametre['metot'],
                'yillik': { str(datetime.now().year): yeni_parametre['fiyat'] }
            })
        save_asgari_fiyatlar(data_file)
        return jsonify({'success': True})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/asgari_fiyatlar/delete', methods=['POST'])
def delete_asgari_fiyat():
    """Asgari fiyat parametresini siler"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        parametre_adi = data.get('parametre', '')
        
        if not parametre_adi:
            return jsonify({'success': False, 'message': 'Parametre adı gerekli'})
        
        # Sabit parametreler (silinemez)
        sabit_parametreler = ['TABAN FİYAT', 'YOL']
        
        if parametre_adi.upper() in [p.upper() for p in sabit_parametreler]:
            return jsonify({'success': False, 'message': 'Bu parametre silinemez'})
        
        # Gerçek uygulamada veritabanından silinir
        return jsonify({
            'success': True, 
            'message': 'Parametre başarıyla silindi'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/teklif/detail/<teklif_id>')
def get_teklif_detail(teklif_id):
    """Teklif detayını getirir"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        teklifler = load_teklif()
        teklif = next((t for t in teklifler if t.get('id') == teklif_id), None)
        
        if not teklif:
            return jsonify({'success': False, 'message': 'Teklif bulunamadı'})
        
        return jsonify({'success': True, 'teklif': teklif})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/teklif/delete', methods=['POST'])
def delete_teklif():
    """Teklifi siler - NUMARA BENZERSİZLİĞİ GARANTİLİ"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        teklif_id = data.get('teklif_id')
        
        if not teklif_id:
            return jsonify({'success': False, 'message': 'Teklif ID gerekli'})
        
        teklifler = load_teklif()
        
        # Silinecek teklifi bul ve numarasını kaydet
        silinen_teklif = next((t for t in teklifler if t.get('id') == teklif_id), None)
        
        if silinen_teklif and silinen_teklif.get('teklif_no'):
            # Kullanılmış numaralar listesine ekle (zaten ekliydi ama emin olmak için)
            used_numbers = load_used_teklif_numbers()
            used_numbers.add(silinen_teklif['teklif_no'])
            save_used_teklif_numbers(used_numbers)
        
        # Teklifi listeden kaldır
        teklifler = [t for t in teklifler if t.get('id') != teklif_id]
        
        if save_teklif(teklifler):
            return jsonify({'success': True, 'message': 'Teklif başarıyla silindi ve numarası korundu'})
        else:
            return jsonify({'success': False, 'message': 'Teklif silinirken hata oluştu'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/teklif/delete_bulk', methods=['POST'])
def delete_teklif_bulk():
    """Birden fazla teklifi siler - NUMARA BENZERSİZLİĞİ GARANTİLİ"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})

    try:
        data = request.get_json() or {}
        teklif_ids = data.get('teklif_ids')

        if not isinstance(teklif_ids, list) or not teklif_ids:
            return jsonify({'success': False, 'message': 'teklif_ids liste olmalı ve boş olmamalı'})

        teklifler = load_teklif()
        id_set = set(str(x) for x in teklif_ids if x is not None)

        # Silinecek teklif numaralarını used_numbers'a ekle
        used_numbers = load_used_teklif_numbers()
        deleted_count = 0

        for t in teklifler:
            try:
                if str(t.get('id')) in id_set:
                    deleted_count += 1
                    teklif_no = t.get('teklif_no')
                    if teklif_no:
                        used_numbers.add(teklif_no)
            except Exception:
                pass

        save_used_teklif_numbers(used_numbers)

        # Teklifleri listeden kaldır
        teklifler_new = [t for t in teklifler if str(t.get('id')) not in id_set]

        if save_teklif(teklifler_new):
            return jsonify({'success': True, 'message': f'{deleted_count} teklif başarıyla silindi ve numaraları korundu'})
        return jsonify({'success': False, 'message': 'Teklifler silinirken hata oluştu'})

    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

# Asgari fiyat tarifesi - yıllara göre
def load_asgari_fiyatlar():
    """Asgari fiyatları yıllara göre yükler"""
    try:
        with open('asgari_fiyatlar.json', 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Asgari fiyatlar yüklenirken hata: {e}")
        return []

def get_parametre_fiyati(parametre_adi, yil):
    """Belirli bir yıl için parametre fiyatını döndürür"""
    asgari_fiyatlar = load_asgari_fiyatlar()
    
    # Parametre adı eşleştirmesi
    parametre_eslestirme = {
        'TOZ': 'TOZ',
        'YG': 'YANMA GAZI', 
        'YANMA GAZI': 'YANMA GAZI',  # Eksik olan eşleştirme eklendi
        'VOC': 'VOC',
        'AĞIR METAL': 'AĞIR METAL',
        'TOC': 'TOC',
        'SO2': 'YANMA GAZI',  # SO2 YANMA GAZI içinde
        'NO': 'YANMA GAZI',   # NO YANMA GAZI içinde  
        'CO': 'YANMA GAZI',   # CO YANMA GAZI içinde
        'O2': 'YANMA GAZI',   # O2 YANMA GAZI içinde
        'PM10': 'PM10',
        'ÇT': 'ÇÖKEN TOZ',
        'Çt': 'ÇÖKEN TOZ',    # Küçük t ile yazım
        'NEM': 'NEM',
        'HF': 'HF',
        'HCL': 'HCL',
        'AMON.': 'AMONYAK',
        'AMONYAK': 'AMONYAK',  # Tam yazım
        'FORM.': 'FORMALDEHİT',
        'FORMALDEHİT': 'FORMALDEHİT',  # Tam yazım
        'CR6': 'CR+6',
        'Cr6': 'CR+6',        # Küçük c ile yazım
        'FOSF.A.': 'FOSFORİK ASİT',
        'FOSFORİK ASİT': 'FOSFORİK ASİT',  # Tam yazım
        'SÜLFÜRİK ASİT': 'SÜLFÜRİK ASİT',  # Eksik olan eşleştirme eklendi
        'SÜLF.A': 'SÜLFÜRİK ASİT',  # Alternatif yazım
        'HCN': 'HCN',
        'PAH': 'PAH',         # Tam yazım
        'DİOKSİN FURAN': 'DİOKSİN FURAN'  # Tam yazım
    }
    
    eslestirilen_parametre = parametre_eslestirme.get(parametre_adi, parametre_adi)
    
    for fiyat_kaydi in asgari_fiyatlar:
        if fiyat_kaydi['parametre'] == eslestirilen_parametre:
            yillik_fiyatlar = fiyat_kaydi.get('yillik', {})
            fiyat = yillik_fiyatlar.get(str(yil), 0)
            return fiyat
    
    return 0

@app.route('/api/parametre-fiyatlari', methods=['GET'])
def get_parametre_fiyatlari():
    """Parametre fiyat tarifesini döndürür"""
    # 2025 yılı için varsayılan fiyatlar
    varsayilan_fiyatlar = {
        'TOZ': 6655,
        'YG': 3855,
        'VOC': 13171,
        'AĞIR METAL': 19000,
        'TOC': 8500,
        'SO2': 4500,
        'NO': 4200,
        'CO': 3800,
        'O2': 3500,
        'PM10': 7200,
        'ÇT': 5500,
        'NEM': 1430,
        'HF': 12000,
        'HCL': 9500,
        'AMON.': 6800,
        'FORM.': 7500,
        'CR6': 15000,
        'FOSF.A.': 11000,
        'HCN': 13000
    }
    
    return jsonify({
        'success': True,
        'fiyatlar': varsayilan_fiyatlar
    })

@app.route('/api/asgari-fiyatlar/<yil>', methods=['GET'])
def get_asgari_fiyatlar_yil(yil):
    """Belirli bir yıl için asgari fiyatları döndürür"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        yil = int(yil)
        asgari_fiyatlar = load_asgari_fiyatlar()
        
        # Yıllık fiyatları filtrele
        yillik_fiyatlar = {}
        for fiyat_kaydi in asgari_fiyatlar:
            parametre = fiyat_kaydi['parametre']
            yillik_fiyat = fiyat_kaydi.get('yillik', {}).get(str(yil), 0)
            yillik_fiyatlar[parametre] = yillik_fiyat
        
        return jsonify({
            'success': True,
            'yil': yil,
            'fiyatlar': yillik_fiyatlar
        })
        
    except ValueError:
        return jsonify({'success': False, 'message': 'Geçersiz yıl formatı'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/otomatik-fiyat-hesapla', methods=['POST'])
def otomatik_fiyat_hesapla():
    """Seçilen parametreler için otomatik fiyat hesaplar"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        secilen_parametreler = data.get('parametreler', [])
        yil = data.get('yil', 2025)  # Varsayılan 2025
        
        if not secilen_parametreler:
            return jsonify({'success': False, 'message': 'Seçilen parametre bulunamadı'})
        
        # Her parametre için fiyat hesapla
        hesaplanan_fiyatlar = []
        toplam_fiyat = 0
        
        for parametre in secilen_parametreler:
            parametre_adi = parametre.get('parametre', '')
            adet = parametre.get('adet', 1)
            
            # Yıllık fiyat tarifesinden fiyatı al
            birim_fiyat = get_parametre_fiyati(parametre_adi, yil)
            top_fiyat = birim_fiyat * adet
            
            hesaplanan_fiyatlar.append({
                'parametre': parametre_adi,
                'metot': parametre.get('metot', ''),
                'birimFiyat': birim_fiyat,
                'adet': adet,
                'topFiyat': top_fiyat
            })
            
            toplam_fiyat += top_fiyat
        
        return jsonify({
            'success': True,
            'parametreler': hesaplanan_fiyatlar,
            'toplam': toplam_fiyat,
            'yil': yil,
            'message': f'{len(secilen_parametreler)} parametre için {yil} yılı toplam {toplam_fiyat:,.2f} TL hesaplandı'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/baca-bilgileri/<firma_adi>', methods=['GET'])
def get_baca_bilgileri(firma_adi):
    """Firma için baca bilgilerini ve parametrelerini döndürür"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Firma adına göre filtrele
        firma_bacalar = [baca for baca in baca_bilgileri if baca.get('firma_adi', '').upper() == firma_adi.upper()]
        
        if not firma_bacalar:
            return jsonify({
                'success': False, 
                'message': f'{firma_adi} firması için baca bilgisi bulunamadı',
                'bacalar': []
            })
        
        # Her baca için parametreleri topla
        tum_parametreler = set()
        baca_detaylari = []
        
        for baca in firma_bacalar:
            baca_adi = baca.get('baca_adi', '')
            parametreler = []
            
            # Baca bilgilerinden parametreleri çıkar
            for key, value in baca.items():
                if key.startswith('parametre_') and value and value.strip():
                    parametre_adi = key.replace('parametre_', '').replace('_', ' ').upper()
                    parametreler.append(parametre_adi)
                    tum_parametreler.add(parametre_adi)
            
            baca_detaylari.append({
                'baca_adi': baca_adi,
                'parametreler': parametreler
            })
        
        return jsonify({
            'success': True,
            'firma_adi': firma_adi,
            'bacalar': baca_detaylari,
            'tum_parametreler': list(tum_parametreler),
            'message': f'{len(firma_bacalar)} baca bulundu, {len(tum_parametreler)} farklı parametre tespit edildi'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/teklif/update', methods=['POST'])
def update_teklif():
    """Teklifi günceller"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        teklif_id = data.get('id')
        
        if not teklif_id:
            return jsonify({'success': False, 'message': 'Teklif ID gerekli'})
        
        teklifler = load_teklif()
        teklif = next((t for t in teklifler if t.get('id') == teklif_id), None)
        
        if not teklif:
            return jsonify({'success': False, 'message': 'Teklif bulunamadı'})
        
        # Teklifi güncelle
        teklif.update({
            'teklif_tipi': data.get('teklif_tipi', teklif.get('teklif_tipi')),
            'firma_adi': data.get('firma_adi', teklif.get('firma_adi')),
            'teklif_tarihi': data.get('teklif_tarihi', teklif.get('teklif_tarihi')),
            'teklif_no': data.get('teklif_no', teklif.get('teklif_no')),
            'indirim_orani': data.get('indirim_orani', teklif.get('indirim_orani')),
            'indirim_tipi': data.get('indirim_tipi', teklif.get('indirim_tipi')),
            'parametreler': data.get('parametreler', teklif.get('parametreler', [])),
            'toplam': data.get('toplam', teklif.get('toplam', 0)),
            'indirim': data.get('indirim', teklif.get('indirim', 0)),
            'netToplam': data.get('netToplam', teklif.get('netToplam', 0)),
            'teklif_giris_metni': data.get('teklif_giris_metni', teklif.get('teklif_giris_metni', '')),
            'genel_hukumler': data.get('genel_hukumler', teklif.get('genel_hukumler', '')),
            'teklif_durumu': data.get('teklif_durumu', teklif.get('teklif_durumu', 'BEKLEMEDE')),
            'updated_at': datetime.now().isoformat()
        })
        
        if save_teklif(teklifler):
            return jsonify({'success': True, 'message': 'Teklif başarıyla güncellendi', 'teklif': teklif})
        else:
            return jsonify({'success': False, 'message': 'Teklif güncellenirken hata oluştu'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/teklif/update_status', methods=['POST'])
def update_teklif_status():
    """Teklif durumunu günceller"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        teklif_id = data.get('teklif_id')
        yeni_durum = data.get('teklif_durumu')
        durum_tarihi = data.get('durum_tarihi')
        
        if not teklif_id:
            return jsonify({'success': False, 'message': 'Teklif ID gerekli'})
        
        if not yeni_durum:
            return jsonify({'success': False, 'message': 'Yeni durum gerekli'})
        
        teklifler = load_teklif()
        teklif = next((t for t in teklifler if t.get('id') == teklif_id), None)
        
        if not teklif:
            return jsonify({'success': False, 'message': 'Teklif bulunamadı'})
        
        # Durumu ve durum tarihini güncelle
        teklif['teklif_durumu'] = yeni_durum
        if durum_tarihi:
            teklif['durum_tarihi'] = durum_tarihi
        teklif['updated_at'] = datetime.now().isoformat()
        
        if save_teklif(teklifler):
            return jsonify({'success': True, 'message': 'Teklif durumu başarıyla güncellendi', 'teklif': teklif})
        else:
            return jsonify({'success': False, 'message': 'Teklif durumu güncellenirken hata oluştu'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/teklif/update_durum_tarihi', methods=['POST'])
def update_teklif_durum_tarihi():
    """Teklif durum tarihini günceller"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        teklif_id = data.get('teklif_id')
        durum_tarihi = data.get('durum_tarihi')
        
        if not teklif_id:
            return jsonify({'success': False, 'message': 'Teklif ID gerekli'})
        
        teklifler = load_teklif()
        teklif = next((t for t in teklifler if t.get('id') == teklif_id), None)
        
        if not teklif:
            return jsonify({'success': False, 'message': 'Teklif bulunamadı'})
        
        # Sadece durum tarihini güncelle
        teklif['durum_tarihi'] = durum_tarihi
        teklif['updated_at'] = datetime.now().isoformat()
        
        if save_teklif(teklifler):
            return jsonify({'success': True, 'message': 'Durum tarihi başarıyla güncellendi', 'teklif': teklif})
        else:
            return jsonify({'success': False, 'message': 'Durum tarihi güncellenirken hata oluştu'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/firma_kayit/add', methods=['POST'])
def add_firma_kayit():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        
        # Yeni firma kaydı oluştur
        yeni_firma = {
            'id': str(uuid.uuid4()),
            'firmaAdi': data.get('firmaAdi', ''),
            'adres': data.get('adres', ''),
            'il': data.get('il', ''),
            'ilce': data.get('ilce', ''),
            'vergiDairesi': data.get('vergiDairesi', ''),
            'vergiNo': data.get('vergiNo', ''),
            'yetkiliAdi': data.get('yetkiliAdi', ''),
            'yetkiliTel': data.get('yetkiliTel', ''),
            'yetkiliMail': data.get('yetkiliMail', ''),
            'danismanAdi': data.get('danismanAdi', ''),
            'danismanMail': data.get('danismanMail', ''),
            'danismanTel': data.get('danismanTel', ''),
            'kayitTarihi': datetime.now().strftime('%d.%m.%Y')
        }
        
        # Mevcut verileri yükle
        firma_kayitlar = load_firma_kayit()
        firma_kayitlar.append(yeni_firma)
        
        # Verileri kaydet
        if save_firma_kayit(firma_kayitlar):
            return jsonify({'success': True, 'message': 'Firma başarıyla kaydedildi', 'firma': yeni_firma})
        else:
            return jsonify({'success': False, 'message': 'Firma kaydedilirken hata oluştu'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/firma_kayit/update', methods=['POST'])
def update_firma_kayit():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        firma_id = data.get('id')
        
        # Mevcut verileri yükle
        firma_kayitlar = load_firma_kayit()
        
        # Firma bul ve güncelle
        for firma in firma_kayitlar:
            if firma.get('id') == firma_id:
                firma.update({
                    'firmaAdi': data.get('firmaAdi', ''),
                    'adres': data.get('adres', ''),
                    'il': data.get('il', ''),
                    'ilce': data.get('ilce', ''),
                    'vergiDairesi': data.get('vergiDairesi', ''),
                    'vergiNo': data.get('vergiNo', ''),
                    'yetkiliAdi': data.get('yetkiliAdi', ''),
                    'yetkiliTel': data.get('yetkiliTel', ''),
                    'yetkiliMail': data.get('yetkiliMail', ''),
                    'danismanAdi': data.get('danismanAdi', ''),
                    'danismanMail': data.get('danismanMail', ''),
                    'danismanTel': data.get('danismanTel', '')
                })
                break
        
        # Verileri kaydet
        if save_firma_kayit(firma_kayitlar):
            return jsonify({'success': True, 'message': 'Firma başarıyla güncellendi'})
        else:
            return jsonify({'success': False, 'message': 'Firma güncellenirken hata oluştu'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/firma_kayit/delete', methods=['POST'])
def delete_firma_kayit():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        firma_id = data.get('id')
        
        # Mevcut verileri yükle
        firma_kayitlar = load_firma_kayit()
        
        # Firma bul ve sil
        firma_kayitlar = [firma for firma in firma_kayitlar if firma.get('id') != firma_id]
        
        # Verileri kaydet
        if save_firma_kayit(firma_kayitlar):
            return jsonify({'success': True, 'message': 'Firma başarıyla silindi'})
        else:
            return jsonify({'success': False, 'message': 'Firma silinirken hata oluştu'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/firma_kayit/import', methods=['POST'])
def import_firma_kayit():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        imported_firmalar = data.get('firmalar', [])
        
        if not imported_firmalar:
            return jsonify({'success': False, 'message': 'İçe aktarılacak firma verisi bulunamadı'})
        
        # Mevcut verileri yükle
        firma_kayitlar = load_firma_kayit()
        
        # Yeni firmaları ekle
        for firma in imported_firmalar:
            # ID kontrolü - eğer aynı ID varsa yeni ID oluştur
            if any(existing_firma.get('id') == firma.get('id') for existing_firma in firma_kayitlar):
                firma['id'] = str(uuid.uuid4())
            firma_kayitlar.append(firma)
        
        # Verileri kaydet
        if save_firma_kayit(firma_kayitlar):
            return jsonify({
                'success': True, 
                'message': f'{len(imported_firmalar)} adet firma başarıyla içe aktarıldı',
                'imported_count': len(imported_firmalar)
            })
        else:
            return jsonify({'success': False, 'message': 'Firmalar içe aktarılırken hata oluştu'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/firma_olcum')
def firma_olcum():
    if not session.get('logged_in') or not can_read(session.get('role')):
        return redirect(url_for('login'))
    
    firma_olcumler = load_firma_olcum()
    
    # Tarihleri GG.AA.YY formatında formatla
    for olcum in firma_olcumler:
        # Tarihleri GG.AA.YY formatında formatla
        if 'baslangic_tarihi' in olcum and olcum['baslangic_tarihi']:
            olcum['baslangic_tarihi'] = format_date_with_day(olcum['baslangic_tarihi'])
        if 'bitis_tarihi' in olcum and olcum['bitis_tarihi']:
            olcum['bitis_tarihi'] = format_date_with_day(olcum['bitis_tarihi'])
    
    # Parametreleri yükle
    parameters = load_parameters()
    
    return render_template('firma_olcum.html', 
                         firma_olcumler=firma_olcumler, 
                         username=session.get('username'), 
                         role=session.get('role'),
                         parameters=parameters)

@app.route('/firma_olcum/add_step1', methods=['GET', 'POST'])
def add_firma_olcum_step1():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        # Form verilerini al
        firma_adi = request.form.get('firma_adi', '').strip()
        olcum_kodu = request.form.get('olcum_kodu', '').strip()
        baslangic_tarihi = request.form.get('baslangic_tarihi', '')
        bitis_tarihi = request.form.get('bitis_tarihi', '')
        il = request.form.get('il', '')
        ilce = request.form.get('ilce', '')
        yetkili = request.form.get('yetkili', '')
        telefon = request.form.get('telefon', '')
        durum = request.form.get('durum', 'Aktif')
        secilen_personel = request.form.getlist('personel')
        
        # Zorunlu alanları kontrol et
        if not firma_adi or not olcum_kodu:
            flash('Firma adı ve ölçüm kodu zorunludur!', 'error')
            return render_template('add_firma_olcum_step1.html', username=session.get('username'), role=session.get('role'))
        
        # Session'a geçici veri kaydet
        session['temp_firma_olcum'] = {
            'firma_adi': firma_adi,
            'olcum_kodu': olcum_kodu,
            'baslangic_tarihi': baslangic_tarihi,
            'bitis_tarihi': bitis_tarihi,
            'il': il,
            'ilce': ilce,
            'yetkili': yetkili,
            'telefon': telefon,
            'durum': durum,
            'secilen_personel': secilen_personel
        }
        
        return redirect(url_for('add_firma_olcum_step2'))
    
    # Saha personeli listesini al (görevi "Saha" olan kullanıcılar)
    users = load_users()
    saha_personeli = []
    for username, user_data in users.items():
        if user_data.get('gorev') == 'Saha':
            saha_personeli.append({
                'username': username,
                'surname': user_data.get('surname', ''),
                'gorev': user_data.get('gorev', '')
            })
    
    # İstediğiniz illeri en üste taşı
    oncelikli_iller = ['KOCAELİ', 'SAKARYA', 'DÜZCE', 'BOLU', 'İSTANBUL', 'BURSA', 'BİLECİK', 'KÜTAHYA']
    
    # İlleri öncelik sırasına göre düzenle
    sirali_cities = []
    
    # Önce öncelikli illeri ekle
    for il_adi in oncelikli_iller:
        for il in CITIES_DATA:
            if il['il_adi'] == il_adi:
                sirali_cities.append(il)
                break
    
    # Sonra diğer illeri ekle
    for il in CITIES_DATA:
        if il['il_adi'] not in oncelikli_iller:
            sirali_cities.append(il)
    
    # Firma listesini yükle ve formatını düzenle
    firma_kayitlar = load_firma_kayit()
    
    # Firma verilerini template için uygun formata çevir
    formatted_firmalar = []
    for firma in firma_kayitlar:
        formatted_firma = {
            'firma_adi': firma.get('firmaAdi', ''),
            'yetkili_adi': firma.get('yetkiliAdi', ''),
            'yetkili_tel': firma.get('yetkiliTel', ''),
            'il': firma.get('il', ''),
            'ilce': firma.get('ilce', ''),
            'adres': firma.get('adres', ''),
            'vergi_dairesi': firma.get('vergiDairesi', ''),
            'vergi_no': firma.get('vergiNo', ''),
            'yetkili_mail': firma.get('yetkiliMail', ''),
            'danisman_adi': firma.get('danismanAdi', ''),
            'danisman_mail': firma.get('danismanMail', ''),
            'danisman_tel': firma.get('danismanTel', ''),
            'kayit_tarihi': firma.get('kayitTarihi', '')
        }
        formatted_firmalar.append(formatted_firma)
    
    return render_template('add_firma_olcum_step1.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         saha_personeli=saha_personeli,
                         cities=sirali_cities,
                         firmalar=formatted_firmalar)

@app.route('/firma_olcum/add_step2', methods=['GET', 'POST'])
def add_firma_olcum_step2():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    # Session'dan geçici veriyi al
    temp_data = session.get('temp_firma_olcum')
    if not temp_data:
        return redirect(url_for('add_firma_olcum_step1'))
    
    if request.method == 'POST':
        try:
            # 2. aşama verilerini al
            baca_sayisi = request.form.get('baca_sayisi', '')
            notlar = request.form.get('notlar', '')
            baca_parametreleri_json = request.form.get('baca_parametreleri', '{}')
            
            # Baca sayısı kontrolü
            if not baca_sayisi or int(baca_sayisi) < 1:
                raise ValueError('Baca sayısı en az 1 olmalıdır')
            
            # JSON'dan baca parametrelerini parse et
            try:
                baca_parametreleri = json.loads(baca_parametreleri_json)
            except json.JSONDecodeError:
                baca_parametreleri = {}
            
            # Yeni kayıt oluştur
            yeni_kayit = {
                'id': str(uuid4()),
                'firma_adi': temp_data['firma_adi'],
                'olcum_kodu': temp_data['olcum_kodu'],
                'baslangic_tarihi': temp_data['baslangic_tarihi'],
                'bitis_tarihi': temp_data['bitis_tarihi'],
                'il': temp_data['il'],
                'ilce': temp_data['ilce'],
                'yetkili': temp_data['yetkili'],
                'telefon': temp_data['telefon'],
                'durum': temp_data['durum'],
                'personel': temp_data['secilen_personel'],
                'baca_sayisi': baca_sayisi,
                'baca_parametreleri': baca_parametreleri,
                'notlar': notlar,
                'olusturma_tarihi': datetime.now().isoformat()
            }
            
            # Veriyi kaydet
            firma_olcumler = load_firma_olcum()
            firma_olcumler.append(yeni_kayit)
            if not save_firma_olcum(firma_olcumler):
                raise Exception('Veriler kaydedilemedi')
            
            # Session'dan geçici veriyi temizle
            session.pop('temp_firma_olcum', None)
            
            # AJAX isteği ise JSON döndür
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': True, 'message': 'Firma ölçüm kaydı başarıyla oluşturuldu!'})
            
            flash('Firma ölçüm kaydı başarıyla oluşturuldu!', 'success')
            return redirect(url_for('firma_olcum'))
            
        except Exception as e:
            print(f"Kayıt hatası: {e}")
            import traceback
            traceback.print_exc()
            
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'error': f'Kayıt hatası: {str(e)}'})
            
            flash(f'Kayıt sırasında hata oluştu: {str(e)}', 'error')
            return redirect(url_for('add_firma_olcum_step2'))
    
    # Parametreleri yükle
    parameters = load_parameters()
    
    return render_template('add_firma_olcum_step2.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         temp_data=temp_data,
                         parameters=parameters)

@app.route('/firma_olcum/detail/<olcum_id>')
def firma_olcum_detail(olcum_id):
    if not session.get('logged_in') or not can_read(session.get('role')):
        return redirect(url_for('login'))
    
    firma_olcumler = load_firma_olcum()
    olcum = next((o for o in firma_olcumler if o['id'] == olcum_id), None)
    
    if not olcum:
        flash('Ölçüm kaydı bulunamadı!', 'error')
        return redirect(url_for('firma_olcum'))
    
    # Parametreleri yükle
    parameters = load_parameters()
    
    # Sadece seçili parametrelerin benzersiz isimlerini al
    selected_param_names = set()
    if olcum.get('baca_parametreleri'):
        for baca_adi, parametreler in olcum['baca_parametreleri'].items():
            for param_name in parametreler:
                # Eğer parametre tam format ise (Parametre Adı (Metot)), sadece parametre adını al
                if ' (' in param_name and param_name.endswith(')'):
                    param_adi = param_name.split(' (')[0]
                    selected_param_names.add(param_adi)
                else:
                    # Kısa isim ise, olduğu gibi kullan
                    selected_param_names.add(param_name)
    
    # Benzersiz parametre listesi oluştur (sadece seçili parametreler)
    unique_parameters = []
    for param_name in sorted(selected_param_names):
        # Her parametre için kaç bacada seçildiğini hesapla
        param_count = 0
        if olcum.get('baca_parametreleri'):
            for baca_adi, parametreler in olcum['baca_parametreleri'].items():
                if param_name in parametreler:
                    param_count += 1
        
        unique_parameters.append({
            'name': param_name,
            'full_name': param_name,
            'count': param_count
        })
    

    
    return render_template('firma_olcum_detail.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         olcum=olcum,
                         parameters=unique_parameters)

@app.route('/firma_olcum/delete/<olcum_id>', methods=['POST'])
def delete_firma_olcum(olcum_id):
    if not session.get('logged_in') or not can_delete(session.get('role')):
        return jsonify({'success': False, 'error': 'Yetkisiz erişim'}), 401
    
    try:
        firma_olcumler = load_firma_olcum()
        olcum = next((o for o in firma_olcumler if o['id'] == olcum_id), None)
        
        if not olcum:
            return jsonify({'success': False, 'error': 'Ölçüm kaydı bulunamadı'}), 404
        
        # Kaydı sil
        firma_olcumler = [o for o in firma_olcumler if o['id'] != olcum_id]
        
        if save_firma_olcum(firma_olcumler):
            return jsonify({'success': True, 'message': 'Ölçüm kaydı başarıyla silindi'})
        else:
            return jsonify({'success': False, 'error': 'Kayıt silinirken hata oluştu'}), 500
            
    except Exception as e:
        print(f"Silme hatası: {e}")
        return jsonify({'success': False, 'error': f'Silme hatası: {str(e)}'}), 500

@app.route('/firma_olcum/delete_selected', methods=['POST'])
def delete_selected_firma_olcum():
    if not session.get('logged_in') or not can_delete(session.get('role')):
        return jsonify({'success': False, 'error': 'Yetkisiz erişim'}), 401
    
    try:
        data = request.get_json()
        ids = data.get('ids', []) if data else []
        
        if not ids:
            return jsonify({'success': False, 'error': 'Silinecek kayıt seçilmedi'}), 400
        
        firma_olcumler = load_firma_olcum()
        original_count = len(firma_olcumler)
        
        # Seçilen kayıtları sil
        firma_olcumler = [o for o in firma_olcumler if o['id'] not in ids]
        
        if save_firma_olcum(firma_olcumler):
            deleted_count = original_count - len(firma_olcumler)
            return jsonify({'success': True, 'message': f'{deleted_count} kayıt başarıyla silindi'})
        else:
            return jsonify({'success': False, 'error': 'Kayıtlar silinirken hata oluştu'}), 500
            
    except Exception as e:
        print(f"Toplu silme hatası: {e}")
        return jsonify({'success': False, 'error': f'Silme hatası: {str(e)}'}), 500

@app.route('/firma_olcum/edit/<olcum_id>', methods=['GET', 'POST'])
def edit_firma_olcum(olcum_id):
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    # Firma ölçüm verilerini yükle
    firma_olcumler = load_firma_olcum()
    olcum = next((o for o in firma_olcumler if o['id'] == olcum_id), None)
    
    if not olcum:
        flash('Ölçüm kaydı bulunamadı!', 'error')
        return redirect(url_for('firma_olcum'))
    
    if request.method == 'POST':
        try:
            # Form verilerini al
            firma_adi = request.form.get('firma_adi', '').strip()
            olcum_kodu = request.form.get('olcum_kodu', '').strip()
            baslangic_tarihi = request.form.get('baslangic_tarihi', '')
            bitis_tarihi = request.form.get('bitis_tarihi', '')
            il = request.form.get('il', '')
            ilce = request.form.get('ilce', '')
            yetkili = request.form.get('yetkili', '')
            telefon = request.form.get('telefon', '')
            durum = request.form.get('durum', 'Aktif')
            secilen_personel = request.form.getlist('personel')
            baca_sayisi = request.form.get('baca_sayisi', '')
            notlar = request.form.get('notlar', '')
            baca_parametreleri_json = request.form.get('baca_parametreleri', '{}')
            
            # Zorunlu alanları kontrol et
            if not firma_adi or not olcum_kodu:
                flash('Firma adı ve ölçüm kodu zorunludur!', 'error')
                return render_template('edit_firma_olcum.html', 
                                     username=session.get('username'), 
                                     role=session.get('role'),
                                     olcum=olcum,
                                     cities=CITIES_DATA)
            
            # JSON'dan baca parametrelerini parse et
            try:
                baca_parametreleri = json.loads(baca_parametreleri_json)
            except json.JSONDecodeError:
                baca_parametreleri = {}
            
            # Kaydı güncelle
            olcum['firma_adi'] = firma_adi
            olcum['olcum_kodu'] = olcum_kodu
            olcum['baslangic_tarihi'] = baslangic_tarihi
            olcum['bitis_tarihi'] = bitis_tarihi
            olcum['il'] = il
            olcum['ilce'] = ilce
            olcum['yetkili'] = yetkili
            olcum['telefon'] = telefon
            olcum['durum'] = durum
            olcum['personel'] = secilen_personel
            olcum['baca_sayisi'] = baca_sayisi
            olcum['baca_parametreleri'] = baca_parametreleri
            olcum['notlar'] = notlar
            
            # Veriyi kaydet
            if save_firma_olcum(firma_olcumler):
                flash('Firma ölçüm kaydı başarıyla güncellendi!', 'success')
                return redirect(url_for('firma_olcum'))
            else:
                flash('Kayıt güncellenirken hata oluştu!', 'error')
                
        except Exception as e:
            print(f"Güncelleme hatası: {e}")
            flash(f'Güncelleme sırasında hata oluştu: {str(e)}', 'error')
    
    # Saha personeli listesini al
    users = load_users()
    saha_personeli = []
    for username, user_data in users.items():
        if user_data.get('gorev') == 'Saha':
            saha_personeli.append({
                'username': username,
                'surname': user_data.get('surname', ''),
                'gorev': user_data.get('gorev', '')
            })
    
    # İstediğiniz illeri en üste taşı
    oncelikli_iller = ['KOCAELİ', 'SAKARYA', 'DÜZCE', 'BOLU', 'İSTANBUL', 'BURSA', 'BİLECİK', 'KÜTAHYA']
    
    # İlleri öncelik sırasına göre düzenle
    sirali_cities = []
    
    # Önce öncelikli illeri ekle
    for il_adi in oncelikli_iller:
        for il in CITIES_DATA:
            if il['il_adi'] == il_adi:
                sirali_cities.append(il)
                break
    
    # Sonra diğer illeri ekle
    for il in CITIES_DATA:
        if il['il_adi'] not in oncelikli_iller:
            sirali_cities.append(il)
    
    # Parametreleri yükle
    parameters = load_parameters()
    
    return render_template('edit_firma_olcum.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         olcum=olcum,
                         saha_personeli=saha_personeli,
                         cities=sirali_cities,
                         parameters=parameters)

@app.route('/saha_olc')
def saha_olc():
    """Saha ölçüm sayfası"""
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    saha_olcumler = load_saha_olc()
    return render_template('saha_olc.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         saha_olcumler=saha_olcumler)

@app.route('/kk_egri')
def kk_egri():
    """Kalite kontrol eğrisi sayfası"""
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    return render_template('kk_egri.html', 
                         username=session.get('username'), 
                         role=session.get('role'))


@app.route('/api/kk_parametreler')
def api_kk_parametreler():
    """KK parametrelerini döndür"""
    try:
        # Sabit KK parametreleri
        kk_parametreler = [
            {'id': 'o2', 'parametre_adi': 'O2', 'kk': '7.0'},
            {'id': 'co', 'parametre_adi': 'CO', 'kk': '500'},
            {'id': 'no', 'parametre_adi': 'NO', 'kk': '500'},
            {'id': 'so2', 'parametre_adi': 'SO2', 'kk': '500'},
            {'id': 'toc', 'parametre_adi': 'TOC', 'kk': '100'}
        ]

        return jsonify(kk_parametreler)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/kk_grafik_olustur', methods=['POST'])
def api_kk_grafik_olustur():
    """Kalite kontrol grafiği oluştur"""
    try:
        data = request.get_json()
        parametre = data.get('parametre')
        baslangic_tarih = data.get('baslangic_tarih')
        bitis_tarih = data.get('bitis_tarih')
        localStorage_verileri = data.get('localStorage_verileri', [])  # localStorage verilerini al
        
        if not parametre:
            return jsonify({'error': 'Parametre seçilmedi'}), 400
        
        # Parametre verilerini yükle
        parameters = load_parameters()
        
        # Parametre yönetiminden KK değerlerini dinamik olarak al
        parametre_eslesme = {
            'SO2': 'So2',
            'NO': 'No', 
            'CO': 'Co',
            'O2': 'O2',
            'TOC': 'TOC'
        }
        
        # Parametre adını eşleştir
        metot_adi = parametre_eslesme.get(parametre, parametre)
        
        # Parametre yönetiminden KK değerlerini bul
        kk_degerleri = None
        for param in parameters:
            if param.get('Metot') == metot_adi and param.get('KK'):
                try:
                    kk_degerleri = {
                        'kk': float(param.get('KK', 0)),
                        'minus_3s': float(param.get('-3S', 0)),
                        'minus_2s': float(param.get('-2S', 0)),
                        'plus_2s': float(param.get('+2S', 0)),
                        'plus_3s': float(param.get('+3S', 0))
                    }
                    break
                except (ValueError, TypeError):
                    continue
        
        # Eğer parametre yönetiminde bulunamazsa, varsayılan değerleri kullan
        if not kk_degerleri:
            varsayilan_degerler = {
                'O2': {'kk': 7.0, 'minus_3s': 6.6, 'minus_2s': 6.8, 'plus_2s': 7.2, 'plus_3s': 7.4},
                'CO': {'kk': 500, 'minus_3s': 400, 'minus_2s': 450, 'plus_2s': 550, 'plus_3s': 600},
                'NO': {'kk': 500, 'minus_3s': 400, 'minus_2s': 450, 'plus_2s': 550, 'plus_3s': 600},
                'SO2': {'kk': 500, 'minus_3s': 400, 'minus_2s': 450, 'plus_2s': 550, 'plus_3s': 600},
                'TOC': {'kk': 100, 'minus_3s': 80, 'minus_2s': 90, 'plus_2s': 110, 'plus_3s': 120}
            }
            kk_degerleri = varsayilan_degerler.get(parametre, {})
        
        if not kk_degerleri:
            return jsonify({'error': f'{parametre} parametresi için KK değerleri bulunamadı'}), 400
        
        # KK değerlerini al
        kk = kk_degerleri['kk']
        minus_3s = kk_degerleri['minus_3s']
        minus_2s = kk_degerleri['minus_2s']
        plus_2s = kk_degerleri['plus_2s']
        plus_3s = kk_degerleri['plus_3s']
        
        # Tarih aralığında veriler oluştur
        from datetime import datetime, timedelta
        
        baslangic = datetime.strptime(baslangic_tarih, '%Y-%m-%d')
        bitis = datetime.strptime(bitis_tarih, '%Y-%m-%d')
        
        # Saha ölçümlerinden KK verilerini al
        parametre_olcumleri = load_parametre_olcum()
        baca_bilgileri = load_baca_bilgileri()
        gercek_tarihler = []
        gercek_degerler = []
        gercek_firmalar = []
        gercek_kodlar = []
        gercek_bacalar = []  # Baca listesi ekle
        gercek_cihazlar = []
        gercek_personeller = []
        
        # localStorage verilerini işle
        for veri in localStorage_verileri:
            try:
                # Tarih formatını kontrol et
                tarih_str = veri.get('tarih', '')
                if tarih_str:
                    olcum_tarih = datetime.strptime(tarih_str, '%Y-%m-%d')
                    
                    if baslangic <= olcum_tarih <= bitis:
                        # Parametre tipine göre değeri al
                        kk_deger = None
                        
                        if parametre == 'O2':
                            kk_deger = veri.get('o2')
                        elif parametre == 'CO':
                            kk_deger = veri.get('co')
                        elif parametre == 'NO':
                            kk_deger = veri.get('no')
                        elif parametre == 'SO2':
                            kk_deger = veri.get('so2')
                        elif parametre == 'TOC':
                            kk_deger = veri.get('toc')
                        
                        if kk_deger is not None and kk_deger != '' and str(kk_deger).strip() != '':
                            try:
                                # Virgülü noktaya çevir (float için gerekli)
                                deger_str = str(kk_deger).replace(',', '.')
                                deger = float(deger_str)
                                # Değeri virgül ile formatla (2 ondalık basamak)
                                deger = round(deger, 2)
                                
                                gercek_tarihler.append(olcum_tarih)
                                gercek_degerler.append(deger)
                                gercek_firmalar.append(veri.get('firma', ''))
                                gercek_kodlar.append(veri.get('kod', ''))
                                gercek_bacalar.append(veri.get('baca', ''))  # Baca bilgisini ekle
                                gercek_cihazlar.append(veri.get('cihaz', ''))
                                gercek_personeller.append(veri.get('personel', 'Admin'))
                            except ValueError:
                                continue
            except:
                continue
        
        # Seçilen parametre için saha ölçümlerini filtrele
        for olcum in parametre_olcumleri:
            # Tarih aralığında mı kontrol et
            try:
                # Tarih formatını kontrol et
                tarih_str = olcum.get('tarih', '')
                if not tarih_str:
                    # parametre_verileri içindeki TARİH'i kontrol et
                    parametre_verileri = olcum.get('parametre_verileri', {})
                    tarih_str = parametre_verileri.get('TARİH', '')
                
                if tarih_str:
                    # GG.AA.YY formatını YYYY-MM-DD'ye çevir
                    if '.' in tarih_str and len(tarih_str) == 8:
                        olcum_tarih = datetime.strptime(tarih_str, '%d.%m.%y')
                    else:
                        olcum_tarih = datetime.strptime(tarih_str, '%Y-%m-%d')
                    
                    if baslangic <= olcum_tarih <= bitis:
                        # Parametre tipine göre KK değerini al
                        kk_deger = None
                        parametre_verileri = olcum.get('parametre_verileri', {})
                    
                        if parametre == 'O2':
                            kk_deger = parametre_verileri.get('KK1-O2')
                        elif parametre == 'CO':
                            kk_deger = parametre_verileri.get('KK1-CO')
                        elif parametre == 'NO':
                            kk_deger = parametre_verileri.get('KK1-NO')
                        elif parametre == 'SO2':
                            kk_deger = parametre_verileri.get('KK1-SO2')
                        elif parametre == 'TOC':
                            kk_deger = parametre_verileri.get('KK1-SPAN')
                        
                        if kk_deger and str(kk_deger).strip() and str(kk_deger) != '':
                            try:
                                # Tire ile ayrılmış değerler varsa (örn: 7,02-7,06-7,04) ilkini al
                                deger_str = str(kk_deger)
                                if '-' in deger_str:
                                    deger_str = deger_str.split('-')[0]
                                
                                # Virgülü noktaya çevir (float için gerekli)
                                deger_str = deger_str.replace(',', '.')
                                deger = float(deger_str)
                                
                                # Baca bilgilerinden personel adını bul
                                personel_adi = 'Admin'  # Varsayılan
                                for baca_bilgi in baca_bilgileri:
                                    if (baca_bilgi.get('firma_adi') == olcum.get('firma_adi') and 
                                        baca_bilgi.get('olcum_kodu') == olcum.get('olcum_kodu') and
                                        baca_bilgi.get('baca_adi') == olcum.get('baca_adi')):
                                        personel_adi = baca_bilgi.get('personel_adi', 'Admin')
                                        break
                                
                                gercek_tarihler.append(olcum_tarih)
                                gercek_degerler.append(deger)
                                gercek_firmalar.append(olcum.get('firma_adi', ''))
                                gercek_kodlar.append(olcum.get('olcum_kodu', ''))
                                gercek_bacalar.append(olcum.get('baca_adi', ''))  # Baca adını ekle
                                gercek_cihazlar.append(olcum.get('cihaz_adi', '-'))  # Cihaz adını doğru alanından al
                                gercek_personeller.append(personel_adi)
                            except ValueError:
                                continue
            except:
                continue
        
        # Standart sapma hesapla (KK değerinden)
        standart_sapma = (plus_3s - kk) / 3  # +3S = KK + 3σ
        
        # Gerçek veriler varsa onları kullan, yoksa boş liste
        if gercek_tarihler and gercek_degerler:
            tarihler = gercek_tarihler
            degerler = gercek_degerler
        else:
            # Veri yoksa sadece başlangıç ve bitiş tarihlerini göster
            tarihler = [baslangic, bitis]
            degerler = [None, None]
        
        # Matplotlib'i lazy loading ile yükle
        plt, np, mdates, Rectangle = load_matplotlib()
        
        # Grafik oluştur (çok daha geniş ve yüksek)
        plt.figure(figsize=(18, 8))
        plt.rcParams['font.size'] = 12  # Font boyutu
        
        # Gerçek verileri mavi nokta olarak çiz
        if gercek_tarihler and gercek_degerler:
            # Tarihleri sırala
            sorted_data = sorted(zip(gercek_tarihler, gercek_degerler))
            sorted_tarihler = [x[0] for x in sorted_data]
            sorted_degerler = [x[1] for x in sorted_data]
            
            plt.plot(sorted_tarihler, sorted_degerler, 'bo', markersize=16, linewidth=4, label='Ölçüm Değerleri')
            # Noktaları birleştiren çizgi
            plt.plot(sorted_tarihler, sorted_degerler, 'b-', linewidth=3, alpha=0.9)
        
        # Ana veri çizgisi (gerçek veri olmadığında çizilmez)
        # plt.plot(tarihler, degerler, 'b-o', linewidth=2, markersize=6)
        
        # Outlier'ı kırmızı ile işaretle (gerçek veri olmadığında çizilmez)
        # if len(degerler) >= 10:
        #     plt.plot(tarihler[9], degerler[9], 'ro', markersize=8)
        
        # Kontrol limitleri (çok daha kalın çizgiler)
        plt.axhline(y=plus_3s, color='red', linestyle='--', linewidth=4)
        plt.axhline(y=kk, color='green', linewidth=4)
        plt.axhline(y=minus_3s, color='red', linestyle='--', linewidth=4)
        
        # 2-sigma limitleri (çok daha kalın mavi tireli çizgiler)
        plt.axhline(y=plus_2s, color='blue', linestyle='--', linewidth=4)
        plt.axhline(y=minus_2s, color='blue', linestyle='--', linewidth=4)
        
        # Kontrol limitlerini Y eksenine yaz (çok daha büyük font)
        plt.text(-0.05, float(plus_3s), '+3S', fontsize=16, ha='right', va='center', color='red', fontweight='bold', transform=plt.gca().get_yaxis_transform())
        plt.text(-0.05, float(plus_2s), '+2S', fontsize=16, ha='right', va='center', color='blue', fontweight='bold', transform=plt.gca().get_yaxis_transform())
        plt.text(-0.05, float(kk), 'KK', fontsize=16, ha='right', va='center', color='green', fontweight='bold', transform=plt.gca().get_yaxis_transform())
        plt.text(-0.05, float(minus_2s), '-2S', fontsize=16, ha='right', va='center', color='blue', fontweight='bold', transform=plt.gca().get_yaxis_transform())
        plt.text(-0.05, float(minus_3s), '-3S', fontsize=16, ha='right', va='center', color='red', fontweight='bold', transform=plt.gca().get_yaxis_transform())
        
        # 1-sigma limitleri (hesaplanmış)
        plus_1s = kk + standart_sapma
        minus_1s = kk - standart_sapma
        plt.axhline(y=plus_1s, color='gray', linestyle=':', linewidth=1, alpha=0.5)
        plt.axhline(y=minus_1s, color='gray', linestyle=':', linewidth=1, alpha=0.5)
        
        # Bölgeleri işaretle (lejant olmadan) - sabit değerlerle
        plt.fill_between(range(len(tarihler)), plus_3s, plus_2s, alpha=0.1, color='red')
        plt.fill_between(range(len(tarihler)), plus_2s, plus_1s, alpha=0.1, color='orange')
        plt.fill_between(range(len(tarihler)), plus_1s, minus_1s, alpha=0.1, color='green')
        plt.fill_between(range(len(tarihler)), minus_1s, minus_2s, alpha=0.1, color='orange')
        plt.fill_between(range(len(tarihler)), minus_2s, minus_3s, alpha=0.1, color='red')
        
        # Grafik ayarları
        # plt.xlabel('Tarih', fontsize=12, fontweight='bold')  # X ekseni etiketi kaldırıldı
        # plt.ylabel('Değer', fontsize=12, fontweight='bold')  # Y ekseni etiketi kaldırıldı
        
        # Başlık ve kontrol limitleri (çok daha büyük font)
        title_text = f'{parametre} KK GRAF.    KK: {kk}    -2S: {minus_2s}    +2S: {plus_2s}    -3S: {minus_3s}    +3S: {plus_3s}'
        plt.title(title_text, fontsize=18, fontweight='bold', loc='left')
        plt.grid(True, alpha=0.3)
        # Lejantı kaldır
        # plt.legend(loc='upper right')
        
        # X ekseni formatı - sadece veri olan tarihleri göster
        plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%d/%m'))
        # Tüm tarihleri tick olarak göster (çok daha büyük font)
        plt.xticks(tarihler, rotation=90, fontsize=14)
        
        # X ekseni limitlerini ayarla - veri olan tarihlerin öncesi ve sonrası için çok daha fazla boşluk
        if tarihler:
            min_date = min(tarihler)
            max_date = max(tarihler)
            # 10 gün öncesi ve sonrası ekle (çok daha geniş görünüm)
            plt.xlim(min_date - timedelta(days=10), max_date + timedelta(days=10))
        
        # Y ekseni limitleri (parametreye göre özelleştirilmiş)
        if parametre == 'TOC':
            # TOC için sabit aralık: 91-109
            y_min = 91
            y_max = 109
        else:
            # Diğer parametreler için dinamik aralık
            # -3S ve +3S değerlerinin üstünde ve altında %30 boşluk
            margin = (plus_3s - minus_3s) * 0.3
            y_min = minus_3s - margin
            y_max = plus_3s + margin
        
        plt.ylim(y_min, y_max)
        
        # Y ekseni tick'lerini dinamik olarak ayarla (çok daha büyük font)
        if parametre == 'O2':
            plt.yticks([6.6, 6.8, 7.0, 7.2, 7.4], fontsize=14)
        elif parametre == 'TOC':
            # TOC için daha dar aralık: 91-109
            plt.yticks([91, 94, 97, 100, 103, 106, 109], fontsize=14)
        else:
            # CO, NO, SO2 için dinamik tick'ler
            step = (y_max - y_min) / 8  # 8 eşit aralık
            ticks = []
            for i in range(9):
                tick_value = y_min + (i * step)
                ticks.append(round(tick_value, 1))
            plt.yticks(ticks, fontsize=14)
        
        # Grafik alanını temizle ve kaydet
        plt.tight_layout(pad=3.0)  # Çok daha fazla boşluk
        
        # Grafiği base64 formatında döndür
        from io import BytesIO
        import base64
        
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)
        img_base64 = base64.b64encode(img_buffer.getvalue()).decode()
        plt.close()
        
        # RAM temizleme
        gc.collect()
        plt.clf()
        plt.cla()
        
        # İstatistiksel veriler (gerçek veriler varsa onları kullan)
        if gercek_degerler:
            min_deger = min(gercek_degerler)
            max_deger = max(gercek_degerler)
            ortalama = sum(gercek_degerler) / len(gercek_degerler)
        else:
            min_deger = '-'
            max_deger = '-'
            ortalama = round(float(kk), 1)
        
        # Gerçek verileri tablo için hazırla
        gercek_veriler = []
        if gercek_tarihler and gercek_degerler:
            for i, tarih in enumerate(gercek_tarihler):
                gercek_veriler.append({
                    'tarih': tarih.strftime('%Y-%m-%d'),
                    'firma': gercek_firmalar[i] if i < len(gercek_firmalar) else '',
                    'kod': gercek_kodlar[i] if i < len(gercek_kodlar) else '',
                    'baca': gercek_bacalar[i] if i < len(gercek_bacalar) else '',
                    'deger': round(gercek_degerler[i], 1),
                    'cihaz': gercek_cihazlar[i] if i < len(gercek_cihazlar) else '',
                    'personel': gercek_personeller[i] if i < len(gercek_personeller) else 'Admin'
                })
        
        istatistikler = {
            'ortalama': round(float(ortalama), 1) if isinstance(ortalama, (int, float)) else ortalama,
            'standart_sapma': round(float(standart_sapma), 1),
            'min_deger': round(float(min_deger), 1) if isinstance(min_deger, (int, float)) else min_deger,
            'max_deger': round(float(max_deger), 1) if isinstance(max_deger, (int, float)) else max_deger,
            'ucl': round(float(plus_3s), 1),
            'cl': round(float(kk), 1),
            'lcl': round(float(minus_3s), 1),
            'gercek_veriler': gercek_veriler
        }
        
        return jsonify({
            'grafik': img_base64,
            'istatistikler': istatistikler,
            'veriler': {
                'tarihler': [t.strftime('%Y-%m-%d') for t in tarihler],
                'degerler': [round(d, 1) if d is not None else None for d in degerler]
            }
        })
        
    except Exception as e:
        print(f"KK grafik oluşturma hatası: {e}")
        return jsonify({'error': f'Grafik oluşturulurken hata oluştu: {str(e)}'}), 500

@app.route('/api/kk_rapor_olustur', methods=['POST'])
def api_kk_rapor_olustur():
    """Kalite kontrol raporu oluştur"""
    try:
        data = request.get_json()
        parametre = data.get('parametre')
        tarih_baslangic = data.get('tarih_baslangic')
        tarih_bitis = data.get('tarih_bitis')
        format_type = data.get('format', 'excel')
        
        if not parametre:
            return jsonify({'error': 'Parametre seçilmedi'}), 400
        
        # Parametre verilerini yükle
        parameters = load_parameters()
        
        # Seçilen parametre için KK değerlerini parametre yönetiminden al
        kk_degerleri = {}
        
        # Parametre yönetiminden KK değerlerini al
        parametre_eslesme = {
            'SO2': 'So2',
            'NO': 'No', 
            'CO': 'Co',
            'O2': 'O2',
            'TOC': 'TOC'
        }
        
        # Parametre adını eşleştir
        metot_adi = parametre_eslesme.get(parametre, parametre)
        
        # Parametre yönetiminden KK değerlerini bul
        for param in parameters:
            if param.get('Metot') == metot_adi and param.get('KK'):
                try:
                    kk_degerleri[parametre] = {
                        'kk': float(param.get('KK', 0)),
                        'minus_3s': float(param.get('-3S', 0)),
                        'minus_2s': float(param.get('-2S', 0)),
                        'plus_2s': float(param.get('+2S', 0)),
                        'plus_3s': float(param.get('+3S', 0))
                    }
                    break
                except (ValueError, TypeError):
                    continue
        
        # Eğer parametre yönetiminde bulunamazsa, varsayılan değerleri kullan
        if parametre not in kk_degerleri:
            varsayilan_degerler = {
                'O2': {'kk': 7.0, 'minus_3s': 6.6, 'minus_2s': 6.8, 'plus_2s': 7.2, 'plus_3s': 7.4},
                'CO': {'kk': 500, 'minus_3s': 400, 'minus_2s': 450, 'plus_2s': 550, 'plus_3s': 600},
                'NO': {'kk': 500, 'minus_3s': 400, 'minus_2s': 450, 'plus_2s': 550, 'plus_3s': 600},
                'SO2': {'kk': 500, 'minus_3s': 400, 'minus_2s': 450, 'plus_2s': 550, 'plus_3s': 600},
                'TOC': {'kk': 100, 'minus_3s': 80, 'minus_2s': 90, 'plus_2s': 110, 'plus_3s': 120}
            }
            kk_degerleri[parametre] = varsayilan_degerler.get(parametre, {})
        
        if parametre not in kk_degerleri:
            return jsonify({'error': f'{parametre} parametresi için KK değerleri tanımlanmamış'}), 400
        
        # KK değerlerini al
        degerler = kk_degerleri[parametre]
        kk = degerler['kk']
        minus_3s = degerler['minus_3s']
        minus_2s = degerler['minus_2s']
        plus_2s = degerler['plus_2s']
        plus_3s = degerler['plus_3s']
        
        # Tarih aralığında veriler oluştur
        from datetime import datetime, timedelta
        
        baslangic = datetime.strptime(tarih_baslangic, '%Y-%m-%d')
        bitis = datetime.strptime(tarih_bitis, '%Y-%m-%d')
        
        # Sağdaki tabloda görünen gerçek verileri al
        gercek_veriler = []
        
        # Saha ölçümlerinden KK verilerini al (api_kk_grafik_olustur ile aynı mantık)
        parametre_olcumleri = load_parametre_olcum()
        baca_bilgileri = load_baca_bilgileri()
        
        # localStorage verilerini de dahil et (manuel elle girilen veriler)
        # Bu veriler frontend'den gönderilmeli, şimdilik boş bırakıyoruz
        # localStorage_verileri = data.get('localStorage_verileri', [])
        
        # Seçilen parametre için saha ölçümlerini filtrele
        for olcum in parametre_olcumleri:
            # Tarih aralığında mı kontrol et
            try:
                # Tarih formatını kontrol et
                tarih_str = olcum.get('tarih', '')
                if not tarih_str:
                    # parametre_verileri içindeki TARİH'i kontrol et
                    parametre_verileri = olcum.get('parametre_verileri', {})
                    tarih_str = parametre_verileri.get('TARİH', '')
                
                if tarih_str:
                    # GG.AA.YY formatını YYYY-MM-DD'ye çevir
                    if '.' in tarih_str and len(tarih_str) == 8:
                        olcum_tarih = datetime.strptime(tarih_str, '%d.%m.%y')
                    else:
                        olcum_tarih = datetime.strptime(tarih_str, '%Y-%m-%d')
                    
                    if baslangic <= olcum_tarih <= bitis:
                        # Parametre tipine göre KK değerini al
                        kk_deger = None
                        parametre_verileri = olcum.get('parametre_verileri', {})
                    
                        if parametre == 'O2':
                            kk_deger = parametre_verileri.get('KK1-O2')
                        elif parametre == 'CO':
                            kk_deger = parametre_verileri.get('KK1-CO')
                        elif parametre == 'NO':
                            kk_deger = parametre_verileri.get('KK1-NO')
                        elif parametre == 'SO2':
                            kk_deger = parametre_verileri.get('KK1-SO2')
                        elif parametre == 'TOC':
                            kk_deger = parametre_verileri.get('KK1-SPAN')
                        
                        if kk_deger and str(kk_deger).strip() and str(kk_deger) != '0':
                            try:
                                # Tire ile ayrılmış değerler varsa (örn: 7,02-7,06-7,04) ilkini al
                                deger_str = str(kk_deger)
                                if '-' in deger_str:
                                    deger_str = deger_str.split('-')[0]
                                
                                # Virgülü noktaya çevir (float için gerekli)
                                deger_str = deger_str.replace(',', '.')
                                deger = float(deger_str)
                                
                                # Baca bilgilerinden personel adını bul
                                personel_adi = 'Admin'  # Varsayılan
                                for baca_bilgi in baca_bilgileri:
                                    if (baca_bilgi.get('firma_adi') == olcum.get('firma_adi') and 
                                        baca_bilgi.get('olcum_kodu') == olcum.get('olcum_kodu') and
                                        baca_bilgi.get('baca_adi') == olcum.get('baca_adi')):
                                        personel_adi = baca_bilgi.get('personel_adi', 'Admin')
                                        break
                                
                                gercek_veriler.append({
                                    'tarih': olcum_tarih.strftime('%Y-%m-%d'),
                                    'firma': olcum.get('firma_adi', ''),
                                    'kod': olcum.get('olcum_kodu', ''),
                                    'deger': round(deger, 1),
                                    'cihaz': olcum.get('baca_adi', ''),
                                    'personel': personel_adi
                                })
                            except ValueError:
                                continue
            except:
                continue
        
        # localStorage verilerini de dahil et (manuel elle girilen veriler)
        localStorage_verileri = data.get('localStorage_verileri', [])
        for veri in localStorage_verileri:
            try:
                # Tarih formatını kontrol et
                tarih_str = veri.get('tarih', '')
                if tarih_str:
                    # YYYY-MM-DD formatını kontrol et
                    if '-' in tarih_str and len(tarih_str) == 10:
                        olcum_tarih = datetime.strptime(tarih_str, '%Y-%m-%d')
                    else:
                        continue
                    
                    if baslangic <= olcum_tarih <= bitis:
                        # localStorage'da parametre adları farklı: CO, NO, SO2, O2, TOC
                        # Saha ölçümlerinde: KK1-CO, KK1-NO, KK1-SO2, KK1-O2, KK1-SPAN
                        kk_deger = None
                        
                        # Parametre adını localStorage formatına çevir (hem büyük hem küçük harf kontrol et)
                        if parametre == 'O2':
                            kk_deger = veri.get('O2') or veri.get('o2')
                        elif parametre == 'CO':
                            kk_deger = veri.get('CO') or veri.get('co')
                        elif parametre == 'NO':
                            kk_deger = veri.get('NO') or veri.get('no')
                        elif parametre == 'SO2':
                            kk_deger = veri.get('SO2') or veri.get('so2')
                        elif parametre == 'TOC':
                            kk_deger = veri.get('TOC') or veri.get('toc')
                        
                        if kk_deger and str(kk_deger).strip() and str(kk_deger) != '':
                            try:
                                # Virgülü noktaya çevir (float için gerekli)
                                deger_str = str(kk_deger).replace(',', '.')
                                deger = float(deger_str)
                                
                                gercek_veriler.append({
                                    'tarih': olcum_tarih.strftime('%Y-%m-%d'),
                                    'firma': veri.get('firma', ''),
                                    'kod': veri.get('kod', ''),
                                    'deger': round(deger, 1),
                                    'baca': veri.get('baca', ''),
                                    'personel': veri.get('personel', '')
                                })
                            except ValueError:
                                continue
            except:
                continue
        
        # İstatistiksel veriler hesapla
        if gercek_veriler:
            degerler_list = [v['deger'] for v in gercek_veriler]
            min_deger = min(degerler_list)
            max_deger = max(degerler_list)
            ortalama = sum(degerler_list) / len(degerler_list)
            standart_sapma = (plus_3s - kk) / 3  # +3S = KK + 3σ
        else:
            min_deger = '-'
            max_deger = '-'
            ortalama = round(float(kk), 1)
            standart_sapma = (plus_3s - kk) / 3
        
        # Format'a göre rapor oluştur
        if format_type == 'excel':
            return create_kk_excel_report(parametre, gercek_veriler, kk, minus_3s, minus_2s, plus_2s, plus_3s, ortalama, standart_sapma, min_deger, max_deger, tarih_baslangic, tarih_bitis)
        elif format_type == 'word':
            return create_kk_word_report(parametre, gercek_veriler, kk, minus_3s, minus_2s, plus_2s, plus_3s, ortalama, standart_sapma, min_deger, max_deger, tarih_baslangic, tarih_bitis)
        elif format_type == 'pdf':
            return create_kk_pdf_report(parametre, gercek_veriler, kk, minus_3s, minus_2s, plus_2s, plus_3s, ortalama, standart_sapma, min_deger, max_deger, tarih_baslangic, tarih_bitis)
        else:
            return jsonify({'error': 'Geçersiz format türü'}), 400
            
    except Exception as e:
        print(f"KK rapor oluşturma hatası: {e}")
        return jsonify({'error': f'Rapor oluşturulurken hata oluştu: {str(e)}'}), 500

def create_kk_excel_report(parametre, gercek_veriler, kk, minus_3s, minus_2s, plus_2s, plus_3s, ortalama, standart_sapma, min_deger, max_deger, tarih_baslangic, tarih_bitis):
    """KK Excel raporu oluştur - sadece veri"""
    try:
        import pandas as pd
        from io import BytesIO
        from datetime import datetime
        
        # Veri oluştur - tarih formatını düzelt
        def format_tarih_for_excel(tarih_str):
            if not tarih_str:
                return ''
            try:
                # YYYY-MM-DD formatından GG.AA.YY formatına çevir
                if '-' in tarih_str and len(tarih_str) == 10:
                    tarih_obj = datetime.strptime(tarih_str, '%Y-%m-%d')
                    return tarih_obj.strftime('%d.%m.%y')
                return tarih_str
            except:
                return tarih_str
        
        data = {
            "Sıra": list(range(1, len(gercek_veriler) + 1)),
            "Tarih": [format_tarih_for_excel(veri.get('tarih', '')) for veri in gercek_veriler],
            "Firma": [veri.get('firma', '') for veri in gercek_veriler],
            "Kod": [veri.get('kod', '') for veri in gercek_veriler],
            "Baca": [veri.get('baca', '') for veri in gercek_veriler],
            "Personel": [veri.get('personel', '') for veri in gercek_veriler],
            "Cihaz": [veri.get('cihaz', '') for veri in gercek_veriler],
            "Değer": [float(veri.get('deger', 0)) if veri.get('deger') else 0 for veri in gercek_veriler]
        }
        
        df = pd.DataFrame(data)
        
        # Excel dosyasını oluştur (sadece veri)
        output = BytesIO()
        df.to_excel(output, index=False, engine='openpyxl')
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'KK_Rapor_{parametre}_{tarih_baslangic}_{tarih_bitis}.xlsx'
        )
        
    except Exception as e:
        print(f"KK Excel rapor oluşturma hatası: {e}")
        return jsonify({'error': f'Excel rapor oluşturulurken hata oluştu: {str(e)}'}), 500

def create_kk_word_report(parametre, gercek_veriler, kk, minus_3s, minus_2s, plus_2s, plus_3s, ortalama, standart_sapma, min_deger, max_deger, tarih_baslangic, tarih_bitis):
    """KK Word raporu oluştur"""
    try:
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            return jsonify({'error': 'Word rapor oluşturma için python-docx kütüphanesi gerekli'}), 500
        
        # Word dosyası oluştur
        doc = Document()
        
        # Sayfa kenar boşluklarını ayarla (1 cm = 28.35 points)
        from docx.shared import Inches
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(1)    # Sol kenar 1 cm
            section.right_margin = Inches(1)   # Sağ kenar 1 cm
            section.top_margin = Inches(1)     # Üst kenar 1 cm
            section.bottom_margin = Inches(1)  # Alt kenar 1 cm
        
        # Başlık
        title = doc.add_heading(f'{parametre} KALİTE KONTROL RAPORU', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Başlık font boyutunu artır
        for run in title.runs:
            run.font.size = Pt(24)
            run.font.bold = True
        
        # Grafik ekle
        doc.add_paragraph()
        # Tarih formatını GG.AA.YY yap
        from datetime import datetime
        baslangic_tarih = datetime.strptime(tarih_baslangic, '%Y-%m-%d').strftime('%d.%m.%y')
        bitis_tarih = datetime.strptime(tarih_bitis, '%Y-%m-%d').strftime('%d.%m.%y')
        
        doc.add_heading(f'{parametre} Kalite Kontrol Grafiği ({baslangic_tarih} - {bitis_tarih})', level=1)
        
        # Grafik oluştur ve Word'e ekle
        try:
            # Grafik oluştur
            from datetime import datetime, timedelta
            
            baslangic = datetime.strptime(tarih_baslangic, '%Y-%m-%d')
            bitis = datetime.strptime(tarih_bitis, '%Y-%m-%d')
            
            # Gerçek verileri hazırla
            gercek_tarihler = []
            gercek_degerler = []
            
            if gercek_veriler:
                for veri in gercek_veriler:
                    try:
                        tarih = datetime.strptime(veri['tarih'], '%Y-%m-%d')
                        gercek_tarihler.append(tarih)
                        gercek_degerler.append(veri['deger'])
                    except:
                        continue
            
            # Grafik oluştur
            plt.figure(figsize=(10, 6))
            plt.rcParams['font.size'] = 12
            
            # Gerçek verileri mavi nokta olarak çiz
            if gercek_tarihler and gercek_degerler:
                # Tarihleri sırala
                sorted_data = sorted(zip(gercek_tarihler, gercek_degerler))
                sorted_tarihler = [x[0] for x in sorted_data]
                sorted_degerler = [x[1] for x in sorted_data]
                
                plt.plot(sorted_tarihler, sorted_degerler, 'bo', markersize=8, linewidth=2, label='Ölçüm Değerleri')
                # Noktaları birleştiren çizgi
                plt.plot(sorted_tarihler, sorted_degerler, 'b-', linewidth=1, alpha=0.7)
            
            # Kontrol limitleri
            plt.axhline(y=plus_3s, color='red', linestyle='--', linewidth=2)
            plt.axhline(y=kk, color='green', linewidth=2)
            plt.axhline(y=minus_3s, color='red', linestyle='--', linewidth=2)
            
            # 2-sigma limitleri
            plt.axhline(y=plus_2s, color='blue', linestyle='--', linewidth=2)
            plt.axhline(y=minus_2s, color='blue', linestyle='--', linewidth=2)
            
            # Kontrol limitlerini grafik üzerine yaz
            plt.text(-0.05, float(plus_3s), '+3S', fontsize=12, ha='right', va='center', color='red', fontweight='bold', transform=plt.gca().get_yaxis_transform())
            plt.text(-0.05, float(plus_2s), '+2S', fontsize=12, ha='right', va='center', color='blue', fontweight='bold', transform=plt.gca().get_yaxis_transform())
            plt.text(-0.05, float(kk), 'KK', fontsize=12, ha='right', va='center', color='green', fontweight='bold', transform=plt.gca().get_yaxis_transform())
            plt.text(-0.05, float(minus_2s), '-2S', fontsize=12, ha='right', va='center', color='blue', fontweight='bold', transform=plt.gca().get_yaxis_transform())
            plt.text(-0.05, float(minus_3s), '-3S', fontsize=12, ha='right', va='center', color='red', fontweight='bold', transform=plt.gca().get_yaxis_transform())
            
            # 1-sigma limitleri
            plus_1s = kk + standart_sapma
            minus_1s = kk - standart_sapma
            plt.axhline(y=plus_1s, color='gray', linestyle=':', linewidth=1, alpha=0.5)
            plt.axhline(y=minus_1s, color='gray', linestyle=':', linewidth=1, alpha=0.5)
            
            # Bölgeleri işaretle
            if gercek_tarihler:
                tarihler = gercek_tarihler
            else:
                tarihler = [baslangic, bitis]
            
            plt.fill_between(range(len(tarihler)), plus_3s, plus_2s, alpha=0.1, color='red')
            plt.fill_between(range(len(tarihler)), plus_2s, plus_1s, alpha=0.1, color='orange')
            plt.fill_between(range(len(tarihler)), plus_1s, minus_1s, alpha=0.1, color='green')
            plt.fill_between(range(len(tarihler)), minus_1s, minus_2s, alpha=0.1, color='orange')
            plt.fill_between(range(len(tarihler)), minus_2s, minus_3s, alpha=0.1, color='red')
            
            # Başlık
            title_text = f'{parametre} KK GRAF.    KK: {kk}    -2S: {minus_2s}    +2S: {plus_2s}    -3S: {minus_3s}    +3S: {plus_3s}'
            plt.title(title_text, fontsize=14, fontweight='bold', loc='left')
            plt.grid(True, alpha=0.3)
            
            # X ekseni formatı
            plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%d/%m'))
            plt.xticks(tarihler, rotation=45, fontsize=10)
            
            # X ekseni limitlerini ayarla
            if tarihler:
                min_date = min(tarihler)
                max_date = max(tarihler)
                plt.xlim(min_date - timedelta(days=2), max_date + timedelta(days=2))
            
            # Y ekseni limitleri
            margin = (plus_3s - minus_3s) * 0.1
            y_min = minus_3s - margin
            y_max = plus_3s + margin
            plt.ylim(y_min, y_max)
            
            # Grafik alanını temizle ve kaydet
            plt.tight_layout()
            
            # Grafiği sabit dosya adıyla kaydet (örnek kod yaklaşımı)
            grafik_path = 'kk_grafik_temp.png'
            plt.savefig(grafik_path, format='png', dpi=150, bbox_inches='tight')
            plt.close()
            
            # Grafiği Word'e ekle
            doc.add_picture(grafik_path, width=Inches(6))
            
            # Geçici dosyayı temizle
            try:
                if os.path.exists(grafik_path):
                    os.remove(grafik_path)
            except Exception as cleanup_error:
                print(f"Grafik dosyası temizleme hatası: {cleanup_error}")
                # Dosya silinemezse devam et
                
        except Exception as e:
            print(f"Grafik ekleme hatası: {e}")
            doc.add_paragraph("Grafik oluşturulamadı.")
        
        # Ölçüm verileri tablosu
        doc.add_heading('Ölçüm Verileri', level=1)
        if gercek_veriler:
            data_table = doc.add_table(rows=1, cols=7)
            data_table.style = 'Table Grid'
            
            # Tabloya sol kenar boşluğu ekle (1 cm)
            data_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            data_table.left_margin = Inches(1)   # Sol kenar 1 cm
            
            # Tablo genişliklerini ayarla - Firma adı 3.1 cm
            data_table.columns[0].width = Inches(0.5)   # Sıra No - dar
            data_table.columns[1].width = Inches(1.0)   # Tarih
            data_table.columns[2].width = Inches(3.1)   # Firma - 3.1 cm
            data_table.columns[3].width = Inches(1.5)   # Kod
            data_table.columns[4].width = Inches(1.5)   # Baca
            data_table.columns[5].width = Inches(1.2)   # Personel
            data_table.columns[6].width = Inches(1.0)   # Değer
            
            data_hdr_cells = data_table.rows[0].cells
            headers = ['Sıra', 'Tarih', 'Firma', 'Kod', 'Baca', 'Personel', 'Değer']
            
            for i, header in enumerate(headers):
                data_hdr_cells[i].text = header
            
            for i, veri in enumerate(gercek_veriler, 1):
                row_cells = data_table.add_row().cells
                row_cells[0].text = str(i)
                
                # Tarih formatını GG.AA.YY yap
                try:
                    from datetime import datetime
                    tarih_str = veri.get('tarih', '')
                    if tarih_str:
                        tarih_obj = datetime.strptime(tarih_str, '%Y-%m-%d')
                        formatted_tarih = tarih_obj.strftime('%d.%m.%y')
                        row_cells[1].text = formatted_tarih
                    else:
                        row_cells[1].text = ''
                except:
                    row_cells[1].text = veri.get('tarih', '')
                
                # Firma adının sadece 2 kelimesini al
                firma_adi = veri.get('firma', '')
                firma_kelimeler = firma_adi.split()
                firma_kisaltilmis = ' '.join(firma_kelimeler[:2]) if len(firma_kelimeler) >= 2 else firma_adi
                row_cells[2].text = firma_kisaltilmis
                row_cells[3].text = veri.get('kod', '')
                row_cells[4].text = veri.get('baca', '')
                row_cells[5].text = veri.get('personel', '')
                row_cells[6].text = str(veri.get('deger', ''))
        else:
            doc.add_paragraph('Bu tarih aralığında ölçüm verisi bulunamadı.')
        
        # Word dosyasını geçici olarak kaydet
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_word:
            doc.save(tmp_word.name)
            tmp_word_path = tmp_word.name
        
        # Word dosyasını oku
        with open(tmp_word_path, 'rb') as f:
            word_content = f.read()
        
        # Geçici dosyayı sil
        os.unlink(tmp_word_path)
        
        # Response oluştur
        response = make_response(word_content)
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="KK_Rapor_{parametre}_{tarih_baslangic}_{tarih_bitis}.docx"'
        
        return response
        
    except Exception as e:
        print(f"KK Word rapor oluşturma hatası: {e}")
        return jsonify({'error': f'Word rapor oluşturulurken hata oluştu: {str(e)}'}), 500

def create_kk_pdf_report(parametre, gercek_veriler, kk, minus_3s, minus_2s, plus_2s, plus_3s, ortalama, standart_sapma, min_deger, max_deger, tarih_baslangic, tarih_bitis):
    """KK PDF raporu oluştur - Word'den PDF'e çevir"""
    try:
        from docx2pdf import convert
        import tempfile
        import os
        
        # Önce Word dosyası oluştur
        word_content = create_kk_word_report(parametre, gercek_veriler, kk, minus_3s, minus_2s, plus_2s, plus_3s, ortalama, standart_sapma, min_deger, max_deger, tarih_baslangic, tarih_bitis)
        
        # Word dosyasını geçici olarak kaydet
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_word:
            tmp_word.write(word_content.data)
            tmp_word_path = tmp_word.name
        
        # PDF dosyası için geçici dosya oluştur
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
            tmp_pdf_path = tmp_pdf.name
        
        # Word'den PDF'e çevir
        convert(tmp_word_path, tmp_pdf_path)
        
        # PDF dosyasını oku
        with open(tmp_pdf_path, 'rb') as f:
            pdf_content = f.read()
        
        # Geçici dosyaları sil
        try:
            os.unlink(tmp_word_path)
            os.unlink(tmp_pdf_path)
        except:
            pass
        
        # Response oluştur
        response = make_response(pdf_content)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename="KK_Rapor_{parametre}_{tarih_baslangic}_{tarih_bitis}.pdf"'
        
        return response
        
    except Exception as e:
        print(f"KK PDF rapor oluşturma hatası: {e}")
        return jsonify({'error': f'PDF rapor oluşturulurken hata oluştu: {str(e)}'}), 500

@app.route('/api/firmalar')
def api_firmalar():
    """Firma listesini döndür"""
    try:
        firma_olcumler = load_firma_olcum()
        firmalar = []
        firma_adi_set = set()
        
        for olcum in firma_olcumler:
            if olcum.get('firma_adi') and olcum.get('firma_adi') not in firma_adi_set:
                firmalar.append({'firma_adi': olcum['firma_adi']})
                firma_adi_set.add(olcum['firma_adi'])
        
        return jsonify(firmalar)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/olcum_kodlari/<firma_adi>')
def api_olcum_kodlari(firma_adi):
    """Belirli firma için ölçüm kodlarını döndür"""
    try:
        firma_olcumler = load_firma_olcum()
        olcum_kodlari = []
        olcum_kodu_set = set()
        
        for olcum in firma_olcumler:
            if (olcum.get('firma_adi') == firma_adi and 
                olcum.get('olcum_kodu') and 
                olcum.get('olcum_kodu') not in olcum_kodu_set):
                olcum_kodlari.append({'olcum_kodu': olcum['olcum_kodu']})
                olcum_kodu_set.add(olcum['olcum_kodu'])
        
        return jsonify(olcum_kodlari)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/baca_listesi/<firma_adi>/<olcum_kodu>')
def api_baca_listesi(firma_adi, olcum_kodu):
    """Belirli firma ve ölçüm kodu için baca listesini döndür"""
    try:
        firma_olcumler = load_firma_olcum()
        baca_listesi = []
        
        for olcum in firma_olcumler:
            if (olcum.get('firma_adi') == firma_adi and 
                olcum.get('olcum_kodu') == olcum_kodu and
                olcum.get('baca_parametreleri')):
                
                baca_parametreleri = olcum['baca_parametreleri']
                if isinstance(baca_parametreleri, dict):
                    for baca_adi in baca_parametreleri.keys():
                        baca_listesi.append({'baca_adi': baca_adi})
        
        return jsonify(baca_listesi)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/parametre_listesi/<firma_adi>/<olcum_kodu>/<baca_adi>')
def api_parametre_listesi(firma_adi, olcum_kodu, baca_adi):
    """Belirli firma, ölçüm kodu ve baca için parametre listesini döndür"""
    try:
        firma_olcumler = load_firma_olcum()
        parametre_listesi = []
        
        for olcum in firma_olcumler:
            if (olcum.get('firma_adi') == firma_adi and 
                olcum.get('olcum_kodu') == olcum_kodu and
                olcum.get('baca_parametreleri')):
                
                baca_parametreleri = olcum['baca_parametreleri']
                if isinstance(baca_parametreleri, dict) and baca_adi in baca_parametreleri:
                    baca_parametreleri_list = baca_parametreleri[baca_adi]
                    if isinstance(baca_parametreleri_list, list):
                        for parametre in baca_parametreleri_list:
                            # Parametre string olarak tutuluyor
                            parametre_listesi.append({'parametre_adi': parametre})
        
        return jsonify(parametre_listesi)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/baca_bilgileri')
def api_baca_bilgileri():
    """Kaydedilen baca bilgileri listesini döndürür."""
    try:
        saved_baca_bilgileri = load_baca_bilgileri()
        return jsonify(saved_baca_bilgileri)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/baca_bilgileri/export_excel', methods=['POST'])
def export_baca_bilgileri_excel():
    try:
        data = request.get_json()
        selected_ids = data.get('selected_ids', [])
        
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Seçili kayıtları filtrele
        if selected_ids:
            filtered_data = [item for item in baca_bilgileri if item.get('id') in selected_ids]
        else:
            filtered_data = baca_bilgileri
        
        if not filtered_data:
            return jsonify({'error': 'Dışa aktarılacak veri bulunamadı'}), 404
        
        # Excel için DataFrame oluştur - baca_bilgileri objesi içindeki verileri doğru şekilde map et
        df_data = []
        for item in filtered_data:
            # Baca bilgileri objesini al
            baca_bilgileri = item.get('baca_bilgileri', {})
            
            # UUID'leri alan adlarına map et - gerçek UUID'leri kullan
            field_mapping = {
                '597fad80-d28f-40ea-bd28-a76c61c5203d': 'BACA NO',
                'ddca398d-0e55-4662-b661-3731e0975bd2': 'YAKIT TÜRÜ',
                '22867c9a-ca3c-4d80-b017-b73dafdd7fef': 'ISIL GÜÇ (MW)',
                '6b3546e0-184c-49de-82e4-e2835e81923b': 'ÇATI ŞEKLİ',
                '98399625-5bbc-465e-8e09-de454f231ae4': 'KAYNAK TÜRÜ',
                'd9958774-43f7-4bc3-8e12-436614a6193a': 'BACA ŞEKLİ',
                'b1b6fc38-98c0-4048-8b8e-795cf7d44c48': 'BACA ÖLÇÜSÜ',
                '6a301d72-f21b-485b-b8fb-116ad5cb223f': 'YERDEN YÜK.',
                '8ec9ecc9-ecda-4bf2-9802-02fa2e3fda4c': 'ÇATI YÜK',
                'ab2b67dd-16a5-4bee-9b6e-b60b8cfc2d0c': 'RÜZGAR HIZ (M/S)',
                'eca60e54-ec39-4412-8884-caa17faed0be': 'ORT. SIC.',
                '64238e0a-6387-4c31-9bf4-d7f800ef17e1': 'ORT. NEM',
                'b09ad69a-e4d4-4219-b055-2cf923ffd499': 'ORT. BAS.',
                '9c8c8bcf-c98e-4109-8b10-63b08b26460e': 'A BACA',
                'af55c55f-f83b-4b90-a655-ee76bf6bb2ac': 'B BACA',
                '20881447-f7c8-4a6b-8583-76c7246082ef': 'C DELİK'
            }
            
            # Temel alanlar
            row = {
                'Firma': item.get('firma_adi', ''),
                'Ölçüm Kodu': item.get('olcum_kodu', ''),
                'Baca': item.get('baca_adi', ''),
            }
            
            # Baca bilgileri objesindeki her UUID için değeri al
            for uuid, field_name in field_mapping.items():
                value = baca_bilgileri.get(uuid, '')
                if value and value.strip() != '':
                    row[field_name] = value
                else:
                    row[field_name] = '*'
            
            # Ek alanlar
            created_at = item.get('created_at', '')
            updated_at = item.get('updated_at', '')
            
            # Tarih formatını düzelt
            if created_at:
                try:
                    created_date = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                    created_at = created_date.strftime('%d.%m.%Y %H:%M')
                except:
                    created_at = created_at
            
            if updated_at:
                try:
                    updated_date = datetime.fromisoformat(updated_at.replace('Z', '+00:00'))
                    updated_at = updated_date.strftime('%d.%m.%Y %H:%M')
                except:
                    updated_at = updated_at
            
            row.update({
                'Fotoğraf': item.get('photo_path', '') or '*',
                'Kayıt Tarihi': created_at or '*',
                'Güncelleme Tarihi': updated_at or '*'
            })
            
            df_data.append(row)
        
        df = pd.DataFrame(df_data)
        
        # Geçici dosya oluştur
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
            # Excel dosyasını oluştur
            with pd.ExcelWriter(tmp_file.name, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Baca Bilgileri')
                
                # Sütun genişliklerini ayarla
                worksheet = writer.sheets['Baca Bilgileri']
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)  # Maksimum 50 karakter
                    worksheet.column_dimensions[column_letter].width = adjusted_width
            
            tmp_file_path = tmp_file.name
        
        # Dosyayı gönder
        return send_file(
            tmp_file_path,
            as_attachment=True,
            download_name=f'baca_bilgileri_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    
    except Exception as e:
        return jsonify({'error': f'Excel dışa aktarma hatası: {str(e)}'}), 500

@app.route('/api/baca_bilgileri/export_pdf', methods=['POST'])
def export_baca_bilgileri_pdf():
    try:
        data = request.get_json()
        selected_ids = data.get('selected_ids', [])
        
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Seçili kayıtları filtrele
        if selected_ids:
            filtered_data = [item for item in baca_bilgileri if item.get('id') in selected_ids]
        else:
            filtered_data = baca_bilgileri
        
        if not filtered_data:
            return jsonify({'error': 'Dışa aktarılacak veri bulunamadı'}), 404
        
        # İlk seçili baca için Word şablonunu kullanarak Word dosyası oluştur
        item = filtered_data[0]  # İlk seçili kayıt
        
        # Bu bacaya ait parametreleri bul
        parametre_olcumleri = load_parametre_olcum()
        baca_parametreleri = []
        
        for parametre in parametre_olcumleri:
            if (parametre.get('firma_adi') == item.get('firma_adi') and 
                parametre.get('olcum_kodu') == item.get('olcum_kodu') and 
                parametre.get('baca_adi') == item.get('baca_adi')):
                baca_parametreleri.append(parametre)
        
        # Form bilgilerini al
        forms = load_forms()
        form_bilgisi = None
        if forms:
            for form in forms:
                if 'EMİSYON' in form.get('form_adi', '').upper():
                    form_bilgisi = form
                    break
            if not form_bilgisi:
                form_bilgisi = forms[0]
        
        # Word dosyasını oluştur
        word_doc = create_baca_word_document_from_template(item, baca_parametreleri, form_bilgisi)
        
        if word_doc is None:
            return jsonify({'error': 'Word dosyası oluşturulamadı'}), 500
        
        # Word dosyasını geçici olarak kaydet
        import tempfile
        import os
        from io import BytesIO
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_word:
            # Document objesini BytesIO'ya kaydet
            doc_buffer = BytesIO()
            word_doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # Geçici dosyaya yaz
            tmp_word.write(doc_buffer.getvalue())
            tmp_word_path = tmp_word.name
        
        # Word dosyasını PDF'e çevir
        try:
            from docx2pdf import convert
            import pythoncom
            
            # COM bileşenlerini başlat
            pythoncom.CoInitialize()
            
            # PDF dosya adını oluştur
            firma_adi = item.get('firma_adi', 'Bilinmeyen')
            olcum_kodu = item.get('olcum_kodu', 'Bilinmeyen')
            baca_adi = item.get('baca_adi', 'Bilinmeyen')
            
            # Dosya adını temizle
            def clean_filename(text):
                tr_to_en = {
                    'ç': 'c', 'ğ': 'g', 'ı': 'i', 'ö': 'o', 'ş': 's', 'ü': 'u',
                    'Ç': 'C', 'Ğ': 'G', 'I': 'I', 'İ': 'I', 'Ö': 'O', 'Ş': 'S', 'Ü': 'U'
                }
                for tr_char, en_char in tr_to_en.items():
                    text = text.replace(tr_char, en_char)
                cleaned = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
                cleaned = cleaned.replace(' ', '-')
                while '--' in cleaned:
                    cleaned = cleaned.replace('--', '-')
                return cleaned
            
            firma_adi_clean = clean_filename(firma_adi)
            olcum_kodu_clean = clean_filename(olcum_kodu)
            baca_adi_clean = clean_filename(baca_adi)
            
            pdf_filename = f"{firma_adi_clean}-{baca_adi_clean}-{olcum_kodu_clean}.pdf"
            pdf_path = os.path.join(tempfile.gettempdir(), pdf_filename)
            
            # Word dosyasını PDF'e çevir
            convert(tmp_word_path, pdf_path)
            
            # PDF dosyasını oku
            with open(pdf_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
            
            # Geçici dosyaları temizle
            os.unlink(tmp_word_path)
            os.unlink(pdf_path)
            
            # PDF dosyasını döndür
            response = make_response(pdf_content)
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = f'attachment; filename="{pdf_filename}"'
            return response
            
        except Exception as pdf_error:
            print(f"PDF dönüştürme hatası: {pdf_error}")
            
            # COM bileşenlerini temizle (hata durumunda da)
            try:
                pythoncom.CoUninitialize()
            except:
                pass
                
            # PDF dönüştürme çalışmazsa, Word dosyasını döndür
            with open(tmp_word_path, 'rb') as word_file:
                word_content = word_file.read()
            
            # Geçici dosyayı temizle
            os.unlink(tmp_word_path)
            
            # Word dosyasını döndür
            response = make_response(word_content)
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = f'attachment; filename="{firma_adi_clean}-{baca_adi_clean}-{olcum_kodu_clean}.docx"'
            return response
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"PDF Export Error: {str(e)}")
        print(f"Error Details: {error_details}")
        return jsonify({'error': f'PDF dışa aktarma hatası: {str(e)}'}), 500

@app.route('/api/baca_bilgileri/import_excel', methods=['POST'])
def import_baca_bilgileri_excel():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Dosya seçilmedi'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Dosya seçilmedi'}), 400
        
        if not file.filename.endswith('.xlsx'):
            return jsonify({'error': 'Sadece Excel (.xlsx) dosyaları desteklenir'}), 400
        
        # Excel dosyasını oku
        df = pd.read_excel(file, engine='openpyxl')
        
        # Gerekli sütunları kontrol et
        required_columns = ['Firma', 'Ölçüm Kodu', 'Baca']
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            return jsonify({'error': f'Eksik sütunlar: {", ".join(missing_columns)}'}), 400
        
        # Mevcut baca bilgilerini yükle
        existing_data = load_baca_bilgileri()
        
        # Yeni kayıtları ekle - baca_bilgileri objesi formatında
        new_records = []
        for index, row in df.iterrows():
            try:
                # UUID mapping'i (export ile aynı)
                field_mapping = {
                    '7425ab86-4cfb-4796-bd2c-a84c3f1c61a4': 'BACA NO',
                    '2517064c-c286-4210-8f60-fa0a6b9e22a9': 'YAKIT TÜRÜ',
                    '2360aa46-81dd-41db-a6b6-5c23e62900fe': 'ISIL GÜÇ (MW)',
                    'f1e7b875-ddf4-4628-87a3-eda6afa775b7': 'ÇATI ŞEKLİ',
                    '7647078a-4f69-4908-99f7-3fff5651cc9b': 'KAYNAK TÜRÜ',
                    '394272bc-a8dc-46c4-a1a8-15ed4c0dbf29': 'BACA ŞEKLİ',
                    '6774fc7e-c124-4272-b826-482d064f3215': 'BACA ÖLÇÜSÜ',
                    '943f1111-1be7-4dd3-b62d-b734963c768e': 'YERDEN YÜK.',
                    '3f622008-8178-460c-a6a1-c5f4eb357231': 'ÇATI YÜK',
                    'eba873d9-1398-4fde-9ef0-851359037990': 'RÜZGAR HIZ (M/S)',
                    '5ab155d9-a52d-4359-abbf-bd3083e216da': 'ORT. SIC.',
                    'dea8eb12-71c2-4af4-a280-acd5eac4660b': 'ORT. NEM',
                    '922ab008-e6ec-442c-a2ee-4159b6ce948e': 'ORT. BAS.',
                    '0c47bc1d-afdf-477a-94b8-abb2c1d6f92c': 'A BACA',
                    'f59ec1f6-f93e-49a7-ac11-c33ce327987a': 'B BACA',
                    '0e9de9b8-d7a4-4125-b299-747e47c71b4b': 'C DELİK'
                }
                
                # Reverse mapping oluştur
                reverse_mapping = {v: k for k, v in field_mapping.items()}
                
                # Baca bilgileri objesi oluştur
                baca_bilgileri = {}
                for field_name, uuid in reverse_mapping.items():
                    value = row.get(field_name, '')
                    if value and str(value).strip() and str(value).lower() != 'nan':  # Boş değilse ve NaN değilse ekle
                        baca_bilgileri[uuid] = str(value)
                
                # Temel alanları kontrol et
                firma_adi = str(row.get('Firma', '')).strip()
                olcum_kodu = str(row.get('Ölçüm Kodu', '')).strip()
                baca_adi = str(row.get('Baca', '')).strip()
                
                if not firma_adi or not olcum_kodu or not baca_adi:
                    return jsonify({'error': f'Satır {index + 1}: Firma, Ölçüm Kodu ve Baca alanları zorunludur'}), 400
                
                # Fotoğraf alanını kontrol et
                photo_path = row.get('Fotoğraf', '')
                if photo_path and str(photo_path).strip() and str(photo_path).lower() != 'nan':
                    photo_path = str(photo_path)
                else:
                    photo_path = None
                
                new_record = {
                    'id': str(uuid4()),
                    'firma_adi': firma_adi,
                    'olcum_kodu': olcum_kodu,
                    'baca_adi': baca_adi,
                    'baca_bilgileri': baca_bilgileri,
                    'photo_path': photo_path,
                    'created_at': str(row.get('Kayıt Tarihi', datetime.now().isoformat())),
                    'updated_at': str(row.get('Güncelleme Tarihi', datetime.now().isoformat()))
                }
                new_records.append(new_record)
                
            except Exception as row_error:
                return jsonify({'error': f'Satır {index + 1} işlenirken hata: {str(row_error)}'}), 400
        
        # Yeni kayıtları mevcut veriye ekle
        existing_data.extend(new_records)
        
        # Veriyi kaydet
        if save_baca_bilgileri(existing_data):
            return jsonify({
                'success': True,
                'message': f'{len(new_records)} kayıt başarıyla içe aktarıldı',
                'imported_count': len(new_records)
            })
        else:
            return jsonify({'error': 'Veriler kaydedilemedi'}), 500
    
    except Exception as e:
        return jsonify({'error': f'Excel içe aktarma hatası: {str(e)}'}), 500

@app.route('/api/baca_bilgileri/<record_id>', methods=['DELETE'])
def api_delete_baca_bilgileri(record_id):
    """Belirtilen ID'ye sahip baca bilgilerini siler."""
    try:
        saved_baca_bilgileri = load_baca_bilgileri()
        
        # Kaydı bul
        record_index = None
        for i, record in enumerate(saved_baca_bilgileri):
            if record.get('id') == record_id:
                record_index = i
                break
        
        if record_index is None:
            return jsonify({'success': False, 'error': 'Kayıt bulunamadı'}), 404
        
        # Kaydı sil
        deleted_record = saved_baca_bilgileri.pop(record_index)
        
        # Dosyaya kaydet
        if save_baca_bilgileri(saved_baca_bilgileri):
            print(f"Baca bilgileri silindi: {deleted_record.get('firma_adi')} - {deleted_record.get('olcum_kodu')} - {deleted_record.get('baca_adi')}")
            return jsonify({'success': True, 'message': 'Baca bilgileri başarıyla silindi'})
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'}), 500
            
    except Exception as e:
        print(f"Baca bilgileri silinirken hata: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/baca_parametreleri')
def api_baca_parametreleri():
    """Baca parametrelerini döndürür."""
    try:
        baca_paralar = load_baca_paralar()
        return jsonify(baca_paralar)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/parametre_bilgileri')
def api_parametre_bilgileri():
    """Tüm parametre bilgilerini döndürür."""
    try:
        parametreler = load_parameters()
        return jsonify(parametreler)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/save_parametre_olcum_old', methods=['POST'])
def save_parametre_olcum_old():
    """Parametre ölçümlerini kaydeder."""
    try:
        firma_adi = request.form.get('firma_adi')
        olcum_kodu = request.form.get('olcum_kodu')
        baca_adi = request.form.get('baca_adi')
        parametre_adi = request.form.get('parametre_adi')
        olcum_sonucu = request.form.get('olcum_sonucu')
        olcum_birimi = request.form.get('olcum_birimi')
        olcum_tarihi = request.form.get('olcum_tarihi')
        olcum_saati = request.form.get('olcum_saati')
        olcum_notlari = request.form.get('olcum_notlari')
        
        if not all([firma_adi, olcum_kodu, baca_adi, parametre_adi, olcum_sonucu, olcum_birimi]):
            return jsonify({'success': False, 'error': 'Eksik bilgi'})
        
        # Saha ölçüm verilerini yükle
        saha_olc_data = load_saha_olc()
        
        # Yeni ölçüm kaydı
        new_olcum = {
            'id': str(uuid4()),
            'firma_adi': firma_adi,
            'olcum_kodu': olcum_kodu,
            'baca_adi': baca_adi,
            'parametre_adi': parametre_adi,
            'olcum_sonucu': olcum_sonucu,
            'olcum_birimi': olcum_birimi,
            'olcum_tarihi': olcum_tarihi,
            'olcum_saati': olcum_saati,
            'olcum_notlari': olcum_notlari,
            'created_at': datetime.now().isoformat()
        }
        
        # Veriyi ekle
        saha_olc_data.append(new_olcum)
        
        # Dosyaya kaydet
        if save_saha_olc(saha_olc_data):
            print(f"Parametre ölçümü kaydedildi: {firma_adi} - {olcum_kodu} - {baca_adi} - {parametre_adi}")
            return jsonify({
                'success': True, 
                'message': 'Parametre ölçümü başarıyla kaydedildi'
            })
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'}), 500
            
    except Exception as e:
        print(f"Parametre ölçümü kaydedilirken hata: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/save_parametre_olcum_saha', methods=['POST'])
def save_parametre_olcum_saha():
    """Saha ölçümü için parametre ölçümlerini kaydeder."""
    try:
        firma_adi = request.form.get('firma_adi')
        olcum_kodu = request.form.get('olcum_kodu')
        baca_adi = request.form.get('baca_adi')
        parametre_adi = request.form.get('parametre_adi')
        parametre_verileri_json = request.form.get('parametre_verileri')
        
        if not all([firma_adi, olcum_kodu, baca_adi, parametre_adi, parametre_verileri_json]):
            return jsonify({'success': False, 'error': 'Eksik bilgi'})
        
        parametre_verileri = json.loads(parametre_verileri_json)
        
        # Mevcut parametre ölçümlerini yükle
        saved_parametre_olcum = load_parametre_olcum()
        
        # İlgili baca bilgisinden personel adını bul (parametre ölçümü için varsayılan)
        personel_adi_default = ''
        try:
            for b in load_baca_bilgileri():
                if (
                    b.get('firma_adi') == firma_adi and
                    b.get('olcum_kodu') == olcum_kodu and
                    b.get('baca_adi') == baca_adi
                ):
                    personel_adi_default = b.get('personel_adi', '')
                    break
        except Exception:
            personel_adi_default = ''
        
        # Kaydedilecek yeni veri
        new_record = {
            'id': str(uuid4()),
            'firma_adi': firma_adi,
            'olcum_kodu': olcum_kodu,
            'baca_adi': baca_adi,
            'parametre_adi': parametre_adi,
            'parametre_verileri': parametre_verileri,
            # Personel adı kayıt içinde bulunmuyorsa baca bilgilerindeki personele düş
            'personel_adi': personel_adi_default,
            'created_at': datetime.now().isoformat(),
            'updated_at': datetime.now().isoformat()
        }
        
        # Aynı firma-ölçüm-baca-parametre kombinasyonu varsa güncelle, yoksa ekle
        record_found = False
        for i, record in enumerate(saved_parametre_olcum):
            if (record.get('firma_adi') == firma_adi and 
                record.get('olcum_kodu') == olcum_kodu and 
                record.get('baca_adi') == baca_adi and
                record.get('parametre_adi') == parametre_adi):
                # Mevcut kaydı güncelle
                new_record['id'] = record.get('id', str(uuid4()))
                new_record['created_at'] = record.get('created_at', datetime.now().isoformat())
                saved_parametre_olcum[i] = new_record
                record_found = True
                break
        
        if not record_found:
            # Yeni kayıt ekle
            saved_parametre_olcum.append(new_record)
        
        # Dosyaya kaydet
        if save_parametre_olcum(saved_parametre_olcum):
            print(f"Parametre ölçümü kaydedildi: {firma_adi} - {olcum_kodu} - {baca_adi} - {parametre_adi}")
            return jsonify({
                'success': True, 
                'message': 'Parametre ölçümü başarıyla kaydedildi'
            })
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'}), 500
            
    except Exception as e:
        print(f"Parametre ölçümü kaydedilirken hata: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/parametre_olcumleri')
def api_parametre_olcumleri():
    """Tüm parametre ölçümlerini döndürür."""
    try:
        parametre_olcumleri = load_parametre_olcum()
        return jsonify(parametre_olcumleri)
    except Exception as e:
        print(f"Parametre ölçümleri yüklenirken hata: {e}")
        return jsonify([])

@app.route('/api/parametre_olcumleri/bulk_delete', methods=['POST'])
def api_bulk_delete_parametre_olcumleri():
    """Seçilen parametre ölçümlerini toplu olarak siler."""
    try:
        data = request.get_json()
        ids_to_delete = data.get('ids', [])
        
        if not ids_to_delete:
            return jsonify({'success': False, 'error': 'Silinecek kayıt seçilmedi'})
        
        # Mevcut parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Silinecek kayıtları filtrele
        original_count = len(parametre_olcumleri)
        parametre_olcumleri = [record for record in parametre_olcumleri if record.get('id') not in ids_to_delete]
        deleted_count = original_count - len(parametre_olcumleri)
        
        # Dosyaya kaydet
        if save_parametre_olcum(parametre_olcumleri):
            print(f"Parametre ölçümleri toplu silindi: {deleted_count} kayıt")
            return jsonify({
                'success': True, 
                'message': f'{deleted_count} adet parametre ölçümü başarıyla silindi',
                'deleted_count': deleted_count
            })
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'})
            
    except Exception as e:
        print(f"Parametre ölçümleri toplu silme hatası: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/save_baca_bilgileri_saha', methods=['POST'])
def save_baca_bilgileri_saha():
    """Saha ölçümü için baca bilgilerini kaydeder."""
    try:
        firma_adi = request.form.get('firma_adi')
        olcum_kodu = request.form.get('olcum_kodu')
        baca_adi = request.form.get('baca_adi')
        baca_bilgileri_json = request.form.get('baca_bilgileri')
        personel_adi = request.form.get('personel_adi', '')
        is_edit = request.form.get('is_edit', 'false') == 'true'
        
        if not all([firma_adi, olcum_kodu, baca_adi, baca_bilgileri_json]):
            return jsonify({'success': False, 'error': 'Eksik bilgi'})
        
        print(f"Gelen baca bilgileri JSON: {baca_bilgileri_json}")
        baca_bilgileri = json.loads(baca_bilgileri_json)
        print(f"Parse edilen baca bilgileri: {baca_bilgileri}")
        
        # Fotoğraf işleme
        photo_path = None
        if 'photo' in request.files:
            photo = request.files['photo']
            if photo and photo.filename:
                # Fotoğraf klasörü oluştur
                photo_dir = os.path.join(app.root_path, 'static', 'uploads', 'photos')
                os.makedirs(photo_dir, exist_ok=True)
                
                # Benzersiz dosya adı oluştur
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"{firma_adi}_{olcum_kodu}_{baca_adi}_{timestamp}.jpg"
                photo_path = os.path.join(photo_dir, filename)
                
                # Fotoğrafı kaydet
                photo.save(photo_path)
                photo_path = f"uploads/photos/{filename}"  # Web erişimi için relative path
        
        # Mevcut baca bilgilerini yükle
        saved_baca_bilgileri = load_baca_bilgileri()
        
        # Kaydedilecek yeni veri
        new_record = {
            'id': str(uuid4()),
            'firma_adi': firma_adi,
            'olcum_kodu': olcum_kodu,
            'baca_adi': baca_adi,
            'baca_bilgileri': baca_bilgileri,
            'personel_adi': personel_adi,
            'photo_path': photo_path,
            'created_at': datetime.now().isoformat(),
            'updated_at': datetime.now().isoformat()
        }
        
        # Eğer güncelleme ise, mevcut kaydı bul ve güncelle
        if is_edit:
            record_found = False
            for i, record in enumerate(saved_baca_bilgileri):
                if (record.get('firma_adi') == firma_adi and 
                    record.get('olcum_kodu') == olcum_kodu and 
                    record.get('baca_adi') == baca_adi):
                    # Mevcut kaydı güncelle
                    new_record['id'] = record.get('id', str(uuid4()))
                    new_record['created_at'] = record.get('created_at', datetime.now().isoformat())
                    # Fotoğraf değişmemişse eski fotoğrafı koru
                    if not photo_path:
                        new_record['photo_path'] = record.get('photo_path')
                    saved_baca_bilgileri[i] = new_record
                    record_found = True
                    break
            
            if not record_found:
                # Kayıt bulunamadıysa yeni kayıt olarak ekle
                saved_baca_bilgileri.append(new_record)
        else:
            # Yeni kayıt ekle
            saved_baca_bilgileri.append(new_record)
        
        # Baca bilgilerini dosyaya kaydet
        if save_baca_bilgileri(saved_baca_bilgileri):
            print(f"Baca bilgileri kaydedildi: {firma_adi} - {olcum_kodu} - {baca_adi}")
            
            # Personel adı değiştiyse, o bacaya ait parametre ölçüm kayıtlarını güncelle
            if personel_adi:
                update_parametre_olcum_personel(firma_adi, olcum_kodu, baca_adi, personel_adi)
            
            return jsonify({
                'success': True, 
                'message': 'Baca bilgileri başarıyla kaydedildi',
                'photo_path': photo_path
            })
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'})
        
    except Exception as e:
        print(f"Baca bilgileri kaydetme hatası: {str(e)}")
        return jsonify({'success': False, 'error': str(e)})

def update_parametre_olcum_personel(firma_adi, olcum_kodu, baca_adi, personel_adi):
    """Baca bilgilerinde personel adı değiştiğinde, o bacaya ait parametre ölçüm kayıtlarını günceller."""
    try:
        # Mevcut parametre ölçüm kayıtlarını yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Güncellenen kayıt sayısı
        updated_count = 0
        
        # O bacaya ait tüm parametre ölçüm kayıtlarını bul ve güncelle
        for record in parametre_olcumleri:
            if (record.get('firma_adi') == firma_adi and 
                record.get('olcum_kodu') == olcum_kodu and 
                record.get('baca_adi') == baca_adi):
                # Personel adını güncelle
                record['personel_adi'] = personel_adi
                updated_count += 1
        
        # Güncellenmiş verileri kaydet
        if updated_count > 0:
            save_parametre_olcum(parametre_olcumleri)
            print(f"Parametre ölçüm kayıtları güncellendi: {updated_count} kayıt - {firma_adi} - {olcum_kodu} - {baca_adi} - Personel: {personel_adi}")
        else:
            print(f"Güncellenecek parametre ölçüm kaydı bulunamadı: {firma_adi} - {olcum_kodu} - {baca_adi}")
            
    except Exception as e:
        print(f"Parametre ölçüm kayıtları güncellenirken hata: {str(e)}")

def sync_all_parametre_personel():
    """Tüm parametre ölçüm kayıtlarını baca bilgileriyle senkronize eder."""
    try:
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Firma + Ölçüm kodu -> Personel adı eşleştirmesi
        personel_map = {}
        for baca in baca_bilgileri:
            firma = baca.get('firma_adi', '').strip()
            olcum = baca.get('olcum_kodu', '').strip()
            personel = baca.get('personel_adi', '').strip()
            
            if firma and olcum and personel:
                key = f"{firma}||{olcum}"
                personel_map[key] = personel
        
        print(f"Personel eşleştirmesi: {personel_map}")
        
        # Güncellenen kayıt sayısı
        updated_count = 0
        
        # Tüm parametre ölçüm kayıtlarını güncelle
        for record in parametre_olcumleri:
            firma = record.get('firma_adi', '').strip()
            olcum = record.get('olcum_kodu', '').strip()
            
            if firma and olcum:
                key = f"{firma}||{olcum}"
                if key in personel_map:
                    # Personel adını güncelle
                    record['personel_adi'] = personel_map[key]
                    updated_count += 1
        
        # Güncellenmiş verileri kaydet
        if updated_count > 0:
            save_parametre_olcum(parametre_olcumleri)
            print(f"Tüm parametre ölçüm kayıtları senkronize edildi: {updated_count} kayıt")
        else:
            print("Senkronize edilecek kayıt bulunamadı")
            
    except Exception as e:
        print(f"Parametre ölçüm kayıtları senkronize edilirken hata: {str(e)}")

def cleanup_orphaned_parametre_personel():
    """Baca bilgilerinde olmayan personel isimlerini parametre ölçümlerinden temizler."""
    try:
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Geçerli personel isimlerini topla
        valid_personel = set()
        for baca in baca_bilgileri:
            personel = baca.get('personel_adi', '').strip()
            if personel:
                valid_personel.add(personel)
        
        print(f"Geçerli personel isimleri: {valid_personel}")
        
        # Parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Temizlenen kayıt sayısı
        cleaned_count = 0
        
        # Geçersiz personel isimlerini temizle
        for record in parametre_olcumleri:
            current_personel = record.get('personel_adi', '').strip()
            if current_personel and current_personel not in valid_personel:
                print(f"Geçersiz personel ismi temizleniyor: {current_personel} -> (boş)")
                record['personel_adi'] = ''
                cleaned_count += 1
        
        # Temizlenmiş verileri kaydet
        if cleaned_count > 0:
            save_parametre_olcum(parametre_olcumleri)
            print(f"Geçersiz personel isimleri temizlendi: {cleaned_count} kayıt")
        else:
            print("Temizlenecek geçersiz personel ismi bulunamadı")
            
    except Exception as e:
        print(f"Geçersiz personel isimleri temizlenirken hata: {str(e)}")

@app.route('/api/sync_parametre_personel', methods=['POST'])
def api_sync_parametre_personel():
    """Tüm parametre ölçüm kayıtlarını baca bilgileriyle senkronize eder."""
    try:
        sync_all_parametre_personel()
        return jsonify({'success': True, 'message': 'Parametre ölçüm kayıtları senkronize edildi'})
    except Exception as e:
        print(f"Senkronizasyon hatası: {str(e)}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/formlar')
def formlar():
    """Formlar sayfası - Ana sekme"""
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    # Kullanıcıları yükle
    users = load_users()
    
    # Parametreleri yükle
    parameters = load_parameters()
    
    # Tekrar eden parametre adlarını kaldır (sadece benzersiz olanları al)
    unique_parameters = []
    seen_names = set()
    
    for param in parameters:
        param_name = param.get('Parametre Adı', '')
        if param_name and param_name not in seen_names:
            seen_names.add(param_name)
            unique_parameters.append(param)
    
    # Baca parametrelerini yükle
    baca_paralar = load_baca_paralar()
    
    return render_template('formlar.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         users=users,
                         parameters=parameters,
                         unique_parameters=unique_parameters,
                         baca_paralar=baca_paralar)

@app.route('/parametre_sahabil')
def parametre_sahabil():
    """Parametre Sahabil sayfası"""
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    # Parametreleri yükle
    parameters = load_parameters()
    
    # Tekrar eden parametre adlarını kaldır (sadece benzersiz olanları al)
    unique_parameters = []
    seen_names = set()
    
    for param in parameters:
        param_name = param.get('Parametre Adı', '')
        if param_name and param_name not in seen_names:
            seen_names.add(param_name)
            unique_parameters.append(param)
    
    return render_template('parametre_sahabil.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         parameters=unique_parameters)

@app.route('/save_parametre_sahabil', methods=['POST'])
def save_parametre_sahabil():
    """Parametre sahabil ölçümlerini kaydeder."""
    try:
        parametre_id = request.form.get('parametre_id')
        parametre_adi = request.form.get('parametre_adi')
        
        if not parametre_id or not parametre_adi:
            return jsonify({'success': False, 'error': 'Eksik parametre bilgisi'})
        
        # Özel parametreler listesi
        specialParams = ['yg', 'voc', 'toc', 'pm10', 'ctoz'];
        isSpecialParam = any(param in parametre_adi.lower() for param in specialParams);
        
        # Ölçüm verilerini hazırla
        olcum_data = {
            'id': str(uuid4()),
            'parametre_id': parametre_id,
            'parametre_adi': parametre_adi,
            'is_special_param': isSpecialParam,
            'created_at': datetime.now().isoformat()
        }
        
        if isSpecialParam:
            # Özel parametreler için basit veri
            olcum_data.update({
                'olcum_tarihi': request.form.get('olcum_tarihi'),
                'olcum_saati': request.form.get('olcum_saati'),
                'olcum_degeri': request.form.get('olcum_degeri'),
                'birim': request.form.get('birim'),
                'notlar': request.form.get('notlar')
            })
        else:
            # Diğer parametreler için detaylı veri
            olcum_data.update({
                'olcum_tarihi': request.form.get('olcum_tarihi'),
                'notlar': request.form.get('notlar'),
                'measurements': {}
            })
            
            # 3 ölçüm için veri topla
            for i in range(1, 4):
                olcum_data['measurements'][f'olcum_{i}'] = {
                    'metot': request.form.get(f'metot_{i}'),
                    'nozzle_cap': request.form.get(f'nozzle_cap_{i}'),
                    'travers': request.form.get(f'travers_{i}'),
                    'b_hiz': request.form.get(f'b_hiz_{i}'),
                    'b_sic': request.form.get(f'b_sic_{i}'),
                    'b_bas': request.form.get(f'b_bas_{i}'),
                    'b_nem_gm3': request.form.get(f'b_nem_gm3_{i}'),
                    'b_nem_yuzde': request.form.get(f'b_nem_yuzde_{i}'),
                    'syc_hac': request.form.get(f'syc_hac_{i}'),
                    'syc_ilk': request.form.get(f'syc_ilk_{i}'),
                    'syc_son': request.form.get(f'syc_son_{i}'),
                    'syc_sic': request.form.get(f'syc_sic_{i}'),
                    'debi': request.form.get(f'debi_{i}'),
                    'isdl': request.form.get(f'isdl_{i}')
                }
        
        # Parametre sahabil verilerini yükle
        parametre_sahabil_data = load_parametre_sahabil()
        parametre_sahabil_data.append(olcum_data)
        
        # Dosyaya kaydet
        if save_parametre_sahabil(parametre_sahabil_data):
            print(f"Parametre sahabil ölçümü kaydedildi: {parametre_adi}")
            return jsonify({
                'success': True, 
                'message': 'Parametre ölçümü başarıyla kaydedildi'
            })
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'}), 500
            
    except Exception as e:
        print(f"Parametre sahabil ölçümü kaydedilirken hata: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/parametre_metodlari')
def api_parametre_metodlari():
    """Parametre metodlarını döndürür."""
    try:
        # Parametreleri yükle
        parameters = load_parameters()
        
        # Parametre adı ve metodunu eşleştir
        parametre_metodlari = {}
        for param in parameters:
            parametre_adi = param.get('Parametre Adı', '')
            metot = param.get('Metot', '')
            if parametre_adi and metot:
                parametre_metodlari[parametre_adi.upper()] = metot
        
        return jsonify(parametre_metodlari)
    except Exception as e:
        print(f"Parametre metodları yüklenirken hata: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/personel_listesi')
def api_personel_listesi():
    """Personel listesini döndürür - sadece görevi 'Saha' olanlar."""
    try:
        # users.json dosyasından personel verilerini yükle
        users_data = load_users()
        
        # Sadece görevi 'Saha' olan personeli filtrele
        saha_personel = []
        for username, user_info in users_data.items():
            # Admin kullanıcısını atla
            if username == 'admin':
                continue
                
            # Görevi 'Saha' olanları al
            if user_info.get('gorev') == 'Saha':
                # Ad ve soyadı birleştir
                ad = username
                soyad = user_info.get('surname', '')
                tam_ad = f"{ad} {soyad}".strip()
                
                saha_personel.append({
                    'id': username,
                    'personel_adi': tam_ad,
                    'gorev': user_info.get('gorev', 'Saha')
                })
        
        return jsonify(saha_personel)
    except Exception as e:
        print(f"Personel listesi yüklenirken hata: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/parametre_sahabil_list')
def api_parametre_sahabil_list():
    """Parametre sahabil ölçümlerini döndürür."""
    try:
        parametre_sahabil_data = load_parametre_sahabil()
        return jsonify(parametre_sahabil_data)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/delete_parametre_sahabil/<record_id>', methods=['DELETE'])
def delete_parametre_sahabil(record_id):
    """Parametre sahabil ölçümünü siler."""
    try:
        parametre_sahabil_data = load_parametre_sahabil()
        
        # Kaydı bul ve sil
        parametre_sahabil_data = [record for record in parametre_sahabil_data if record.get('id') != record_id]
        
        if save_parametre_sahabil(parametre_sahabil_data):
            return jsonify({'success': True, 'message': 'Kayıt başarıyla silindi'})
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'}), 500
            
    except Exception as e:
        print(f"Parametre sahabil silme hatası: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/delete_selected_parametre_sahabil', methods=['POST'])
def delete_selected_parametre_sahabil():
    """Seçilen parametre sahabil ölçümlerini siler."""
    try:
        data = request.get_json()
        selected_ids = data.get('ids', [])
        
        if not selected_ids:
            return jsonify({'success': False, 'error': 'Seçilen kayıt bulunamadı'})
        
        parametre_sahabil_data = load_parametre_sahabil()
        
        # Seçilen kayıtları sil
        parametre_sahabil_data = [record for record in parametre_sahabil_data if record.get('id') not in selected_ids]
        
        if save_parametre_sahabil(parametre_sahabil_data):
            return jsonify({'success': True, 'message': f'{len(selected_ids)} kayıt başarıyla silindi'})
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'}), 500
            
    except Exception as e:
        print(f"Seçilen parametre sahabil silme hatası: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/import_parametre_sahabil', methods=['POST'])
def import_parametre_sahabil():
    """Parametre sahabil verilerini Excel dosyasından içe aktarır."""
    # Session kontrolünü geçici olarak kaldırıyoruz
    # if not session.get('logged_in') or not can_write(session.get('role')):
    #     return jsonify({'success': False, 'error': 'Yetkisiz erişim'}), 403
    
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'Dosya seçilmedi'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'Dosya seçilmedi'})
        
        if not file.filename.endswith(('.xlsx', '.xls')):
            return jsonify({'success': False, 'error': 'Sadece Excel dosyaları (.xlsx, .xls) kabul edilir'})
        
        # Excel dosyasını oku
        df = pd.read_excel(file)
        
        # Mevcut parametre sahabil verilerini yükle
        parametre_sahabil_data = load_parametre_sahabil()
        
        # Excel verilerini parametre sahabil formatına dönüştür
        imported_count = 0
        print(f"Excel dosyası okundu. Toplam satır: {len(df)}")
        
        for index, row in df.iterrows():
            # Debug: İlk 10 satırı yazdır
            if index < 10:
                print(f"Satır {index}: {[str(val) if pd.notna(val) else 'NaN' for val in row]}")
            
            # Boş satırları atla
            if pd.isna(row.iloc[1]) or str(row.iloc[1]).strip() == '':
                continue
            
            parametre_turu = str(row.iloc[1]).strip()  # Sütun 2 (Unnamed: 1)
            parametre_adi = str(row.iloc[2]).strip() if len(row) > 2 and pd.notna(row.iloc[2]) else ''  # Sütun 3 (Unnamed: 2)
            
            # Ölçüm verilerini topla (sütun 4, 5, 6)
            olcum_verileri = {}
            for i in range(3, min(6, len(row))):  # Sütun 4, 5, 6 (Unnamed: 3, 4, 5)
                if pd.notna(row.iloc[i]) and str(row.iloc[i]).strip() != '':
                    olcum_verileri[f'olcum_{i-2}'] = str(row.iloc[i]).strip()
            
            # Sadece parametre türü ve adı olan satırları kaydet
            if parametre_turu and parametre_adi and parametre_turu not in ['1.ÖLÇ.', '2.ÖLÇ', '3.ÖLÇ.']:
                print(f"Kayıt ekleniyor: {parametre_turu} - {parametre_adi}")
                # Yeni kayıt oluştur
                new_record = {
                    'id': str(uuid4()),
                    'parametre_turu': parametre_turu,
                    'parametre_adi': parametre_adi,
                    'olcum_verileri': olcum_verileri,
                    'imported_from_excel': True,
                    'created_at': datetime.now().isoformat()
                }
                
                parametre_sahabil_data.append(new_record)
                imported_count += 1
        
        # Verileri kaydet
        if save_parametre_sahabil(parametre_sahabil_data):
            return jsonify({
                'success': True, 
                'message': f'{imported_count} adet parametre ölçümü başarıyla içe aktarıldı'
            })
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'}), 500
            
    except Exception as e:
        print(f"Parametre sahabil import hatası: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/export_all_parametre_sahabil')
def export_all_parametre_sahabil():
    """Tüm parametre sahabil verilerini Excel olarak dışa aktarır."""
    # Session kontrolünü geçici olarak kaldırıyoruz
    # if not session.get('logged_in') or not can_read(session.get('role')):
    #     return jsonify({'success': False, 'error': 'Yetkisiz erişim'}), 403
    
    try:
        parametre_sahabil_data = load_parametre_sahabil()
        
        if not parametre_sahabil_data:
            return jsonify({'success': False, 'error': 'Dışa aktarılacak veri bulunamadı'})
        
        # DataFrame oluştur
        export_data = []
        for record in parametre_sahabil_data:
            row = {
                'Parametre Türü': record.get('parametre_turu', ''),
                'Parametre Adı': record.get('parametre_adi', ''),
                '1.ÖLÇ': record.get('olcum_verileri', {}).get('olcum_1', ''),
                '2.ÖLÇ': record.get('olcum_verileri', {}).get('olcum_2', ''),
                '3.ÖLÇ': record.get('olcum_verileri', {}).get('olcum_3', ''),
                'Oluşturma Tarihi': record.get('created_at', '')
            }
            export_data.append(row)
        
        df = pd.DataFrame(export_data)
        
        # Excel dosyası oluştur
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
            df.to_excel(tmp.name, index=False, sheet_name='Parametre Sahabil', engine='openpyxl')
            tmp.flush()
            
            return send_file(
                tmp.name,
                as_attachment=True,
                download_name=f'parametre_sahabil_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx',
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            
    except Exception as e:
        print(f"Parametre sahabil export hatası: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/save_parametre_field', methods=['POST'])
def save_parametre_field():
    """Yeni parametre alanını kaydeder."""
    try:
        data = request.get_json()
        parametre_type = data.get('parametreType')
        field_name = data.get('fieldName')
        
        if not parametre_type or not field_name:
            return jsonify({'success': False, 'message': 'Parametre tipi ve alan adı gerekli!'}), 400
        
        # Mevcut alanları yükle
        fields = load_parametre_fields()
        
        # Parametre tipine göre alanları organize et
        if parametre_type not in fields:
            fields[parametre_type] = []
        
        # Alan zaten var mı kontrol et
        if field_name in fields[parametre_type]:
            return jsonify({'success': False, 'message': f'"{field_name}" alanı zaten mevcut!'}), 400
        
        # Yeni alanı ekle
        fields[parametre_type].append(field_name)
        
        # Kaydet
        if save_parametre_fields(fields):
            return jsonify({
                'success': True, 
                'message': f'"{field_name}" alanı {parametre_type.upper()} parametresine başarıyla eklendi!',
                'fieldName': field_name,
                'parametreType': parametre_type
            })
        else:
            return jsonify({'success': False, 'message': 'Alan kaydedilirken hata oluştu!'}), 500
            
    except Exception as e:
        print(f"Parametre alanı kaydedilirken hata: {e}")
        return jsonify({'success': False, 'message': 'Sunucu hatası!'}), 500

@app.route('/get_parametre_fields/<parametre_type>')
def get_parametre_fields(parametre_type):
    """Belirli bir parametre tipinin alanlarını getirir."""
    try:
        fields = load_parametre_fields()
        parametre_fields = fields.get(parametre_type, [])
        return jsonify({'success': True, 'fields': parametre_fields})
    except Exception as e:
        print(f"Parametre alanları getirilirken hata: {e}")
        return jsonify({'success': False, 'message': 'Sunucu hatası!'}), 500

@app.route('/update_parametre_field', methods=['POST'])
def update_parametre_field():
    """Parametre alanını günceller."""
    try:
        data = request.get_json()
        parametre_type = data.get('parametreType')
        old_field_name = data.get('oldFieldName')
        new_field_name = data.get('newFieldName')
        
        if not parametre_type or not old_field_name or not new_field_name:
            return jsonify({'success': False, 'message': 'Tüm alanlar gerekli!'}), 400
        
        # Mevcut alanları yükle
        fields = load_parametre_fields()
        
        # Parametre tipi var mı kontrol et
        if parametre_type not in fields:
            return jsonify({'success': False, 'message': 'Parametre tipi bulunamadı!'}), 404
        
        # Eski alan var mı kontrol et
        if old_field_name not in fields[parametre_type]:
            return jsonify({'success': False, 'message': f'"{old_field_name}" alanı bulunamadı!'}), 404
        
        # Yeni alan adı zaten var mı kontrol et
        if new_field_name in fields[parametre_type] and new_field_name != old_field_name:
            return jsonify({'success': False, 'message': f'"{new_field_name}" alanı zaten mevcut!'}), 400
        
        # Alanı güncelle
        field_index = fields[parametre_type].index(old_field_name)
        fields[parametre_type][field_index] = new_field_name
        
        # Kaydet
        if save_parametre_fields(fields):
            return jsonify({
                'success': True, 
                'message': f'"{old_field_name}" alanı "{new_field_name}" olarak güncellendi!'
            })
        else:
            return jsonify({'success': False, 'message': 'Alan güncellenirken hata oluştu!'}), 500
            
    except Exception as e:
        print(f"Parametre alanı güncellenirken hata: {e}")
        return jsonify({'success': False, 'message': 'Sunucu hatası!'}), 500

@app.route('/delete_parametre_field', methods=['POST'])
def delete_parametre_field():
    """Parametre alanını siler."""
    try:
        data = request.get_json()
        parametre_type = data.get('parametreType')
        field_name = data.get('fieldName')
        
        if not parametre_type or not field_name:
            return jsonify({'success': False, 'message': 'Parametre tipi ve alan adı gerekli!'}), 400
        
        # Mevcut alanları yükle
        fields = load_parametre_fields()
        
        # Parametre tipi var mı kontrol et
        if parametre_type not in fields:
            return jsonify({'success': False, 'message': 'Parametre tipi bulunamadı!'}), 404
        
        # Alan var mı kontrol et
        if field_name not in fields[parametre_type]:
            return jsonify({'success': False, 'message': f'"{field_name}" alanı bulunamadı!'}), 404
        
        # Alanı sil
        fields[parametre_type].remove(field_name)
        
        # Kaydet
        if save_parametre_fields(fields):
            return jsonify({
                'success': True, 
                'message': f'"{field_name}" alanı başarıyla silindi!'
            })
        else:
            return jsonify({'success': False, 'message': 'Alan silinirken hata oluştu!'}), 500
            
    except Exception as e:
        print(f"Parametre alanı silinirken hata: {e}")
        return jsonify({'success': False, 'message': 'Sunucu hatası!'}), 500

@app.route('/import_parametre_fields', methods=['POST'])
def import_parametre_fields():
    """Tablodaki formatı Excel dosyasından içe aktarır."""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': 'Dosya seçilmedi!'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': 'Dosya seçilmedi!'}), 400
        
        if not file.filename.endswith(('.xlsx', '.xls')):
            return jsonify({'success': False, 'message': 'Sadece Excel dosyaları (.xlsx, .xls) kabul edilir!'}), 400
        
        # Excel dosyasını header olmadan oku
        df = pd.read_excel(file, header=None)
        
        # Tablodaki formatı kontrol et
        if len(df.columns) < 8:
            return jsonify({'success': False, 'message': 'Dosya formatı uygun değil! En az 8 sütun olmalı.'}), 400
        
        # Başlık satırını kontrol et
        first_row = df.iloc[0].tolist()
        expected_headers = [
            'TOZ,A.MET,SÜLF.A,HF,HCL,AMON.,FORM.,CR6,FOSF.A.,HCN',
            'NEM', 'YG', 'TOC', 'VOC', 'PM10', 'ÇT', 'boş', 'boş'
        ]
        
        # Başlık kontrolü (ilk 6 sütun için)
        for i in range(6):
            if str(first_row[i]).strip() != expected_headers[i]:
                return jsonify({'success': False, 'message': f'Başlık formatı uygun değil! Sütun {i+1}: {first_row[i]}'}), 400
        
        # TARİH satırını kontrol et
        second_row = df.iloc[1].tolist()
        for i in range(6):
            if str(second_row[i]).strip() != 'TARİH':
                return jsonify({'success': False, 'message': f'TARİH satırı uygun değil! Sütun {i+1}: {second_row[i]}'}), 400
        
        # Veri satırlarını al (3. satırdan itibaren)
        data_rows = []
        for i in range(2, len(df)):
            row = df.iloc[i].tolist()
            # Boş satırları atla
            if any(str(cell).strip() for cell in row[:6]):
                data_rows.append([str(cell).strip() if pd.notna(cell) else '' for cell in row])
        
        # Yeni alanlar için dictionary oluştur
        new_fields = {
            'genel': [], 'nem': [], 'yg': [], 'toc': [], 'voc': [], 'pm10': [], 'ct': []
        }
        
        # Her sütun için alanları topla
        for row in data_rows:
            if len(row) >= 8:
                # 1. sütun: genel parametreler
                if row[0]:
                    new_fields['genel'].append(row[0])
                
                # 2. sütun: nem parametreleri
                if row[1]:
                    new_fields['nem'].append(row[1])
                
                # 3. sütun: yg parametreleri
                if row[2]:
                    new_fields['yg'].append(row[2])
                
                # 4. sütun: toc parametreleri
                if row[3]:
                    new_fields['toc'].append(row[3])
                
                # 5. sütun: voc parametreleri
                if row[4]:
                    new_fields['voc'].append(row[4])
                
                # 6. sütun: pm10 parametreleri
                if row[5]:
                    new_fields['pm10'].append(row[5])
                
                # 7. sütun: ct parametreleri
                if row[6]:
                    new_fields['ct'].append(row[6])
        
        # Boş listeleri temizle
        new_fields = {k: v for k, v in new_fields.items() if v}
        
        # Kaydet
        save_parametre_fields(new_fields)
        
        total_fields = sum(len(fields) for fields in new_fields.values())
        return jsonify({
            'success': True, 
            'message': f'{total_fields} alan başarıyla içe aktarıldı!'
        })
            
    except Exception as e:
        print(f"Parametre formatı içe aktarılırken hata: {e}")
        return jsonify({'success': False, 'message': f'Dosya okuma hatası: {str(e)}'}), 500

@app.route('/export_parametre_fields')
def export_parametre_fields():
    """Tablodaki formatı Excel dosyasına dışa aktarır."""
    try:
        # Tabloda gördüğünüz formatı tam olarak oluştur
        excel_data = []
        
        # Başlık satırı
        excel_data.append([
            'TOZ,A.MET,SÜLF.A,HF,HCL,AMON.,FORM.,CR6,FOSF.A.,HCN',
            'NEM', 'YG', 'TOC', 'VOC', 'PM10', 'ÇT', 'boş', 'boş'
        ])
        
        # TARİH satırı
        excel_data.append(['TARİH', 'TARİH', 'TARİH', 'TARİH', 'TARİH', 'TARİH', 'TARİH', '', ''])
        
        # Veri satırları - tablodaki gördüğünüz sırayla
        data_rows = [
            ['METOT', '1İMP-İ', 'B.SIC', 'B.NEM', 'GAZ HAC.', 'METOT', 'T.İÇİ-1', '', ''],
            ['NOZZLE ÇAP', '1-İMP-S', 'O2', 'B.SIC', 'GAZ.SIC.', 'ORT.SIC', 'T.İÇİ-2', '', ''],
            ['TRAVERS', '2-İMP-İ', 'CO', 'TOC(PPM)', 'SEY.GAZ.HAC', 'ORT.NEM', 'T.İÇİ-3', '', ''],
            ['B.HIZ', '2-İMP-S', 'NO', 'KK1-SPAN', 'SEY.GAZ.SIC', 'ORT.RUZ.HIZ', 'T.İÇİ-4', '', ''],
            ['B.SIC', '3-İMP-İ', 'NOX', 'KK1-SPAN', '', 'ÇEK.HACİM', 'T-DIŞ-1', '', ''],
            ['B.BAS(KPA)', '3-İMP-S', 'SO2', 'KK2-SPAN', '', 'SYC.İLK', 'T.DIŞ-2', '', ''],
            ['B.NEM(G/M3)', 'HAC.', 'KK1-O2', 'KK2-SPAN', '', 'SYC.SON', 'T.DIŞ-3', '', ''],
            ['B.NEM(%)', '', 'KK1-CO', '', '', 'ISDL', 'TDİS-4', '', ''],
            ['SYC.HAC.', '', 'KK1-NO', '', '', '', 'İLK KURULUM', '', ''],
            ['SYC.İLK', '', 'KK1-SO2', '', '', '', '2. KURULUM', '', ''],
            ['SYC.SON', '', 'KK2-O2', '', '', '', 'TOPLAMA', '', ''],
            ['SYC.SIC', '', 'KK2-CO', '', '', '', '', '', ''],
            ['DEBİ', '', 'KK2-NO', '', '', '', '', '', ''],
            ['ISDL', '', 'KK2-SO2', '', '', '', '', '', ''],
            ['', '', 'T90', '', '', '', '', '', '']
        ]
        
        excel_data.extend(data_rows)
        
        # DataFrame oluştur
        df = pd.DataFrame(excel_data)
        
        # Excel dosyası oluştur
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
            df.to_excel(tmp.name, index=False, header=False, engine='openpyxl')
            
            return send_file(
                tmp.name,
                as_attachment=True,
                download_name=f'parametre_format_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx',
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            
    except Exception as e:
        print(f"Parametre formatı dışa aktarılırken hata: {e}")
        return jsonify({'success': False, 'message': 'Dışa aktarma hatası!'}), 500


@app.route('/clear_parametre_fields', methods=['POST'])
def clear_parametre_fields():
    """Tüm parametre alanlarını siler."""
    try:
        # Boş dictionary oluştur
        empty_fields = {}
        
        # Kaydet
        save_parametre_fields(empty_fields)
        
        return jsonify({
            'success': True, 
            'message': 'Tüm parametre alanları başarıyla silindi!'
        })
            
    except Exception as e:
        print(f"Parametre alanları silinirken hata: {e}")
        return jsonify({'success': False, 'message': f'Silme hatası: {str(e)}'}), 500


@app.route('/export_selected_parametre_sahabil')
def export_selected_parametre_sahabil():
    """Seçilen parametre sahabil verilerini Excel olarak dışa aktarır."""
    # Session kontrolünü geçici olarak kaldırıyoruz
    # if not session.get('logged_in') or not can_read(session.get('role')):
    #     return jsonify({'success': False, 'error': 'Yetkisiz erişim'}), 403
    
    try:
        selected_ids = request.args.get('ids', '').split(',')
        if not selected_ids or selected_ids[0] == '':
            return jsonify({'success': False, 'error': 'Seçilen kayıt bulunamadı'})
        
        parametre_sahabil_data = load_parametre_sahabil()
        
        # Seçilen kayıtları filtrele
        selected_data = [record for record in parametre_sahabil_data if record.get('id') in selected_ids]
        
        if not selected_data:
            return jsonify({'success': False, 'error': 'Seçilen kayıtlar bulunamadı'})
        
        # DataFrame oluştur
        export_data = []
        for record in selected_data:
            row = {
                'Parametre Türü': record.get('parametre_turu', ''),
                'Parametre Adı': record.get('parametre_adi', ''),
                '1.ÖLÇ': record.get('olcum_verileri', {}).get('olcum_1', ''),
                '2.ÖLÇ': record.get('olcum_verileri', {}).get('olcum_2', ''),
                '3.ÖLÇ': record.get('olcum_verileri', {}).get('olcum_3', ''),
                'Oluşturma Tarihi': record.get('created_at', '')
            }
            export_data.append(row)
        
        df = pd.DataFrame(export_data)
        
        # Excel dosyası oluştur
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
            df.to_excel(tmp.name, index=False, sheet_name='Parametre Sahabil', engine='openpyxl')
            tmp.flush()
            
            return send_file(
                tmp.name,
                as_attachment=True,
                download_name=f'parametre_sahabil_selected_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx',
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            
    except Exception as e:
        print(f"Seçilen parametre sahabil export hatası: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/baca_bilgileri')
def baca_bilgileri():
    if not session.get('logged_in') or not can_read(session.get('role')):
        return redirect(url_for('login'))
    
    baca_bilgileri_list = load_baca_bilgileri()
    return render_template('baca_bilgileri.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         baca_bilgileri=baca_bilgileri_list)

@app.route('/baca_bilgileri/add', methods=['POST'])
def add_baca_bilgisi():
    if not session.get('logged_in') or not can_write(session.get('role')):
        return redirect(url_for('login'))
    
    try:
        baca_bilgileri_list = load_baca_bilgileri()
        
        new_baca = {
            'id': str(uuid4()),
            'firma_adi': request.form.get('firma_adi', ''),
            'olcum_kodu': request.form.get('olcum_kodu', ''),
            'baca_no': request.form.get('baca_no', '').strip(),
            'yakit_turu': request.form.get('yakit_turu', ''),
            'isil_guc': request.form.get('isil_guc', ''),
            'cati_sekli': request.form.get('cati_sekli', ''),
            'kaynak_turu': request.form.get('kaynak_turu', ''),
            'baca_sekli': request.form.get('baca_sekli', ''),
            'baca_olcusu': request.form.get('baca_olcusu', ''),
            'yerden_yuk': request.form.get('yerden_yuk', ''),
            'cati_yuk': request.form.get('cati_yuk', ''),
            'ruzgar_hiz': request.form.get('ruzgar_hiz', ''),
            'ort_sic': request.form.get('ort_sic', ''),
            'ort_nem': request.form.get('ort_nem', ''),
            'ort_bas': request.form.get('ort_bas', ''),
            'a_baca': request.form.get('a_baca', ''),
            'b_baca': request.form.get('b_baca', ''),
            'c_delik': request.form.get('c_delik', ''),
            'foto': ''  # Fotoğraf işlemi daha sonra eklenecek
        }
        
        if not new_baca['baca_no']:
            flash('Baca numarası zorunludur!', 'error')
            return redirect(url_for('baca_bilgileri'))
        
        baca_bilgileri_list.append(new_baca)
        
        if save_baca_bilgileri(baca_bilgileri_list):
            flash('Baca bilgisi başarıyla eklendi!', 'success')
        else:
            flash('Baca bilgisi eklenirken hata oluştu!', 'error')
            
    except Exception as e:
        print(f"Baca bilgisi ekleme hatası: {e}")
        flash(f'Baca bilgisi eklenirken hata oluştu: {str(e)}', 'error')
    
    return redirect(url_for('baca_bilgileri'))

@app.route('/baca_bilgileri/edit/<baca_id>', methods=['POST'])
def edit_baca_bilgisi(baca_id):
    if not session.get('logged_in') or not can_edit(session.get('role')):
        return redirect(url_for('login'))
    
    try:
        baca_bilgileri_list = load_baca_bilgileri()
        baca = next((b for b in baca_bilgileri_list if b['id'] == baca_id), None)
        
        if not baca:
            flash('Baca bilgisi bulunamadı!', 'error')
            return redirect(url_for('baca_bilgileri'))
        
        # Baca bilgilerini güncelle
        baca['firma_adi'] = request.form.get('firma_adi', '')
        baca['olcum_kodu'] = request.form.get('olcum_kodu', '')
        baca['baca_no'] = request.form.get('baca_no', '').strip()
        baca['yakit_turu'] = request.form.get('yakit_turu', '')
        baca['isil_guc'] = request.form.get('isil_guc', '')
        baca['cati_sekli'] = request.form.get('cati_sekli', '')
        baca['kaynak_turu'] = request.form.get('kaynak_turu', '')
        baca['baca_sekli'] = request.form.get('baca_sekli', '')
        baca['baca_olcusu'] = request.form.get('baca_olcusu', '')
        baca['yerden_yuk'] = request.form.get('yerden_yuk', '')
        baca['cati_yuk'] = request.form.get('cati_yuk', '')
        baca['ruzgar_hiz'] = request.form.get('ruzgar_hiz', '')
        baca['ort_sic'] = request.form.get('ort_sic', '')
        baca['ort_nem'] = request.form.get('ort_nem', '')
        baca['ort_bas'] = request.form.get('ort_bas', '')
        baca['a_baca'] = request.form.get('a_baca', '')
        baca['b_baca'] = request.form.get('b_baca', '')
        baca['c_delik'] = request.form.get('c_delik', '')
        
        if not baca['baca_no']:
            flash('Baca numarası zorunludur!', 'error')
            return redirect(url_for('baca_bilgileri'))
        
        if save_baca_bilgileri(baca_bilgileri_list):
            flash('Baca bilgisi başarıyla güncellendi!', 'success')
        else:
            flash('Baca bilgisi güncellenirken hata oluştu!', 'error')
            
    except Exception as e:
        print(f"Baca bilgisi güncelleme hatası: {e}")
        flash(f'Baca bilgisi güncellenirken hata oluştu: {str(e)}', 'error')
    
    return redirect(url_for('baca_bilgileri'))

@app.route('/baca_bilgileri/delete/<baca_id>', methods=['POST'])
def delete_baca_bilgisi(baca_id):
    if not session.get('logged_in') or not can_delete(session.get('role')):
        return redirect(url_for('login'))
    
    try:
        baca_bilgileri_list = load_baca_bilgileri()
        baca_bilgileri_to_keep = [b for b in baca_bilgileri_list if b['id'] != baca_id]
        
        if len(baca_bilgileri_list) == len(baca_bilgileri_to_keep):
            flash('Silinecek baca bilgisi bulunamadı.', 'error')
        else:
            if save_baca_bilgileri(baca_bilgileri_to_keep):
                flash('Baca bilgisi başarıyla silindi.', 'success')
            else:
                flash('Baca bilgisi silinirken hata oluştu.', 'error')
                
    except Exception as e:
        print(f"Baca bilgisi silme hatası: {e}")
        flash(f'Baca bilgisi silinirken hata oluştu: {str(e)}', 'error')
    
    return redirect(url_for('baca_bilgileri'))

# Baca Para Yönetimi Route'ları
@app.route('/add_baca_para', methods=['POST'])
def add_baca_para():
    """Yeni baca parametresi ekler."""
    if not can_write(session.get('role')):
        flash('Bu işlem için yetkiniz yok!', 'error')
        return redirect(url_for('formlar'))
    
    try:
        baca_par_adi = request.form.get('baca_par_adi', '').strip()
        liste_icerigi = request.form.get('liste_icerigi', '').strip()
        
        if not baca_par_adi:
            flash('Baca parametre adı zorunludur!', 'error')
            return redirect(url_for('formlar'))
        
        # Baca parametrelerini yükle
        baca_paralar = load_baca_paralar()
        
        # Yeni parametre oluştur
        yeni_para = {
            'id': str(uuid4()),
            'baca_par_adi': baca_par_adi,
            'liste_icerigi': liste_icerigi,
            'created_at': datetime.now().isoformat()
        }
        
        baca_paralar.append(yeni_para)
        save_baca_paralar(baca_paralar)
        
        flash(f'Baca parametresi başarıyla eklendi: {baca_par_adi}', 'success')
        
    except Exception as e:
        flash(f'Hata oluştu: {str(e)}', 'error')
    
    # Başarılı ekleme sonrası baca bilgileri sekmesi ile yönlendir
    return redirect(url_for('formlar', tab='baca-bilgileri'))

@app.route('/edit_baca_para/<para_id>', methods=['POST'])
def edit_baca_para(para_id):
    """Baca parametresini düzenler."""
    if not can_edit(session.get('role')):
        flash('Bu işlem için yetkiniz yok!', 'error')
        return redirect(url_for('formlar'))
    
    try:
        baca_par_adi = request.form.get('baca_par_adi', '').strip()
        liste_icerigi = request.form.get('liste_icerigi', '').strip()
        
        if not baca_par_adi:
            flash('Baca parametre adı zorunludur!', 'error')
            return redirect(url_for('formlar'))
        
        # Baca parametrelerini yükle
        baca_paralar = load_baca_paralar()
        
        # Parametreyi bul ve güncelle
        for para in baca_paralar:
            if str(para.get('id')) == str(para_id):
                para['baca_par_adi'] = baca_par_adi
                para['liste_icerigi'] = liste_icerigi
                para['updated_at'] = datetime.now().isoformat()
                break
        
        save_baca_paralar(baca_paralar)
        flash(f'Baca parametresi başarıyla güncellendi: {baca_par_adi}', 'success')
        
    except Exception as e:
        flash(f'Hata oluştu: {str(e)}', 'error')
    
    # Başarılı güncelleme sonrası baca bilgileri sekmesi ile yönlendir
    return redirect(url_for('formlar', tab='baca-bilgileri'))

@app.route('/delete_baca_para/<para_id>', methods=['POST'])
def delete_baca_para(para_id):
    """Baca parametresini siler."""
    if not can_delete(session.get('role')):
        return jsonify({'success': False, 'error': 'Yetkiniz yok'})
    
    try:
        baca_paralar = load_baca_paralar()
        
        # Parametreyi bul ve sil
        for i, para in enumerate(baca_paralar):
            if str(para.get('id')) == str(para_id):
                deleted_para = baca_paralar.pop(i)
                save_baca_paralar(baca_paralar)
                return jsonify({'success': True, 'message': f'Baca parametresi silindi: {deleted_para.get("baca_par_adi", "Bilinmeyen")}'})
        
        return jsonify({'success': False, 'error': 'Baca parametresi bulunamadı'})
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Hata: {str(e)}'})

@app.route('/delete_selected_baca_para', methods=['POST'])
def delete_selected_baca_para():
    """Seçilen baca parametrelerini toplu olarak siler."""
    if not can_delete(session.get('role')):
        return jsonify({'success': False, 'error': 'Yetkiniz yok'})
    
    try:
        data = request.get_json()
        selected_ids = data.get('ids', [])
        
        if not selected_ids:
            return jsonify({'success': False, 'error': 'Silinecek parametre seçilmedi'})
        
        baca_paralar = load_baca_paralar()
        deleted_count = 0
        deleted_names = []
        
        # Seçilen parametreleri sil
        for para_id in selected_ids:
            for i, para in enumerate(baca_paralar):
                if str(para.get('id')) == str(para_id):
                    deleted_para = baca_paralar.pop(i)
                    deleted_count += 1
                    deleted_names.append(deleted_para.get('baca_par_adi', 'Bilinmeyen'))
                    break
        
        if deleted_count > 0:
            save_baca_paralar(baca_paralar)
            return jsonify({
                'success': True, 
                'message': f'{deleted_count} baca parametresi başarıyla silindi: {", ".join(deleted_names)}'
            })
        else:
            return jsonify({'success': False, 'error': 'Silinecek parametre bulunamadı'})
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Hata: {str(e)}'})

@app.route('/export_baca_bilgileri')
def export_baca_bilgileri():
    """Baca bilgilerini dışa aktarır."""
    if not can_read(session.get('role')):
        flash('Bu işlem için yetkiniz yok!', 'error')
        return redirect(url_for('formlar'))
    
    try:
        baca_paralar = load_baca_paralar()
        
        if not baca_paralar:
            flash('Dışa aktarılacak baca parametresi bulunamadı!', 'warning')
            return redirect(url_for('formlar'))
        
        # DataFrame oluştur
        df = pd.DataFrame(baca_paralar)
        
        # Geçici dosya oluştur
        with tempfile.NamedTemporaryFile(mode='w', suffix='.xlsx', delete=False, encoding='utf-8') as tmp_file:
            df.to_excel(tmp_file.name, index=False, engine='openpyxl')
            tmp_file_path = tmp_file.name
        
        # Dosyayı gönder
        return send_file(
            tmp_file_path,
            as_attachment=True,
            download_name=f'baca_parametreleri_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        
    except Exception as e:
        flash(f'Dışa aktarma hatası: {str(e)}', 'error')
        return redirect(url_for('formlar'))

def load_baca_paralar():
    """Baca parametrelerini JSON dosyasından yükler."""
    BACA_PARALAR_FILE = 'baca_paralar.json'
    if not os.path.exists(BACA_PARALAR_FILE):
        # Varsayılan parametreler - 26 alan sırasıyla
        default_paralar = [
            {
                'id': str(uuid4()),
                'baca_par_adi': 'BACA NO',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'YAKIT TÜRÜ',
                'liste_icerigi': 'BİOKÜTLE, DOĞAL GAZ, KÖMÜR, SIVI YAKIT',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'ISIL GÜÇ (MW)',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'ÇATI ŞEKLİ',
                'liste_icerigi': 'BAĞIMSIZ, DÜZ, EĞİK',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KAYNAK TÜRÜ',
                'liste_icerigi': 'YAKMA, PROSES',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'BACA ŞEKLİ',
                'liste_icerigi': 'DÖRTGEN, DAİRESEL',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'BACA ÖLÇÜSÜ',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'YERDEN YÜK.',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'ÇATI YÜK',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'RÜZGAR HIZ (M/S)',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'ORT SIC',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'ORT NEM',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'ORT BAS.',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'A-BACA',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'B-BACA',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'C-DELİK',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'FOTO',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KK1-O2',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KK1-CO',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KK1-NO',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KK1-SO2',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KK2-O2',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KK2-CO',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KK2-NO',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KK2-SO2',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'T90',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            }
        ]
        save_baca_paralar(default_paralar)
        return default_paralar
    
    try:
        with open(BACA_PARALAR_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Baca parametreleri yüklenirken hata: {e}")
        return []

def save_baca_paralar(baca_paralar_data):
    """Baca parametrelerini JSON dosyasına kaydeder."""
    # Yol başta tanımlandı: BACA_PARALAR_FILE
    try:
        with open(BACA_PARALAR_FILE, 'w', encoding='utf-8') as f:
            json.dump(baca_paralar_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Baca parametreleri kaydedilirken hata: {e}")
        return False

def load_parametre_sahabil():
    """Parametre sahabil verilerini JSON dosyasından yükler."""
    # Yol başta tanımlandı: PARAMETRE_SAHABIL_FILE
    if not os.path.exists(PARAMETRE_SAHABIL_FILE):
        return []
    try:
        with open(PARAMETRE_SAHABIL_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Parametre sahabil verileri yüklenirken hata: {e}")
        return []

def save_parametre_sahabil(parametre_sahabil_data):
    """Parametre sahabil verilerini JSON dosyasına kaydeder."""
    try:
        with open(PARAMETRE_SAHABIL_FILE, 'w', encoding='utf-8') as f:
            json.dump(parametre_sahabil_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Parametre sahabil verileri kaydedilirken hata: {e}")
        return False

def load_parametre_fields():
    """Parametre alanlarını JSON dosyasından yükler."""
    try:
        if not os.path.exists(PARAMETRE_FIELDS_FILE):
            return {}
        with open(PARAMETRE_FIELDS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Parametre alanları yüklenirken hata: {e}")
        return {}

def load_asgari_fiyatlar():
    try:
        if not os.path.exists(ASGARI_FIYATLAR_FILE):
            return []
        with open(ASGARI_FIYATLAR_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Asgari fiyatlar yüklenirken hata: {e}")
        return []

def save_asgari_fiyatlar(data):
    ok = _atomic_write_json(ASGARI_FIYATLAR_FILE, data, indent=2, ensure_ascii=False)
    if not ok:
        try:
            print("Asgari fiyatlar kaydedilemedi")
        except Exception:
            pass
    return bool(ok)

def load_par_saha_headers():
    """PAR_SAHA başlıklarını yükler."""
    try:
        if not os.path.exists(PAR_SAHA_HEADERS_FILE):
            return {"groups": []}
        with open(PAR_SAHA_HEADERS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"PAR SAHA header yüklenirken hata: {e}")
        return {"groups": []}

def save_par_saha_headers(data):
    """PAR_SAHA başlıklarını kaydeder."""
    try:
        return bool(_atomic_write_json(PAR_SAHA_HEADERS_FILE, data, indent=2, ensure_ascii=False))
    except Exception as e:
        print(f"PAR SAHA header kaydedilirken hata: {e}")
        return False

def save_parametre_fields(parametre_fields_data):
    """Parametre alanlarını JSON dosyasına kaydeder."""
    return bool(_atomic_write_json(PARAMETRE_FIELDS_FILE, parametre_fields_data, indent=2, ensure_ascii=False))

def load_forms():
    """Form verilerini JSON dosyasından yükler."""
    if not os.path.exists(FORMS_FILE):
        return []
    try:
        with open(FORMS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Form verileri yüklenirken hata: {e}")
        return []

def save_forms(forms_data):
    """Form verilerini JSON dosyasına kaydeder."""
    return bool(_atomic_write_json(FORMS_FILE, forms_data, indent=4, ensure_ascii=False))

@app.route('/api/baca_bilgileri/bulk_delete', methods=['POST'])
def api_bulk_delete_baca_bilgileri():
    """Seçili baca bilgilerini toplu olarak siler."""
    try:
        data = request.get_json()
        if not data or 'ids' not in data:
            return jsonify({'success': False, 'error': 'Silinecek kayıt ID\'leri belirtilmedi'}), 400
        
        record_ids = data['ids']
        if not isinstance(record_ids, list) or len(record_ids) == 0:
            return jsonify({'success': False, 'error': 'Geçersiz kayıt ID listesi'}), 400
        
        saved_baca_bilgileri = load_baca_bilgileri()
        original_count = len(saved_baca_bilgileri)
        
        # Seçili kayıtları filtrele
        filtered_records = []
        deleted_records = []
        
        for record in saved_baca_bilgileri:
            if record.get('id') in record_ids:
                deleted_records.append(record)
            else:
                filtered_records.append(record)
        
        # Silinen kayıt sayısını kontrol et
        if len(deleted_records) == 0:
            return jsonify({'success': False, 'error': 'Silinecek kayıt bulunamadı'}), 404
        
        # Dosyaya kaydet
        if save_baca_bilgileri(filtered_records):
            deleted_info = [f"{r.get('firma_adi')} - {r.get('olcum_kodu')} - {r.get('baca_adi')}" for r in deleted_records]
            print(f"Baca bilgileri toplu silindi: {', '.join(deleted_info)}")
            return jsonify({
                'success': True, 
                'message': f'{len(deleted_records)} kayıt başarıyla silindi',
                'deleted_count': len(deleted_records)
            })
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'}), 500
            
    except Exception as e:
        print(f"Baca bilgileri toplu silinirken hata: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/forms')
def api_forms():
    """Form listesini döndürür."""
    try:
        forms = load_forms()
        return jsonify(forms)
    except Exception as e:
        print(f"Formlar yüklenirken hata: {e}")
        return jsonify({'error': str(e)}), 500

# PAR_SAHA başlıklarını getir
@app.route('/api/par_saha_headers', methods=['GET'])
def api_get_par_saha_headers():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'}), 401
    data = load_par_saha_headers()
    return jsonify({'success': True, **data})

# PAR_SAHA başlıklarını kaydet
@app.route('/api/par_saha_headers', methods=['POST'])
def api_save_par_saha_headers():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'}), 401
    try:
        data = request.get_json() or {}
        if not isinstance(data, dict) or 'groups' not in data or not isinstance(data['groups'], list):
            return jsonify({'success': False, 'message': 'Geçersiz veri'}), 400
        ok = save_par_saha_headers({'groups': data['groups']})
        return jsonify({'success': ok})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/add_form', methods=['POST'])
def add_form():
    """Yeni form ekler."""
    try:
        formAdi = request.form.get('formAdi')
        formKodu = request.form.get('formKodu')
        yayinTarihi = request.form.get('yayinTarihi')
        revizyonTarihi = request.form.get('revizyonTarihi')
        revizyonNo = request.form.get('revizyonNo')
        
        if not all([formAdi, formKodu, yayinTarihi, revizyonTarihi, revizyonNo]):
            flash('Tüm alanları doldurun!', 'error')
            return redirect(url_for('formlar'))
        
        forms = load_forms()
        new_form = {
            'id': str(uuid4()),
            'formAdi': formAdi,
            'formKodu': formKodu,
            'yayinTarihi': yayinTarihi,
            'revizyonTarihi': revizyonTarihi,
            'revizyonNo': revizyonNo,
            'created_at': datetime.now().isoformat()
        }
        
        forms.append(new_form)
        
        if save_forms(forms):
            flash('Form başarıyla eklendi!', 'success')
        else:
            flash('Form eklenirken hata oluştu!', 'error')
            
        return redirect(url_for('formlar'))
        
    except Exception as e:
        print(f"Form eklenirken hata: {e}")
        flash('Form eklenirken hata oluştu!', 'error')
        return redirect(url_for('formlar'))

@app.route('/edit_form/<form_id>', methods=['POST'])
def edit_form(form_id):
    """Form düzenler."""
    try:
        formAdi = request.form.get('formAdi')
        formKodu = request.form.get('formKodu')
        yayinTarihi = request.form.get('yayinTarihi')
        revizyonTarihi = request.form.get('revizyonTarihi')
        revizyonNo = request.form.get('revizyonNo')
        
        if not all([formAdi, formKodu, yayinTarihi, revizyonTarihi, revizyonNo]):
            flash('Tüm alanları doldurun!', 'error')
            return redirect(url_for('formlar'))
        
        forms = load_forms()
        form_index = None
        
        for i, form in enumerate(forms):
            if form.get('id') == form_id:
                form_index = i
                break
        
        if form_index is not None:
            forms[form_index].update({
                'formAdi': formAdi,
                'formKodu': formKodu,
                'yayinTarihi': yayinTarihi,
                'revizyonTarihi': revizyonTarihi,
                'revizyonNo': revizyonNo,
                'updated_at': datetime.now().isoformat()
            })
            
            if save_forms(forms):
                flash('Form başarıyla güncellendi!', 'success')
            else:
                flash('Form güncellenirken hata oluştu!', 'error')
        else:
            flash('Form bulunamadı!', 'error')
            
        return redirect(url_for('formlar'))
        
    except Exception as e:
        print(f"Form güncellenirken hata: {e}")
        flash('Form güncellenirken hata oluştu!', 'error')
        return redirect(url_for('formlar'))

@app.route('/delete_form/<form_id>', methods=['POST'])
def delete_form(form_id):
    """Form siler."""
    try:
        forms = load_forms()
        form_to_delete = None
        
        for form in forms:
            if form.get('id') == form_id:
                form_to_delete = form
                break
        
        if form_to_delete:
            forms.remove(form_to_delete)
            
            if save_forms(forms):
                return jsonify({
                    'success': True,
                    'message': f'"{form_to_delete.get("formAdi")}" formu başarıyla silindi'
                })
            else:
                return jsonify({
                    'success': False,
                    'error': 'Form silinirken hata oluştu'
                })
        else:
            return jsonify({
                'success': False,
                'error': 'Form bulunamadı'
            })
            
    except Exception as e:
        print(f"Form silinirken hata: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        })

@app.route('/export_forms', methods=['POST'])
def export_forms():
    """Formları Excel, Word veya PDF formatında dışa aktarır."""
    try:
        data = request.get_json()
        format_type = data.get('format', 'excel')
        
        forms = load_forms()
        
        if not forms:
            return jsonify({'error': 'Dışa aktarılacak form bulunamadı'}), 404
        
        if format_type == 'excel':
            return export_forms_excel(forms)
        elif format_type == 'word':
            return export_forms_word(forms)
        elif format_type == 'pdf':
            return export_forms_pdf(forms)
        else:
            return jsonify({'error': 'Geçersiz format'}), 400
            
    except Exception as e:
        print(f"Form dışa aktarma hatası: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/firma_rapor_export', methods=['POST'])
def api_firma_rapor_export():
    """Bir firmanın tüm bacalarında yapılan tüm parametreleri içeren rapor oluşturur."""
    try:
        print("=== FIRMA RAPOR EXPORT BAŞLADI ===")
        data = request.get_json()
        print(f"Gelen data: {data}")
        
        firma_adi = data.get('firma_adi')
        olcum_kodu = data.get('olcum_kodu')
        print(f"Firma adı: {firma_adi}")
        print(f"Ölçüm kodu: {olcum_kodu}")
        
        if not firma_adi or not olcum_kodu:
            print("Firma adı veya ölçüm kodu eksik!")
            return jsonify({'error': 'Firma adı ve ölçüm kodu gerekli!'}), 400
        
        # Şablon kullanımını kontrol et (varsayılan olarak şablon kullan)
        use_template = data.get('use_template', True)
        print(f"Şablon kullanılacak: {use_template}")
        
        # Word raporu oluştur
        if use_template:
            print("create_firma_raporu_from_template çağrılıyor...")
            result = create_firma_raporu_from_template(firma_adi, olcum_kodu)
            print(f"create_firma_raporu_from_template sonucu: {type(result)}")
            return result
        else:
            print("create_firma_raporu çağrılıyor...")
            return create_firma_raporu(firma_adi, olcum_kodu)
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Firma rapor export hatası: {e}")
        print(f"Hata detayı: {error_details}")
        return jsonify({'error': f'Rapor oluşturma hatası: {str(e)}'}), 500

@app.route('/api/firma_rapor_pdf_export', methods=['POST'])
def api_firma_rapor_pdf_export():
    """Bir firmanın tüm bacalarında yapılan tüm parametreleri içeren PDF rapor oluşturur."""
    try:
        data = request.get_json()
        firma_adi = data.get('firma_adi')
        olcum_kodu = data.get('olcum_kodu')
        
        if not firma_adi or not olcum_kodu:
            return jsonify({'error': 'Firma adı ve ölçüm kodu gerekli!'}), 400
        
        # Word raporu oluştur
        word_response = create_firma_raporu_from_template(firma_adi, olcum_kodu)
        
        # Hata kontrolü
        if hasattr(word_response, 'get_data'):
            # Flask Response objesi
            word_data = word_response.get_data()
        elif isinstance(word_response, tuple):
            # Hata tuple'ı
            return word_response
        else:
            return jsonify({'error': 'Bilinmeyen response tipi'}), 500
        
        # Word dosyasını geçici olarak kaydet
        import tempfile
        import os
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_word:
            tmp_word.write(word_data)
            tmp_word_path = tmp_word.name
        
        # Word dosyasını PDF'e çevir
        try:
            from docx2pdf import convert
            
            # PDF dosya adını oluştur
            def clean_filename(text):
                tr_to_en = {
                    'ç': 'c', 'ğ': 'g', 'ı': 'i', 'ö': 'o', 'ş': 's', 'ü': 'u',
                    'Ç': 'C', 'Ğ': 'G', 'I': 'I', 'İ': 'I', 'Ö': 'O', 'Ş': 'S', 'Ü': 'U'
                }
                for tr_char, en_char in tr_to_en.items():
                    text = text.replace(tr_char, en_char)
                cleaned = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
                cleaned = cleaned.replace(' ', '-')
                while '--' in cleaned:
                    cleaned = cleaned.replace('--', '-')
                return cleaned
            
            firma_adi_clean = clean_filename(firma_adi)
            olcum_kodu_clean = clean_filename(olcum_kodu)
            
            pdf_filename = f"Firma_Raporu_{firma_adi_clean}_{olcum_kodu_clean}.pdf"
            pdf_path = os.path.join(tempfile.gettempdir(), pdf_filename)
            
            # Word dosyasını PDF'e çevir
            convert(tmp_word_path, pdf_path)
            
            # PDF dosyasını oku
            with open(pdf_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
            
            # Geçici dosyaları temizle
            os.unlink(tmp_word_path)
            os.unlink(pdf_path)
            
            # PDF dosyasını döndür
            response = make_response(pdf_content)
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = f'attachment; filename="{pdf_filename}"'
            return response
            
        except Exception as pdf_error:
            # PDF dönüştürme başarısız olursa Word dosyasını döndür
            print(f"PDF dönüştürme hatası: {pdf_error}")
            
            # Geçici Word dosyasını temizle
            if os.path.exists(tmp_word_path):
                os.unlink(tmp_word_path)
            
            # Word dosyasını döndür
            return word_response
        
    except Exception as e:
        print(f"PDF Firma rapor export hatası: {e}")
        return jsonify({'error': f'PDF Rapor oluşturma hatası: {str(e)}'}), 500

@app.route('/api/firma_olcum_detail_word_export', methods=['POST'])
def api_firma_olcum_detail_word_export():
    """Firma ölçüm detay bilgilerini Word formatında dışa aktarır."""
    try:
        data = request.get_json()
        olcum_id = data.get('olcum_id')
        
        if not olcum_id:
            return jsonify({'error': 'Ölçüm ID gerekli!'}), 400
        
        # Firma ölçüm verilerini yükle
        firma_olcumler = load_firma_olcum()
        
        # Seçili ölçümü bul
        olcum = None
        for item in firma_olcumler:
            if str(item.get('id')) == str(olcum_id):
                olcum = item
                break
        
        if not olcum:
            return jsonify({'error': 'Ölçüm bulunamadı'}), 404
        
        # Firma kayıt bilgilerini al
        firma_kayitlar = load_firma_kayit()
        firma_bilgisi = None
        for firma in firma_kayitlar:
            if firma.get('firmaAdi') == olcum.get('firma_adi'):
                firma_bilgisi = firma
                break
        
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Bu ölçüme ait bacaları bul
        olcum_bacalar = []
        for baca in baca_bilgileri:
            if (baca.get('firma_adi') == olcum.get('firma_adi') and 
                baca.get('olcum_kodu') == olcum.get('olcum_kodu')):
                olcum_bacalar.append(baca)
        
        # Parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Bu ölçüme ait parametreleri bul
        olcum_parametreleri = []
        for parametre in parametre_olcumleri:
            if (parametre.get('firma_adi') == olcum.get('firma_adi') and 
                parametre.get('olcum_kodu') == olcum.get('olcum_kodu')):
                olcum_parametreleri.append(parametre)
        
        # Word dokümanı oluştur
        doc = create_firma_olcum_detail_word_document(olcum, firma_bilgisi, olcum_bacalar, olcum_parametreleri)
        return doc
        
    except Exception as e:
        print(f"Firma ölçüm detay Word export hatası: {e}")
        import traceback
        print(f"DEBUG: Hata detayı: {traceback.format_exc()}")
        return jsonify({'error': f'Word export hatası: {str(e)}'}), 500

@app.route('/api/firma_olcum_word_export', methods=['POST'])
def api_firma_olcum_word_export():
    """Firma ölçüm bilgilerini Word formatında dışa aktarır."""
    try:
        data = request.get_json()
        olcum_ids = data.get('olcum_ids', [])
        
        if not olcum_ids:
            return jsonify({'error': 'En az bir ölçüm ID gerekli!'}), 400
        
        # Firma ölçüm verilerini yükle
        firma_olcumler = load_firma_olcum()
        
        # Seçili ölçümleri filtrele
        if olcum_ids:
            filtered_olcumler = [item for item in firma_olcumler if str(item.get('id')) in olcum_ids]
        else:
            filtered_olcumler = firma_olcumler
        
        if not filtered_olcumler:
            return jsonify({'error': 'Dışa aktarılacak veri bulunamadı'}), 404
        
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Her ölçüm için Word dokümanı oluştur
        documents = []
        for olcum in filtered_olcumler:
            # Bu ölçüme ait bacaları bul
            olcum_bacalar = []
            for baca in baca_bilgileri:
                if (baca.get('firma_adi') == olcum.get('firma_adi') and 
                    baca.get('olcum_kodu') == olcum.get('olcum_kodu')):
                    olcum_bacalar.append(baca)
            
            # Bu ölçüme ait parametreleri bul
            olcum_parametreleri = []
            for parametre in parametre_olcumleri:
                if (parametre.get('firma_adi') == olcum.get('firma_adi') and 
                    parametre.get('olcum_kodu') == olcum.get('olcum_kodu')):
                    olcum_parametreleri.append(parametre)
            
            # Word dokümanı oluştur
            doc = create_firma_olcum_word_document(olcum, olcum_bacalar, olcum_parametreleri)
            documents.append(doc)
        
        # Eğer tek doküman varsa onu döndür, birden fazla varsa birleştir
        if len(documents) == 1:
            return documents[0]
        else:
            # Birden fazla dokümanı birleştir
            merged_doc = merge_word_documents(documents)
            return merged_doc
        
    except Exception as e:
        print(f"Firma ölçüm Word export hatası: {e}")
        import traceback
        print(f"DEBUG: Hata detayı: {traceback.format_exc()}")
        return jsonify({'error': f'Word export hatası: {str(e)}'}), 500

@app.route('/api/firma_olcum_pdf_export', methods=['POST'])
def api_firma_olcum_pdf_export():
    """Firma ölçüm bilgilerini PDF formatında dışa aktarır."""
    try:
        data = request.get_json()
        olcum_ids = data.get('olcum_ids', [])
        
        if not olcum_ids:
            return jsonify({'error': 'En az bir ölçüm ID gerekli!'}), 400
        
        # Firma ölçüm verilerini yükle
        firma_olcumler = load_firma_olcum()
        
        # Seçili ölçümleri filtrele
        if olcum_ids:
            filtered_olcumler = [item for item in firma_olcumler if str(item.get('id')) in olcum_ids]
        else:
            filtered_olcumler = firma_olcumler
        
        if not filtered_olcumler:
            return jsonify({'error': 'Dışa aktarılacak veri bulunamadı'}), 404
        
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Her ölçüm için Word dokümanı oluştur
        documents = []
        for olcum in filtered_olcumler:
            # Bu ölçüme ait bacaları bul
            olcum_bacalar = []
            for baca in baca_bilgileri:
                if (baca.get('firma_adi') == olcum.get('firma_adi') and 
                    baca.get('olcum_kodu') == olcum.get('olcum_kodu')):
                    olcum_bacalar.append(baca)
            
            # Bu ölçüme ait parametreleri bul
            olcum_parametreleri = []
            for parametre in parametre_olcumleri:
                if (parametre.get('firma_adi') == olcum.get('firma_adi') and 
                    parametre.get('olcum_kodu') == olcum.get('olcum_kodu')):
                    olcum_parametreleri.append(parametre)
            
            # Word dokümanı oluştur
            doc = create_firma_olcum_word_document(olcum, olcum_bacalar, olcum_parametreleri)
            documents.append(doc)
        
        # Eğer tek doküman varsa onu kullan, birden fazla varsa birleştir
        if len(documents) == 1:
            word_doc = documents[0]
        else:
            # Birden fazla dokümanı birleştir
            word_doc = merge_word_documents(documents)
        
        # Word dosyasını geçici olarak kaydet
        from io import BytesIO
        import tempfile
        import os
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_word:
            tmp_word.write(word_doc.get_data())
            tmp_word_path = tmp_word.name
        
        # Word dosyasını PDF'e çevir
        try:
            from docx2pdf import convert
            
            # PDF dosya adını oluştur
            def clean_filename(text):
                tr_to_en = {
                    'ç': 'c', 'ğ': 'g', 'ı': 'i', 'ö': 'o', 'ş': 's', 'ü': 'u',
                    'Ç': 'C', 'Ğ': 'G', 'I': 'I', 'İ': 'I', 'Ö': 'O', 'Ş': 'S', 'Ü': 'U'
                }
                for tr_char, en_char in tr_to_en.items():
                    text = text.replace(tr_char, en_char)
                cleaned = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
                cleaned = cleaned.replace(' ', '-')
                while '--' in cleaned:
                    cleaned = cleaned.replace('--', '-')
                return cleaned
            
            # Birden fazla ölçüm varsa genel bir isim kullan
            if len(filtered_olcumler) == 1:
                olcum = filtered_olcumler[0]
                firma_adi = olcum.get('firma_adi', '')
                firma_kelimeleri = firma_adi.split()
                firma_ilk_2_kelime = ' '.join(firma_kelimeleri[:2]) if len(firma_kelimeleri) >= 2 else firma_adi
                olcum_kodu_clean = clean_filename(olcum.get('olcum_kodu', ''))
                pdf_filename = f"{clean_filename(firma_ilk_2_kelime)}_{olcum_kodu_clean}_SAHAFORM.pdf"
            else:
                # Birden fazla ölçüm için genel isim
                pdf_filename = f"Firma_Olcum_{len(filtered_olcumler)}_kayit_SAHAFORM.pdf"
            
            pdf_path = os.path.join(tempfile.gettempdir(), pdf_filename)
            
            # Word dosyasını PDF'e çevir
            convert(tmp_word_path, pdf_path)
            
            # PDF dosyasını oku
            with open(pdf_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
            
            # Geçici dosyaları temizle
            os.unlink(tmp_word_path)
            os.unlink(pdf_path)
            
            # PDF dosyasını döndür
            response = make_response(pdf_content)
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = f'attachment; filename="{pdf_filename}"'
            return response
            
        except Exception as pdf_error:
            # PDF dönüştürme başarısız olursa Word dosyasını döndür
            print(f"PDF dönüştürme hatası: {pdf_error}")
            
            # Geçici Word dosyasını temizle
            if os.path.exists(tmp_word_path):
                os.unlink(tmp_word_path)
            
            # Word dosyasını döndür
            return word_doc
        
    except Exception as e:
        print(f"Firma ölçüm PDF export hatası: {e}")
        return jsonify({'error': f'PDF export hatası: {str(e)}'}), 500

@app.route('/api/firma_olcum_detail_excel_export', methods=['POST'])
def api_firma_olcum_detail_excel_export():
    """Firma ölçüm detay bilgilerini Excel formatında dışa aktarır."""
    try:
        data = request.get_json()
        olcum_id = data.get('olcum_id')
        
        if not olcum_id:
            return jsonify({'error': 'Ölçüm ID gerekli!'}), 400
        
        # Firma ölçüm verilerini yükle
        firma_olcumler = load_firma_olcum()
        
        # Seçili ölçümü bul
        olcum = None
        for item in firma_olcumler:
            if str(item.get('id')) == str(olcum_id):
                olcum = item
                break
        
        if not olcum:
            return jsonify({'error': 'Ölçüm bulunamadı'}), 404
        
        # Firma kayıt bilgilerini al
        firma_kayitlar = load_firma_kayit()
        firma_bilgisi = None
        for firma in firma_kayitlar:
            if firma.get('firmaAdi') == olcum.get('firma_adi'):
                firma_bilgisi = firma
                break
        
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Bu ölçüme ait bacaları bul
        olcum_bacalar = []
        for baca in baca_bilgileri:
            if (baca.get('firma_adi') == olcum.get('firma_adi') and 
                baca.get('olcum_kodu') == olcum.get('olcum_kodu')):
                olcum_bacalar.append(baca)
        
        # Excel için DataFrame oluştur
        df_data = []
        
        # Temel bilgiler
        df_data.append(['TEMEL BİLGİLER', ''])
        df_data.append(['Firma Adı', firma_bilgisi.get('firmaAdi', '') if firma_bilgisi else olcum.get('firma_adi', '')])
        df_data.append(['Ölçüm Kodu', olcum.get('olcum_kodu', '')])
        df_data.append(['Başlangıç Tarihi', olcum.get('baslangic_tarihi', '')])
        df_data.append(['Bitiş Tarihi', olcum.get('bitis_tarihi', '')])
        df_data.append(['Baca Sayısı', olcum.get('baca_sayisi', '')])
        df_data.append(['Durum', olcum.get('durum', '')])
        df_data.append(['', ''])
        
        # İletişim bilgileri
        df_data.append(['İLETİŞİM BİLGİLERİ', ''])
        df_data.append(['İl', firma_bilgisi.get('il', '') if firma_bilgisi else olcum.get('il', '')])
        df_data.append(['İlçe', firma_bilgisi.get('ilce', '') if firma_bilgisi else olcum.get('ilce', '')])
        df_data.append(['Yetkili', firma_bilgisi.get('yetkiliAdi', '') if firma_bilgisi else olcum.get('yetkili', '')])
        df_data.append(['Telefon', str(firma_bilgisi.get('yetkiliTel', '')) if firma_bilgisi else olcum.get('telefon', '')])
        df_data.append(['Personel', ', '.join(olcum.get('personel', []))])
        df_data.append(['', ''])
        
        # Baca-Parametre matrisi
        df_data.append(['BACA-PARAMETRE MATRİSİ', ''])
        
        # Parametre başlıkları
        baca_parametreleri = olcum.get('baca_parametreleri', {})
        all_params = set()
        for baca_params in baca_parametreleri.values():
            for param in baca_params:
                all_params.add(param)
        
        param_headers = ['Baca Adı'] + list(all_params)
        df_data.append(param_headers)
        
        # Baca satırları
        for baca_adi, parametreler in baca_parametreleri.items():
            row = [baca_adi]
            for param in all_params:
                if param in parametreler:
                    row.append('X')
                else:
                    row.append('')
            df_data.append(row)
        
        df = pd.DataFrame(df_data)
        
        # Geçici dosya oluştur
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
            # Excel dosyasını oluştur
            with pd.ExcelWriter(tmp_file.name, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, header=False, sheet_name='Firma Ölçüm Detayı')
                
                # Sütun genişliklerini ayarla
                worksheet = writer.sheets['Firma Ölçüm Detayı']
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
            
            # Dosya adını oluştur
            firma_adi = firma_bilgisi.get('firmaAdi', '') if firma_bilgisi else olcum.get('firma_adi', '')
            olcum_kodu = olcum.get('olcum_kodu', '')
            filename = f"{firma_adi}_{olcum_kodu}_DETAY.xlsx"
            
            return send_file(
                tmp_file.name,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
        
    except Exception as e:
        print(f"Firma ölçüm detay Excel export hatası: {e}")
        import traceback
        print(f"DEBUG: Hata detayı: {traceback.format_exc()}")
        return jsonify({'error': f'Excel export hatası: {str(e)}'}), 500

@app.route('/api/firma_olcum_excel_export', methods=['POST'])
def api_firma_olcum_excel_export():
    """Firma ölçüm bilgilerini Excel formatında dışa aktarır - Basit format."""
    try:
        # Pandas kütüphanesini import et
        try:
            import pandas as pd
            print(f"DEBUG: pandas başarıyla import edildi: {pd}")
        except ImportError as e:
            print(f"DEBUG: pandas import hatası: {e}")
            return jsonify({'error': 'pandas kütüphanesi yüklü değil'}), 500
        data = request.get_json()
        olcum_ids = data.get('olcum_ids', [])
        
        if not olcum_ids:
            return jsonify({'error': 'En az bir ölçüm ID gerekli!'}), 400
        
        # Firma ölçümlerini yükle
        firma_olcumler = load_firma_olcum()
        
        # Seçili ölçümleri filtrele
        if olcum_ids:
            filtered_olcumler = [item for item in firma_olcumler if str(item.get('id')) in olcum_ids]
        else:
            filtered_olcumler = firma_olcumler
        
        if not filtered_olcumler:
            return jsonify({'error': 'Dışa aktarılacak veri bulunamadı'}), 404
        
        # Firma kayıt bilgilerini al
        firma_kayitlar = load_firma_kayit()
        
        # Excel için DataFrame oluştur - basit format
        df_data = []
        for index, olcum in enumerate(filtered_olcumler, 1):
            # Firma bilgilerini bul
            firma_bilgisi = None
            for firma in firma_kayitlar:
                if firma.get('firmaAdi') == olcum.get('firma_adi'):
                    firma_bilgisi = firma
                    break
            
            # Tarih formatını düzelt
            baslangic_tarihi = olcum.get('baslangic_tarihi', '')
            bitis_tarihi = olcum.get('bitis_tarihi', '')
            
            # Tarihleri DD.MM.YY formatına çevir
            if baslangic_tarihi:
                try:
                    baslangic_date = datetime.strptime(baslangic_tarihi, '%Y-%m-%d')
                    baslangic_tarihi = baslangic_date.strftime('%d.%m.%y')
                except:
                    pass
            
            if bitis_tarihi:
                try:
                    bitis_date = datetime.strptime(bitis_tarihi, '%Y-%m-%d')
                    bitis_tarihi = bitis_date.strftime('%d.%m.%y')
                except:
                    pass
            
            # Personel listesini string'e çevir
            personel = olcum.get('personel', [])
            if isinstance(personel, list):
                personel_str = ', '.join(personel)
            else:
                personel_str = str(personel) if personel else ''
            
            # Parametreleri al ve sayılarını hesapla
            baca_parametreleri = olcum.get('baca_parametreleri', {})
            parametre_sayilari = {}
            
            for baca_params in baca_parametreleri.values():
                for parametre in baca_params:
                    if parametre in parametre_sayilari:
                        parametre_sayilari[parametre] += 1
                    else:
                        parametre_sayilari[parametre] = 1
            
            # Parametreleri sayılarıyla birlikte formatla
            parametre_listesi = []
            for parametre, sayi in parametre_sayilari.items():
                parametre_listesi.append(f"{parametre} ({sayi})")
            
            parametre_str = ', '.join(parametre_listesi) if parametre_listesi else ''
            
            row = {
                'RA': index,
                'FIRMA': firma_bilgisi.get('firmaAdi', '') if firma_bilgisi else olcum.get('firma_adi', ''),
                'OLC_KOD': olcum.get('olcum_kodu', ''),
                'BAS TRH': baslangic_tarihi,
                'BIT TAR': bitis_tarihi,
                'BACA_SAY': olcum.get('baca_sayisi', ''),
                'PARAMETRE': parametre_str,
                'PER.': personel_str,
                'IL': firma_bilgisi.get('il', '') if firma_bilgisi else olcum.get('il', ''),
                'ILCE': firma_bilgisi.get('ilce', '') if firma_bilgisi else olcum.get('ilce', ''),
                'YETK': firma_bilgisi.get('yetkiliAdi', '') if firma_bilgisi else olcum.get('yetkili', ''),
                'TEL': str(firma_bilgisi.get('yetkiliTel', '')) if firma_bilgisi else olcum.get('telefon', ''),
                'DURU': olcum.get('durum', '')
            }
            
            df_data.append(row)
        
        print(f"DEBUG: df_data hazırlandı: {df_data}")
        df = pd.DataFrame(df_data)
        print(f"DEBUG: DataFrame oluşturuldu: {df}")
        
        # Geçici dosya oluştur
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
            # Excel dosyasını oluştur
            with pd.ExcelWriter(tmp_file.name, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Firma Ölçümleri')
                
                # Sütun genişliklerini ayarla
                worksheet = writer.sheets['Firma Ölçümleri']
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
            
            # Dosya adını oluştur
            if len(filtered_olcumler) == 1:
                olcum = filtered_olcumler[0]
                filename = f"{olcum.get('firma_adi', 'bilinmeyen')}_{olcum.get('olcum_kodu', 'bilinmeyen')}_OLCUM.xlsx"
            else:
                filename = f"TUM_FIRMA_OLCUMLERI_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            
            return send_file(
                tmp_file.name,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
        
    except Exception as e:
        print(f"Firma ölçüm Excel export hatası: {e}")
        return jsonify({'error': f'Excel export hatası: {str(e)}'}), 500

def create_firma_olcum_detail_word_document(olcum, firma_bilgisi, olcum_bacalar, olcum_parametreleri):
    """Firma ölçüm detay bilgilerini Word şablonu kullanarak oluşturur."""
    try:
        print(f"DEBUG: create_firma_olcum_detail_word_document başladı")
        
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            print(f"DEBUG: python-docx kütüphanesi yüklü değil")
            return jsonify({'error': 'python-docx kütüphanesi yüklü değil'}), 500
        
        # Şablon dosyasını yükle
        template_path = os.path.join(os.path.dirname(__file__), 'static', 'images', 'EMISYON_OLCUM_BILGI_FORMU.docx')
        print(f"DEBUG: Template path = {template_path}")
        print(f"DEBUG: Template exists = {os.path.exists(template_path)}")
        
        if not os.path.exists(template_path):
            print(f"DEBUG: Word şablonu bulunamadı")
            return jsonify({'error': 'Word şablonu bulunamadı'}), 500
        
        # Şablonu yükle
        print(f"DEBUG: Document yükleniyor...")
        doc = Document(template_path)
        print(f"DEBUG: Document başarıyla yüklendi")
        
        # Parametre listesini oluştur (firma_olcum verilerinden)
        baca_parametreleri = olcum.get('baca_parametreleri', {})
        
        # Tüm parametreleri topla
        all_params = set()
        for baca_params in baca_parametreleri.values():
            for param in baca_params:
                all_params.add(param)
        
        # Parametreleri sırala
        sorted_params = sorted(all_params)
        
        # Parametre listesini formatla (parametre adı + sayı)
        parametre_listesi = []
        for param in sorted_params:
            # Bu parametrenin kaç kez kullanıldığını say
            count = sum(1 for baca_params in baca_parametreleri.values() for p in baca_params if p == param)
            parametre_listesi.append(f"{param} ({count})")
        
        # Debug: Parametre listesi oluşturma
        print(f"DEBUG: Baca parametreleri: {baca_parametreleri}")
        print(f"DEBUG: Tüm parametreler: {all_params}")
        print(f"DEBUG: Sıralanmış parametreler: {sorted_params}")
        print(f"DEBUG: Oluşturulan parametre listesi: {parametre_listesi}")
        
        # Şablon verilerini hazırla - Firma kayıt bilgilerini kullan
        data = {
            'FIRMA_ADI': firma_bilgisi.get('firmaAdi', '') if firma_bilgisi else olcum.get('firma_adi', ''),
            'OLCUM_KODU': olcum.get('olcum_kodu', ''),
            'BASLANGIC_TARIHI': olcum.get('baslangic_tarihi', ''),
            'BITIS_TARIHI': olcum.get('bitis_tarihi', ''),
            'BACA_SAYISI': str(olcum.get('baca_sayisi', '')),
            'PARAMETRELER': ', '.join(parametre_listesi),
            'PERSONEL': ', '.join(olcum.get('personel', [])),
            'IL': firma_bilgisi.get('il', '') if firma_bilgisi else olcum.get('il', ''),
            'ILCE': firma_bilgisi.get('ilce', '') if firma_bilgisi else olcum.get('ilce', ''),
            'YETKILI': firma_bilgisi.get('yetkiliAdi', '') if firma_bilgisi else olcum.get('yetkili', ''),
            'TELEFON': str(firma_bilgisi.get('yetkiliTel', '')) if firma_bilgisi else olcum.get('telefon', ''),
            'DURUM': olcum.get('durum', '')
        }
        
        print(f"DEBUG: Data hazırlandı: {data}")
        
        # Şablondaki yer tutucuları değiştir
        print(f"DEBUG: Placeholder'lar değiştiriliyor...")
        replace_placeholders_in_document(doc, data)
        print(f"DEBUG: Placeholder'lar değiştirildi")
        
        # Baca listesi tablosunu ekle
        print(f"DEBUG: Baca listesi tablosu ekleniyor...")
        add_baca_listesi_detail_table(doc, olcum, baca_parametreleri)
        print(f"DEBUG: Baca listesi tablosu eklendi")
        
        # Dosyayı kaydet
        print(f"DEBUG: Dosya kaydediliyor...")
        from io import BytesIO
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        print(f"DEBUG: Dosya başarıyla kaydedildi")
        
        # Dosya adını oluştur
        print(f"DEBUG: Dosya adı oluşturuluyor...")
        def clean_filename(text):
            tr_to_en = {
                'ç': 'c', 'ğ': 'g', 'ı': 'i', 'ö': 'o', 'ş': 's', 'ü': 'u',
                'Ç': 'C', 'Ğ': 'G', 'I': 'I', 'İ': 'I', 'Ö': 'O', 'Ş': 'S', 'Ü': 'U'
            }
            for tr_char, en_char in tr_to_en.items():
                text = text.replace(tr_char, en_char)
            cleaned = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
            cleaned = cleaned.replace(' ', '-')
            while '--' in cleaned:
                cleaned = cleaned.replace('--', '-')
            return cleaned
        
        # Firma adının ilk 2 kelimesini al
        firma_adi = firma_bilgisi.get('firmaAdi', '') if firma_bilgisi else olcum.get('firma_adi', '')
        firma_kelimeleri = firma_adi.split()
        firma_ilk_2_kelime = ' '.join(firma_kelimeleri[:2]) if len(firma_kelimeleri) >= 2 else firma_adi
        
        # Dosya adını oluştur: "FIRMA ADI İKİ KELİME _OLCUM KODU_SAHAFORM"
        firma_ilk_2_kelime_clean = clean_filename(firma_ilk_2_kelime)
        olcum_kodu_clean = clean_filename(olcum.get('olcum_kodu', ''))
        filename = f"{firma_ilk_2_kelime_clean}_{olcum_kodu_clean}_SAHAFORM.docx"
        
        # Response oluştur
        print(f"DEBUG: Response oluşturuluyor...")
        response = make_response(docx_io.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        print(f"DEBUG: Response başarıyla oluşturuldu, filename: {filename}")
        return response
        
    except Exception as e:
        print(f"Firma ölçüm detay Word doküman oluşturma hatası: {e}")
        import traceback
        print(f"DEBUG: Hata detayı: {traceback.format_exc()}")
        return jsonify({'error': f'Word doküman oluşturma hatası: {str(e)}'}), 500

def create_firma_olcum_word_document(olcum, olcum_bacalar, olcum_parametreleri):
    """Firma ölçüm bilgilerini Word şablonu kullanarak oluşturur."""
    try:
        print(f"DEBUG: create_firma_olcum_word_document başladı")
        print(f"DEBUG: DOCX_AVAILABLE = {DOCX_AVAILABLE}")
        
        # Docx kütüphanesini yükle
        print(f"DEBUG: load_docx çağrılıyor...")
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        print(f"DEBUG: load_docx sonrası DOCX_AVAILABLE = {DOCX_AVAILABLE}")
        print(f"DEBUG: Document = {Document}")
        
        if not Document:
            print(f"DEBUG: python-docx kütüphanesi yüklü değil")
            return jsonify({'error': 'python-docx kütüphanesi yüklü değil'}), 500
        
        # Şablon dosyasını yükle
        template_path = os.path.join(os.path.dirname(__file__), 'static', 'images', 'EMISYON_OLCUM_BILGI_FORMU.docx')
        print(f"DEBUG: Template path = {template_path}")
        print(f"DEBUG: Template exists = {os.path.exists(template_path)}")
        
        if not os.path.exists(template_path):
            print(f"DEBUG: Word şablonu bulunamadı")
            return jsonify({'error': 'Word şablonu bulunamadı'}), 500
        
        # Şablonu yükle
        print(f"DEBUG: Document yükleniyor...")
        doc = Document(template_path)
        print(f"DEBUG: Document başarıyla yüklendi")
        
        # Parametre listesini oluştur (firma_olcum verilerinden)
        priority_params = ['Toz', 'Yg', 'Voc', 'Toc']  # Büyük/küçük harf düzeltildi
        
        # Firma ölçüm verilerinden parametreleri al
        baca_parametreleri = olcum.get('baca_parametreleri', {})
        
        # Tüm parametreleri topla
        all_params = set()
        for baca_params in baca_parametreleri.values():
            for param in baca_params:
                all_params.add(param)
        
        # Parametreleri sırala: önce öncelikli olanlar, sonra diğerleri
        sorted_params = []
        
        # Önce öncelikli parametreleri ekle
        for param in priority_params:
            if param in all_params:
                sorted_params.append(param)
                all_params.remove(param)
        
        # Sonra diğer parametreleri ekle
        sorted_params.extend(sorted(all_params))
        
        # Parametre listesini formatla (parametre adı + sayı)
        parametre_listesi = []
        for param in sorted_params:
            # Bu parametrenin kaç kez kullanıldığını say
            count = sum(1 for baca_params in baca_parametreleri.values() for p in baca_params if p == param)
            parametre_listesi.append(f"{param} ({count})")
        
        # Debug: Parametre listesi oluşturma
        print(f"DEBUG: Baca parametreleri: {baca_parametreleri}")
        print(f"DEBUG: Tüm parametreler: {all_params}")
        print(f"DEBUG: Sıralanmış parametreler: {sorted_params}")
        print(f"DEBUG: Oluşturulan parametre listesi: {parametre_listesi}")
        
        # Şablon verilerini hazırla - Şablondaki gerçek yer tutucuları kullan
        data = {
            'FIRMA_ADI': olcum.get('firma_adi', ''),
            'OLCUM_KODU': olcum.get('olcum_kodu', ''),
            'BASLANGIC_TARIHI': olcum.get('baslangic_tarihi', ''),
            'BITIS_TARIHI': olcum.get('bitis_tarihi', ''),
            'BACA_SAYISI': str(olcum.get('baca_sayisi', '')),
            'PARAMETRELER': ', '.join(parametre_listesi),
            'PERSONEL': ', '.join(olcum.get('personel', [])),
            'IL': olcum.get('il', ''),
            'ILCE': olcum.get('ilce', ''),
            'YETKILI': olcum.get('yetkili', ''),
            'TELEFON': olcum.get('telefon', ''),
            'DURUM': olcum.get('durum', '')
        }
        
        print(f"DEBUG: Data hazırlandı: {data}")
        
        # Şablondaki yer tutucuları değiştir
        print(f"DEBUG: Placeholder'lar değiştiriliyor...")
        replace_placeholders_in_document(doc, data)
        print(f"DEBUG: Placeholder'lar değiştirildi")
        
        # "Buraya baca listesi tablosu eklenecek..." placeholder metnini kaldır
        for paragraph in doc.paragraphs:
            if "Buraya baca listesi tablosu eklenecek" in paragraph.text:
                # Paragrafı temizle
                paragraph.clear()
                break
        
        # Baca listesi tablosunu ekle - firma_olcum verilerinden baca ve parametre bilgilerini geç
        print(f"DEBUG: Baca listesi tablosu ekleniyor...")
        add_baca_listesi_table(doc, olcum, olcum_parametreleri)
        print(f"DEBUG: Baca listesi tablosu eklendi")
        
        # "Buraya baca listesi tablosu eklenecek..." placeholder metnini kaldır
        print(f"DEBUG: Placeholder metni temizleniyor...")
        for paragraph in doc.paragraphs:
            if "Buraya baca listesi tablosu eklenecek" in paragraph.text:
                # Paragrafı temizle
                paragraph.clear()
                break
        print(f"DEBUG: Placeholder metni temizlendi")
        
        # Dosyayı kaydet
        print(f"DEBUG: Dosya kaydediliyor...")
        from io import BytesIO
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        print(f"DEBUG: Dosya başarıyla kaydedildi")
        
        # Dosya adını oluştur
        print(f"DEBUG: Dosya adı oluşturuluyor...")
        def clean_filename(text):
            tr_to_en = {
                'ç': 'c', 'ğ': 'g', 'ı': 'i', 'ö': 'o', 'ş': 's', 'ü': 'u',
                'Ç': 'C', 'Ğ': 'G', 'I': 'I', 'İ': 'I', 'Ö': 'O', 'Ş': 'S', 'Ü': 'U'
            }
            for tr_char, en_char in tr_to_en.items():
                text = text.replace(tr_char, en_char)
            cleaned = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
            cleaned = cleaned.replace(' ', '-')
            while '--' in cleaned:
                cleaned = cleaned.replace('--', '-')
            return cleaned
        
        # Firma adının ilk 2 kelimesini al
        firma_adi = olcum.get('firma_adi', '')
        firma_kelimeleri = firma_adi.split()
        firma_ilk_2_kelime = ' '.join(firma_kelimeleri[:2]) if len(firma_kelimeleri) >= 2 else firma_adi
        
        # Dosya adını oluştur: "FIRMA ADI İKİ KELİME _OLCUM KODU_SAHAFORM"
        firma_ilk_2_kelime_clean = clean_filename(firma_ilk_2_kelime)
        olcum_kodu_clean = clean_filename(olcum.get('olcum_kodu', ''))
        filename = f"{firma_ilk_2_kelime_clean}_{olcum_kodu_clean}_SAHAFORM.docx"
        
        # Response oluştur
        print(f"DEBUG: Response oluşturuluyor...")
        response = make_response(docx_io.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        print(f"DEBUG: Response başarıyla oluşturuldu, filename: {filename}")
        return response
        
    except Exception as e:
        print(f"Firma ölçüm Word doküman oluşturma hatası: {e}")
        import traceback
        print(f"DEBUG: Hata detayı: {traceback.format_exc()}")
        return jsonify({'error': f'Word doküman oluşturma hatası: {str(e)}'}), 500

def add_baca_listesi_detail_table(doc, olcum, baca_parametreleri):
    """Baca listesi tablosunu ekler (detay sayfası için)."""
    try:
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            print("Word rapor oluşturma için python-docx kütüphanesi gerekli")
            return
        
        # Tüm parametreleri topla
        all_params = set()
        for baca_params in baca_parametreleri.values():
            for param in baca_params:
                all_params.add(param)
        
        # Parametreleri sırala
        sorted_params = sorted(all_params)
        
        # Tablo oluştur
        if baca_parametreleri:
            # Başlık satırı
            doc.add_paragraph()
            title = doc.add_paragraph('BACA-PARAMETRE MATRİSİ')
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.bold = True
                run.font.size = Pt(14)
            
            # Tablo oluştur
            table = doc.add_table(rows=1, cols=len(sorted_params) + 1)
            table.style = 'Table Grid'
            
            # Başlık satırı
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Baca Adı'
            header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_cells[0].paragraphs[0].runs[0].font.bold = True
            
            for i, param in enumerate(sorted_params):
                header_cells[i + 1].text = param
                header_cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                header_cells[i + 1].paragraphs[0].runs[0].font.bold = True
            
            # Veri satırları
            for baca_adi, parametreler in baca_parametreleri.items():
                row_cells = table.add_row().cells
                row_cells[0].text = baca_adi
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[0].paragraphs[0].runs[0].font.bold = True
                
                for i, param in enumerate(sorted_params):
                    if param in parametreler:
                        row_cells[i + 1].text = 'X'
                        row_cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        row_cells[i + 1].paragraphs[0].runs[0].font.bold = True
                        row_cells[i + 1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Yeşil
                    else:
                        row_cells[i + 1].text = ''
                        row_cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Tablo genişliğini ayarla
            table.autofit = True
            table.allow_autofit = True
            
    except Exception as e:
        print(f"Baca listesi detay tablosu ekleme hatası: {e}")

def add_baca_listesi_table(doc, olcum, olcum_parametreleri):
    """Baca listesi tablosunu ekler (özel parametre sıralaması ile)."""
    try:
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            print("Word rapor oluşturma için python-docx kütüphanesi gerekli")
            return
        # Özel parametre sıralaması: Toz-Yg-Voc-Toc (ilk 4 sıra) - büyük/küçük harf düzeltildi
        priority_params = ['Toz', 'Yg', 'Voc', 'Toc']
        
        # Firma ölçüm verilerinden baca ve parametre bilgilerini al
        baca_parametreleri = olcum.get('baca_parametreleri', {})
        
        # Baca isimlerini topla
        baca_isimleri = list(baca_parametreleri.keys())
        
        # Tüm parametreleri topla
        all_params = set()
        for parametreler in baca_parametreleri.values():
            for parametre in parametreler:
                all_params.add(parametre)
        
        # Parametreleri sırala: önce öncelikli olanlar, sonra diğerleri
        sorted_params = []
        
        # Önce öncelikli parametreleri ekle
        for param in priority_params:
            if param in all_params:
                sorted_params.append(param)
                all_params.remove(param)
        
        # Sonra diğer parametreleri ekle
        sorted_params.extend(sorted(all_params))
        
        if not baca_isimleri or not sorted_params:
            return
        
        # Şablonda zaten başlık olabilir, bu yüzden başlık eklemeyelim
        # Sadece tablo ekleyelim - boşluk azaltıldı
        
        # Tablo oluştur (sıra numarası + baca + parametreler)
        table = doc.add_table(rows=len(baca_isimleri) + 1, cols=len(sorted_params) + 2)
        table.style = 'Table Grid'
        
        # Başlık satırı
        header_row = table.rows[0]
        header_row.cells[0].text = "SIRA"
        header_row.cells[0].paragraphs[0].runs[0].font.bold = True
        header_row.cells[1].text = "BACA"
        header_row.cells[1].paragraphs[0].runs[0].font.bold = True
        
        for i, param in enumerate(sorted_params):
            # Parametre sayısını hesapla
            param_count = 0
            for baca_params in baca_parametreleri.values():
                if param in baca_params:
                    param_count += 1
            
            # Başlığa sayıyı ekle
            header_text = f"{param} ({param_count})"
            header_row.cells[i + 2].text = header_text
            header_row.cells[i + 2].paragraphs[0].runs[0].font.bold = True
        
        # Veri satırları
        for i, baca_adi in enumerate(baca_isimleri):
            row = table.rows[i + 1]
            
            # Sıra numarası
            row.cells[0].text = str(i + 1)
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Baca adı
            row.cells[1].text = baca_adi
            row.cells[1].paragraphs[0].runs[0].font.bold = True
            
            # Her parametre için X veya boş kontrol et
            for j, param in enumerate(sorted_params):
                # Bu bacaya ait bu parametre seçili mi kontrol et
                baca_parametreleri = olcum.get('baca_parametreleri', {}).get(baca_adi, [])
                
                if param in baca_parametreleri:
                    # Parametre seçili - X yaz
                    row.cells[j + 2].text = "X"
                    row.cells[j + 2].paragraphs[0].runs[0].font.bold = True
                else:
                    # Parametre seçili değil - boş bırak
                    row.cells[j + 2].text = ""
        
        # Sütun genişliklerini ayarla
        # SIRA sütunu - 0.5 cm
        table.columns[0].width = Inches(0.5 / 2.54)
        
        # BACA sütunu - 3 cm
        table.columns[1].width = Inches(3.0 / 2.54)
        
        # Parametre sütunları - 1.5 cm
        for i in range(2, len(table.columns)):
            table.columns[i].width = Inches(1.5 / 2.54)
        
        # Tablo formatını ayarla
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    except Exception as e:
        print(f"Baca listesi tablosu ekleme hatası: {e}")

@app.route('/api/birlesik_veriler/<firma_adi>/<olcum_kodu>')
def api_birlesik_veriler(firma_adi, olcum_kodu):
    """Bir firmanın baca bilgileri ve parametre ölçümlerini birleştirerek döndürür."""
    try:
        # Verileri yükle
        baca_bilgileri = load_baca_bilgileri()
        parametre_olcumleri = load_parametre_olcum()
        
        # Firma verilerini filtrele
        firma_bacalar = [b for b in baca_bilgileri if b['firma_adi'] == firma_adi and b['olcum_kodu'] == olcum_kodu]
        firma_parametreleri = [p for p in parametre_olcumleri if p['firma_adi'] == firma_adi and p['olcum_kodu'] == olcum_kodu]
        
        # Birleştirilmiş veri yapısı oluştur
        birlesik_veriler = []
        
        for baca in firma_bacalar:
            baca_verisi = {
                'baca_id': baca['id'],
                'baca_adi': baca['baca_adi'],
                'baca_bilgileri': baca['baca_bilgileri'],
                'created_at': baca['created_at'],
                'updated_at': baca['updated_at'],
                'parametreler': []
            }
            
            # Bu bacaya ait parametreleri bul
            baca_parametreleri = [p for p in firma_parametreleri if p['baca_adi'] == baca['baca_adi']]
            
            for parametre in baca_parametreleri:
                baca_verisi['parametreler'].append({
                    'parametre_id': parametre['id'],
                    'parametre_adi': parametre['parametre_adi'],
                    'parametre_verileri': parametre['parametre_verileri'],
                    'created_at': parametre['created_at'],
                    'updated_at': parametre['updated_at']
                })
            
            birlesik_veriler.append(baca_verisi)
        
        return jsonify({
            'firma_adi': firma_adi,
            'olcum_kodu': olcum_kodu,
            'baca_sayisi': len(birlesik_veriler),
            'toplam_parametre_sayisi': len(firma_parametreleri),
            'veriler': birlesik_veriler
        })
        
    except Exception as e:
        print(f"Birleştirilmiş veri hatası: {e}")
        return jsonify({'error': f'Veri birleştirme hatası: {str(e)}'}), 500

@app.route('/api/baca_word_export', methods=['POST'])
def api_baca_word_export():
    """Seçili baca bilgilerini ve ölçüm parametrelerini Word formatında dışa aktarır."""
    try:
        print("=== Word Export Başladı ===")
        data = request.get_json()
        print(f"Gelen veri: {data}")
        
        firma = data.get('firma')
        olcum_kodu = data.get('olcum_kodu')
        baca_adi = data.get('baca_adi')
        
        print(f"Firma: {firma}")
        print(f"Ölçüm Kodu: {olcum_kodu}")
        print(f"Baca Adı: {baca_adi}")
        
        if not all([firma, olcum_kodu, baca_adi]):
            print("Eksik parametreler!")
            return jsonify({'error': 'Eksik parametreler'}), 400
        
        # Baca bilgilerini al
        print("Baca bilgileri yükleniyor...")
        baca_bilgileri = load_baca_bilgileri()
        print(f"Toplam {len(baca_bilgileri)} baca bilgisi bulundu")
        
        baca_bilgisi = None
        
        for baca in baca_bilgileri:
            print(f"Kontrol edilen baca: {baca.get('firma_adi')} - {baca.get('olcum_kodu')} - {baca.get('baca_adi')}")
            if (baca.get('firma_adi') == firma and 
                baca.get('olcum_kodu') == olcum_kodu and 
                baca.get('baca_adi') == baca_adi):
                baca_bilgisi = baca
                print("Baca bilgisi bulundu!")
                break
        
        if not baca_bilgisi:
            print("Baca bilgisi bulunamadı!")
            return jsonify({'error': 'Baca bilgisi bulunamadı'}), 404
        
        # Parametre ölçümlerini al
        parametre_olcumleri = load_parametre_olcum()
        baca_parametreleri = []
        
        for parametre in parametre_olcumleri:
            if (parametre.get('firma_adi') == firma and 
                parametre.get('olcum_kodu') == olcum_kodu and 
                parametre.get('baca_adi') == baca_adi):
                baca_parametreleri.append(parametre)
        
        # Form bilgilerini al (logo ve üst bilgiler için)
        forms = load_forms()
        form_bilgisi = None
        if forms:
            # İlk formu kullan (veya EMİSYON FORMU olanı)
            for form in forms:
                if 'EMİSYON' in form.get('form_adi', '').upper():
                    form_bilgisi = form
                    break
            if not form_bilgisi:
                form_bilgisi = forms[0]
        
        # Word dosyası oluştur (sadece şablon kullan)
        print("Word dosyası oluşturuluyor...")
        doc = create_baca_word_document_from_template(baca_bilgisi, baca_parametreleri, form_bilgisi)
        
        # Eğer doc None ise, hata döndür
        if doc is None:
            return jsonify({'error': 'Word dosyası oluşturulurken hata oluştu'}), 500
        
        # Eğer doc bir Document objesi ise, Flask response'a çevir
        if hasattr(doc, 'save'):
            # Geçici dosya oluştur
            import tempfile
            import os
            from io import BytesIO
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                doc.save(tmp_file.name)
                tmp_file_path = tmp_file.name
            
            # Dosyayı oku
            with open(tmp_file_path, 'rb') as f:
                doc_content = f.read()
            
            # Geçici dosyayı sil
            os.unlink(tmp_file_path)
            
            # Dosya adını oluştur
            firma_adi = baca_bilgisi.get('firma_adi', '')
            olcum_kodu = baca_bilgisi.get('olcum_kodu', '')
            baca_adi = baca_bilgisi.get('baca_adi', '')
            
            def clean_filename(text):
                tr_to_en = {
                    'ç': 'c', 'ğ': 'g', 'ı': 'i', 'ö': 'o', 'ş': 's', 'ü': 'u',
                    'Ç': 'C', 'Ğ': 'G', 'I': 'I', 'İ': 'I', 'Ö': 'O', 'Ş': 'S', 'Ü': 'U'
                }
                for tr_char, en_char in tr_to_en.items():
                    text = text.replace(tr_char, en_char)
                cleaned = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
                cleaned = cleaned.replace(' ', '-')
                while '--' in cleaned:
                    cleaned = cleaned.replace('--', '-')
                return cleaned
            
            firma_adi_clean = clean_filename(firma_adi)
            olcum_kodu_clean = clean_filename(olcum_kodu)
            baca_adi_clean = clean_filename(baca_adi)
            
            filename = f"{firma_adi_clean}-{baca_adi_clean}-{olcum_kodu_clean}.docx"
            
            # Flask response oluştur
            response = make_response(doc_content)
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
        else:
            # Eğer doc bir hata response'u ise, direkt döndür
            return doc
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Baca Word export hatası: {e}")
        print(f"Hata detayı: {error_details}")
        return jsonify({'error': f'Word dosyası oluşturulurken hata oluştu: {str(e)}', 'details': error_details}), 500

def export_forms_excel(forms):
    """Formları Excel formatında dışa aktarır."""
    try:
        # DataFrame oluştur
        df_data = []
        for form in forms:
            df_data.append({
                'FORM ADI': form.get('formAdi', ''),
                'FORM KODU': form.get('formKodu', ''),
                'YAYIN TARİHİ': form.get('yayinTarihi', ''),
                'REVİZYON TARİHİ': form.get('revizyonTarihi', ''),
                'REVİZYON NO': form.get('revizyonNo', ''),
                'OLUŞTURMA TARİHİ': form.get('created_at', ''),
                'GÜNCELLEME TARİHİ': form.get('updated_at', '')
            })
        
        df = pd.DataFrame(df_data)
        
        # Excel dosyası oluştur
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Formlar', index=False)
            
            # Sütun genişliklerini ayarla
            worksheet = writer.sheets['Formlar']
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'Formlar_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'
        )
        
    except Exception as e:
        print(f"Excel dışa aktarma hatası: {e}")
        raise e

def export_forms_word(forms):
    """Formları Word formatında dışa aktarır."""
    try:
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            return jsonify({'error': 'python-docx kütüphanesi yüklü değil'}), 500
        
        doc = Document()
        
        # Başlık
        title = doc.add_heading('FORM LİSTESİ', 0)
        title.alignment = 1  # Ortalı
        
        # Tarih
        doc.add_paragraph(f'Oluşturma Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph('')
        
        # Tablo oluştur
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Başlık satırı
        header_cells = table.rows[0].cells
        headers = ['SIRA', 'FORM ADI', 'FORM KODU', 'YAYIN TARİHİ', 'REVİZYON TARİHİ', 'REVİZYON NO']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Veri satırları
        for i, form in enumerate(forms, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = form.get('formAdi', '')
            row_cells[2].text = form.get('formKodu', '')
            row_cells[3].text = form.get('yayinTarihi', '')
            row_cells[4].text = form.get('revizyonTarihi', '')
            row_cells[5].text = form.get('revizyonNo', '')
        
        # Dosyayı kaydet
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'Formlar_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
        )
        
    except Exception as e:
        print(f"Word dışa aktarma hatası: {e}")
        raise e

def export_forms_pdf(forms):
    return 'PDF özelliği bu ortamda devre dışı (WeasyPrint eksik veya desteklenmiyor).', 501

def create_firma_raporu(firma_adi, olcum_kodu):
    """
    Bir firmanın tüm bacalarında yapılan tüm parametreleri içeren rapor oluşturur.
    JSON şablonunu kullanarak esnek rapor formatı sağlar.
    """
    try:
        # Şablonu yükle
        with open('rapor_sablonu.json', 'r', encoding='utf-8') as f:
            sablon = json.load(f)
        
        # Verileri yükle
        baca_bilgileri = load_baca_bilgileri()
        parametre_olcumleri = load_parametre_olcum()
        
        # Firma verilerini filtrele
        firma_bacalar = [b for b in baca_bilgileri if b['firma_adi'] == firma_adi and b['olcum_kodu'] == olcum_kodu]
        firma_parametreleri = [p for p in parametre_olcumleri if p['firma_adi'] == firma_adi and p['olcum_kodu'] == olcum_kodu]
        
        if not firma_bacalar:
            return "Firma verisi bulunamadı!", 404
        
        # Toplam sayfa sayısını hesapla (sadece header var)
        toplam_sayfa = 1  # Sadece 1 sayfa (header ile)
        
        # Debug için sayfa hesaplamasını yazdır
        print(f"Sayfa hesaplaması: Sadece header, Toplam={toplam_sayfa} sayfa")
        
        # Emisyon formu verilerini yükle
        emisyon_formu = None
        try:
            with open('forms.json', 'r', encoding='utf-8') as f:
                forms_data = json.load(f)
                emisyon_formu = next((form for form in forms_data if form['formAdi'] == 'EMİSYON ÖLÇÜM FORMU'), None)
        except Exception as e:
            print(f"Form verileri yüklenirken hata: {e}")
        
        # Tarih formatını GG.AA.YY şeklinde düzenle
        def format_tarih(tarih_str):
            try:
                if not tarih_str:
                    return ""
                # YYYY-MM-DD formatını GG.AA.YY'ye çevir
                dt = datetime.strptime(tarih_str, '%Y-%m-%d')
                return dt.strftime('%d.%m.%y')
            except:
                return tarih_str
        
        # Form bilgilerini hazırla
        if emisyon_formu:
            yayin_tarihi = format_tarih(emisyon_formu['yayinTarihi'])
            revizyon_tarihi = format_tarih(emisyon_formu['revizyonTarihi'])
            form_kodu = emisyon_formu['formKodu']
            revizyon_no = emisyon_formu['revizyonNo']
        else:
            yayin_tarihi = "01.08.15"
            revizyon_tarihi = "29.02.24"
            form_kodu = "AÇ.F.52"
            revizyon_no = "03"
        
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            return "Word rapor oluşturma için python-docx kütüphanesi gerekli", 500
        
        # Word dokümanı oluştur
        doc = Document()
        
        # Sayfa ayarları
        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        # Header oluştur - Sadece 1 satır
        header = section.header
        header_table = header.add_table(rows=1, cols=3, width=Inches(8))
        header_table.style = 'Table Grid'
        
        # Logo hücresi - 3cm genişlik
        logo_cell = header_table.cell(0, 0)
        logo_cell.width = Inches(1.2)  # 3cm
        
        # Logo ekle - FORM54.png
        try:
            logo_path = "static/images/FORM54.png"
            if os.path.exists(logo_path):
                logo_cell.text = ""
                logo_paragraph = logo_cell.paragraphs[0]
                logo_run = logo_paragraph.add_run()
                logo_run.add_picture(logo_path, width=Inches(0.8), height=Inches(0.6))
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Logo eklenirken hata: {e}")
        
        # Başlık hücresi - 9cm genişlik
        baslik_cell = header_table.cell(0, 1)
        baslik_cell.width = Inches(3.5)  # 9cm
        baslik_cell.text = "EMİSYON ÖLÇÜM FORMU"
        baslik_paragraph = baslik_cell.paragraphs[0]
        baslik_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        baslik_run = baslik_paragraph.runs[0]
        baslik_run.font.size = Pt(16)
        baslik_run.font.bold = True
        baslik_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Form bilgileri hücresi (sağ sütun) - Dinamik sayfa numarası ile
        form_cell = header_table.cell(0, 2)
        form_cell.text = f"Form Kodu: {form_kodu}\n"
        form_cell.text += f"Yayın Tarihi: {yayin_tarihi}\n"
        form_cell.text += f"Revizyon No: {revizyon_no}\n"
        form_cell.text += f"Revizyon Tarihi: {revizyon_tarihi}\n"
        form_cell.text += f"Sayfa No: 1/{toplam_sayfa}"  # Dinamik sayfa numarası
        
        # Header satır yüksekliğini ayarla - 1 satır 0.4cm
        for row in header_table.rows:
            row.height = Inches(0.16)  # 0.4cm
        
        # Header ile sayfa üstü arasında 3mm boşluk
        section.top_margin = Inches(0.12)  # 3mm
        section.header_distance = Inches(0.12)  # 3mm
        
        # Ana içerik - Baca bilgileri tablosu
        doc.add_paragraph()  # Boşluk
        
        # Baca bilgileri başlığı
        baca_baslik = doc.add_heading("BACA BİLGİLERİ", level=1)
        baca_baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if firma_bacalar:
            # Baca bilgileri tablosu - 4 sütun x 4 satır
            baca_table = doc.add_table(rows=4, cols=4)
            baca_table.style = 'Table Grid'
            
            # Tablo genişliğini ayarla
            baca_table.autofit = False
            for row in baca_table.rows:
                row.height = Inches(0.3)  # Satır yüksekliği
            
            # Sütun genişliklerini ayarla
            for i, col in enumerate(baca_table.columns):
                if i == 0:  # İlk sütun (etiketler)
                    col.width = Inches(2.5)
                else:  # Diğer sütunlar (değerler)
                    col.width = Inches(1.5)
            
            # 1. Satır
            baca_table.cell(0, 0).text = "FİRMA ADI:"
            baca_table.cell(0, 1).text = firma_adi
            baca_table.cell(0, 2).text = "ÖLÇÜM KODU:"
            baca_table.cell(0, 3).text = olcum_kodu
            
            # 2. Satır
            baca_table.cell(1, 0).text = "BACA ADI:"
            baca_table.cell(1, 1).text = firma_bacalar[0]['baca_adi'] if firma_bacalar else ""
            baca_table.cell(1, 2).text = "BACA NO:"
            baca_table.cell(1, 3).text = firma_bacalar[0]['baca_bilgileri'].get('6774fc7e-c124-4272-b826-482d064f3215', '') if firma_bacalar else ""
            
            # 3. Satır
            baca_table.cell(2, 0).text = "ISIL GÜÇ:"
            baca_table.cell(2, 1).text = firma_bacalar[0]['baca_bilgileri'].get('2360aa46-81dd-41db-a6b6-5c23e62900fe', '') if firma_bacalar else ""
            baca_table.cell(2, 2).text = "KAYNAK TÜRÜ:"
            baca_table.cell(2, 3).text = firma_bacalar[0]['baca_bilgileri'].get('7647078a-4f69-4908-99f7-3fff5651cc9b', '') if firma_bacalar else ""
            
            # 4. Satır
            baca_table.cell(3, 0).text = "YAKIT TÜRÜ:"
            baca_table.cell(3, 1).text = firma_bacalar[0]['baca_bilgileri'].get('2517064c-c286-4210-8f60-fa0a6b9e22a9', '') if firma_bacalar else ""
            baca_table.cell(3, 2).text = "ÇATI ŞEKLİ:"
            baca_table.cell(3, 3).text = firma_bacalar[0]['baca_bilgileri'].get('f1e7b875-ddf4-4628-87a3-eda6afa775b7', '') if firma_bacalar else ""
            
            # Tüm hücrelerde metin hizalaması
            for row in baca_table.rows:
                for cell in row.cells:
                    # Paragraf hizalaması
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # İlk karakteri kalın yap (etiketler için)
                        if paragraph.text and paragraph.text.endswith(':'):
                            if paragraph.runs:
                                paragraph.runs[0].font.bold = True
                    
                    # Hücre dikey hizalaması
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Dosyayı kaydet
        from io import BytesIO
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        
        # Response oluştur
        response = make_response(docx_io.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename={firma_adi}_{olcum_kodu}_Raporu.docx'
        
        return response
        
    except Exception as e:
        print(f"Rapor oluşturma hatası: {e}")
        return f"Rapor oluşturma hatası: {str(e)}", 500



def create_baca_word_document_from_template(baca_bilgisi, baca_parametreleri, form_bilgisi):
    """Word şablonunu kullanarak baca raporu oluşturur."""
    try:
        print("=== Şablon Word Dosyası Oluşturma Başladı ===")
        print(f"Baca bilgisi: {baca_bilgisi}")
        print(f"Parametre sayısı: {len(baca_parametreleri) if baca_parametreleri else 0}")
        
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            print("python-docx kütüphanesi yüklü değil!")
            return jsonify({'error': 'python-docx kütüphanesi yüklü değil'}), 500
        
        # Şablon dosyasını yükle
        template_path = os.path.join(os.path.dirname(__file__), 'static', 'images', 'F52 EMISYON.docx')
        print(f"Şablon yolu: {template_path}")
        print(f"Şablon dosyası var mı: {os.path.exists(template_path)}")
        
        if not os.path.exists(template_path):
            print("Word şablonu bulunamadı!")
            return jsonify({'error': 'Word şablonu bulunamadı'}), 500
        
        # Şablonu kopyala
        doc = Document(template_path)
        
        # Sayfa kenar boşluklarını ayarla - Sol ve sağ 5 mm
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(5/25.4)  # 5mm = 5/25.4 inç
            section.right_margin = Inches(5/25.4)  # 5mm = 5/25.4 inç
        
        # Boş değerleri "-" ile değiştiren yardımcı fonksiyon
        def get_value_or_dash(value):
            if value is None or value == '' or str(value).strip() == '':
                return '-'
            return str(value).strip()
        
        # Veri hazırla - Doğru ID'leri kullan ve boş değerleri "*" ile değiştir
        data = {
            'FIRMA_ADI': baca_bilgisi.get('firma_adi', '').title(),  # İlk karakterler büyük
            'OLCUM_KODU': baca_bilgisi.get('olcum_kodu', ''),
            'BACA_ADI': baca_bilgisi.get('baca_adi', '').title(),  # İlk karakterler büyük
            'BACA_NO': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('597fad80-d28f-40ea-bd28-a76c61c5203d', '')),
            'ISIL_GUC': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('22867c9a-ca3c-4d80-b017-b73dafdd7fef', '')),
            'KAYNAK_TURU': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('98399625-5bbc-465e-8e09-de454f231ae4', '')),
            'YAKIT_TURU': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('ddca398d-0e55-4662-b661-3731e0975bd2', '')),
            'CATI_SEKLI': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('6b3546e0-184c-49de-82e4-e2835e81923b', '')),
            'BACA_SEKLI': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('d9958774-43f7-4bc3-8e12-436614a6193a', '')),
            'BACA_OLCUSU': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('b1b6fc38-98c0-4048-8b8e-795cf7d44c48', '')),
            'YERDEN_YUKSEKLIK': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('6a301d72-f21b-485b-b8fb-116ad5cb223f', '')),
            'CATI_YUK': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('8ec9ecc9-ecda-4bf2-9802-02fa2e3fda4c', '')),
            'RUZGAR_HIZ': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('ab2b67dd-16a5-4bee-9b6e-b60b8cfc2d0c', '')),
            'ORT_SIC': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('eca60e54-ec39-4412-8884-caa17faed0be', '')),
            'ORT_NEM': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('64238e0a-6387-4c31-9bf4-d7f800ef17e1', '')),
            'ORT_BAS': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('b09ad69a-e4d4-4219-b055-2cf923ffd499', '')),
            'A_BACA': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('9c8c8bcf-c98e-4109-8b10-63b08b26460e', '')),
            'B_BACA': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('af55c55f-f83b-4b90-a655-ee76bf6bb2ac', '')),
            'C_DELIK': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('20881447-f7c8-4a6b-8583-76c7246082ef', '')),
            'PERSONEL': get_value_or_dash(baca_bilgisi.get('personel_adi', '')),
            'TARIH': datetime.now().strftime('%d.%m.%Y')
        }
        
        # Şablondaki yer tutucuları değiştir
        replace_placeholders_in_document(doc, data)
        
        # Parametre ölçümlerini ekle
        if baca_parametreleri:
            add_parametre_measurements_to_document(doc, baca_parametreleri)
        
        # Dosya adını temizle
        def clean_filename(text):
            # Türkçe karakterleri İngilizce karşılıklarıyla değiştir
            tr_to_en = {
                'ç': 'c', 'ğ': 'g', 'ı': 'i', 'ö': 'o', 'ş': 's', 'ü': 'u',
                'Ç': 'C', 'Ğ': 'G', 'I': 'I', 'İ': 'I', 'Ö': 'O', 'Ş': 'S', 'Ü': 'U'
            }
            for tr_char, en_char in tr_to_en.items():
                text = text.replace(tr_char, en_char)
            # Sadece alfanumerik karakterler ve bazı özel karakterler
            cleaned = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
            # Boşlukları tire ile değiştir
            cleaned = cleaned.replace(' ', '-')
            # Birden fazla tireyi tek tireye çevir
            while '--' in cleaned:
                cleaned = cleaned.replace('--', '-')
            return cleaned
        
        # Dosya adını oluştur
        firma_adi = baca_bilgisi.get('firma_adi', '')
        olcum_kodu = baca_bilgisi.get('olcum_kodu', '')
        baca_adi = baca_bilgisi.get('baca_adi', '')
        
        firma_adi_clean = clean_filename(firma_adi)
        olcum_kodu_clean = clean_filename(olcum_kodu)
        baca_adi_clean = clean_filename(baca_adi)
        
        filename = f"{firma_adi_clean}-{baca_adi_clean}-{olcum_kodu_clean}.docx"
        
        # Document objesini döndür (sayfa sonu ekleme işlemi merge_word_documents'da yapılacak)
        return doc
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Şablon rapor oluşturma hatası: {e}")
        print(f"Hata detayı: {error_details}")
        # Hata durumunda jsonify döndür
        return jsonify({'error': f'Şablon rapor oluşturma hatası: {str(e)}'}), 500

def add_parametre_measurements_to_document(doc, baca_parametreleri):
    """Word dokümanına parametre ölçümlerini ekler."""
    try:
        if baca_parametreleri:
            # Boşluk ekle
            doc.add_paragraph()
            
            # Parametre tablolarını ekle
            baca_adi = baca_parametreleri[0].get('baca_adi', '') if baca_parametreleri else ''
            for parametre in baca_parametreleri:
                add_parametre_details_to_document(doc, parametre, baca_adi)
            
    except Exception as e:
        print(f"Parametre ölçümleri ekleme hatası: {e}")

def add_parametre_details_to_document(doc, parametre, baca_adi=""):
    """Her parametre için detay bilgileri ekler."""
    try:
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            print("Word rapor oluşturma için python-docx kütüphanesi gerekli")
            return
        parametre_adi = parametre.get('parametre_adi', '')
        parametre_verileri = parametre.get('parametre_verileri', {})
        
        if not parametre_verileri:
            return
        
        # Parametre başlığı - 2-3 satır sola kaydır
        if baca_adi:
            subtitle_para = doc.add_paragraph(f"    {parametre_adi.upper()} PARAMETRE ÖLÇÜM DETAYLARI ({baca_adi})")
        else:
            subtitle_para = doc.add_paragraph(f"    {parametre_adi.upper()} PARAMETRE ÖLÇÜM DETAYLARI")
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        subtitle_run = subtitle_para.runs[0]
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.bold = True
        subtitle_run.font.name = 'Arial'
        try:
            subtitle_run.font.color.rgb = RGBColor(0, 0, 0)
        except:
            pass
        
                    # Detay bilgileri tablosu
        if parametre_verileri:
            # Boş değerleri "-" ile değiştiren yardımcı fonksiyon
            def get_value_or_dash(value):
                if value is None or value == '' or str(value).strip() == '':
                    return '-'
                return str(value).strip()
            
            # Parametre verilerini al ve sırala
            details = []
            for key, value in parametre_verileri.items():
                key_name = key
                details.append([key_name, get_value_or_dash(value)])
            
            # Detayları mantıklı sıraya koy
            priority_order = [
                'TARİH', 'METOT', 'NOZZLE ÇAP', 'TRAVERS', 'B.HIZ', 'B.SIC', 'B.BAS(KPA)', 
                'B.NEM(G/M3)', 'B.NEM(%)', 'SYC.HAC.', 'SYC.İLK', 'SYC.SON', 'SYC.SIC', 
                'DEBİ', 'ISDL', 'AA', 'SS', 'DD'
            ]
            
            # Öncelikli sıraya göre sırala
            sorted_details = []
            for priority_key in priority_order:
                for detail in details:
                    if detail[0] == priority_key:
                        sorted_details.append(detail)
                        break
            
            # Sıralanmamış olanları da ekle
            for detail in details:
                if detail not in sorted_details:
                    sorted_details.append(detail)
            
            if sorted_details:
                # 6 sütunlu detay tablosu oluştur (daha dar tablo)
                detail_table = doc.add_table(rows=1, cols=6)
                detail_table.style = 'Table Grid'
                detail_table.autofit = True
                
                # Tablo genişliğini manuel olarak ayarla
                detail_table.allow_autofit = False
                detail_table.autofit = False
                
                # Tabloyu 5 mm sola kaydır (Word şablonu gibi)
                detail_table.alignment = WD_ALIGN_PARAGRAPH.LEFT
                detail_table.left_indent = Inches(0.5/2.54)  # 5 mm sola kaydır
                
                # Her sütun için genişlik ayarla - 3,5 cm = 3.5/2.54 inç (daha geniş)
                for i, column in enumerate(detail_table.columns):
                    column.width = Inches(3.5/2.54)  # Her sütun 3,5 cm genişlik
                
                # Başlık satırı
                header_row = detail_table.rows[0]
                header_row.cells[0].text = "Ölç.Param."
                header_row.cells[1].text = "Değer"
                header_row.cells[2].text = "Ölç.Param."
                header_row.cells[3].text = "Değer"
                header_row.cells[4].text = "Ölç.Param."
                header_row.cells[5].text = "Değer"
                
                # Başlık formatla
                for cell in header_row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].runs[0].font.size = Pt(11)  # 11 punto
                        cell.paragraphs[0].runs[0].font.name = 'Arial'
                    
                    # Başlık hücre genişliğini ayarla - 3,5 cm
                    cell.width = Inches(3.5/2.54)
                
                # Detay verilerini 6 sütunlu olarak ekle (3 parametre-değer çifti)
                for i in range(0, len(sorted_details), 3):
                    row = detail_table.add_row()
                    
                    # İlk parametre
                    row.cells[0].text = sorted_details[i][0]
                    row.cells[1].text = sorted_details[i][1]
                    
                    # İkinci parametre (varsa)
                    if i + 1 < len(sorted_details):
                        row.cells[2].text = sorted_details[i + 1][0]
                        row.cells[3].text = sorted_details[i + 1][1]
                    else:
                        row.cells[2].text = ""
                        row.cells[3].text = ""
                    
                    # Üçüncü parametre (varsa)
                    if i + 2 < len(sorted_details):
                        row.cells[4].text = sorted_details[i + 2][0]
                        row.cells[5].text = sorted_details[i + 2][1]
                    else:
                        row.cells[4].text = ""
                        row.cells[5].text = ""
                    
                    # Hücre formatla
                    for j in range(6):
                        if j % 2 == 0:  # Parametre adı sütunları
                            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:  # Değer sütunları
                            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Run kontrolü yap
                        if row.cells[j].paragraphs[0].runs:
                            row.cells[j].paragraphs[0].runs[0].font.size = Pt(11)  # 11 punto
                            row.cells[j].paragraphs[0].runs[0].font.name = 'Arial'
                        
                        # Hücre genişliğini ayarla - 3,5 cm
                        row.cells[j].width = Inches(3.5/2.54)
 
        # Tablo sonrası 1 boşluk
        doc.add_paragraph()
          
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Parametre detayları ekleme hatası: {e}")
        print(f"Hata detayı: {error_details}")
        # Hatayı yukarı fırlatma, sadece logla
        raise e

def replace_placeholders_in_document(doc, data):
    """Word dokümanındaki yer tutucuları gerçek verilerle değiştirir."""
    # Docx kütüphanesini yükle
    Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
    
    if not Document:
        print("Word rapor oluşturma için python-docx kütüphanesi gerekli")
        return
    
    # Paragraflardaki yer tutucuları değiştir
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        
        for key, value in data.items():
            # Hem {{KEY}} hem de {KEY} formatlarını kontrol et
            placeholder1 = f"{{{{{key}}}}}"
            placeholder2 = f"{{{key}}}"
            if placeholder1 in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder1, str(value))
            elif placeholder2 in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder2, str(value))
        
        # Eğer metin değiştiyse ve başlık satırıysa formatla
        if paragraph.text != original_text:
            # Başlık satırlarını kontrol et
            if any(key in original_text for key in ['FIRMA_ADI', 'OLCUM_KODU', 'BACA_ADI']):
                # Başlık formatlaması
                for run in paragraph.runs:
                    run.font.size = Pt(14)  # 14 punto
                    run.font.bold = True
                    run.font.name = 'Arial'
                    # RGBColor kullanımını güvenli hale getir
                    try:
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Siyah renk
                    except:
                        # RGBColor kullanılamıyorsa varsayılan rengi kullan
                        pass
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Sola dayalı
    
    # Tablolardaki yer tutucuları değiştir
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        # Hem {{KEY}} hem de {KEY} formatlarını kontrol et
                        placeholder1 = f"{{{{{key}}}}}"
                        placeholder2 = f"{{{key}}}"
                        if placeholder1 in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder1, str(value))
                        elif placeholder2 in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder2, str(value))
    
    # Header ve footer'daki yer tutucuları değiştir
    for section in doc.sections:
        # Header
        if section.header:
            for paragraph in section.header.paragraphs:
                for key, value in data.items():
                    # Hem {{KEY}} hem de {KEY} formatlarını kontrol et
                    placeholder1 = f"{{{{{key}}}}}"
                    placeholder2 = f"{{{key}}}"
                    if placeholder1 in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder1, str(value))
                    elif placeholder2 in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder2, str(value))
            # Header tablolarını da kontrol et
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in data.items():
                                placeholder1 = f"{{{{{key}}}}}"
                                placeholder2 = f"{{{key}}}"
                                if placeholder1 in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder1, str(value))
                                elif placeholder2 in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder2, str(value))
        
        # Footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                for key, value in data.items():
                    # Hem {{KEY}} hem de {KEY} formatlarını kontrol et
                    placeholder1 = f"{{{{{key}}}}}"
                    placeholder2 = f"{{{key}}}"
                    if placeholder1 in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder1, str(value))
                    elif placeholder2 in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder2, str(value))
            # Footer tablolarını da kontrol et
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in data.items():
                                placeholder1 = f"{{{{{key}}}}}"
                                placeholder2 = f"{{{key}}}"
                                if placeholder1 in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder1, str(value))
                                elif placeholder2 in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder2, str(value))

def merge_word_documents(documents):
    """Birden fazla Word dokümanını tek bir dokümanda birleştirir."""
    if not documents:
        return None
    
    # İlk dokümanı ana doküman olarak kullan
    merged_doc = documents[0]
    
    # Diğer dokümanları ekle
    for i in range(1, len(documents)):
        doc = documents[i]
        # Dokümanın tüm içeriğini ana dokümana ekle
        for element in doc.element.body:
            merged_doc.element.body.append(element)
        
        # Sayfa sonu ekle (son doküman hariç)
        if i < len(documents) - 1:
            merged_doc.add_page_break()
    
    return merged_doc

def create_firma_raporu_from_template(firma_adi, olcum_kodu):
    """Her baca için ayrı Word dosyası oluşturup birleştirir."""
    try:
        print("=== CREATE_FIRMA_RAPORU_FROM_TEMPLATE BAŞLADI ===")
        
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        print(f"load_docx sonucu - Document: {Document is not None}")
        
        if not Document:
            print("python-docx kütüphanesi yüklü değil!")
            return jsonify({'error': 'python-docx kütüphanesi yüklü değil'}), 500
        
        # Firma verilerini al
        print("Baca bilgileri yükleniyor...")
        baca_bilgileri = load_baca_bilgileri()
        print(f"Toplam baca bilgisi sayısı: {len(baca_bilgileri)}")
        
        firma_bacalar = [baca for baca in baca_bilgileri if baca.get('firma_adi') == firma_adi and baca.get('olcum_kodu') == olcum_kodu]
        print(f"Firma için bulunan baca sayısı: {len(firma_bacalar)}")
        
        if not firma_bacalar:
            print("Firma için baca bilgisi bulunamadı!")
            return jsonify({'error': 'Firma için baca bilgisi bulunamadı'}), 404
        
        # Parametre ölçümlerini al
        parametre_olcumleri = load_parametre_olcum()
        
        # Form bilgilerini al
        forms = load_forms()
        form_bilgisi = None
        for form in forms:
            if form.get('formKodu') == 'AÇ.F.52':
                form_bilgisi = form
                break
        if not form_bilgisi:
            form_bilgisi = forms[0] if forms else {}
        
        # Her baca için ayrı Word dosyası oluştur
        baca_documents = []
        
        for baca_bilgisi in firma_bacalar:
            # Baca parametrelerini al
            baca_adi = baca_bilgisi.get('baca_adi', '')
            baca_parametreleri = [param for param in parametre_olcumleri 
                                 if param.get('firma_adi') == firma_adi 
                                 and param.get('olcum_kodu') == olcum_kodu 
                                 and param.get('baca_adi') == baca_adi]
            
            print(f"Baca: {baca_adi}, Parametre sayısı: {len(baca_parametreleri)}")
            for param in baca_parametreleri:
                print(f"  - {param.get('parametre_adi', '')}: {len(param.get('parametre_verileri', {}))} veri")
            
            # Her baca için ayrı Word dosyası oluştur (mevcut fonksiyonu kullan)
            baca_doc = create_baca_word_document_from_template(baca_bilgisi, baca_parametreleri, form_bilgisi)
            
            # Hata kontrolü
            if baca_doc is None:
                print(f"Baca {baca_adi} için doküman oluşturulamadı!")
                continue
            
            # Dokümanı listeye ekle
            baca_documents.append(baca_doc)
        
        # Tüm dokümanları birleştir
        if baca_documents:
            # İlk dokümanı ana doküman olarak kullan
            merged_doc = baca_documents[0]
            
            # Diğer dokümanları ekle (her biri yeni sayfadan başlasın)
            for i in range(1, len(baca_documents)):
                # Sayfa sonu ekle
                merged_doc.add_page_break()
                
                # Dokümanın tüm içeriğini ana dokümana ekle
                doc = baca_documents[i]
                for element in doc.element.body:
                    merged_doc.element.body.append(element)
        else:
            return jsonify({'error': 'Hiç baca dokümanı oluşturulamadı'}), 500
        
        # Dosya adını temizle (Türkçe karakterleri ve geçersiz karakterleri kaldır)
        def clean_filename(text):
            # Türkçe karakterleri İngilizce karşılıklarıyla değiştir
            tr_to_en = {
                'ç': 'c', 'ğ': 'g', 'ı': 'i', 'ö': 'o', 'ş': 's', 'ü': 'u',
                'Ç': 'C', 'Ğ': 'G', 'I': 'I', 'İ': 'I', 'Ö': 'O', 'Ş': 'S', 'Ü': 'U'
            }
            for tr_char, en_char in tr_to_en.items():
                text = text.replace(tr_char, en_char)
            # Sadece alfanumerik karakterler ve bazı özel karakterler
            cleaned = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
            # Boşlukları tire ile değiştir
            cleaned = cleaned.replace(' ', '-')
            # Birden fazla tireyi tek tireye çevir
            while '--' in cleaned:
                cleaned = cleaned.replace('--', '-')
            return cleaned
        
        # Dosya adını oluştur
        firma_adi_clean = clean_filename(firma_adi)
        olcum_kodu_clean = clean_filename(olcum_kodu)
        filename = f"F54-{firma_adi_clean}-{olcum_kodu_clean}-FirmaRaporu.docx"
        
        # Dosyayı kaydet
        from io import BytesIO
        docx_io = BytesIO()
        merged_doc.save(docx_io)
        docx_io.seek(0)
        
        # Response oluştur
        response = make_response(docx_io.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        print(f"Firma şablon rapor oluşturma hatası: {e}")
        return jsonify({'error': f'Firma şablon rapor oluşturma hatası: {str(e)}'}), 500

# Parametre ölçümleri export API'leri
@app.route('/api/parametre_olcumleri_excel_export', methods=['POST'])
def api_parametre_olcumleri_excel_export():
    """Parametre ölçümlerini Excel formatında dışa aktarır."""
    try:
        print("Excel export başlatıldı...")
        data = request.get_json()
        olcum_ids = data.get('olcum_ids', [])
        
        print(f"Gelen olcum_ids: {olcum_ids}")
        
        if not olcum_ids:
            return jsonify({'error': 'En az bir ölçüm ID gerekli!'}), 400
        
        # Parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        print(f"Toplam parametre ölçümü sayısı: {len(parametre_olcumleri)}")
        
        # Seçili ölçümleri filtrele
        if olcum_ids:
            filtered_olcumler = [item for item in parametre_olcumleri if str(item.get('id')) in olcum_ids]
        else:
            filtered_olcumler = parametre_olcumleri
        
        print(f"Filtrelenmiş ölçüm sayısı: {len(filtered_olcumler)}")
        
        if not filtered_olcumler:
            return jsonify({'error': 'Dışa aktarılacak veri bulunamadı'}), 404
        
        # Tüm parametre alanlarını tanımla (frontend ile aynı sırada)
        allParametreFields = [
            'TARİH', 'METOT', 'NOZZLE ÇAP', 'TRAVERS', 'B.HIZ', 'B.SIC', 'B.BAS(KPA)', 
            'B.NEM(G/M3)', 'B.NEM(%)', 'SYC.HAC.', 'SYC.İLK', 'SYC.SON', 'SYC.SIC', 'DEBİ', 'ISDL',
            '1İMP-İ', '1-İMP-S', '2-İMP-İ', '2-İMP-S', '3-İMP-İ', '3-İMP-S', 'HAC.',
            'O2', 'CO', 'NO', 'NOX', 'SO2', 'KK1-O2', 'KK1-CO', 'KK1-NO', 'KK1-SO2', 
            'KK2-O2', 'KK2-CO', 'KK2-NO', 'KK2-SO2', 'T90',
                            'TOC(PPM)', 'KK1-SPAN', 'KK1-SPAN', 'KK2-O', 'KK2-SPAN',
            'GAZ HAC.', 'GAZ.SIC.', 'SEY.GAZ.HAC', 'SEY.GAZ.SIC',
            'ORT.SIC', 'ORT.NEM', 'ORT.RUZ.HIZ', 'ÇEK.HACİM',
            'T.İÇİ-1', 'T.İÇİ-2', 'T.İÇİ-3', 'T.İÇİ-4', 'T-DIŞ-1', 'T.DIŞ-2', 
            'T.DIŞ-3', 'T.DIŞ-4', 'İLK KURULUM', '2. KURULUM'
        ]
        
        print("DataFrame oluşturuluyor...")
        
        # Excel için DataFrame oluştur - ekranda gördüğünüz format
        df_data = []
        for index, olcum in enumerate(filtered_olcumler, 1):
            # Temel bilgiler (ekranda gördüğünüz sütunlar)
            row = {
                '#': index,
                'Firma': olcum.get('firma_adi', ''),
                'Ölçüm Kodu': olcum.get('olcum_kodu', ''),
                'Baca': olcum.get('baca_adi', ''),
                'Parametre': olcum.get('parametre_adi', ''),
                'Kayıt Tarihi': olcum.get('created_at', '')[:19].replace('T', ' ') if olcum.get('created_at') else ''
            }
            
            # Parametre verilerini ekle (ekranda gördüğünüz parametre sütunları)
            parametre_verileri = olcum.get('parametre_verileri', {})
            for alan in allParametreFields:
                row[alan] = parametre_verileri.get(alan, '*')
            
            df_data.append(row)
        
        print(f"DataFrame verisi hazırlandı, satır sayısı: {len(df_data)}")
        
        df = pd.DataFrame(df_data)
        print(f"DataFrame oluşturuldu, boyut: {df.shape}")
        
        # Geçici dosya oluştur
        print("Geçici dosya oluşturuluyor...")
        tmp_file_path = None
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
                tmp_file_path = tmp_file.name
                print(f"Geçici dosya yolu: {tmp_file_path}")
            
            # Excel dosyasını oluştur
            with pd.ExcelWriter(tmp_file_path, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Parametre Ölçümleri')
                
                # Sütun genişliklerini ayarla
                worksheet = writer.sheets['Parametre Ölçümleri']
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
            
            print("Excel dosyası oluşturuldu, okunuyor...")
            
            # Dosyayı oku ve response oluştur
            with open(tmp_file_path, 'rb') as f:
                file_content = f.read()
            
            print(f"Dosya okundu, boyut: {len(file_content)} bytes")
            
        finally:
            # Geçici dosyayı sil (hata olsa bile)
            if tmp_file_path and os.path.exists(tmp_file_path):
                try:
                    os.unlink(tmp_file_path)
                    print("Geçici dosya silindi")
                except Exception as delete_error:
                    print(f"Geçici dosya silinirken hata: {delete_error}")
                    # Dosya silinemezse sorun değil, sistem temizleyecek
        
        # Response oluştur
        response = make_response(file_content)
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        
        # Dosya adını oluştur
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"Parametre_Olcumleri_{len(filtered_olcumler)}_kayit_{timestamp}.xlsx"
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        print(f"Response hazırlandı, dosya adı: {filename}")
        return response
        
    except Exception as e:
        print(f"Parametre ölçümleri Excel export hatası: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Excel export hatası: {str(e)}'}), 500

@app.route('/api/parametre_olcumleri_word_export', methods=['POST'])
def api_parametre_olcumleri_word_export():
    """Parametre ölçümlerini Word formatında dışa aktarır."""
    try:
        data = request.get_json()
        olcum_ids = data.get('olcum_ids', [])
        
        if not olcum_ids:
            return jsonify({'error': 'En az bir ölçüm ID gerekli!'}), 400
        
        # Parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Seçili ölçümleri filtrele
        if olcum_ids:
            filtered_olcumler = [item for item in parametre_olcumleri if str(item.get('id')) in olcum_ids]
        else:
            filtered_olcumler = parametre_olcumleri
        
        if not filtered_olcumler:
            return jsonify({'error': 'Dışa aktarılacak veri bulunamadı'}), 404
        
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            return jsonify({'error': 'Word rapor oluşturma için python-docx kütüphanesi gerekli'}), 500
        
        # Word dokümanı oluştur
        doc = Document()
        
        # Başlık
        title = doc.add_heading('Parametre Ölçümleri Raporu', 0)
        title.alignment = 1  # Ortalı
        
        # Tarih
        doc.add_paragraph(f'Oluşturma Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph('')
        
        # Her ölçüm için tablo oluştur
        for i, olcum in enumerate(filtered_olcumler, 1):
            # Ölçüm başlığı
            doc.add_heading(f'{i}. Ölçüm Kaydı', level=1)
            
            # Temel bilgiler tablosu
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Alan'
            hdr_cells[1].text = 'Değer'
            
            # Temel bilgileri ekle
            basic_info = [
                ('Firma Adı', olcum.get('firma_adi', '')),
                ('Ölçüm Kodu', olcum.get('olcum_kodu', '')),
                ('Baca Adı', olcum.get('baca_adi', '')),
                ('Parametre Adı', olcum.get('parametre_adi', '')),
                ('Kayıt Tarihi', olcum.get('created_at', '')[:19].replace('T', ' ') if olcum.get('created_at') else '')
            ]
            
            for field, value in basic_info:
                row_cells = table.add_row().cells
                row_cells[0].text = field
                row_cells[1].text = str(value)
            
            # Parametre verileri tablosu
            parametre_verileri = olcum.get('parametre_verileri', {})
            if parametre_verileri:
                doc.add_paragraph('')
                doc.add_heading('Parametre Verileri', level=2)
                
                param_table = doc.add_table(rows=1, cols=2)
                param_table.style = 'Table Grid'
                param_hdr_cells = param_table.rows[0].cells
                param_hdr_cells[0].text = 'Parametre'
                param_hdr_cells[1].text = 'Değer'
                
                for key, value in parametre_verileri.items():
                    param_row_cells = param_table.add_row().cells
                    param_row_cells[0].text = str(key)
                    param_row_cells[1].text = str(value)
            
            # Sayfa sonu ekle (son ölçüm hariç)
            if i < len(filtered_olcumler):
                doc.add_page_break()
        
        # Dosyayı kaydet
        from io import BytesIO
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        
        # Response oluştur
        response = make_response(docx_io.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        # Dosya adını oluştur
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"Parametre_Olcumleri_{len(filtered_olcumler)}_kayit_{timestamp}.docx"
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        print(f"Parametre ölçümleri Word export hatası: {e}")
        return jsonify({'error': f'Word export hatası: {str(e)}'}), 500

@app.route('/api/parametre_olcumleri_pdf_export', methods=['POST'])
def api_parametre_olcumleri_pdf_export():
    """Parametre ölçümlerini PDF formatında dışa aktarır."""
    try:
        data = request.get_json()
        olcum_ids = data.get('olcum_ids', [])
        
        if not olcum_ids:
            return jsonify({'error': 'En az bir ölçüm ID gerekli!'}), 400
        
        # Parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Seçili ölçümleri filtrele
        if olcum_ids:
            filtered_olcumler = [item for item in parametre_olcumleri if str(item.get('id')) in olcum_ids]
        else:
            filtered_olcumler = parametre_olcumleri
        
        if not filtered_olcumler:
            return jsonify({'error': 'Dışa aktarılacak veri bulunamadı'}), 404
        
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            return jsonify({'error': 'Word rapor oluşturma için python-docx kütüphanesi gerekli'}), 500
        
        # Word dokümanı oluştur (Word export fonksiyonunu kullan)
        doc = Document()
        
        # Başlık
        title = doc.add_heading('Parametre Ölçümleri Raporu', 0)
        title.alignment = 1  # Ortalı
        
        # Tarih
        doc.add_paragraph(f'Oluşturma Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph('')
        
        # Her ölçüm için tablo oluştur
        for i, olcum in enumerate(filtered_olcumler, 1):
            # Ölçüm başlığı
            doc.add_heading(f'{i}. Ölçüm Kaydı', level=1)
            
            # Temel bilgiler tablosu
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Alan'
            hdr_cells[1].text = 'Değer'
            
            # Temel bilgileri ekle
            basic_info = [
                ('Firma Adı', olcum.get('firma_adi', '')),
                ('Ölçüm Kodu', olcum.get('olcum_kodu', '')),
                ('Baca Adı', olcum.get('baca_adi', '')),
                ('Parametre Adı', olcum.get('parametre_adi', '')),
                ('Kayıt Tarihi', olcum.get('created_at', '')[:19].replace('T', ' ') if olcum.get('created_at') else '')
            ]
            
            for field, value in basic_info:
                row_cells = table.add_row().cells
                row_cells[0].text = field
                row_cells[1].text = str(value)
            
            # Parametre verileri tablosu
            parametre_verileri = olcum.get('parametre_verileri', {})
            if parametre_verileri:
                doc.add_paragraph('')
                doc.add_heading('Parametre Verileri', level=2)
                
                param_table = doc.add_table(rows=1, cols=2)
                param_table.style = 'Table Grid'
                param_hdr_cells = param_table.rows[0].cells
                param_hdr_cells[0].text = 'Parametre'
                param_hdr_cells[1].text = 'Değer'
                
                for key, value in parametre_verileri.items():
                    param_row_cells = param_table.add_row().cells
                    param_row_cells[0].text = str(key)
                    param_row_cells[1].text = str(value)
            
            # Sayfa sonu ekle (son ölçüm hariç)
            if i < len(filtered_olcumler):
                doc.add_page_break()
        
        # Word dosyasını geçici olarak kaydet
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_word:
            doc.save(tmp_word.name)
            tmp_word_path = tmp_word.name
        
        # Word dosyasını PDF'e çevir
        try:
            from docx2pdf import convert
            
            # PDF dosya adını oluştur
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            pdf_filename = f"Parametre_Olcumleri_{len(filtered_olcumler)}_kayit_{timestamp}.pdf"
            pdf_path = os.path.join(tempfile.gettempdir(), pdf_filename)
            
            # Word'ü PDF'e çevir
            convert(tmp_word_path, pdf_path)
            
            # PDF dosyasını oku
            with open(pdf_path, 'rb') as f:
                pdf_content = f.read()
            
            # Geçici dosyaları sil
            os.unlink(tmp_word_path)
            os.unlink(pdf_path)
            
            # Response oluştur
            response = make_response(pdf_content)
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = f'attachment; filename="{pdf_filename}"'
            
            return response
            
        except Exception as pdf_error:
            print(f"PDF dönüştürme hatası: {pdf_error}")
            # PDF dönüştürme başarısız olursa Word dosyasını döndür
            with open(tmp_word_path, 'rb') as f:
                word_content = f.read()
            
            os.unlink(tmp_word_path)
            
            response = make_response(word_content)
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = f'attachment; filename="Parametre_Olcumleri_{len(filtered_olcumler)}_kayit_{timestamp}.docx"'
            
            return response
        
    except Exception as e:
        print(f"Parametre ölçümleri PDF export hatası: {e}")
        return jsonify({'error': f'PDF export hatası: {str(e)}'}), 500

@app.route('/api/teklif/yazdir/<teklif_id>', methods=['POST'])
def yazdir_teklif(teklif_id):
    """Teklifi Word formatında yazdırır"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        format_type = data.get('format', 'word')  # word veya pdf
        
        # Teklif verilerini yükle
        teklifler = load_teklif()
        teklif = next((t for t in teklifler if t.get('id') == teklif_id), None)
        
        if not teklif:
            return jsonify({'success': False, 'message': 'Teklif bulunamadı'})
        
        # Firma bilgilerini al
        firma_kayitlar = load_firma_kayit()
        firma = next((f for f in firma_kayitlar if f.get('firmaAdi') == teklif.get('firma_adi')), None)
        
        # Eğer firma bulunamazsa, teklif verisindeki firma_adi'yi kullan
        if not firma:
            print(f"Firma bulunamadı: {teklif.get('firma_adi')}")
            # Geçici firma objesi oluştur
            firma = {'firmaAdi': teklif.get('firma_adi', '')}
        
        if format_type == 'word':
            return create_word_teklif(teklif, firma)
        else:
            return create_pdf_teklif(teklif, firma)
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

def create_word_teklif(teklif, firma, return_file_info: bool = False):
    try:
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        if not DOCX_AVAILABLE:
            return jsonify({'success': False, 'message': 'Word dosyası oluşturma için gerekli kütüphane yüklü değil'})

        from docx.oxml.shared import qn
        from docx.oxml import OxmlElement

        def set_labeled_table_value(doc_obj, label, value):
            label_norm = (label or '').strip().rstrip(':')
            value_str = '' if value is None else str(value)
            for table in doc_obj.tables:
                try:
                    for row in table.rows:
                        cells = row.cells
                        if not cells:
                            continue
                        # Find label in first cell
                        if label_norm.lower() in (cells[0].text or '').strip().rstrip(':').lower():
                            # Templates use: [0]=label, [1]=short code, [2]=value
                            if len(cells) >= 3:
                                target = cells[2]
                            elif len(cells) == 2:
                                target = cells[1]
                            else:
                                target = cells[-1]
                            target.text = value_str
                            return True
                except Exception:
                    continue
            return False

        def _norm(s: str) -> str:
            return (s or '').strip().lower()

        def _is_break_or_empty_block(el) -> bool:
            """True if block is an empty paragraph or paragraph containing only page/section break markers."""
            try:
                if not el.tag.endswith('}p'):
                    return False
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                texts = [t.text for t in el.xpath('.//w:t', namespaces=ns) if getattr(t, 'text', None)]
                if any((tx or '').strip() for tx in texts):
                    return False
                if el.xpath('.//w:drawing', namespaces=ns) or el.xpath('.//w:pict', namespaces=ns):
                    return False
                # treat all-empty paragraph as removable
                return True
            except Exception:
                return False

        def _strip_break_markers_in_paragraph(p_el):
            try:
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                for pb in list(p_el.xpath('.//w:pPr/w:pageBreakBefore', namespaces=ns)):
                    try:
                        pb.getparent().remove(pb)
                    except Exception:
                        pass
                for sp in list(p_el.xpath('.//w:pPr/w:sectPr', namespaces=ns)):
                    try:
                        sp.getparent().remove(sp)
                    except Exception:
                        pass
                for br in list(p_el.xpath('.//w:br[@w:type="page"]', namespaces=ns)):
                    try:
                        br.getparent().remove(br)
                    except Exception:
                        pass
                for lrp in list(p_el.xpath('.//w:lastRenderedPageBreak', namespaces=ns)):
                    try:
                        lrp.getparent().remove(lrp)
                    except Exception:
                        pass
            except Exception:
                pass

        def _clean_template_breaks(doc_obj):
            """Aggressively clean leading/trailing empty/break-only paragraphs and remove break markers
            that can produce visible 'Sayfa Sonu' and extra blank pages when merged.
            """
            try:
                body = doc_obj.element.body
                children = list(body.iterchildren())

                # Remove leading empty paragraphs
                for child in children:
                    if child.tag.endswith('}sectPr'):
                        continue
                    if _is_break_or_empty_block(child):
                        try:
                            body.remove(child)
                        except Exception:
                            pass
                        continue
                    # first real block: strip break markers inside it
                    if child.tag.endswith('}p'):
                        _strip_break_markers_in_paragraph(child)
                    break

                # Remove trailing empty paragraphs
                while True:
                    last = None
                    for child in reversed(list(body.iterchildren())):
                        if child.tag.endswith('}sectPr'):
                            continue
                        last = child
                        break
                    if last is None:
                        break
                    if last.tag.endswith('}p'):
                        _strip_break_markers_in_paragraph(last)
                        if _is_break_or_empty_block(last):
                            try:
                                body.remove(last)
                            except Exception:
                                pass
                            continue
                    break
            except Exception:
                pass

        def _apply_header_image_only(doc_obj):
            header_img_path = os.path.join(app.root_path, 'static', 'images', 'tek_ust1.jpg')
            if not os.path.exists(header_img_path):
                return
            for section in doc_obj.sections:
                try:
                    header = section.header
                    if header is None:
                        continue
                    try:
                        for el in list(header._element):
                            header._element.remove(el)
                    except Exception:
                        pass
                    try:
                        header.is_linked_to_previous = False
                    except Exception:
                        pass
                    hp = header.add_paragraph('')
                    hp.paragraph_format.space_before = Pt(0)
                    hp.paragraph_format.space_after = Pt(0)
                    hp.paragraph_format.line_spacing = 1.0
                    hr = hp.add_run()
                    hr.add_picture(header_img_path, width=Inches(7.5))
                    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    spacer = header.add_paragraph('')
                    spacer.paragraph_format.space_before = Pt(0)
                    spacer.paragraph_format.space_after = Pt(0)
                    spacer.paragraph_format.line_spacing = 1.0
                    spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass

        def _bump_doc_font_sizes(doc_obj, delta_pt: float = 1.0):
            try:
                def _bump_runs_in_paragraphs(paragraphs):
                    for p in paragraphs or []:
                        for r in getattr(p, 'runs', []) or []:
                            try:
                                if r.font is None or r.font.size is None:
                                    continue
                                cur = r.font.size.pt
                                if cur is None:
                                    continue
                                r.font.size = Pt(float(cur) + float(delta_pt))
                            except Exception:
                                pass

                _bump_runs_in_paragraphs(getattr(doc_obj, 'paragraphs', []) or [])

                for t in getattr(doc_obj, 'tables', []) or []:
                    try:
                        for row in t.rows:
                            for cell in row.cells:
                                _bump_runs_in_paragraphs(getattr(cell, 'paragraphs', []) or [])
                    except Exception:
                        continue

                for section in getattr(doc_obj, 'sections', []) or []:
                    try:
                        hdr = getattr(section, 'header', None)
                        ftr = getattr(section, 'footer', None)
                        if hdr is not None:
                            _bump_runs_in_paragraphs(getattr(hdr, 'paragraphs', []) or [])
                            for t in getattr(hdr, 'tables', []) or []:
                                for row in t.rows:
                                    for cell in row.cells:
                                        _bump_runs_in_paragraphs(getattr(cell, 'paragraphs', []) or [])
                        if ftr is not None:
                            _bump_runs_in_paragraphs(getattr(ftr, 'paragraphs', []) or [])
                            for t in getattr(ftr, 'tables', []) or []:
                                for row in t.rows:
                                    for cell in row.cells:
                                        _bump_runs_in_paragraphs(getattr(cell, 'paragraphs', []) or [])
                    except Exception:
                        continue
            except Exception:
                pass

        def _set_specific_font_size(doc_obj, from_pt: float, to_pt: float):
            try:
                def _apply(paragraphs):
                    for p in paragraphs or []:
                        for r in getattr(p, 'runs', []) or []:
                            try:
                                if r.font is None or r.font.size is None:
                                    continue
                                cur = r.font.size.pt
                                if cur is None:
                                    continue
                                if abs(float(cur) - float(from_pt)) < 0.25:
                                    r.font.size = Pt(float(to_pt))
                            except Exception:
                                pass

                _apply(getattr(doc_obj, 'paragraphs', []) or [])

                for t in getattr(doc_obj, 'tables', []) or []:
                    try:
                        for row in t.rows:
                            for cell in row.cells:
                                _apply(getattr(cell, 'paragraphs', []) or [])
                    except Exception:
                        continue

                for section in getattr(doc_obj, 'sections', []) or []:
                    try:
                        hdr = getattr(section, 'header', None)
                        ftr = getattr(section, 'footer', None)
                        if hdr is not None:
                            _apply(getattr(hdr, 'paragraphs', []) or [])
                            for t in getattr(hdr, 'tables', []) or []:
                                for row in t.rows:
                                    for cell in row.cells:
                                        _apply(getattr(cell, 'paragraphs', []) or [])
                        if ftr is not None:
                            _apply(getattr(ftr, 'paragraphs', []) or [])
                            for t in getattr(ftr, 'tables', []) or []:
                                for row in t.rows:
                                    for cell in row.cells:
                                        _apply(getattr(cell, 'paragraphs', []) or [])
                    except Exception:
                        continue
            except Exception:
                pass

        def _set_footer_distance_cm(doc_obj, cm: float = 0.8):
            try:
                dist_in = float(cm) / 2.54
                for section in getattr(doc_obj, 'sections', []) or []:
                    try:
                        section.footer_distance = Inches(dist_in)
                    except Exception:
                        pass
            except Exception:
                pass

        def _apply_footer_paging(doc_obj, teklif_no: str = '', teklif_tarihi: str = ''):
            """Resimdeki gibi footer: 1 satır tablo, 3 sütun (Sol: Sayı, Orta: Firma, Sağ: Form kodu + sayfa)"""
            try:
                teklif_no = (teklif_no or '').strip()
                
                for section in doc_obj.sections:
                    try:
                        footer = section.footer
                        footer.is_linked_to_previous = False
                        
                        # Mevcut içeriği temizle
                        for p in list(footer.paragraphs):
                            p._element.getparent().remove(p._element)
                        for t in list(footer.tables):
                            t._element.getparent().remove(t._element)
                        
                        # 1 satır, 3 sütun tablo
                        table = footer.add_table(rows=1, cols=3)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
                        # Sütun genişlikleri
                        widths = [Inches(2.0), Inches(3.5), Inches(2.0)]
                        for idx, cell in enumerate(table.rows[0].cells):
                            cell.width = widths[idx]
                        
                        # Sol hücre: Sayı
                        cell_left = table.rows[0].cells[0]
                        p_left = cell_left.paragraphs[0]
                        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run_left = p_left.add_run(f"Sayı:{teklif_no}")
                        run_left.font.bold = True
                        run_left.font.color.rgb = RGBColor(0, 0, 128)
                        run_left.font.size = Pt(9)
                        
                        # Orta hücre: Firma bilgileri (3 satır)
                        cell_center = table.rows[0].cells[1]
                        p_center = cell_center.paragraphs[0]
                        p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        run1 = p_center.add_run("AKARE ÇEVRE LABORATUVAR VE DAN. HİZM. TİC.LTD.ŞTİ Kirazlıyalı Mah. Süleyman Demirel Cad. No:28/A")
                        run1.font.bold = True
                        run1.font.color.rgb = RGBColor(0, 0, 128)
                        run1.font.size = Pt(9)
                        
                        p_center.add_run("\n")
                        
                        run2 = p_center.add_run("Körfez V.D 013 065 1290 Körfez-KOCAELİ")
                        run2.font.bold = True
                        run2.font.color.rgb = RGBColor(0, 0, 128)
                        run2.font.size = Pt(9)
                        
                        p_center.add_run("\n")
                        
                        run3 = p_center.add_run("info@akarecevre.com  www.akarecevre.com")
                        run3.font.bold = True
                        run3.font.color.rgb = RGBColor(0, 0, 128)
                        run3.font.size = Pt(9)
                        
                        # Sağ hücre: Form kodu ve sayfa numarası
                        cell_right = table.rows[0].cells[2]
                        p_right = cell_right.paragraphs[0]
                        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        
                        run_form = p_right.add_run("AÇ.F.102/Rev04/14.08.2025")
                        run_form.font.bold = True
                        run_form.font.color.rgb = RGBColor(0, 0, 128)
                        run_form.font.size = Pt(9)
                        
                        p_right.add_run("  ")
                        
                        # Sayfa numarası
                        run_page = p_right.add_run()
                        run_page.font.bold = True
                        run_page.font.color.rgb = RGBColor(0, 0, 128)
                        run_page.font.size = Pt(9)
                        
                        fldChar1 = OxmlElement('w:fldChar')
                        fldChar1.set(qn('w:fldCharType'), 'begin')
                        run_page._r.append(fldChar1)
                        instrText = OxmlElement('w:instrText')
                        instrText.text = 'PAGE'
                        run_page._r.append(instrText)
                        fldChar2 = OxmlElement('w:fldChar')
                        fldChar2.set(qn('w:fldCharType'), 'end')
                        run_page._r.append(fldChar2)
                        
                        run_page.add_text('/')
                        
                        fldChar3 = OxmlElement('w:fldChar')
                        fldChar3.set(qn('w:fldCharType'), 'begin')
                        run_page._r.append(fldChar3)
                        instrText2 = OxmlElement('w:instrText')
                        instrText2.text = 'NUMPAGES'
                        run_page._r.append(instrText2)
                        fldChar4 = OxmlElement('w:fldChar')
                        fldChar4.set(qn('w:fldCharType'), 'end')
                        run_page._r.append(fldChar4)
                        
                        # Tablo kenarlıklarını kaldır
                        for cell in table.rows[0].cells:
                            tcPr = cell._element.get_or_add_tcPr()
                            tcBorders = OxmlElement('w:tcBorders')
                            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                                border = OxmlElement(f'w:{border_name}')
                                border.set(qn('w:val'), 'none')
                                tcBorders.append(border)
                            tcPr.append(tcBorders)
                        
                    except Exception as e:
                        continue
            except Exception:
                pass

        def _remove_all_section_breaks(doc_obj):
            """Remove section properties that can create unexpected blank pages when merging."""
            try:
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                body = doc_obj.element.body
                # body-level sectPr
                for child in list(body.iterchildren()):
                    if child.tag.endswith('}sectPr'):
                        try:
                            body.remove(child)
                        except Exception:
                            pass
                # paragraph-level sectPr
                for p in list(body.iterchildren()):
                    if not p.tag.endswith('}p'):
                        continue
                    for sp in list(p.xpath('.//w:pPr/w:sectPr', namespaces=ns)):
                        try:
                            sp.getparent().remove(sp)
                        except Exception:
                            pass
            except Exception:
                pass

        def _fill_teklif2_table(doc_obj, teklif_obj):
            try:
                params = (teklif_obj or {}).get('parametreler', []) or []
                if not params:
                    return

                def _find_pricing_table():
                    for t in doc_obj.tables:
                        try:
                            if not t.rows:
                                continue
                            # Header row may not be the first row (some templates have a title row)
                            for r in t.rows[:3]:
                                hdr = '|'.join(_norm(c.text) for c in r.cells)
                                if 'parametre' in hdr and ('metodu' in hdr or 'metot' in hdr or 'metod' in hdr):
                                    return t
                        except Exception:
                            continue
                    return None

                table = _find_pricing_table()
                if table is None:
                    return

                header_row_idx = 0
                header_cells = table.rows[0].cells if table.rows else []
                header_texts = [_norm(c.text) for c in header_cells]
                try:
                    for ri in range(min(3, len(table.rows))):
                        ht = [_norm(c.text) for c in table.rows[ri].cells]
                        joined = ' '.join(ht)
                        if 'parametre' in joined and ('metodu' in joined or 'metot' in joined or 'metod' in joined):
                            header_row_idx = ri
                            header_cells = table.rows[ri].cells
                            header_texts = ht
                            break
                except Exception:
                    pass

                def _find_col(*needles):
                    for i, ht in enumerate(header_texts):
                        if all(n in ht for n in needles):
                            return i
                    return None

                col_param = _find_col('parametre')
                col_metot = _find_col('metod')
                if col_metot is None:
                    col_metot = _find_col('metot')
                if col_metot is None and len(header_texts) >= 2:
                    col_metot = 1
                col_adet = _find_col('adet')

                col_birim = _find_col('birim')
                if col_birim is None:
                    col_birim = _find_col('birim', 'fiyat')
                if col_birim is None:
                    col_birim = _find_col('fiyat')

                col_toplam = _find_col('toplam')
                if col_toplam is None:
                    # Templates may use 'TOP. FİYAT' instead of 'Toplam'
                    col_toplam = _find_col('top', 'fiyat')
                if col_toplam is None:
                    col_toplam = _find_col('top')
                
                # İndirim veya İskonto sütununu bul
                col_indirim = _find_col('indirim')
                if col_indirim is None:
                    col_indirim = _find_col('iskonto')

                # Fallback: if header-based detection fails, assume last 2 columns are money columns
                try:
                    ncols = len(header_texts)
                    if ncols >= 2:
                        if col_toplam is None:
                            col_toplam = ncols - 1
                        if col_birim is None:
                            col_birim = ncols - 2
                except Exception:
                    pass

                totals_start_idx = None
                for i, row in enumerate(table.rows):
                    row_text = ' '.join(_norm(c.text) for c in row.cells)
                    if 'toplam' in row_text and i > header_row_idx:
                        totals_start_idx = i
                        break

                data_start = header_row_idx + 1
                data_end = totals_start_idx if totals_start_idx is not None else len(table.rows)
                capacity = max(0, data_end - data_start)

                if len(params) > capacity:
                    for _ in range(len(params) - capacity):
                        if totals_start_idx is None:
                            table.add_row()
                        else:
                            last_data_row = table.rows[data_end - 1]._tr
                            new_tr = deepcopy(last_data_row)
                            table._tbl.insert(totals_start_idx, new_tr)
                            data_end += 1
                            totals_start_idx += 1

                def _as_int(v, default=0):
                    try:
                        return int(v)
                    except Exception:
                        try:
                            return int(float(v))
                        except Exception:
                            return default

                def _as_float(v, default=0.0):
                    try:
                        return float(v)
                    except Exception:
                        return default

                # Word formatting helpers (alignment + row height)
                try:
                    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN_PARAGRAPH
                    from docx.enum.table import WD_ALIGN_VERTICAL as _WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE as _WD_ROW_HEIGHT_RULE
                    from docx.shared import Pt as _Pt
                except Exception:
                    _WD_ALIGN_PARAGRAPH = None
                    _WD_ALIGN_VERTICAL = None
                    _WD_ROW_HEIGHT_RULE = None
                    _Pt = None

                def _align_cell(cell, right: bool = False):
                    try:
                        if _WD_ALIGN_VERTICAL is not None:
                            cell.vertical_alignment = _WD_ALIGN_VERTICAL.CENTER
                        if _WD_ALIGN_PARAGRAPH is not None:
                            for p in cell.paragraphs:
                                p.alignment = _WD_ALIGN_PARAGRAPH.RIGHT if right else _WD_ALIGN_PARAGRAPH.LEFT
                        # Add a bit of right padding for numeric cells so text doesn't touch the border
                        if right:
                            try:
                                tcPr = cell._tc.get_or_add_tcPr()
                                tcMar = tcPr.find(qn('w:tcMar'))
                                if tcMar is None:
                                    tcMar = OxmlElement('w:tcMar')
                                    tcPr.append(tcMar)
                                right_el = tcMar.find(qn('w:right'))
                                if right_el is None:
                                    right_el = OxmlElement('w:right')
                                    tcMar.append(right_el)
                                # twips: 1/20 pt. ~180 twips ≈ 0.125 in ≈ ~2 characters feel
                                right_el.set(qn('w:w'), '180')
                                right_el.set(qn('w:type'), 'dxa')
                            except Exception:
                                pass
                    except Exception:
                        pass

                def _set_row_height(row_obj, points: float):
                    try:
                        if _Pt is None or _WD_ROW_HEIGHT_RULE is None:
                            return
                        row_obj.height = _Pt(points)
                        row_obj.height_rule = _WD_ROW_HEIGHT_RULE.AT_LEAST
                    except Exception:
                        pass

                for idx, pr in enumerate(params):
                    row = table.rows[data_start + idx]
                    _set_row_height(row, 18)
                    adet_val = _as_int(pr.get('adet', 0), 0)
                    birim_val = _as_float(pr.get('birimFiyat', 0), 0.0)
                    toplam_val = _as_float(pr.get('topFiyat', birim_val * adet_val), birim_val * adet_val)

                    def _set_cell(ci, v, is_money: bool = False, align_right: bool = False):
                        if ci is None or ci >= len(row.cells):
                            return
                        try:
                            cell = row.cells[ci]
                            txt = '' if v is None else str(v)
                            if is_money:
                                try:
                                    txt = f"{float(v):.2f}"
                                except Exception:
                                    pass
                            cell.text = ''
                            cell.text = txt
                            _align_cell(cell, right=bool(is_money or align_right))
                        except Exception:
                            pass

                    _set_cell(col_param, pr.get('parametre', ''))
                    # Guarantee metot column is filled
                    metot_value = pr.get('metot', '')
                    if (metot_value is None) or (str(metot_value).strip() == ''):
                        metot_value = pr.get('metod', '')
                    _set_cell(col_metot, metot_value)
                    _set_cell(col_adet, adet_val, align_right=True)
                    _set_cell(col_birim, birim_val, is_money=True)
                    _set_cell(col_toplam, toplam_val, is_money=True)
                    # İskonto sütunu parametre satırlarında yok, sadece özet bölümünde var

                # Remove extra template rows so row count matches param count
                try:
                    desired = len(params)
                    # Recompute bounds after any row insertions
                    if totals_start_idx is None:
                        data_end = len(table.rows)
                    else:
                        data_end = totals_start_idx
                    # Delete from bottom of data area to avoid index shifts
                    last_keep = data_start + desired - 1
                    for ridx in range(data_end - 1, last_keep, -1):
                        if ridx < data_start:
                            break
                        try:
                            table._tbl.remove(table.rows[ridx]._tr)
                            if totals_start_idx is not None:
                                totals_start_idx -= 1
                        except Exception:
                            pass
                except Exception:
                    pass

                toplam = teklif_obj.get('toplam', None)
                indirim = teklif_obj.get('indirim', None)
                net = teklif_obj.get('netToplam', None)
                
                print(f"✓ Değerler: Toplam={toplam}, İndirim={indirim}, Net={net}, totals_start_idx={totals_start_idx}")

                if totals_start_idx is not None:
                    def _row_key(row_obj) -> str:
                        return ' '.join(_norm(c.text) for c in row_obj.cells)

                    def _write_total(row_idx: int, value):
                        if value is None:
                            return
                        try:
                            txt = f"{float(value):.2f}"
                        except Exception:
                            txt = str(value)
                        try:
                            r = table.rows[row_idx]
                            _set_row_height(r, 18)
                            r.cells[-1].text = txt
                            _align_cell(r.cells[-1], right=True)
                        except Exception:
                            pass

                    # Özet satırlarını doldur (yeni şablon TOPLAM, İSKONTO, TOPLAM TUTAR satırlarını içeriyor)
                    for i in range(totals_start_idx, len(table.rows)):
                        # Satırın ilk hücresinin içeriğini al (normalize etmeden)
                        try:
                            first_cell_text = table.rows[i].cells[0].text.strip().upper()
                            print(f"  → Satır {i}: '{first_cell_text}'")
                            
                            # İskonto satırını bul - büyük harfe çevirip kontrol et
                            if 'İSKONTO' in first_cell_text or 'ISKONTO' in first_cell_text or 'İNDİRİM' in first_cell_text or 'INDIRIM' in first_cell_text:
                                print(f"    ✓ İskonto satırı bulundu, değer yazılıyor: {indirim}")
                                _write_total(i, indirim)
                            # Toplam tutar satırını bul
                            elif 'TOPLAM TUTAR' in first_cell_text or 'NET TOPLAM' in first_cell_text:
                                print(f"    ✓ Toplam tutar satırı bulundu, değer yazılıyor: {net}")
                                _write_total(i, net)
                            # Toplam satırını bul (ama toplam tutar veya iskonto değil)
                            elif first_cell_text.startswith('TOPLAM') and 'TUTAR' not in first_cell_text and 'İSKONTO' not in first_cell_text and 'ISKONTO' not in first_cell_text:
                                print(f"    ✓ Toplam satırı bulundu, değer yazılıyor: {toplam}")
                                _write_total(i, toplam)
                        except Exception as e:
                            print(f"Satır {i} işlenirken hata: {e}")
                            continue

            except Exception as _e:
                print(f"TEKLİF-2 tablo doldurma hatası: {_e}")

        teklif1_path = os.path.join(app.root_path, 'static', 'images', 'TEKLİF-1 GİRİŞ.docx')
        teklif2_path = os.path.join(app.root_path, 'static', 'images', 'TEKLİF - 2 FİYAT.docx')
        teklif3_path = os.path.join(app.root_path, 'static', 'images', 'TEKLİF - 3 GENEL HUKUM.docx')

        if not os.path.exists(teklif1_path):
            return jsonify({'success': False, 'message': 'TEKLİF-1 GİRİŞ.docx şablonu bulunamadı'})

        # 1) Prepare TEKLİF-1 (master)
        doc1 = Document(teklif1_path)
        tel_email = f"{firma.get('telefon', '') if firma else ''} / {firma.get('email', '') if firma else ''}"
        teklif_tarihi_raw = teklif.get('teklif_tarihi', '')
        try:
            teklif_tarihi = format_tarih_gg_aa_yyyy(teklif_tarihi_raw)
        except Exception:
            teklif_tarihi = teklif_tarihi_raw
        set_labeled_table_value(doc1, 'Firma Adı', firma.get('firmaAdi', '') if firma else '')
        set_labeled_table_value(doc1, 'Firma Yetkili', firma.get('yetkiliAdi', '') if firma else '')
        set_labeled_table_value(doc1, 'Firma Adresi', firma.get('adres', '') if firma else '')
        set_labeled_table_value(doc1, 'Tel', tel_email)
        set_labeled_table_value(doc1, 'Tel / E-Posta', tel_email)
        set_labeled_table_value(doc1, 'Faks', firma.get('faks', '') if firma else '')
        set_labeled_table_value(doc1, 'E-Posta', firma.get('email', '') if firma else '')
        set_labeled_table_value(doc1, 'Talep', teklif.get('teklif_tipi', ''))
        set_labeled_table_value(doc1, 'Teklif Kodu', teklif.get('teklif_no', ''))
        set_labeled_table_value(doc1, 'Sayfa Adedi', '')
        set_labeled_table_value(doc1, 'Tarih', teklif_tarihi)
        _clean_template_breaks(doc1)
        _apply_header_image_only(doc1)
        _set_footer_distance_cm(doc1, 0.8)
        # Footer placeholder'ları şablonda tanımlı: {{TEKLIF_NO}}, {{SAYFA_NO}}
        # replace_placeholders_in_document ile değiştirilecek
        _bump_doc_font_sizes(doc1, 1.0)
        _set_specific_font_size(doc1, 14.0, 12.0)

        # Merge with docxcompose when available (more stable than altChunk/body-copy)
        final_path = None
        try:
            try:
                from docxcompose.composer import Composer
            except Exception:
                Composer = None

            master = doc1
            composer = Composer(master) if Composer is not None else None

            def _append_with_page_break(src_doc):
                # Ensure no leftover section breaks from the source
                _remove_all_section_breaks(src_doc)
                _clean_template_breaks(src_doc)
                # page break at end of current doc (not at start of next)
                master.add_page_break()
                if composer is not None:
                    composer.append(src_doc)
                else:
                    # Fallback: append body elements (least preferred)
                    src_body = src_doc.element.body
                    for child in list(src_body.iterchildren()):
                        if child.tag.endswith('}sectPr'):
                            continue
                        master.element.body.append(deepcopy(child))

            def _fill_teklif3_acceptance_sentence(doc_obj, tarih_text: str, teklif_kodu: str):
                """TEKLİF-3 şablonundaki 'Xxxxx tarih ve xxxx sayılı...' cümlesini doldur."""
                try:
                    # Requested final text format
                    # Example: "gg.aa.yyyy Tarih ve TEYY-SSS Sayılı Teklifinizi Kabul Ediyoruz."
                    tarih_text = (tarih_text or '').strip() or 'gg.aa.yyyy'
                    teklif_kodu = (teklif_kodu or '').strip() or 'TEYY-SSS'
                    target_sentence = f"{tarih_text} Tarih ve {teklif_kodu} Sayılı Teklifinizi Kabul Ediyoruz."
                    keywords = ['tarih', 'sayılı', 'teklifinizi', 'kabul']

                    try:
                        from docx.shared import RGBColor
                    except Exception:
                        RGBColor = None

                    def _matches(p_text: str) -> bool:
                        t = (p_text or '').strip().lower()
                        if not t:
                            return False
                        return all(k in t for k in keywords)

                    def _set_paragraph(par):
                        try:
                            if not _matches(par.text):
                                return False
                            par.text = ''
                            run = par.add_run(target_sentence)
                            # Remove hyperlink-like underline/blue styling coming from template
                            try:
                                run.font.underline = False
                            except Exception:
                                pass
                            try:
                                if RGBColor is not None:
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                            except Exception:
                                pass
                            return True
                        except Exception:
                            return False

                    # Paragraphs outside tables
                    for p in getattr(doc_obj, 'paragraphs', []) or []:
                        if _set_paragraph(p):
                            return

                    # Paragraphs inside tables
                    for t in getattr(doc_obj, 'tables', []) or []:
                        for r in t.rows:
                            for c in r.cells:
                                for p in c.paragraphs:
                                    if _set_paragraph(p):
                                        return
                except Exception:
                    return

            if os.path.exists(teklif2_path):
                d2 = Document(teklif2_path)
                _fill_teklif2_table(d2, teklif)
                _set_footer_distance_cm(d2, 0.8)
                # Footer placeholder'ları şablonda tanımlı
                _bump_doc_font_sizes(d2, 1.0)
                _set_specific_font_size(d2, 14.0, 12.0)
                _append_with_page_break(d2)

            if os.path.exists(teklif3_path):
                d3 = Document(teklif3_path)
                _fill_teklif3_acceptance_sentence(d3, teklif_tarihi, teklif.get('teklif_no', ''))
                _set_footer_distance_cm(d3, 0.8)
                # Footer placeholder'ları şablonda tanımlı
                _bump_doc_font_sizes(d3, 1.0)
                _set_specific_font_size(d3, 14.0, 12.0)
                _append_with_page_break(d3)

            # Apply header/footer on the final merged document (all sections)
            _apply_header_image_only(master)
            _set_footer_distance_cm(master, 0.8)
            
            # Footer placeholder'ları değiştir
            placeholder_data = {
                'TEKLIF_NO': teklif.get('teklif_no', '')
            }
            replace_placeholders_in_document(master, placeholder_data)
            
            _bump_doc_font_sizes(master, 1.0)
            _set_specific_font_size(master, 14.0, 12.0)

            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as fout:
                final_path = fout.name
            if composer is not None:
                composer.save(final_path)
            else:
                master.save(final_path)

            firma_adi = firma.get('firmaAdi', '') if firma else ''
            if not firma_adi:
                firma_adi = teklif.get('firma_adi', '')
            firma_kisa = '_'.join((firma_adi or '').strip().split()[:2]) or 'Firma'
            teklif_tarihi = teklif.get('teklif_tarihi', '')
            if teklif_tarihi:
                try:
                    tarih_obj = datetime.strptime(teklif_tarihi, '%Y-%m-%d')
                    tarih_format = tarih_obj.strftime('%d%m%y')
                except Exception:
                    tarih_format = datetime.now().strftime('%d%m%y')
            else:
                tarih_format = datetime.now().strftime('%d%m%y')
            tutar_kdv_haric = teklif.get('netToplam', 0)
            dosya_adi = f"{firma_kisa}_{teklif.get('teklif_no', '')}_{tarih_format}_{tutar_kdv_haric:.0f}TL.docx"

            if return_file_info:
                return final_path, dosya_adi

            return send_file(
                final_path,
                as_attachment=True,
                download_name=dosya_adi,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        except Exception as save_error:
            if final_path and os.path.exists(final_path):
                try:
                    os.unlink(final_path)
                except Exception:
                    pass
            raise save_error

    except Exception as e:
        return jsonify({'success': False, 'message': f'Word dosyası oluşturma hatası: {str(e)}'})

def create_pdf_teklif(teklif, firma):
    """PDF formatında teklif oluşturur (Word'den PDF'e çevir)"""
    try:
        # Önce DOCX üret
        word_path, word_name = create_word_teklif(teklif, firma, return_file_info=True)

        pdf_path = None
        try:
            from docx2pdf import convert

            base_name = os.path.splitext(word_name)[0]
            pdf_name = f"{base_name}.pdf"
            pdf_path = os.path.join(tempfile.gettempdir(), pdf_name)

            # Dönüştür
            convert(word_path, pdf_path)

            with open(pdf_path, 'rb') as f:
                pdf_bytes = f.read()

            response = make_response(pdf_bytes)
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = f'attachment; filename="{pdf_name}"'
            return response

        except Exception as pdf_error:
            print(f"Teklif PDF dönüştürme hatası: {pdf_error}")
            # PDF başarısızsa Word döndür
            with open(word_path, 'rb') as f:
                word_bytes = f.read()
            response = make_response(word_bytes)
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = f'attachment; filename="{word_name}"'
            return response
        finally:
            # Geçici dosyaları temizle
            try:
                if word_path and os.path.exists(word_path):
                    os.unlink(word_path)
            except Exception:
                pass
            try:
                if pdf_path and os.path.exists(pdf_path):
                    os.unlink(pdf_path)
            except Exception:
                pass
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'PDF dosyası oluşturma hatası: {str(e)}'})

@app.route('/export_graph_and_data', methods=['POST'])
def export_graph_and_data():
    """Grafiği ve veri tablosunu Excel dosyasına dışa aktarır."""
    try:
        # Grafiği oluştur ve PNG olarak kaydet
        plt, np, mdates, Rectangle = load_matplotlib()
        plt.plot([1, 2, 3, 4], [10, 20, 25, 30])
        plt.title("Örnek Grafik")
        graph_path = tempfile.NamedTemporaryFile(delete=False, suffix='.png').name
        plt.savefig(graph_path)
        plt.close()

        # Veri tablosunu oluştur
        data = [
            ["Sıra", "Tarih", "Firma", "Kod", "Baca", "Personel", "Cihaz", "Değer"],
            [1, "22.08.25", "AKARE ÇE", "E-250281-02", "SSS", "Admin", "515"],
            [2, "11.08.25", "ATALAY T", "EE-250811-02", "DDDD", "Admin", "490"],
        ]

        # Excel dosyasını oluştur
        output = BytesIO()
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Grafik ve Veri"

        # PNG'yi ekle
        from openpyxl.drawing.image import Image
        img = Image(graph_path)
        ws.add_image(img, "A1")

        # PNG'nin altına veri tablosunu ekle (örneğin A20'den başlat)
        start_row = 20
        for i, row in enumerate(data):
            for j, value in enumerate(row):
                ws.cell(row=start_row + i, column=1 + j, value=value)

        wb.save(output)
        output.seek(0)

        return send_file(
            output,
            as_attachment=True,
            download_name="grafik_ve_veri.xlsx",
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    except Exception as e:
        print(f"Grafik ve veri dışa aktarılırken hata: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/kk_excel_with_graph', methods=['POST'])
def kk_excel_with_graph():
    try:
        data = request.get_json()
        parametre = data.get('parametre')
        tarih_baslangic = data.get('tarih_baslangic')
        tarih_bitis = data.get('tarih_bitis')
        table_data = data.get('localStorage_verileri', [])
        graph_png = data.get('graph_png')

        # PNG'yi base64'ten dosyaya çevir
        import base64
        from openpyxl import Workbook
        from openpyxl.drawing.image import Image
        import tempfile
        import re
        img_data = re.sub('^data:image/.+;base64,', '', graph_png)
        img_bytes = base64.b64decode(img_data)
        img_path = tempfile.NamedTemporaryFile(delete=False, suffix='.png').name
        with open(img_path, 'wb') as f:
            f.write(img_bytes)

        # Excel oluştur
        wb = Workbook()
        ws = wb.active
        ws.title = "Grafik ve Veri"
        # PNG'yi ekle
        img = Image(img_path)
        ws.add_image(img, "A1")
        # Tabloyu ekle (10. satırdan itibaren)
        headers = ["Sıra", "Tarih", "Firma", "Kod", "Baca", "Personel", "Cihaz", "Değer"]
        for col, val in enumerate(headers, start=1):
            ws.cell(row=10, column=col, value=val)
        for i, row in enumerate(table_data, start=11):
            for col, key in enumerate(["sira", "tarih", "firma", "kod", "baca", "personel", "cihaz", "deger"], start=1):
                ws.cell(row=i, column=col, value=row.get(key, ''))
        # Dosyayı kaydet
        output = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
        wb.save(output.name)
        return send_file(
            output.name,
            as_attachment=True,
            download_name=f'KK_Rapor_{parametre}_{tarih_baslangic}_{tarih_bitis}.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    except Exception as e:
        print(f"Excel export with graph error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/pivot/compare')
def api_pivot_compare():
    """PIVOT karşılaştırma API endpoint'i - yıllara göre baca sayısı, parametre sayısı, fiyat ve personel performansını döndürür."""
    try:
        years = request.args.get('years', '')
        if not years:
            return jsonify({'success': False, 'error': 'Yıl parametresi gerekli'}), 400
        
        year_list = [int(y.strip()) for y in years.split(',') if y.strip()]
        if not year_list:
            return jsonify({'success': False, 'error': 'Geçerli yıl listesi gerekli'}), 400
        
        # Verileri yükle
        baca_bilgileri = load_baca_bilgileri()
        parametre_olcumleri = load_parametre_olcum()
        teklifler = load_teklif()

        def _strip_scope_prefix(pname: str) -> str:
            try:
                s = (pname or '').strip()
                if not s:
                    return ''
                up = s.upper()
                if up.startswith('(E)') or up.startswith('(İ)') or up.startswith('(I)'):
                    s = s[s.find(')') + 1:].strip()
                if s.startswith('-'):
                    s = s[1:].strip()
                return s
            except Exception:
                return (pname or '').strip()

        def _norm_param_key(pname: str) -> str:
            """Normalize parameter names so prefixes/diacritics/punctuation differences don't split keys."""
            try:
                s = _strip_scope_prefix(pname)
                s = (s or '').strip().lower()
                if not s:
                    return ''
                tr_map = str.maketrans({'ı':'i','İ':'i','ş':'s','Ş':'s','ç':'c','Ç':'c','ğ':'g','Ğ':'g','ü':'u','Ü':'u','ö':'o','Ö':'o'})
                s = s.translate(tr_map)
                # keep letters/numbers/spaces only
                import re
                s = re.sub(r'[^a-z0-9\s]+', ' ', s)
                s = re.sub(r'\s+', ' ', s).strip()
                return s
            except Exception:
                return _strip_scope_prefix(pname).strip().lower()

        def _is_non_baca_item(pname: str) -> bool:
            try:
                k = _norm_param_key(pname)
                if not k:
                    return True
                # Offer/fee rows that must not contribute to baca estimation
                if 'yol' in k:
                    return True
                if 'raporlama' in k:
                    return True
                if 'taban fiyat' in k:
                    return True
                if 'tesis tabani' in k:
                    return True
                return False
            except Exception:
                return False

        def _is_non_parametre_total_item(pname: str) -> bool:
            try:
                k = _norm_param_key(pname)
                if not k:
                    return True
                if 'yol' in k:
                    return True
                if 'raporlama' in k:
                    return True
                if 'taban fiyat' in k:
                    return True
                if 'tesis tabani' in k:
                    return True
                return False
            except Exception:
                return False
        
        result_data = {}
        
        for year in year_list:
            # O yıl için teklifleri filtrele
            year_teklifler = []
            for t in teklifler:
                tdate = t.get('teklif_tarihi', '')
                if tdate and tdate[:4] == str(year):
                    year_teklifler.append(t)
            
            # Toplam teklif sayısı ve tutarı
            toplam_teklif_adedi = len(year_teklifler)
            toplam_teklif_tutari = sum(float(t.get('netToplam', 0) or 0) for t in year_teklifler)
            
            # Kabul olan teklifleri filtrele
            def _norm(sval: str) -> str:
                m = (sval or '').strip().lower()
                tr_map = str.maketrans({'ı':'i','İ':'i','ş':'s','Ş':'s','ç':'c','Ç':'c','ğ':'g','Ğ':'g','ü':'u','Ü':'u','ö':'o','Ö':'o'})
                return m.translate(tr_map)
            
            kabul_teklifler = []
            for t in year_teklifler:
                status = _norm(t.get('teklif_durumu', ''))
                if any(k in status for k in ['kabul', 'onay']):
                    kabul_teklifler.append(t)
            
            kabul_adet = len(kabul_teklifler)
            kabul_tutari = sum(float(t.get('netToplam', 0) or 0) for t in kabul_teklifler)
            
            # Kapsam içi/dışı teklifleri hesapla - TÜM teklifler arasından (sadece kabul olanlar değil)
            kapsam_ici_teklifler = [t for t in year_teklifler if 'kapsam' in _norm(t.get('teklif_tipi', '')) and ('ici' in _norm(t.get('teklif_tipi', '')) or 'i̇ci̇' in _norm(t.get('teklif_tipi', '')))]
            kapsam_disi_teklifler = [t for t in year_teklifler if 'kapsam' in _norm(t.get('teklif_tipi', '')) and ('disi' in _norm(t.get('teklif_tipi', '')) or 'dis' in _norm(t.get('teklif_tipi', '')))]
            is_birligi_teklifler = [t for t in year_teklifler if ('is' in _norm(t.get('teklif_tipi', '')) or 'i̇s' in _norm(t.get('teklif_tipi', ''))) and ('birlik' in _norm(t.get('teklif_tipi', '')) or 'bi̇rli̇k' in _norm(t.get('teklif_tipi', '')) or 'bi̇rli̇gi̇' in _norm(t.get('teklif_tipi', '')))]
            
            # Kapsam içi/dışı tutarları - sadece KABUL olanların tutarını hesapla
            kapsam_ici_kabul_teklifler = [t for t in kapsam_ici_teklifler if any(k in _norm(t.get('teklif_durumu', '')) for k in ['kabul', 'onay'])]
            kapsam_disi_kabul_teklifler = [t for t in kapsam_disi_teklifler if any(k in _norm(t.get('teklif_durumu', '')) for k in ['kabul', 'onay'])]
            is_birligi_kabul_teklifler = [t for t in is_birligi_teklifler if any(k in _norm(t.get('teklif_durumu', '')) for k in ['kabul', 'onay'])]
            
            # Adet sayıları - sadece KABUL olanların sayısı
            kapsam_ici_adet = len(kapsam_ici_kabul_teklifler)
            kapsam_disi_adet = len(kapsam_disi_kabul_teklifler)
            is_birligi_adet = len(is_birligi_kabul_teklifler)
            
            kapsam_ici_tutar = sum(float(t.get('netToplam', 0) or 0) for t in kapsam_ici_kabul_teklifler)
            kapsam_disi_tutar = sum(float(t.get('netToplam', 0) or 0) for t in kapsam_disi_kabul_teklifler)
            is_birligi_tutar = sum(float(t.get('netToplam', 0) or 0) for t in is_birligi_kabul_teklifler)
            
            # Baca sayısını hesapla - Firma ölçümlerindeki BACA SAY değerlerinin toplamı
            firma_olcumler = load_firma_olcum()
            toplam_baca_adedi = 0
            for olcum in firma_olcumler:
                olcum_date = olcum.get('baslangic_tarihi', '')
                if olcum_date and olcum_date[:4] == str(year):
                    baca_sayisi = olcum.get('baca_sayisi', '0')
                    try:
                        toplam_baca_adedi += int(baca_sayisi)
                    except (ValueError, TypeError):
                        # Baca sayısı sayı değilse 0 olarak kabul et
                        pass

            # Fallback: if firma_olcum doesn't provide bacas for the year, estimate from offers
            # Each offer: take max(adet) among its parameters (Toz 10, Toc 5 => 10 bacas)
            teklif_baca_adedi = 0
            try:
                for t in year_teklifler:
                    max_adet = 0
                    for p in (t.get('parametreler', []) or []):
                        if not isinstance(p, dict):
                            continue
                        pname = p.get('parametre')
                        if not pname:
                            continue
                        if _is_non_baca_item(pname):
                            continue
                        try:
                            a = int(p.get('adet', 1) or 1)
                        except Exception:
                            a = 1
                        if a > max_adet:
                            max_adet = a
                    if max_adet > 0:
                        teklif_baca_adedi += max_adet
            except Exception:
                teklif_baca_adedi = 0

            baca_sayisi_kaynak = 'FIRMA_OLCUM'
            if toplam_baca_adedi == 0 and teklif_baca_adedi > 0:
                toplam_baca_adedi = teklif_baca_adedi
                baca_sayisi_kaynak = 'TEKLIF_MAX_ADET'
            
            # Parametre isim eşleştirme tablosu - önce tanımla
            parametre_eslesme = {
                'YG': 'YANMA GAZI',
                'YANMA GAZI': 'YANMA GAZI',  # Eksik olan eşleştirme eklendi
                'AMET': 'AĞIR METAL', 
                'SÜLF.A': 'SÜLFÜRİK ASİT',
                'SÜLFÜRİK ASİT': 'SÜLFÜRİK ASİT',  # Eksik olan eşleştirme eklendi
                'TOZ': 'TOZ',
                'VOC': 'VOC',
                'TOC': 'TOC',
                'NEM': 'NEM',
                'PM10': 'PM10',
                'HIZ': 'HIZ',
                'HF': 'HF',
                'HCL': 'HCL',
                'AMONYAK': 'AMONYAK',
                'FORMALDEHİT': 'FORMALDEHİT',
                'CR+6': 'CR+6',
                'CR6': 'CR+6',  # Alternatif yazım
                'FOSFORİK ASİT': 'FOSFORİK ASİT',
                'HCN': 'HCN',
                'DİOKSİN FURAN': 'DİOKSİN FURAN',
                'PAH': 'PAH',
                'ÇÖKEN TOZ': 'ÇÖKEN TOZ',
                'ÇT': 'ÇÖKEN TOZ',  # Alternatif yazım
                'MODELLEME': 'MODELLEME'
            }

            # Normalize mapping keys for more tolerant matching
            parametre_eslesme_norm = {}
            try:
                for k, v in parametre_eslesme.items():
                    parametre_eslesme_norm[_norm_param_key(k)] = v
            except Exception:
                parametre_eslesme_norm = {}
            
            # Firma ölçümlerinden parametre sayılarını hesapla
            from collections import defaultdict
            parametre_sayilari_firma = defaultdict(int)
            parametre_sayilari_teklif = defaultdict(int)
            toplam_parametre_adedi = 0
            
            # Firma ölçümlerindeki parametreleri say
            for olcum in firma_olcumler:
                olcum_date = olcum.get('baslangic_tarihi', '')
                if olcum_date and olcum_date[:4] == str(year):
                    baca_parametreleri = olcum.get('baca_parametreleri', {})
                    for baca_adi, parametre_listesi in baca_parametreleri.items():
                        for parametre in parametre_listesi:
                            if parametre:  # Boş parametre değilse
                                norm_key = _norm_param_key(parametre)
                                # Parametre eşleştirmesi yap (tolerant)
                                mapped = parametre_eslesme_norm.get(norm_key)
                                eslesen_param = mapped if mapped else _strip_scope_prefix(parametre).upper().strip()
                                parametre_sayilari_firma[eslesen_param] += 1
                                toplam_parametre_adedi += 1

            # Tekliflerden parametre sayılarını hesapla (kabul olan teklifler)
            try:
                for t in kabul_teklifler:
                    for p in (t.get('parametreler', []) or []):
                        pname = p.get('parametre') if isinstance(p, dict) else None
                        if not pname:
                            continue
                        # Use adet if present, otherwise count as 1
                        try:
                            adet_val = int(p.get('adet', 1)) if isinstance(p, dict) else 1
                        except Exception:
                            adet_val = 1
                        if adet_val <= 0:
                            adet_val = 1
                        norm_key = _norm_param_key(pname)
                        mapped = parametre_eslesme_norm.get(norm_key)
                        eslesen_param = mapped if mapped else _strip_scope_prefix(pname).upper().strip()
                        parametre_sayilari_teklif[eslesen_param] += adet_val
            except Exception:
                pass

            # Decide which source to use primarily: if firma_olcum has no rows for that year, fallback to teklifler
            use_teklif = False
            try:
                use_teklif = (sum(parametre_sayilari_firma.values()) == 0 and sum(parametre_sayilari_teklif.values()) > 0)
            except Exception:
                use_teklif = False

            parametre_sayilari = parametre_sayilari_teklif if use_teklif else parametre_sayilari_firma
            try:
                toplam_parametre_adedi = 0
                for k, v in parametre_sayilari.items():
                    if _is_non_parametre_total_item(k):
                        continue
                    try:
                        toplam_parametre_adedi += int(v)
                    except Exception:
                        pass
            except Exception:
                pass
            
            # Parametre tutarlarını hesapla: Parametre Adedi × Asgari Fiyat
            parametre_tl = defaultdict(float)
            asgari_fiyatlar = load_asgari_fiyatlar()
            
            # Asgari fiyatları yıla göre filtrele
            for param_adi, adet in parametre_sayilari.items():
                if adet > 0:  # Sadece adedi olan parametreler için hesapla
                    # Parametre ismini eşleştir (zaten eşleştirilmiş)
                    eslesen_param = param_adi  # Artık zaten eşleştirilmiş halde geliyor
                    
                    # Asgari fiyat tablosundan o yılın fiyatını bul
                    fiyat = 0
                    try:
                        target_key = _norm_param_key(eslesen_param)
                    except Exception:
                        target_key = (eslesen_param or '').strip().lower()
                    for asg_fiyat in asgari_fiyatlar:
                        try:
                            src_key = _norm_param_key(asg_fiyat.get('parametre', ''))
                        except Exception:
                            src_key = (asg_fiyat.get('parametre', '') or '').strip().lower()
                        if src_key and target_key and src_key == target_key:
                            yillik_fiyatlar = asg_fiyat.get('yillik', {})
                            try:
                                fiyat = float(yillik_fiyatlar.get(str(year), 0) or 0)
                            except Exception:
                                fiyat = 0
                            break
                    
                    # Tutar = Adet × Fiyat
                    parametre_tl[param_adi] = adet * fiyat
                    print(f"DEBUG PIVOT: {param_adi} -> {eslesen_param} -> {adet} adet × {fiyat} TL = {adet * fiyat} TL")

            toplam_parametre_tutari = 0.0
            try:
                for k, v in parametre_tl.items():
                    if _is_non_parametre_total_item(k):
                        continue
                    try:
                        toplam_parametre_tutari += float(v or 0)
                    except Exception:
                        pass
            except Exception:
                toplam_parametre_tutari = 0.0
            
            # Personel performansını hesapla (baca bilgileri kayıtlarından)
            personel_performans = defaultdict(int)
            personel_tutarlar = defaultdict(float)  # Personel başına toplam tutar
            personel_parametre_performans = defaultdict(int)  # Personel-Parametre kombinasyonu
            
            for baca_kayit in baca_bilgileri:
                created_at = baca_kayit.get('created_at', '')
                if created_at and created_at[:4] == str(year):
                    personel_adi = baca_kayit.get('personel_adi', '')
                    if personel_adi:
                        # Sadece ad kısmını al (boşluktan önceki kısım)
                        personel_adi_kisa = personel_adi.split(' ')[0] if personel_adi else ''
                        # Baca sayısını artır - kısa ad ile
                        personel_performans[personel_adi_kisa] += 1
                        
                        # Bu personelin girdiği baca için parametreleri bul ve fiyat hesapla
                        firma_adi = baca_kayit.get('firma_adi', '')
                        olcum_kodu = baca_kayit.get('olcum_kodu', '')
                        baca_adi = baca_kayit.get('baca_adi', '')
                        
                        # Dinamik olarak firma_olcum.json'dan parametreleri oku
                        baca_parametreleri = []
                        
                        # Firma ölçümlerinden bu baca için parametreleri bul
                        for olcum in firma_olcumler:
                            if (olcum.get('firma_adi') == firma_adi and 
                                olcum.get('olcum_kodu') == olcum_kodu and
                                olcum.get('baca_parametreleri')):
                                
                                baca_parametreleri_dict = olcum['baca_parametreleri']
                                if isinstance(baca_parametreleri_dict, dict) and baca_adi in baca_parametreleri_dict:
                                    baca_parametreleri_list = baca_parametreleri_dict[baca_adi]
                                    if isinstance(baca_parametreleri_list, list):
                                        # Parametreleri normalize et (büyük harf, trim)
                                        baca_parametreleri = [p.upper().strip() for p in baca_parametreleri_list if p and p.strip()]
                                        break
                        
                        # Eğer parametre bulunamadıysa, eski sabit kodlanmış listeyi kullan (fallback)
                        if not baca_parametreleri:
                            # HAFİZE A bacası için: TOZ, YG, VOC, AĞIR METAL
                            if firma_adi == "HAFİZE" and baca_adi == "a":
                                baca_parametreleri = ["TOZ", "YANMA GAZI", "VOC", "AĞIR METAL"]
                            # HAFİZE B bacası için: TOZ, YG, SÜLFÜRİK ASİT, PM10  
                            elif firma_adi == "HAFİZE" and baca_adi == "b":
                                baca_parametreleri = ["TOZ", "YANMA GAZI", "SÜLFÜRİK ASİT", "PM10"]
                            # AKARE ÇEVRE FÜZYON BACASI için: TOZ, NEM, YG (sadece bu 3 parametre!)
                            elif "AKARE ÇEVRE" in firma_adi and "FÜZYON" in baca_adi:
                                baca_parametreleri = ["TOZ", "NEM", "YANMA GAZI"]
                            # ATALAY TOKER adem bacası için: TOZ, YANMA GAZI, VOC, AĞIR METAL
                            elif firma_adi == "ATALAY TOKER" and baca_adi == "adem":
                                baca_parametreleri = ["TOZ", "YANMA GAZI", "VOC", "AĞIR METAL"]
                        
                        # Bu parametreler için o yılın asgari fiyatlarını hesapla
                        toplam_fiyat = 0
                        for parametre in baca_parametreleri:
                            fiyat = get_parametre_fiyati(parametre, year)
                            toplam_fiyat += fiyat
                        
                        # Debug log
                        print(f"Personel: {personel_adi} ({personel_adi_kisa}), Firma: {firma_adi}, Baca: {baca_adi}")
                        print(f"Bulunan parametreler: {baca_parametreleri}")
                        print(f"Toplam Fiyat: {toplam_fiyat}")
                        
                        # Kısa ad ile grupla
                        personel_tutarlar[personel_adi_kisa] += toplam_fiyat
                        
                        # Personel-Parametre kombinasyonu hesapla
                        for parametre in baca_parametreleri:
                            kombinasyon_adi = f"{parametre}-{personel_adi_kisa}"
                            personel_parametre_performans[kombinasyon_adi] += 1
            
            summary = {
                'toplam_baca_adedi': toplam_baca_adedi,
                'toplam_parametre_adedi': toplam_parametre_adedi,
                'toplam_parametre_tutari': round(float(toplam_parametre_tutari or 0), 2),
                'toplam_teklif_adedi': toplam_teklif_adedi,
                'kabul_adet': kabul_adet,
                'kapsam_ici_adet': kapsam_ici_adet,
                'kapsam_disi_adet': kapsam_disi_adet,
                'is_birligi_adet': is_birligi_adet,
                'toplam_teklif_tutari': round(toplam_teklif_tutari, 2),
                'kabul_tutari': round(kabul_tutari, 2),
                'kapsam_ici_tutar': round(kapsam_ici_tutar, 2),
                'kapsam_disi_tutar': round(kapsam_disi_tutar, 2),
                'is_birligi_tutar': round(is_birligi_tutar, 2),
                'parametre_sayilari': dict(parametre_sayilari),
                'parametre_sayilari_kaynak': 'TEKLIF' if use_teklif else 'FIRMA_OLCUM',
                'baca_sayisi_kaynak': baca_sayisi_kaynak,
                'personel_performans': dict(personel_performans),
                'personel_tutarlar': dict(personel_tutarlar),
                'personel_parametre_performans': dict(personel_parametre_performans)
            }
            
            result_data[str(year)] = {
                'summary': summary,
                'parametre_tl': parametre_tl
            }
        
        return jsonify({
            'success': True,
            'data': result_data
        })
        
    except Exception as e:
        print(f"Pivot compare API hatası: {e}")
        import traceback
        print(f"Hata detayı: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/pivot/export-xlsx')
def api_pivot_export_xlsx():
    """PIVOT verilerini XLSX formatında export et - sayılar sayı formatında, tarihler tarih formatında"""
    try:
        years = request.args.get('years', '')
        if not years:
            return jsonify({'success': False, 'error': 'Yıl parametresi gerekli'}), 400
        
        year_list = [int(y.strip()) for y in years.split(',') if y.strip()]
        if not year_list:
            return jsonify({'success': False, 'error': 'Geçerli yıl listesi gerekli'}), 400
        
        # Pivot verilerini al (mevcut API'yi kullan)
        import requests
        try:
            response = requests.get(f'http://localhost:{os.environ.get("PORT", 5001)}/api/pivot/compare?years={",".join(map(str, year_list))}')
            if response.status_code != 200:
                return jsonify({'success': False, 'error': 'Pivot verileri alınamadı'}), 500
            data = response.json()
            if not data.get('success'):
                return jsonify({'success': False, 'error': 'Pivot verileri başarısız'}), 500
        except Exception as e:
            # Direct API call başarısız olursa, verileri doğrudan hesapla
            baca_bilgileri = load_baca_bilgileri()
            parametre_olcumleri = load_parametre_olcum()
            teklifler = load_teklif()
            firma_olcumler = load_firma_olcum()
            asgari_fiyatlar = load_asgari_fiyatlar()
            
            # Basit veri yapısı oluştur
            data = {'success': True, 'data': {}}
            for year in year_list:
                data['data'][str(year)] = {
                    'summary': {
                        'toplam_teklif_adedi': 0,
                        'kabul_adet': 0,
                        'kapsam_ici_adet': 0,
                        'kapsam_disi_adet': 0,
                        'is_birligi_adet': 0,
                        'toplam_teklif_tutari': 0,
                        'kabul_tutari': 0,
                        'kapsam_ici_tutar': 0,
                        'kapsam_disi_tutar': 0,
                        'is_birligi_tutar': 0,
                        'toplam_baca_adedi': 0,
                        'toplam_parametre_adedi': 0,
                        'parametre_sayilari': {},
                        'personel_performans': {},
                        'personel_tutarlar': {},
                        'personel_parametre_performans': {}
                    }
                }
        
        # XLSX dosyası oluştur
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter
        from datetime import datetime
        
        wb = Workbook()
        ws = wb.active
        ws.title = "PIVOT Karşılaştırma"
        
        # Stil tanımlamaları
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        section_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Başlık satırı
        headers = ['METRİK']
        for year in year_list:
            headers.extend([f'{year} Adet', f'{year} TL'])
        
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border
        
        row_num = 2
        
        # TEKLİF TAKİP bölümü
        section_cell = ws.cell(row=row_num, column=1, value='TEKLİF TAKİP')
        section_cell.font = Font(bold=True)
        section_cell.fill = section_fill
        section_cell.border = border
        ws.merge_cells(start_row=row_num, start_column=1, end_row=row_num, end_column=len(headers))
        row_num += 1
        
        # Teklif satırları
        teklif_rows = [
            ('TOP. TEKLİF', lambda d: d.get('summary', {}).get('toplam_teklif_adedi', 0), lambda d: d.get('summary', {}).get('toplam_teklif_tutari', 0)),
            ('TOP KABUL OLAN TEK.', lambda d: d.get('summary', {}).get('kabul_adet', 0), lambda d: d.get('summary', {}).get('kabul_tutari', 0)),
            ('KAPSAM İÇİ TEKLİF', lambda d: d.get('summary', {}).get('kapsam_ici_adet', 0), lambda d: d.get('summary', {}).get('kapsam_ici_tutar', 0)),
            ('KAPSAM DIŞI TEKLİF', lambda d: d.get('summary', {}).get('kapsam_disi_adet', 0), lambda d: d.get('summary', {}).get('kapsam_disi_tutar', 0)),
            ('İŞ BİRLİĞİ TEKLİF', lambda d: d.get('summary', {}).get('is_birligi_adet', 0), lambda d: d.get('summary', {}).get('is_birligi_tutar', 0))
        ]
        
        for title, adet_func, tutar_func in teklif_rows:
            ws.cell(row=row_num, column=1, value=title).border = border
            
            for col, year in enumerate(year_list, 2):
                year_data = data['data'].get(str(year), {})
                
                # Adet sütunu
                adet_value = adet_func(year_data)
                adet_cell = ws.cell(row=row_num, column=col, value=adet_value)
                adet_cell.number_format = '0'  # Sayı formatı
                adet_cell.alignment = Alignment(horizontal='right')
                adet_cell.border = border
                
                # TL sütunu
                tutar_value = tutar_func(year_data)
                tutar_cell = ws.cell(row=row_num, column=col + len(year_list), value=tutar_value)
                tutar_cell.number_format = '#,##0.00'  # Para formatı
                tutar_cell.alignment = Alignment(horizontal='right')
                tutar_cell.border = border
            
            row_num += 1
        
        # PARAMETRE bölümü
        section_cell = ws.cell(row=row_num, column=1, value='PARAMETRE')
        section_cell.font = Font(bold=True)
        section_cell.fill = section_fill
        section_cell.border = border
        ws.merge_cells(start_row=row_num, start_column=1, end_row=row_num, end_column=len(headers))
        row_num += 1
        
        # Parametre satırları
        parametre_rows = [
            ('TOP. BACA SAYISI', lambda d: d.get('summary', {}).get('toplam_baca_adedi', 0), lambda d: d.get('summary', {}).get('kabul_tutari', 0)),
            ('TOP. PARAMETRE SAYISI', lambda d: d.get('summary', {}).get('toplam_parametre_adedi', 0), lambda d: 0)
        ]
        
        for title, adet_func, tutar_func in parametre_rows:
            ws.cell(row=row_num, column=1, value=title).border = border
            
            for col, year in enumerate(year_list, 2):
                year_data = data['data'].get(str(year), {})
                
                # Adet sütunu
                adet_value = adet_func(year_data)
                adet_cell = ws.cell(row=row_num, column=col, value=adet_value)
                adet_cell.number_format = '0'  # Sayı formatı
                adet_cell.alignment = Alignment(horizontal='right')
                adet_cell.border = border
                
                # TL sütunu
                tutar_value = tutar_func(year_data)
                tutar_cell = ws.cell(row=row_num, column=col + len(year_list), value=tutar_value)
                tutar_cell.number_format = '#,##0.00'  # Para formatı
                tutar_cell.alignment = Alignment(horizontal='right')
                tutar_cell.border = border
            
            row_num += 1
        
        # Dinamik parametre satırları
        all_params = set()
        for year in year_list:
            year_data = data['data'].get(str(year), {})
            params = year_data.get('summary', {}).get('parametre_sayilari', {})
            all_params.update(params.keys())
        
        # Asgari fiyat fallback
        asgari_fiyatlar_fallback = {
            'TOZ': 6655, 'YG': 3855, 'YANMA GAZI': 3855, 'AMET': 19000, 'AĞIR METAL': 19000,
            'SÜLF.A': 6290, 'SÜLFÜRİK ASİT': 6290, 'VOC': 13170, 'TOC': 13170, 'HIZ': 1430,
            'NEM': 1430, 'HF': 7505, 'HCL': 7505, 'AMONYAK': 6645, 'FORMALDEHİT': 6685,
            'CR+6': 6290, 'FOSFORİK ASİT': 6290, 'HCN': 6810, 'DİOKSİN FURAN': 83950,
            'PAH': 42220, 'PM10': 7370, 'ÇÖKEN TOZ': 6640, 'MODELLEME': 10560
        }
        
        for param in sorted(all_params):
            title = f'TOP. {param} SAYISI'
            ws.cell(row=row_num, column=1, value=title).border = border
            
            for col, year in enumerate(year_list, 2):
                year_data = data['data'].get(str(year), {})
                
                # Adet sütunu
                adet_value = year_data.get('summary', {}).get('parametre_sayilari', {}).get(param, 0)
                adet_cell = ws.cell(row=row_num, column=col, value=adet_value)
                adet_cell.number_format = '0'  # Sayı formatı
                adet_cell.alignment = Alignment(horizontal='right')
                adet_cell.border = border
                
                # TL sütunu (parametre sayısı × asgari fiyat)
                fiyat = asgari_fiyatlar_fallback.get(param, 0)
                tutar_value = adet_value * fiyat
                tutar_cell = ws.cell(row=row_num, column=col + len(year_list), value=tutar_value)
                tutar_cell.number_format = '#,##0.00'  # Para formatı
                tutar_cell.alignment = Alignment(horizontal='right')
                tutar_cell.border = border
            
            row_num += 1
        
        # PERSONEL PERFORMANS bölümü
        section_cell = ws.cell(row=row_num, column=1, value='PERSONEL PERFORMANS')
        section_cell.font = Font(bold=True)
        section_cell.fill = section_fill
        section_cell.border = border
        ws.merge_cells(start_row=row_num, start_column=1, end_row=row_num, end_column=len(headers))
        row_num += 1
        
        # Personel performans satırları
        all_personel = set()
        for year in year_list:
            year_data = data['data'].get(str(year), {})
            personel = year_data.get('summary', {}).get('personel_performans', {})
            all_personel.update(personel.keys())
        
        for personel in sorted(all_personel):
            title = f'{personel} BACA SAYISI'
            ws.cell(row=row_num, column=1, value=title).border = border
            
            for col, year in enumerate(year_list, 2):
                year_data = data['data'].get(str(year), {})
                
                # Adet sütunu
                adet_value = year_data.get('summary', {}).get('personel_performans', {}).get(personel, 0)
                adet_cell = ws.cell(row=row_num, column=col, value=adet_value)
                adet_cell.number_format = '0'  # Sayı formatı
                adet_cell.alignment = Alignment(horizontal='right')
                adet_cell.border = border
                
                # TL sütunu
                tutar_value = year_data.get('summary', {}).get('personel_tutarlar', {}).get(personel, 0)
                tutar_cell = ws.cell(row=row_num, column=col + len(year_list), value=tutar_value)
                tutar_cell.number_format = '#,##0.00'  # Para formatı
                tutar_cell.alignment = Alignment(horizontal='right')
                tutar_cell.border = border
            
            row_num += 1
        
        # Sütun genişliklerini ayarla
        ws.column_dimensions['A'].width = 30
        for col in range(2, len(headers) + 1):
            ws.column_dimensions[get_column_letter(col)].width = 15
        
        # Dosyayı memory'de oluştur
        from io import BytesIO
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        
        # Dosya adı oluştur
        filename = f'PIVOT_Karsilastirma_{"_".join(map(str, year_list))}.xlsx'
        
        return send_file(
            output,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        
    except Exception as e:
        print(f"Pivot XLSX export hatası: {e}")
        import traceback
        print(f"Hata detayı: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': str(e)}), 500

# Admin Backup/Restore endpoints
@app.route('/api/admin/backup', methods=['POST'])
def api_admin_backup():
    """Admin için veri yedekleme"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'error': 'Oturum açmanız gerekiyor'}), 401
    
    # Sadece admin kullanıcısı erişebilir
    if session.get('username') != 'admin':
        return jsonify({'success': False, 'error': 'Bu işlem için admin yetkisi gerekiyor'}), 403
    
    try:
        from datetime import datetime
        import shutil
        
        # Yedekleme dizini oluştur
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_dir = f'backups/backup_{timestamp}'
        os.makedirs(backup_dir, exist_ok=True)
        
        # Yedeklenecek dosyalar
        data_files = [
            'firma_kayit.json',
            'teklif.json', 
            'firma_olcum.json',
            'saha_olc.json',
            'parameters.json',
            'baca_bilgileri.json',
            'parametre_olcum.json',
            'parametre_sahabil.json',
            'asgari_fiyatlar.json',
            'forms.json',
            'users.json',
            'used_teklif_numbers.json',
            'par_saha_header_groups.json'
        ]
        
        backed_up_files = []
        for file in data_files:
            if os.path.exists(file):
                shutil.copy2(file, backup_dir)
                backed_up_files.append(file)
        
        # Eski yedekleri temizle (7 günden eski)
        cleanup_old_backups()
        
        return jsonify({
            'success': True, 
            'message': f'{len(backed_up_files)} dosya yedeklendi. Yedek klasörü: {backup_dir}',
            'backup_dir': backup_dir,
            'files': backed_up_files
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Yedekleme hatası: {str(e)}'}), 500

@app.route('/api/admin/restore', methods=['POST'])
def api_admin_restore():
    """Admin için veri geri yükleme"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'error': 'Oturum açmanız gerekiyor'}), 401
    
    # Sadece admin kullanıcısı erişebilir
    if session.get('username') != 'admin':
        return jsonify({'success': False, 'error': 'Bu işlem için admin yetkisi gerekiyor'}), 403
    
    try:
        import shutil
        from datetime import datetime
        
        # En son yedek klasörünü bul
        backup_base_dir = 'backups'
        if not os.path.exists(backup_base_dir):
            return jsonify({'success': False, 'error': 'Yedek klasörü bulunamadı'}), 404
        
        # Yedek klasörlerini listele
        backup_dirs = [d for d in os.listdir(backup_base_dir) if d.startswith('backup_')]
        if not backup_dirs:
            return jsonify({'success': False, 'error': 'Yedek dosyası bulunamadı'}), 404
        
        # En son yedek klasörünü seç
        latest_backup = sorted(backup_dirs)[-1]
        latest_backup_path = os.path.join(backup_base_dir, latest_backup)
        
        # Mevcut dosyaları yedekle (güvenlik için)
        current_backup_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        current_backup_dir = f'backups/current_backup_{current_backup_time}'
        os.makedirs(current_backup_dir, exist_ok=True)
        
        data_files = [
            'firma_kayit.json',
            'teklif.json', 
            'firma_olcum.json',
            'saha_olc.json',
            'parameters.json',
            'baca_bilgileri.json',
            'parametre_olcum.json',
            'parametre_sahabil.json',
            'asgari_fiyatlar.json',
            'forms.json',
            'users.json',
            'used_teklif_numbers.json',
            'par_saha_header_groups.json'
        ]
        
        # Mevcut dosyaları yedekle
        for file in data_files:
            if os.path.exists(file):
                shutil.copy2(file, current_backup_dir)
        
        # Yedek dosyalarını geri yükle
        restored_files = []
        for file in data_files:
            backup_file = os.path.join(latest_backup_path, file)
            if os.path.exists(backup_file):
                shutil.copy2(backup_file, file)
                restored_files.append(file)
        
        return jsonify({
            'success': True, 
            'message': f'{len(restored_files)} dosya geri yüklendi. Yedek: {latest_backup}',
            'restored_files': restored_files,
            'backup_used': latest_backup
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Geri yükleme hatası: {str(e)}'}), 500

def cleanup_old_backups():
    """7 günden eski yedekleri temizler"""
    import shutil
    backup_dir = 'backups'
    if not os.path.exists(backup_dir):
        return
    
    from datetime import datetime
    current_time = datetime.now()
    for item in os.listdir(backup_dir):
        item_path = os.path.join(backup_dir, item)
        if os.path.isdir(item_path):
            # Dizin oluşturma zamanını kontrol et
            creation_time = datetime.fromtimestamp(os.path.getctime(item_path))
            if (current_time - creation_time).days > 7:
                shutil.rmtree(item_path)
                print(f"🗑️ Eski yedek silindi: {item}")

@app.route('/api/admin/backup-download', methods=['POST'])
def api_admin_backup_download():
    """Admin için veri yedekleme - ZIP dosyası olarak indir"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'error': 'Oturum açmanız gerekiyor'}), 401
    
    # Sadece admin kullanıcısı erişebilir
    if session.get('username') != 'admin':
        return jsonify({'success': False, 'error': 'Bu işlem için admin yetkisi gerekiyor'}), 403
    
    try:
        import zipfile
        import io
        from datetime import datetime
        
        # Yedeklenecek dosyalar
        data_files = [
            'firma_kayit.json',
            'teklif.json', 
            'firma_olcum.json',
            'saha_olc.json',
            'parameters.json',
            'baca_bilgileri.json',
            'parametre_olcum.json',
            'parametre_sahabil.json',
            'asgari_fiyatlar.json',
            'forms.json',
            'users.json',
            'used_teklif_numbers.json',
            'par_saha_header_groups.json'
        ]
        
        # ZIP dosyası oluştur
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for file in data_files:
                if os.path.exists(file):
                    zip_file.write(file, file)
        
        zip_buffer.seek(0)
        
        # Dosya adı oluştur
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'emisyon_backup_{timestamp}.zip'
        
        return zip_buffer.getvalue(), 200, {
            'Content-Type': 'application/zip',
            'Content-Disposition': f'attachment; filename="{filename}"'
        }
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Yedekleme hatası: {str(e)}'}), 500

@app.route('/api/forms/genel-hukum-docx', methods=['GET'])
def api_forms_genel_hukum_docx():
    """GENEL_HUKUM.docx dosyasının içeriğini okuyup HTML formatında döndürür"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'error': 'Oturum açmanız gerekiyor'}), 401
    
    try:
        # GENEL_HUKUM.docx dosyasını oku
        genel_hukum_doc_path = os.path.join(app.root_path, 'static', 'images', 'GENEL_HUKUM.docx')
        if not os.path.exists(genel_hukum_doc_path):
            return jsonify({'success': False, 'error': 'GENEL_HUKUM.docx dosyası bulunamadı'}), 404
        
        from docx import Document as DocxDocument
        
        # Word belgesini aç ve içeriği al
        genel_hukum_doc = DocxDocument(genel_hukum_doc_path)
        
        # HTML içeriği oluştur
        html_content = ""
        
        for paragraph in genel_hukum_doc.paragraphs:
            if paragraph.text.strip():
                paragraph_text = paragraph.text.strip()
                
                # İlk paragraf "1. GENEL HÜKÜMLER" ise başlık olarak işle
                if 'GENEL HÜKÜMLER' in paragraph_text and paragraph_text.startswith('1.'):
                    html_content += f'<strong>{paragraph_text}</strong>\n'
                else:
                    # Normal madde olarak ekle
                    html_content += f'<li><strong>{paragraph_text}</strong></li>\n'
        
        # HTML'i <ul> tag'i ile sar
        if html_content:
            html_content = f'<ul>\n{html_content}</ul>'
        
        return jsonify({
            'success': True,
            'content': html_content,
            'message': 'GENEL_HUKUM.docx içeriği başarıyla okundu'
        })
        
    except Exception as e:
        print(f"GENEL_HUKUM.docx okuma hatası: {e}")
        return jsonify({'success': False, 'error': f'Dosya okuma hatası: {str(e)}'}), 500

@app.route('/api/admin/restore-upload', methods=['POST'])
def api_admin_restore_upload():
    """Admin için veri geri yükleme - Yüklenen dosyalardan"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'error': 'Oturum açmanız gerekiyor'}), 401
    
    # Sadece admin kullanıcısı erişebilir
    if session.get('username') != 'admin':
        return jsonify({'success': False, 'error': 'Bu işlem için admin yetkisi gerekiyor'}), 403
    
    try:
        import zipfile
        import io
        import shutil
        from datetime import datetime
        
        if 'files' not in request.files:
            return jsonify({'success': False, 'error': 'Dosya bulunamadı'}), 400
        
        files = request.files.getlist('files')
        if not files or files[0].filename == '':
            return jsonify({'success': False, 'error': 'Dosya seçilmedi'}), 400
        
        # Mevcut dosyaları yedekle (güvenlik için)
        current_backup_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        current_backup_dir = f'backups/current_backup_{current_backup_time}'
        os.makedirs(current_backup_dir, exist_ok=True)
        
        data_files = [
            'firma_kayit.json',
            'teklif.json', 
            'firma_olcum.json',
            'saha_olc.json',
            'parameters.json',
            'baca_bilgileri.json',
            'parametre_olcum.json',
            'parametre_sahabil.json',
            'asgari_fiyatlar.json',
            'forms.json',
            'users.json',
            'used_teklif_numbers.json',
            'par_saha_header_groups.json'
        ]
        
        # Mevcut dosyaları yedekle
        for file in data_files:
            if os.path.exists(file):
                shutil.copy2(file, current_backup_dir)
        
        restored_files = []
        
        # ZIP dosyası kontrolü
        for file in files:
            if file.filename.endswith('.zip'):
                # ZIP dosyasını çıkar
                zip_data = file.read()
                with zipfile.ZipFile(io.BytesIO(zip_data), 'r') as zip_file:
                    for zip_info in zip_file.infolist():
                        if zip_info.filename in data_files:
                            # Dosyayı çıkar ve kaydet
                            file_data = zip_file.read(zip_info.filename)
                            with open(zip_info.filename, 'wb') as f:
                                f.write(file_data)
                            restored_files.append(zip_info.filename)
            else:
                # Tekil JSON dosyası
                filename = file.filename
                if filename in data_files:
                    file.save(filename)
                    restored_files.append(filename)
        
        return jsonify({
            'success': True, 
            'message': f'{len(restored_files)} dosya geri yüklendi.',
            'restored_files': restored_files
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Geri yükleme hatası: {str(e)}'}), 500

if __name__ == '__main__':
    # Teklif numarası benzersizliği için migration çalıştır
    migrate_existing_teklif_numbers()
    
    # Eski teklif numaralarını yeni formata dönüştür (TE26-001 -> 2026/TE-001)
    convert_teklif_numbers_to_new_format()
    
    # Render için port ayarı (Render'ın verdiği PORT değişkenini kullan, yoksa 5001 kullan)
    port = int(os.environ.get('PORT', 5001))
    
    # Geliştirme modu kontrolü - her zaman debug=True
    debug = True
    
    # Uygulamayı başlat
    if __name__ == '__main__':
        app.run(
            host='0.0.0.0',
            port=port,
            debug=debug,
            threaded=True,
            use_reloader=True,
            use_debugger=True
        )