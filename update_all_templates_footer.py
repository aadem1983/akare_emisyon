"""TÃ¼m teklif ÅŸablonlarÄ±na footer ekler (TEKLÄ°F-1, TEKLÄ°F-2, TEKLÄ°F-3)"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def add_footer_to_template(template_path):
    """Åablona footer ekler"""
    try:
        doc = Document(template_path)
        
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            
            # Mevcut footer paragraflarÄ±nÄ± temizle
            for p in list(footer.paragraphs):
                p.clear()
            
            # Ä°lk paragraf varsa kullan, yoksa ekle
            if footer.paragraphs:
                p = footer.paragraphs[0]
            else:
                p = footer.add_paragraph()
            
            # Sol: SayÄ±
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run1 = p.add_run("SayÄ±:{{TEKLIF_NO}}")
            run1.font.bold = True
            run1.font.color.rgb = RGBColor(0, 0, 128)
            run1.font.size = Pt(8)
            
            # Orta: Firma bilgileri (yeni paragraf)
            p2 = footer.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run2a = p2.add_run("AKARE Ã‡EVRE LABORATUVAR VE DAN. HÄ°ZM. TÄ°C.LTD.ÅTÄ° KirazlÄ±yalÄ± Mah. SÃ¼leyman Demirel Cad. No:28/A")
            run2a.font.bold = True
            run2a.font.color.rgb = RGBColor(0, 0, 128)
            run2a.font.size = Pt(8)
            
            p2.add_run("\n")
            
            run2b = p2.add_run("KÃ¶rfez V.D 013 065 1290 KÃ¶rfez-KOCAELÄ°")
            run2b.font.bold = True
            run2b.font.color.rgb = RGBColor(0, 0, 128)
            run2b.font.size = Pt(8)
            
            p2.add_run("\n")
            
            run2c = p2.add_run("info@akarecevre.com  www.akarecevre.com")
            run2c.font.bold = True
            run2c.font.color.rgb = RGBColor(0, 0, 128)
            run2c.font.size = Pt(8)
            
            # SaÄŸ: Form kodu ve sayfa (yeni paragraf)
            p3 = footer.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            run3a = p3.add_run("AÃ‡.F.102/Rev04/14.08.2025  ")
            run3a.font.bold = True
            run3a.font.color.rgb = RGBColor(0, 0, 128)
            run3a.font.size = Pt(8)
            
            # Sayfa numarasÄ±
            run3b = p3.add_run()
            run3b.font.bold = True
            run3b.font.color.rgb = RGBColor(0, 0, 128)
            run3b.font.size = Pt(8)
            
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run3b._r.append(fldChar1)
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            run3b._r.append(instrText)
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run3b._r.append(fldChar2)
            
            run3b.add_text('/')
            
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'begin')
            run3b._r.append(fldChar3)
            instrText2 = OxmlElement('w:instrText')
            instrText2.text = 'NUMPAGES'
            run3b._r.append(instrText2)
            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')
            run3b._r.append(fldChar4)
        
        doc.save(template_path)
        print(f"âœ… Footer eklendi: {template_path}")
        return True
        
    except Exception as e:
        print(f"âŒ Hata ({template_path}): {e}")
        return False

if __name__ == '__main__':
    templates_dir = os.path.join('static', 'images')
    
    templates = [
        'TEKLÄ°F-1 GÄ°RÄ°Å.docx',
        'TEKLÄ°F - 2 FÄ°YAT.docx',
        'TEKLÄ°F - 3 GENEL HUKUM.docx'
    ]
    
    success_count = 0
    for template_name in templates:
        template_path = os.path.join(templates_dir, template_name)
        if os.path.exists(template_path):
            if add_footer_to_template(template_path):
                success_count += 1
        else:
            print(f"âš ï¸ Åablon bulunamadÄ±: {template_path}")
    
    print(f"\nâœ… {success_count}/{len(templates)} ÅŸablon gÃ¼ncellendi!")
    print("ğŸ“ Footer formatÄ±: 8pt, bold, koyu mavi")
    print("ğŸ“ Placeholder: {{TEKLIF_NO}}")
