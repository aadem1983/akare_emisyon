from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.shared import qn
from docx.shared import Inches, Pt, RGBColor
import os
import shutil


def _clear_header_footer(hf):
    try:
        for el in list(hf._element):
            hf._element.remove(el)
    except Exception:
        pass


def _add_page_fields(paragraph):
    run = paragraph.add_run('')

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    fldChar1.set(qn('w:dirty'), 'true')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

    run.add_text(' / ')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    fldChar3.set(qn('w:dirty'), 'true')
    run._r.append(fldChar3)

    instrText2 = OxmlElement('w:instrText')
    instrText2.text = 'NUMPAGES'
    run._r.append(instrText2)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar4)



def _set_run_size(paragraph, pt: float):
    for r in paragraph.runs:
        try:
            r.font.size = Pt(pt)
        except Exception:
            pass


def _set_run_blue(paragraph):
    try:
        for r in paragraph.runs:
            try:
                r.font.color.rgb = RGBColor(0, 0, 128)
            except Exception:
                pass
    except Exception:
        pass


def _set_paragraph_bottom_border(paragraph, color_hex: str = '000080', size: str = '12'):
    """Add a bottom border to a paragraph (acts as a horizontal line)."""
    try:
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is None:
            pBdr = OxmlElement('w:pBdr')
            pPr.append(pBdr)

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(size))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color_hex)
        pBdr.append(bottom)
    except Exception:
        pass



def apply_footer_layout(doc_path: str, out_path: str):
    doc = Document(doc_path)

    for section in doc.sections:
        for attr in ('footer', 'first_page_footer', 'even_page_footer'):
            hf = getattr(section, attr, None)
            if hf is None:
                continue
            try:
                hf.is_linked_to_previous = False
            except Exception:
                pass

            _clear_header_footer(hf)

            # 1) Blue horizontal line + right top form code line
            p_top = hf.add_paragraph('')
            try:
                p_top.paragraph_format.space_before = Pt(0)
                p_top.paragraph_format.space_after = Pt(2)
            except Exception:
                pass
            try:
                pf = p_top.paragraph_format
                try:
                    pf.tab_stops.clear_all()
                except Exception:
                    pass
                pf.tab_stops.add_tab_stop(Inches(7.5), alignment=WD_TAB_ALIGNMENT.RIGHT)
            except Exception:
                pass
            p_top.add_run('\t')
            p_top.add_run('{{FORM_KODU}}/Rev{{REVIZYON_NO}}/{{REVIZYON_TARIHI}}')
            _set_run_size(p_top, 8)
            _set_run_blue(p_top)
            _set_paragraph_bottom_border(p_top, '000080', '12')

            # 2) Middle company info (fixed)
            p_mid = hf.add_paragraph('')
            p_mid.text = (
                'AKARE ÇEVRE LABORATUVAR VE DAN. HİZM. TİC.LTD.ŞTİ\n'
                'Kirazlıyalı Mah. Süleyman Demirel Cad. No:28/A\n'
                'Körfez V.D 013 065 1290 Körfez-KOCAELİ\n'
                'info@akarecevre.com  www.akarecevre.com'
            )
            try:
                p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_mid.paragraph_format.space_before = Pt(0)
                p_mid.paragraph_format.space_after = Pt(0)
            except Exception:
                pass
            _set_run_size(p_mid, 8)
            _set_run_blue(p_mid)

            # 3) Bottom: left Sayı, right page x/y
            p_bot = hf.add_paragraph('')
            try:
                p_bot.paragraph_format.space_before = Pt(0)
                p_bot.paragraph_format.space_after = Pt(0)
            except Exception:
                pass
            try:
                pf = p_bot.paragraph_format
                try:
                    pf.tab_stops.clear_all()
                except Exception:
                    pass
                pf.tab_stops.add_tab_stop(Inches(7.5), alignment=WD_TAB_ALIGNMENT.RIGHT)
            except Exception:
                pass
            p_bot.add_run('Sayı:{{TEKLIF_NO}}')
            p_bot.add_run('\t')
            _add_page_fields(p_bot)
            _set_run_size(p_bot, 8)
            _set_run_blue(p_bot)

    doc.save(out_path)



def main():
    base = os.path.dirname(os.path.abspath(__file__))
    templates = [
        os.path.join(base, 'static', 'images', 'TEKLİF-1 GİRİŞ.docx'),
        os.path.join(base, 'static', 'images', 'TEKLİF - 2 FİYAT.docx'),
        os.path.join(base, 'static', 'images', 'TEKLİF - 3 GENEL HUKUM.docx'),
    ]

    for p in templates:
        if not os.path.exists(p):
            raise FileNotFoundError(p)

    backup_dir = os.path.join(base, 'static', 'images', 'template_backups')
    os.makedirs(backup_dir, exist_ok=True)

    images_dir = os.path.join(base, 'static', 'images')

    for p in templates:
        base_name = os.path.basename(p)
        dst_backup = os.path.join(backup_dir, base_name)
        shutil.copy2(p, dst_backup)

        root, ext = os.path.splitext(base_name)
        new_path = os.path.join(images_dir, f"{root}__NEW{ext}")
        apply_footer_layout(p, out_path=new_path)
        print(f"OK: wrote NEW template: {new_path}")


if __name__ == '__main__':
    main()
