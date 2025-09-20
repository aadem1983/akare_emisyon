from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
import os

def create_emisyon_olcum_bilgi_formu_template():
    """EMISYON_OLCUM_BILGI_FORMU.docx şablonunu oluşturur."""
    
    # Yeni doküman oluştur
    doc = Document()
    
    # Sayfa kenar boşluklarını ayarla (5mm sol ve sağ)
    section = doc.sections[0]
    section.left_margin = Inches(0.2)  # 5mm
    section.right_margin = Inches(0.2)  # 5mm
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    
    # Header oluştur
    header = section.header
    header_table = header.add_table(rows=1, cols=3, width=Inches(7.5))
    header_table.style = 'Table Grid'
    
    # Logo hücresi (sol)
    logo_cell = header_table.cell(0, 0)
    logo_cell.width = Inches(2.0)
    logo_cell.text = "AKare ÇEVRE LABORATUVARI"
    logo_paragraph = logo_cell.paragraphs[0]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_paragraph.runs[0]
    logo_run.font.size = Pt(10)
    logo_run.font.bold = True
    
    # Başlık hücresi (orta)
    baslik_cell = header_table.cell(0, 1)
    baslik_cell.width = Inches(3.5)
    baslik_cell.text = "EMİSYON ÖLÇÜM BİLGİ FORMU"
    baslik_paragraph = baslik_cell.paragraphs[0]
    baslik_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    baslik_run = baslik_paragraph.runs[0]
    baslik_run.font.size = Pt(16)
    baslik_run.font.bold = True
    
    # Form bilgileri hücresi (sağ)
    form_cell = header_table.cell(0, 2)
    form_cell.width = Inches(2.0)
    form_cell.text = "Form Kodu: AÇ.F.99\n"
    form_cell.text += "Yayın Tarihi: 01.08.2015\n"
    form_cell.text += "Revizyon No: 00\n"
    form_cell.text += "Revizyon Tarihi: 29.02.2024\n"
    form_cell.text += "Sayfa No: 1/1"
    
    # Header satır yüksekliğini ayarla
    for row in header_table.rows:
        row.height = Inches(0.4)
    
    # Ana içerik - Boşluk
    doc.add_paragraph()
    
    # Firma bilgileri tablosu - 6 satır x 4 sütun (2x2 düzen)
    firma_table = doc.add_table(rows=6, cols=4)
    firma_table.style = 'Table Grid'
    firma_table.autofit = False
    
    # Tablo genişliğini ayarla (satır yüksekliği 9.5 mm)
    for row in firma_table.rows:
        row.height = Inches(0.37)  # 9.5 mm ≈ 0.37 inç
    
    # Sütun genişliklerini ayarla
    firma_table.columns[0].width = Inches(2.0)  # Etiketler (sol)
    firma_table.columns[1].width = Inches(2.75)  # Değerler (sol)
    firma_table.columns[2].width = Inches(2.0)  # Etiketler (sağ)
    firma_table.columns[3].width = Inches(2.75)  # Değerler (sağ)
    
    # Tablo verilerini doldur (yer tutucular ile)
    firma_data = [
        ("Firma Adı", ": {{FIRMA_ADI}}"),
        ("Ölçüm Kodu", ": {{OLCUM_KODU}}"),
        ("Başl. Tar", ": {{BASLANGIC_TARIHI}}"),
        ("Bitiş Tarih", ": {{BITIS_TARIHI}}"),
        ("Baca Say", ": {{BACA_SAYISI}}"),
        ("Parametre", ": {{PARAMETRELER}}"),
        ("Per.", ": {{PERSONEL}}"),
        ("İl", ": {{IL}}"),
        ("İlçe/il", ": {{ILCE}}"),
        ("Yetkili", ": {{YETKILI}}"),
        ("Telefon", ": {{TELEFON}}"),
        ("Durum", ": {{DURUM}}")
    ]
    
    # Verileri tabloya ekle (2x2 düzen)
    for i, (label, value) in enumerate(firma_data):
        row = i // 2  # 2 sütunlu tablo için
        col = (i % 2) * 2  # 0 veya 2 (sol veya sağ sütun)
        
        if row < len(firma_table.rows):  # Tablo sınırları içinde
            # Etiket hücresi
            label_cell = firma_table.cell(row, col)
            label_cell.text = label
            label_paragraph = label_cell.paragraphs[0]
            label_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            label_run = label_paragraph.runs[0]
            label_run.font.bold = True
            label_run.font.size = Pt(11)
            
            # Değer hücresi
            value_cell = firma_table.cell(row, col + 1)
            value_cell.text = value
            value_paragraph = value_cell.paragraphs[0]
            value_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            value_run = value_paragraph.runs[0]
            value_run.font.size = Pt(11)
    
    # Tüm hücrelerde dikey hizalama
    for row in firma_table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Boşluk
    doc.add_paragraph()
    
    # Baca listesi tablosu için yer tutucu
    baca_baslik = doc.add_paragraph("BACA LİSTESİ VE PARAMETRE ÖLÇÜMLERİ")
    baca_baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
    baca_baslik.runs[0].font.bold = True
    baca_baslik.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    
    # Baca listesi tablosu için yer tutucu metin
    baca_placeholder = doc.add_paragraph("Buraya baca listesi tablosu eklenecek...")
    baca_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    baca_placeholder.runs[0].font.italic = True
    baca_placeholder.runs[0].font.size = Pt(10)
    
    # Dosyayı kaydet
    template_path = os.path.join('static', 'images', 'EMISYON_OLCUM_BILGI_FORMU.docx')
    doc.save(template_path)
    
    print(f"Şablon başarıyla oluşturuldu: {template_path}")
    print("Yer tutucular:")
    print("- {{FIRMA_ADI}}")
    print("- {{OLCUM_KODU}}")
    print("- {{BASLANGIC_TARIHI}}")
    print("- {{BITIS_TARIHI}}")
    print("- {{BACA_SAYISI}}")
    print("- {{PARAMETRELER}}")
    print("- {{PERSONEL}}")
    print("- {{IL}}")
    print("- {{ILCE}}")
    print("- {{YETKILI}}")
    print("- {{TELEFON}}")
    print("- {{DURUM}}")

if __name__ == "__main__":
    create_emisyon_olcum_bilgi_formu_template() 