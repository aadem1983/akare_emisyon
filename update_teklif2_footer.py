"""TEKLİF-2 FİYAT şablonuna footer ekler"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

template_path = os.path.join('static', 'images', 'TEKLİF - 2 FİYAT.docx')

try:
    doc = Document(template_path)
    
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # Mevcut footer paragraflarını temizle
        for p in list(footer.paragraphs):
            p.clear()
        
        # İlk paragraf varsa kullan, yoksa ekle
        if footer.paragraphs:
            p = footer.paragraphs[0]
        else:
            p = footer.add_paragraph()
        
        # Sol: Sayı
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run1 = p.add_run("Sayı:{{TEKLIF_NO}}")
        run1.font.bold = True
        run1.font.color.rgb = RGBColor(0, 0, 128)
        run1.font.size = Pt(8)
        
        # Orta: Firma bilgileri (yeni paragraf)
        p2 = footer.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run2a = p2.add_run("AKARE ÇEVRE LABORATUVAR VE DAN. HİZM. TİC.LTD.ŞTİ Kirazlıyalı Mah. Süleyman Demirel Cad. No:28/A")
        run2a.font.bold = True
        run2a.font.color.rgb = RGBColor(0, 0, 128)
        run2a.font.size = Pt(8)
        
        p2.add_run("\n")
        
        run2b = p2.add_run("Körfez V.D 013 065 1290 Körfez-KOCAELİ")
        run2b.font.bold = True
        run2b.font.color.rgb = RGBColor(0, 0, 128)
        run2b.font.size = Pt(8)
        
        p2.add_run("\n")
        
        run2c = p2.add_run("info@akarecevre.com  www.akarecevre.com")
        run2c.font.bold = True
        run2c.font.color.rgb = RGBColor(0, 0, 128)
        run2c.font.size = Pt(8)
        
        # Sağ: Form kodu ve sayfa (yeni paragraf)
        p3 = footer.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run3a = p3.add_run("AÇ.F.102/Rev04/14.08.2025  ")
        run3a.font.bold = True
        run3a.font.color.rgb = RGBColor(0, 0, 128)
        run3a.font.size = Pt(8)
        
        # Sayfa numarası
        run3b = p3.add_run()
        run3b.font.bold = True
        run3b.font.color.rgb = RGBColor(0, 0, 128)
        run3b.font.size = Pt(8)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run3b._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        run3b._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3b._r.append(fldChar2)
        
        run3b.add_text('/')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        run3b._r.append(fldChar3)
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = 'NUMPAGES'
        run3b._r.append(instrText2)
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run3b._r.append(fldChar4)
    
    doc.save(template_path)
    print(f"✅ Footer eklendi (8pt, bold, koyu mavi): {template_path}")
    
except Exception as e:
    print(f"❌ Hata: {e}")
    import traceback
    traceback.print_exc()
